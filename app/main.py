"""
Scout AgentOS
======

Production deployment entry point for Scout.

Run:
    python -m app.main
"""

from os import getenv
from pathlib import Path
from datetime import datetime
import json
import jwt
import bcrypt

from agno.os import AgentOS
from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse

from scout.agent import scout  # single agent (fallback)
from db import get_postgres_db
from app.model_config import (get_model, get_all_models, save_models, clear_cache as clear_model_cache,
    get_timezone, save_timezone, get_current_datetime, get_current_date, OPENROUTER_BASE_URL)
from db.connection import get_db_conn
from app.s3_storage import s3_upload_async, s3_delete_async, s3_download, s3_test, s3_sync_all, s3_list, is_s3_enabled, save_s3_config, _get_s3_config, _local_to_s3_key
from app.slot_contract import normalise_mapping, sanitise_mapping
from app import training_jobs

# ---------------------------------------------------------------------------
# Production-safe host configuration
# ---------------------------------------------------------------------------
API_HOST = getenv("API_HOST", "")
FRONTEND_HOST = getenv("FRONTEND_HOST", "")


# ---------------------------------------------------------------------------
# CORS Middleware for agent-ui
# ---------------------------------------------------------------------------
def add_cors_middleware(app: FastAPI):
    cors_origins = getenv("CORS_ORIGINS")

    if cors_origins and cors_origins != "*":
        allow_origins = [origin.strip() for origin in cors_origins.split(",")]
    elif cors_origins == "*":
        # Explicitly reject wildcard CORS in production
        import logging
        logging.getLogger("legalscout.security").warning("CORS_ORIGINS='*' is insecure — using same-origin only")
        allow_origins = []
    else:
        # Single-port architecture: same-origin requests don't need CORS
        frontend = getenv("FRONTEND_HOST", "")
        api = getenv("API_HOST", "")
        allow_origins = [o for o in [frontend, api] if o]

    app.add_middleware(
        CORSMiddleware,
        allow_origins=allow_origins,
        allow_credentials=True,
        allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
        allow_headers=["Authorization", "Content-Type", "Accept", "X-Requested-With"],
    )


# ---------------------------------------------------------------------------
# Create AgentOS
# ---------------------------------------------------------------------------
agent_os = AgentOS(
    name="Scout",
    agents=[scout],
    tracing=True,
    scheduler=True,
    db=get_postgres_db(),
    config=str(Path(__file__).parent / "config.yaml"),
)

app = agent_os.get_app()

# ---------------------------------------------------------------------------
# Override AgentOS default routes (/, /health) with our own
# AgentOS registers these automatically but we need custom behavior:
#   / → serve frontend instead of AgentOS JSON
#   /health → serve DB status instead of simple {"status": "ok"}
# ---------------------------------------------------------------------------
_agenos_override_paths = {"/", "/health"}
app.routes[:] = [r for r in app.routes if not (hasattr(r, "path") and r.path in _agenos_override_paths)]

# Mount frontend static assets early (before other routes can intercept)
#
# ★★★This mount is why the `immutable` branch in serve_frontend() below never
# fired for the assets it was written for: EVERY /_next/* request is answered
# here by StaticFiles and never reaches the catch-all. StaticFiles sends no
# Cache-Control at all, so the biggest, most cacheable files in the app —
# content-hashed JS and CSS — were re-downloaded on every hard load while an
# unhashed 1 MB pdf worker sat in the browser cache forever. The header belongs
# on the thing that actually serves the bytes.
class _HashedStatic(StaticFiles):
    """StaticFiles + immutable caching.

    Safe because Next.js puts a content hash in every filename under
    /_next/static — a changed file is a changed URL, so a cached copy can never
    be stale. Anything under /_next WITHOUT a hash would be wrong to freeze, so
    the header is applied by path, not blanket.
    """

    def file_response(self, *args, **kwargs):  # type: ignore[override]
        resp = super().file_response(*args, **kwargs)
        try:
            path = str(args[0]) if args else ""
            if "/static/" in path.replace("\\", "/"):
                resp.headers["Cache-Control"] = "public, max-age=31536000, immutable"
            else:
                resp.headers["Cache-Control"] = "no-cache"
        except Exception:  # noqa: BLE001 — never fail a static asset over a header
            pass
        return resp


_frontend_dir_early = Path("/app/static-frontend")
_next_dir_early = _frontend_dir_early / "_next"
if _next_dir_early.exists():
    app.mount("/_next", _HashedStatic(directory=str(_next_dir_early)), name="frontend-next")

# Add CORS after app is created
add_cors_middleware(app)


# ---------------------------------------------------------------------------
# Prometheus Metrics — /metrics endpoint for monitoring
# ---------------------------------------------------------------------------
from prometheus_fastapi_instrumentator import Instrumentator
from prometheus_client import Counter, Histogram

# Auto-instrument all HTTP requests (latency, status codes, in-progress)
_instrumentator = Instrumentator(
    should_group_status_codes=True,
    should_ignore_untemplated=True,
    excluded_handlers=["/health", "/metrics"],
)
_instrumentator.instrument(app).expose(app, endpoint="/metrics", include_in_schema=False)

# Custom business metrics
DOCUMENTS_GENERATED = Counter(
    "legalscout_documents_generated_total",
    "Total number of documents generated",
    ["template_type", "company"],
)
AGENT_CALLS = Counter(
    "legalscout_agent_calls_total",
    "Total number of AI agent invocations",
)
DOCUMENT_GENERATION_DURATION = Histogram(
    "legalscout_document_generation_seconds",
    "Time spent generating a document",
    buckets=[0.5, 1, 2, 5, 10, 30, 60],
)


# ---------------------------------------------------------------------------
# Graceful Shutdown
# ---------------------------------------------------------------------------
import signal


def _handle_shutdown(signum, frame):
    """Log shutdown signal for observability."""
    import logging
    sig_name = signal.Signals(signum).name
    logging.getLogger("legalscout").info(f"Received {sig_name} — shutting down gracefully")


signal.signal(signal.SIGTERM, _handle_shutdown)
signal.signal(signal.SIGINT, _handle_shutdown)


# ---------------------------------------------------------------------------
# Security Headers Middleware
# ---------------------------------------------------------------------------
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request as StarletteRequest
from starlette.responses import Response as StarletteResponse


class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: StarletteRequest, call_next):
        # Reject oversized chat requests (backend enforcement of input limit)
        #
        # The 50KB limit is about what a PERSON can type. It must not be applied
        # to /runs/{id}/continue, whose body is not a message at all: resuming a
        # paused run posts the entire tool payload back, and that payload is
        # generated by us — picker candidates, preview results, document state.
        # Measured on an ordinary Corporate Shareholder Consent flow, a resume
        # body runs 11KB-25KB, and a longer run crossed 50,000 and came back
        # "Message too large (max 50KB)". To the user that is a card that
        # silently refuses to submit, on a document they have already answered
        # every question for.
        #
        # So the resume path keeps a bound — an unbounded body is still a denial
        # of service — but one sized for a tool payload rather than a sentence.
        if request.method == "POST" and "/agents/" in str(request.url):
            is_resume = str(request.url.path).rstrip("/").endswith("/continue")
            limit = 2_000_000 if is_resume else 50_000
            content_length = request.headers.get("content-length")
            if content_length and int(content_length) > limit:
                return JSONResponse(
                    {
                        "error": (
                            "Resume payload too large (max 2MB)"
                            if is_resume
                            else "Message too large (max 50KB)"
                        )
                    },
                    status_code=413,
                )
        # ---- Bind the conversation for the whole resume request ------------
        #
        # A `requires_user_input=True` tool CANNOT read its session: agno builds
        # `user_input_schema` from `sig.parameters` with no exclusions, so
        # declaring `run_context` puts it in the schema, the frontend echoes the
        # schema back on resume, and the entrypoint then receives it twice
        # (TypeError) — the card never executes. See people_picker._session_of.
        #
        # The pickers work around it by having their `lookup_*` companion publish
        # the session id inside the payload. `ask_questions` has no companion, so
        # its answers had nowhere to be stored and survived only if the model
        # chose to forward them into the next call's custom_data. Measured on the
        # client's tracker: the same AGM Minutes template kept the meeting date
        # on one run and dropped it on the next, purely on that choice.
        #
        # So the session is bound HERE instead, from the request that carries it.
        # Everything downstream reads it through slot_resolver.current_session_scope().
        #
        # Body is read only for the resume path, which is already size-capped
        # above, and is replayed to the app afterwards — consuming it without
        # putting it back would leave agno with an empty form.
        if request.method == "POST" and str(request.url.path).rstrip("/").endswith("/continue"):
            try:
                raw = await request.body()

                async def _replay() -> dict:
                    return {"type": "http.request", "body": raw, "more_body": False}

                request._receive = _replay  # noqa: SLF001 — the standard replay
                sid = ""
                for part in raw.split(b"&"):
                    if part.startswith(b"session_id="):
                        from urllib.parse import unquote_plus

                        sid = unquote_plus(part.split(b"=", 1)[1].decode("utf-8", "replace"))
                        break
                if sid:
                    from scout.tools.slot_resolver import session_scope

                    with session_scope(sid):
                        return await self._finish(request, call_next)
            except Exception as e:  # noqa: BLE001 — never fail a resume for this
                logger.warning(f"[SCOPE] could not bind session on resume: {e}")

        return await self._finish(request, call_next)

    async def _finish(self, request: StarletteRequest, call_next):
        response: StarletteResponse = await call_next(request)

        # GET /agents ships the agent's ENTIRE system prompt to the browser.
        # Measured: 125,390 bytes total, of which system_message is 105,113 and
        # tools another 22,252. Nothing in this UI reads either field (grepped:
        # zero references to `system_message` anywhere in agent-ui/src) — and
        # until the fix in Sessions.tsx the sidebar's history fetch was queued
        # BEHIND this download, so every page load waited on 125KB it had no
        # use for.
        #
        # It is also the whole prompt, handed to any logged-in browser.
        #
        # The route belongs to agno (mounted by agent_os.get_app()), so it is
        # trimmed here on the way out rather than by redefining it. Only the
        # two unused keys are dropped; every field the UI reads — id, name,
        # db_id, model, sessions, memory, knowledge — is untouched. Anything
        # unexpected (wrong status, wrong content type, unparsable body) is
        # passed through exactly as agno produced it.
        if (
            request.method == "GET"
            and request.url.path.rstrip("/") == "/agents"
            and response.status_code == 200
            and response.headers.get("content-type", "").startswith("application/json")
        ):
            body = b""
            async for chunk in response.body_iterator:
                body += chunk
            try:
                payload = json.loads(body)
                if isinstance(payload, list):
                    for entry in payload:
                        if isinstance(entry, dict):
                            entry.pop("system_message", None)
                            entry.pop("tools", None)
                    body = json.dumps(payload).encode()
            except Exception as e:  # noqa: BLE001 — never break /agents for a trim
                logger.warning(f"[AGENTS] trim skipped: {e}")
            headers = dict(response.headers)
            headers.pop("content-length", None)
            # StarletteResponse, not `Response` — the bare name is NOT bound at
            # module scope in this file (line 184 imports it aliased). A bare
            # `Response` here would be a NameError on every /agents request.
            return StarletteResponse(
                content=body,
                status_code=response.status_code,
                headers=headers,
                media_type=response.media_type,
            )

        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "SAMEORIGIN"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        return response


app.add_middleware(SecurityHeadersMiddleware)


# ---------------------------------------------------------------------------
# Startup: Auto-sync templates + cleanup stale knowledge
# ---------------------------------------------------------------------------
def startup_sync():
    """Validate config, ensure directories, rebuild agent knowledge, resume training.

    ★ NOT registered with @app.on_event("startup"). `app = agent_os.get_app()`
    installs an Agno lifespan, and Starlette SILENTLY IGNORES every on_event
    handler once a lifespan exists — so this function never ran. Directory
    creation, the migration check, the template-knowledge and legal-skills
    refreshes and the training-job resume were all dead on arrival.

    It is invoked at module-import time instead (bottom of this file), the same
    pattern training_jobs.py already uses for its watchdog. The body is
    deliberately synchronous and every step is individually guarded, because an
    import-time failure here would take the whole app down.
    """
    # Ensure all required document directories exist (defense-in-depth beyond Docker)
    import logging as _log
    for _d in [
        "/documents/legal/templates", "/documents/legal/data",
        "/documents/legal/output", "/documents/legal/uploads",
        "/documents/legal/previews", "/documents/legal/knowledge",
        "/documents/legal/extracts",
    ]:
        Path(_d).mkdir(parents=True, exist_ok=True)

    # Version
    from pathlib import Path as _P
    _vf = _P("/app/VERSION")
    _ver = _vf.read_text().strip() if _vf.exists() else "unknown"
    print(f"[STARTUP] Legal Scout v{_ver}")

    # Validate critical env vars
    _api_key = getenv("OPENROUTER_API_KEY", "")
    if not _api_key:
        print("[STARTUP] WARNING: OPENROUTER_API_KEY not set — chat and training will not work!")
    else:
        print(f"[STARTUP] OpenRouter API key: ...{_api_key[-4:]}")

    _jwt = getenv("JWT_SECRET_KEY", "")
    if not _jwt:
        print("[STARTUP] WARNING: JWT_SECRET_KEY not set — authentication will not work!")

    # Check migration status
    try:
        from db.connection import get_db_conn as _sconn
        _sc = _sconn()
        _scur = _sc.cursor()
        _scur.execute("SELECT COUNT(*) FROM schema_migrations")
        _applied = _scur.fetchone()[0]
        _scur.close(); _sc.close()

        import glob
        _total = len(glob.glob("/app/db/migration_*.sql"))
        if _applied < _total:
            print(f"[STARTUP] WARNING: {_total - _applied} pending migration(s) — run 'python -m db.migrate'")
        else:
            print(f"[STARTUP] Migrations: {_applied} applied (up to date)")
    except Exception as e:
        _log.getLogger("legalscout").warning(f"Migration check failed: {e}")

    # Rebuild agent template knowledge from DB
    try:
        _refresh_agent_knowledge()
    except Exception as e:
        print(f"[STARTUP] Agent refresh warning: {e}")

    # Rebuild agent legal-skills L1 block from DB
    try:
        _refresh_legal_skills()
    except Exception as e:
        print(f"[STARTUP] Legal skills refresh warning: {e}")

    # Resume any server-side template-training job left running/queued.
    # The advisory lock inside training_jobs guarantees only ONE of the 2
    # uvicorn workers actually runs it; never block startup on failure.
    try:
        training_jobs.resume_incomplete()
    except Exception as e:
        print(f"[STARTUP] Training-job resume warning: {e}")


def _refresh_agent_knowledge():
    """Reload template knowledge into the agent's system prompt."""
    import scout.agent as _am
    _am.TEMPLATE_KNOWLEDGE = _am._build_template_knowledge()
    # Rebuild INSTRUCTIONS f-string with updated TEMPLATE_KNOWLEDGE
    # The INSTRUCTIONS template uses {TEMPLATE_KNOWLEDGE} which was captured at import
    # We need to re-evaluate it by rebuilding from scratch
    _old = _am.INSTRUCTIONS
    _marker = "## Your Template Knowledge (auto-loaded from database)"
    _end_marker = "\n═══"
    if _marker in _old:
        _before = _old[:_old.index(_marker)]
        _after_idx = _old.index(_end_marker, _old.index(_marker))
        _after = _old[_after_idx:]
        _am.INSTRUCTIONS = _before + _marker + "\n" + _am.TEMPLATE_KNOWLEDGE + "\n" + _after
    _am.scout.instructions = _am.INSTRUCTIONS
    print(f"[AGENT] Template knowledge refreshed ({len(_am.TEMPLATE_KNOWLEDGE)} chars)")


def _refresh_legal_skills():
    """Reload the L1 legal-skills block into the agent's system prompt.

    Mirrors _refresh_agent_knowledge but splices between the legal-skills
    markers ("## Legal Skills (playbooks — load on demand)" .. "\n■■■") so it
    never overlaps the template-knowledge span (marker .. "\n═══").
    """
    import scout.agent as _am
    _am.LEGAL_SKILLS_BLOCK = _am._build_legal_skills_block()
    _old = _am.INSTRUCTIONS
    _marker = "## Legal Skills (playbooks — load on demand)"
    _end_marker = "\n■■■"
    if _marker in _old and _end_marker in _old:
        _before = _old[:_old.index(_marker)]
        _after_idx = _old.index(_end_marker, _old.index(_marker)) + len(_end_marker)
        _after = _old[_after_idx:]
        # _build_legal_skills_block() already includes the trailing end marker.
        _am.INSTRUCTIONS = _before + _am.LEGAL_SKILLS_BLOCK + _after
        _am.scout.instructions = _am.INSTRUCTIONS
    print(f"[AGENT] Legal skills refreshed ({len(_am.LEGAL_SKILLS_BLOCK)} chars)")


# ---------------------------------------------------------------------------
# Health Check Endpoint — for load balancers and monitoring
# ---------------------------------------------------------------------------
import time as _startup_time

_APP_START_TIME = _startup_time.time()


def _parse_follow_ups(content: str) -> list:
    """Pull the suggestion list out of whatever shape the model returned.

    The old code did `json.loads(content.strip().strip("`"))` and required a
    BARE array. Anything else — an object, a ```json fence, a sentence before
    the JSON — raised, and the caller's bare `except` turned that into an empty
    list with no trace. That is how follow-ups sat dead behind their hardcoded
    fallback after a model change, which is recorded in the comment below.

    So: accept the object we ask for, accept a bare array, and accept either one
    embedded in prose. Return [] only when there is genuinely nothing usable.
    """
    # `re` is NOT bound at module scope in this file — line 630 imports it as
    # `_re`. Referencing a bare `re` here would raise NameError, and the
    # caller's bare `except` would turn that into an empty list with no trace:
    # the very failure this function exists to end.
    import re

    text = (content or "").strip()
    if not text:
        return []
    # Strip a ```json ... ``` fence if the model added one.
    fence = re.match(r"^```[a-zA-Z]*\s*(.*?)\s*```$", text, re.S)
    if fence:
        text = fence.group(1).strip()

    def _clean(seq) -> list:
        out = []
        for item in seq if isinstance(seq, list) else []:
            if isinstance(item, str) and item.strip():
                out.append(item.strip())
            elif isinstance(item, dict):
                # Some models answer [{"question": "..."}] despite the schema.
                for key in ("question", "text", "follow_up", "suggestion"):
                    v = item.get(key)
                    if isinstance(v, str) and v.strip():
                        out.append(v.strip())
                        break
        return out

    candidates = [text]
    # Last resort: the first {...} or [...] found anywhere in the reply.
    for pattern in (r"\{.*\}", r"\[.*\]"):
        m = re.search(pattern, text, re.S)
        if m:
            candidates.append(m.group(0))

    for raw in candidates:
        try:
            parsed = json.loads(raw)
        except (ValueError, TypeError):
            continue
        if isinstance(parsed, dict):
            for key in ("follow_ups", "suggestions", "questions"):
                cleaned = _clean(parsed.get(key))
                if cleaned:
                    return cleaned
        cleaned = _clean(parsed)
        if cleaned:
            return cleaned
    return []


@app.post("/api/session-title")
async def set_session_title(request: Request):
    """Give a chat a short title, replacing the raw first message.

    Agno names a session after whatever the user typed first, verbatim. That
    makes the sidebar a wall of half-sentences — and, as the comment on
    AGENTOS_PROTECTED_ROOTS notes, a session title here is real company and
    director names.

    The title is written to `session_data.session_name`, which is what
    GET /sessions reads (verified against the live table: setting that key
    changes the API response). Ownership is checked against agno_sessions.user_id
    so one user cannot rename — or probe the existence of — another's chat.
    """
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")

    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Body must be JSON")

    session_id = str(body.get("session_id") or "").strip()
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")

    title = " ".join(str(body.get("title") or "").split())
    if not title:
        raise HTTPException(status_code=400, detail="title is required")
    # Long enough for "Shareholder Consent · Golden Delta Agri Limited", short
    # enough that no caller can push a document into the sidebar.
    title = title[:120]

    uid = str(user.get("user_id") or user.get("id") or "")
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        # Ownership and the write in ONE statement: a check-then-update would
        # race, and a row whose user_id is NULL (pre-migration-027 sessions)
        # must not be claimable by whoever asks first.
        cur.execute(
            """
            UPDATE ai.agno_sessions
               SET session_data = COALESCE(session_data, '{}'::jsonb)
                                  || jsonb_build_object('session_name', %s::text)
             WHERE session_id = %s
               AND user_id = %s
            """,
            (title, session_id, uid),
        )
        updated = cur.rowcount
        conn.commit()
        cur.close()
    except Exception as e:
        logger.warning(f"[TITLE] rename failed for {session_id}: {e}")
        raise HTTPException(status_code=500, detail="Could not set title")
    finally:
        if conn:
            conn.close()

    if not updated:
        # Same response whether the session belongs to someone else or does not
        # exist — telling them apart is an existence oracle over other users'
        # chats.
        raise HTTPException(status_code=404, detail="Session not found")
    return {"success": True, "title": title}


@app.post("/api/suggest-followups")
async def suggest_followups(request: Request):
    """LLM-powered follow-up suggestions based on conversation context."""
    try:
        body = await request.json()
        question = body.get("question", "")
        answer = body.get("answer", "")
        if not question or not answer:
            return {"suggestions": []}

        api_key = getenv("OPENROUTER_API_KEY", "")
        if not api_key:
            return {"suggestions": []}

        import httpx
        # Shape and wording follow Open WebUI's task-model template, which has
        # three rules this prompt used to lack, each of which produced a wrong
        # suggestion here:
        #   * "from the user's point of view" — without it the model writes the
        #     ASSISTANT's next line ("I can prepare that for you") as a
        #     suggestion the user is then asked to click.
        #   * "do not repeat" — this endpoint only ever sees the last exchange,
        #     so it will happily re-suggest something already answered.
        #   * primary language — this is a Myanmar product. A Burmese question
        #     must not get English follow-ups.
        prompt = f"""Suggest 3 natural follow-up questions the user might ask next.

Q: {question}
A: {answer[:500]}

Context: a legal document automation system for Myanmar corporate law — AGM
minutes, director consents, shareholder resolutions, company registers.

Rules:
- Write every question from the USER's point of view, as something they would
  type. Never write what the assistant would say or offer.
- Do not suggest anything the answer above already covers, and do not repeat
  the user's question back to them.
- Keep each one short — one line, no more than about 12 words.
- Use the conversation's primary language. If the question above is in Burmese,
  reply in Burmese. Default to English if it is mixed or unclear.

Respond with ONLY this JSON object and nothing else:
{{"follow_ups": ["Question 1?", "Question 2?", "Question 3?"]}}"""

        resp = httpx.post(
            f"{OPENROUTER_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            # max_tokens must cover the model's REASONING as well as its answer.
            # gemini-3.6-flash cannot have reasoning disabled ("Reasoning is
            # mandatory for this endpoint"), and `exclude` only hides it from the
            # response — the tokens are still spent. At 200 the reasoning ate the
            # whole budget, content came back as `[\n  "`, json.loads threw, and
            # the bare except below returned an empty list. Follow-ups were dead
            # from the moment the chat model changed, silently. 800 leaves room.
            json={"model": get_model("task"), "messages": [{"role": "user", "content": prompt}], "max_tokens": 800, "temperature": 0.3},
            timeout=10,
        )
        resp.raise_for_status()
        result = resp.json()
        content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
        suggestions = _parse_follow_ups(content)
        return {"suggestions": suggestions[:3]}
    except Exception as e:
        # Logged, not swallowed: this returning [] silently is exactly why the
        # follow-ups sat on their hardcoded fallback for so long without anyone
        # being able to see why.
        logger.warning(f"[FOLLOWUPS] suggestion generation failed: {e}")
        return {"suggestions": []}


@app.get("/health")
async def health_check():
    """Health check with full dependency status."""
    from psycopg import OperationalError
    from pathlib import Path

    checks = {}

    # 1. Database
    db_latency_ms = None
    try:
        start = _startup_time.time()
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT 1")
        cur.close(); conn.close()
        db_latency_ms = round((_startup_time.time() - start) * 1000, 1)
        checks["database"] = {"status": "connected", "latency_ms": db_latency_ms}
    except Exception:
        checks["database"] = {"status": "disconnected", "latency_ms": None}

    # 2. API key
    api_key = getenv("OPENROUTER_API_KEY", "")
    checks["api_key"] = {"status": "set" if api_key else "missing"}

    # 3. Documents directory
    docs_dir = Path("/documents/legal/templates")
    checks["documents"] = {"status": "writable" if docs_dir.exists() else "missing"}

    # 4. Templates loaded
    try:
        import scout.agent as _ag
        tpl_count = _ag.TEMPLATE_KNOWLEDGE.count("**") // 2 if _ag.TEMPLATE_KNOWLEDGE else 0
        checks["templates"] = {"status": "loaded", "count": tpl_count}
    except Exception:
        checks["templates"] = {"status": "not loaded", "count": 0}

    uptime_seconds = round(_startup_time.time() - _APP_START_TIME)
    is_healthy = checks["database"]["status"] == "connected"

    from fastapi.responses import JSONResponse
    return JSONResponse(
        status_code=200 if is_healthy else 503,
        content={
            "status": "healthy" if is_healthy else "degraded",
            "uptime_seconds": uptime_seconds,
            "environment": "production",
            "checks": checks,
            # Backward compatible
            "database": checks["database"],
        },
    )


# ---------------------------------------------------------------------------
# Startup Security Checks — reject weak defaults in production
# ---------------------------------------------------------------------------
import logging as _logging

_logger = _logging.getLogger("legalscout.security")

_JWT_SECRET = getenv("JWT_SECRET_KEY", "")
_ADMIN_PASS = getenv("ADMIN_PASSWORD", "")

_WEAK_JWT_SECRETS = {"legal-scout-default-secret", "secret", "changeme", "dev-only-change-in-production", "dev-secret-not-for-production", ""}
_WEAK_PASSWORDS = {"admin123", "password", "123456", "admin", "change-me-12chars", ""}

_security_warnings = []
if _JWT_SECRET in _WEAK_JWT_SECRETS or len(_JWT_SECRET) < 32:
    _security_warnings.append("JWT_SECRET_KEY is weak. Generate with: openssl rand -hex 32")
if _ADMIN_PASS in _WEAK_PASSWORDS or len(_ADMIN_PASS) < 10:
    _security_warnings.append("ADMIN_PASSWORD is weak. Use 10+ characters.")
if not getenv("OPENROUTER_API_KEY"):
    _security_warnings.append("OPENROUTER_API_KEY is not set — chat will not work.")

for w in _security_warnings:
    _logger.warning(f"[SECURITY] {w}")

# Block startup with weak auth secrets
_critical = [w for w in _security_warnings if "JWT_SECRET_KEY" in w or "ADMIN_PASSWORD" in w]
if _critical:
    print(f"\n[SECURITY] FATAL: {len(_critical)} security issue(s):")
    for c in _critical:
        print(f"  - {c}")
    print("Fix these in .env before deploying. Exiting.")
    import sys; sys.exit(1)


# ---------------------------------------------------------------------------
# Rate Limiting — prevent brute force and API abuse
# ---------------------------------------------------------------------------
import time as _time
from collections import defaultdict

_rate_limits: dict = defaultdict(list)  # IP -> [timestamps]
RATE_LIMIT_WINDOW = 60  # seconds
RATE_LIMIT_LOGIN = 5    # max login attempts per minute
RATE_LIMIT_API = 120    # max API calls per minute

def _check_rate_limit(ip: str, limit: int) -> bool:
    """Returns True if rate limit exceeded."""
    now = _time.time()
    key = f"{ip}:{limit}"
    _rate_limits[key] = [t for t in _rate_limits[key] if now - t < RATE_LIMIT_WINDOW]
    if len(_rate_limits[key]) >= limit:
        return True
    _rate_limits[key].append(now)
    return False


# ---------------------------------------------------------------------------
# Simple In-Memory Cache
# ---------------------------------------------------------------------------
_cache: dict = {}
_cache_ttl: dict = {}
CACHE_TTL_SECONDS = 30  # Cache for 30 seconds

def cached_response(key: str):
    """Get cached response if still valid."""
    if key in _cache and key in _cache_ttl:
        if _time.time() - _cache_ttl[key] < CACHE_TTL_SECONDS:
            return _cache[key]
    return None

def set_cache(key: str, value):
    """Set cache value."""
    _cache[key] = value
    _cache_ttl[key] = _time.time()

def clear_cache(prefix: str = ""):
    """Clear cache entries matching prefix."""
    keys_to_remove = [k for k in _cache if k.startswith(prefix)] if prefix else list(_cache.keys())
    for k in keys_to_remove:
        _cache.pop(k, None)
        _cache_ttl.pop(k, None)


# ---------------------------------------------------------------------------
# Input Validation Helpers
# ---------------------------------------------------------------------------
import re as _re

def sanitize_string(value: str, max_length: int = 500) -> str:
    """Sanitize string input — strip, truncate, remove dangerous chars."""
    if not isinstance(value, str):
        return str(value)[:max_length] if value else ""
    return value.strip()[:max_length]

def validate_email(email: str) -> bool:
    """Basic email validation."""
    return bool(_re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email))

def validate_filename(name: str) -> bool:
    """Validate filename — no path traversal, no absolute paths, safe chars only."""
    if not name or not isinstance(name, str):
        return False
    # Block path traversal and absolute paths
    if '..' in name or '/' in name or '\\' in name or name.startswith(('.', '~')):
        return False
    # Only allow safe characters: letters, numbers, spaces, hyphens, underscores, dots, parentheses
    return bool(_re.match(r'^[\w\s\-\.\(\)]+$', name))


def _highlight_placeholders_in_docx(source_path, dest_path):
    """Create a copy of DOCX with yellow-highlighted placeholders.
    Handles placeholders split across multiple Word runs by merging first."""
    import re as _hlre
    from docx import Document as _HlDoc
    from docx.oxml.ns import qn as _hlqn
    from docx.oxml import OxmlElement as _HlEl
    import copy as _hlcopy

    doc = _HlDoc(str(source_path))
    patt = _hlre.compile(r'\{\{[^}]+\}\}|\[[^\]]+_[^\]]+\]|\{[^}]+\}')

    def _add_yellow(run):
        rPr = run._r.get_or_add_rPr()
        for old in rPr.findall(_hlqn('w:highlight')):
            rPr.remove(old)
        hl = _HlEl('w:highlight')
        hl.set(_hlqn('w:val'), 'yellow')
        rPr.append(hl)

    def _merge_and_highlight(paragraph):
        """Merge split runs that form a placeholder, then highlight."""
        runs = paragraph.runs
        if not runs:
            return

        # Build full text and map char positions to runs
        full_text = ''.join(r.text or '' for r in runs)
        if not patt.search(full_text):
            return

        # Find all placeholder spans in the full text
        matches = list(patt.finditer(full_text))
        if not matches:
            return

        # Build char→run index mapping
        char_to_run = []
        for ri, run in enumerate(runs):
            for _ in (run.text or ''):
                char_to_run.append(ri)

        # For each match, find which runs it spans
        runs_to_highlight = set()
        for m in matches:
            for ci in range(m.start(), m.end()):
                if ci < len(char_to_run):
                    runs_to_highlight.add(char_to_run[ci])

        # Highlight those runs
        for ri in runs_to_highlight:
            if ri < len(runs):
                _add_yellow(runs[ri])

    for para in doc.paragraphs:
        _merge_and_highlight(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _merge_and_highlight(para)

    doc.save(str(dest_path))


# ---------------------------------------------------------------------------
# Structured Logging
# ---------------------------------------------------------------------------
import logging
import traceback
from app.logging_config import setup_logging, get_logger

logger = setup_logging()


def log_error(context: str, error: Exception):
    """Log error with context and traceback."""
    logger.error(f"[{context}] {type(error).__name__}: {error}")
    logger.debug(traceback.format_exc())


# ---------------------------------------------------------------------------
# Authentication Middleware — protects ALL /api/ endpoints + rate limiting
# ---------------------------------------------------------------------------
from starlette.middleware.base import BaseHTTPMiddleware
from fastapi import Request, HTTPException

# Routes that don't require authentication
PUBLIC_ROUTES = [
    "/api/auth/login",
    "/api/auth/logout",      # clearing your own cookie needs no token
    "/api/version",
    # PDF preview endpoints validate token via query param themselves
    "/api/templates/preview-pdf/",
    "/api/documents/preview-pdf/",
    "/docs",
    "/openapi.json",
    "/redoc",
    "/chat",
    "/dashboard",
    # ★ "/documents/legal/" was listed here as "Static file serving". It was dead
    # text: this list is only consulted after the `path.startswith("/api/")` gate
    # below, and /documents/... never gets that far. Removing it changes no
    # behaviour — the gating now happens in STATIC_PROTECTED_ROOTS, above that
    # early return, which is the only place it can happen at all.
]

# ---------------------------------------------------------------------------
# AgentOS (Agno-mounted) route protection
# ---------------------------------------------------------------------------
# `app = agent_os.get_app()` mounts Agno's own routers at the ROOT of the app
# (/agents, /sessions, /teams, ...). They never went behind the JWT check —
# PUBLIC_ROUTES used to list them with the comment "have their own auth". They
# do not. `GET /sessions` returned every session id, title and timestamp with
# no Authorization header, and a session title is the user's literal first
# message, i.e. real company and director names.
#
# Matched on the FIRST PATH SEGMENT only, so a frontend route that merely
# starts with the same letters can never be caught by accident. Everything not
# listed here keeps its current behaviour — in particular `/`, `/health`,
# `/login`, `/admin/*`, `/_next/*`, `/favicon.ico` and every other static
# frontend path fall through untouched, so a logged-out user can still render
# the login page.
AGENTOS_PROTECTED_ROOTS = frozenset({
    "agents",               # incl. /agents/{id}/runs and .../continue (chat + HITL resume)
    "sessions",             # leaks chat history — the reason this block exists
    "teams",
    "workflows",
    "eval-runs",
    "evals",
    "schedules",
    "approvals",
    "components",
    "config",               # AgentOS config dump (agent ids, models, db ids)
    "registry",
    "models",
    "memories",
    "memory",
    "memory_topics",
    "optimize-memories",
    "user_memory_stats",
    "traces",               # full agent run traces incl. tool arguments
    "trace_session_stats",
    "knowledge",            # bare /knowledge is Agno's; ours is /api/knowledge/*
    "databases",            # /databases/{id}/migrate
    "metrics",              # Agno usage metrics + the Prometheus scrape endpoint
})

# ---------------------------------------------------------------------------
# Static file trees mounted outside /api
# ---------------------------------------------------------------------------
# ★★★ Every client document was downloadable without a password.
#
# `app.mount("/documents", StaticFiles(...))` serves the whole documents tree,
# and the auth middleware returns early on anything that does not start with
# "/api/" — so generated documents, uploaded DICA filings, the firm's templates
# and the cached previews were all served to anyone who asked. Measured
# 2026-08-06 against the running app: a real AGM minutes .docx came back 200,
# 29,313 bytes, with no token of any kind. The container publishes on
# 0.0.0.0, so that was the whole local network, and would be the whole internet
# on a public host.
#
# PUBLIC_ROUTES carried a "/documents/legal/" entry commented "Static file
# serving", which read like a deliberate policy but never executed — that list
# is consulted AFTER the /api/ early return. The gate has to sit before it,
# which is here.
#
# Matched on the first path segment, like the AgentOS roots above, and checked
# with _request_jwt — which reads the Authorization header, then `?token=`, then
# the ls_session cookie. The cookie is what matters: a plain <a href> download
# link and an <img>/canvas fetch cannot set a header, but the browser sends the
# cookie automatically on same-origin requests. So every existing download link
# keeps working unchanged, and no token has to be put in a URL.
STATIC_PROTECTED_ROOTS = frozenset({
    "documents",            # generated docs, uploaded DICA filings, templates, previews
})

# The JWT is also issued as an HttpOnly cookie at login so that same-origin
# requests carry it automatically. See `auth_login` for why.
SESSION_COOKIE_NAME = "ls_session"


def _request_jwt(request: Request) -> str:
    """The raw JWT off a request, from any transport the app already uses.

    Header first (what /api/* uses), then `?token=` (what the PDF/DOCX viewers
    use because an <img>/canvas fetch cannot set headers), then the session
    cookie (set at login, sent automatically on same-origin requests).
    """
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        return auth[7:].strip()
    qp = request.query_params.get("token", "")
    if qp:
        return qp.strip()
    try:
        return (request.cookies.get(SESSION_COOKIE_NAME) or "").strip()
    except Exception:
        return ""


class AuthMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        from starlette.responses import JSONResponse
        path = request.url.path
        ip = request.client.host if request.client else "unknown"

        # Skip for OPTIONS (CORS preflight)
        if request.method == "OPTIONS":
            return await call_next(request)

        # AgentOS routes mounted at the app root — gate on the first segment.
        root = path.split("/", 2)[1] if path.startswith("/") else ""
        if root in AGENTOS_PROTECTED_ROOTS or root in STATIC_PROTECTED_ROOTS:
            token = _request_jwt(request)
            if not token:
                return JSONResponse(status_code=401, content={"detail": "Authentication required"})
            try:
                payload = jwt.decode(token, getenv("JWT_SECRET_KEY", ""), algorithms=["HS256"])
            except jwt.ExpiredSignatureError:
                return JSONResponse(status_code=401, content={"detail": "Token expired. Please login again."})
            except jwt.InvalidTokenError:
                return JSONResponse(status_code=401, content={"detail": "Invalid token."})

            # ★ Pin every Agno route to the caller's own identity.
            #
            # Agno's routers all begin with the same three lines — session.py:122,
            # agents/router.py:229, memory.py:94, teams/router.py:177 and ~12 more:
            #
            #     if hasattr(request.state, "user_id") and request.state.user_id is not None:
            #         user_id = request.state.user_id
            #
            # i.e. request.state.user_id OVERRIDES whatever the client sent. Nothing
            # in this app ever set it, so that override was inert and Agno used the
            # client's value verbatim. Two consequences, both measured on the running
            # app before this line existed:
            #
            #   * GET /sessions with no user_id returned EVERY session (70 rows),
            #     including other users'. The JWT gate above proves you are someone;
            #     it did not scope you to your own data. `?user_id=999` returning 0
            #     rows and `?user_id=1` returning 57 shows the DB filter was working
            #     perfectly — it was simply the CLIENT's choice whether to apply it,
            #     and a client's choice is not an authorization boundary.
            #   * POST /agents/{id}/runs took user_id from a form field the browser
            #     filled in from localStorage (useAIStreamHandler.tsx:1234). A session
            #     could therefore be filed under any user id the caller typed.
            #
            # Setting it here fixes both at once, for every Agno route, without
            # touching vendored code. str() because agno_sessions.user_id is VARCHAR
            # and the JWT carries an int — "1" != 1 would silently match nothing.
            #
            # `request.state` is backed by the ASGI scope, so this survives into the
            # endpoint even though BaseHTTPMiddleware hands it a different Request
            # object. Agno logs a warning (it does NOT 403) when the client sent a
            # different value, so an out-of-date frontend keeps working.
            uid = payload.get("user_id")
            if uid is not None:
                request.state.user_id = str(uid)
            return await call_next(request)

        # Skip for non-API routes
        if not path.startswith("/api/"):
            return await call_next(request)

        # Skip for public routes
        for public in PUBLIC_ROUTES:
            if path.startswith(public):
                # Rate limit login specifically
                if path == "/api/auth/login" and request.method == "POST":
                    if _check_rate_limit(ip, RATE_LIMIT_LOGIN):
                        logger.warning(f"Rate limit exceeded for login from {ip}")
                        return JSONResponse(status_code=429, content={"detail": "Too many login attempts. Try again in 1 minute."})
                return await call_next(request)

        # General API rate limiting
        if _check_rate_limit(ip, RATE_LIMIT_API):
            logger.warning(f"API rate limit exceeded from {ip}")
            return JSONResponse(status_code=429, content={"detail": "Too many requests. Slow down."})

        # Check JWT token
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            token = auth[7:]
            try:
                jwt.decode(token, getenv("JWT_SECRET_KEY", ""), algorithms=["HS256"])
                return await call_next(request)
            except jwt.ExpiredSignatureError:
                return JSONResponse(status_code=401, content={"detail": "Token expired. Please login again."})
            except jwt.InvalidTokenError:
                return JSONResponse(status_code=401, content={"detail": "Invalid token."})

        # No valid token
        return JSONResponse(status_code=401, content={"detail": "Authentication required"})

app.add_middleware(AuthMiddleware)


# ---------------------------------------------------------------------------
# Request Logging Middleware — correlation IDs + timing
# ---------------------------------------------------------------------------
import uuid
import time as _req_time

_request_logger = get_logger("request")


class EffectTurnMiddleware(BaseHTTPMiddleware):
    """Opens one effects turn per agent run, so side effects can be grouped.

    A "turn" is one thing the user asked for. Only agent-run requests get one:
    a static asset or a poll has no side effects worth grouping, and opening a
    turn for every request would file a turn nobody took.

    `turn_scope` is a true no-op while SCOUT_EFFECTS_LEDGER is off — it yields
    None without minting an id or touching the database — so this middleware
    costs a flag read on the default path. It never fails a request: the scope
    swallows its own errors and resets its ContextVar in a finally, so a turn
    cannot leak into the next request served on the same task.
    """

    def _is_agent_run(self, request: Request) -> bool:
        path = request.url.path
        return request.method == "POST" and "/runs" in path and (
            path.startswith("/agents/") or path.startswith("/teams/")
        )

    async def dispatch(self, request: Request, call_next):
        if not self._is_agent_run(request):
            return await call_next(request)

        actor = None
        try:
            user = get_current_user(request)
            actor = (user or {}).get("email")
        except Exception:  # noqa: BLE001 — auth shape must never fail a run
            actor = None

        from scout.effects import turn_scope

        # The session id is not knowable here (agno reads it from the form
        # body); the tools layer binds it to the turn on the first effect.
        with turn_scope(actor_email=actor, meta={"path": request.url.path}):
            return await call_next(request)


app.add_middleware(EffectTurnMiddleware)


class RequestLoggingMiddleware(BaseHTTPMiddleware):
    """Logs every request with method, path, status, duration, and a unique request ID."""

    # Paths to skip logging (noisy/health)
    SKIP_PATHS = {"/health", "/favicon.ico"}

    async def dispatch(self, request: Request, call_next):
        path = request.url.path

        # Skip noisy endpoints
        if path in self.SKIP_PATHS:
            return await call_next(request)

        # Generate correlation ID
        request_id = request.headers.get("X-Request-ID", str(uuid.uuid4())[:8])
        start = _req_time.time()

        response = await call_next(request)

        duration_ms = round((_req_time.time() - start) * 1000, 1)

        # Add request ID to response headers
        response.headers["X-Request-ID"] = request_id

        # Log the request
        ip = request.client.host if request.client else "unknown"
        # Demote noisy expected 4xx (new sessions, stale cached preview-pdf) to INFO
        _noisy_404 = response.status_code == 404 and ("/sessions/" in path or "/runs" in path)
        _stale_preview = response.status_code == 401 and "/preview-pdf/" in path
        if _noisy_404 or _stale_preview:
            log_level = logging.INFO
        else:
            log_level = logging.WARNING if response.status_code >= 400 else logging.INFO

        _request_logger.log(
            log_level,
            f"{request.method} {path} → {response.status_code} ({duration_ms}ms)",
            extra={
                "request_id": request_id,
                "method": request.method,
                "path": path,
                "status_code": response.status_code,
                "duration_ms": duration_ms,
                "ip": ip,
            },
        )

        return response


app.add_middleware(RequestLoggingMiddleware)


# ---------------------------------------------------------------------------
# Authentication Helpers
# ---------------------------------------------------------------------------

JWT_SECRET = getenv("JWT_SECRET_KEY", "")
JWT_ALGORITHM = "HS256"
JWT_EXPIRY_HOURS = 24

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed.encode())

def create_token(user_id: int, email: str, role: str) -> str:
    from datetime import datetime, timedelta
    payload = {
        "user_id": user_id,
        "email": email,
        "role": role,
        "exp": datetime.utcnow() + timedelta(hours=JWT_EXPIRY_HOURS),
        "iat": datetime.utcnow(),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def get_current_user(request: Request) -> dict:
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        return None
    token = auth[7:]
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        return None
    except jwt.InvalidTokenError:
        return None

def require_admin(request: Request) -> dict:
    user = get_current_user(request)
    if not user or user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return user

def log_activity(user_id: int, email: str, action: str, details: str = "", ip: str = ""):
    try:
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("INSERT INTO activity_logs (user_id, user_email, action, details, ip_address) VALUES (%s, %s, %s, %s, %s)",
            (user_id, email, action, details, ip))
        conn.commit(); cur.close(); conn.close()
    except Exception as e:
        logging.getLogger("legalscout").warning(f"Activity log failed: {e}")


def generate_embedding(text: str) -> list[float] | None:
    """Generate embedding via OpenRouter using configured model."""
    import httpx
    from app.model_config import get_model
    text = text[:8000]

    model = get_model("embedding")
    openrouter_key = getenv("OPENROUTER_API_KEY")
    if not openrouter_key:
        logger.warning("OPENROUTER_API_KEY not set — cannot generate embeddings")
        return None

    try:
        res = httpx.post(
            f"{OPENROUTER_BASE_URL}/embeddings",
            headers={"Authorization": f"Bearer {openrouter_key}", "Content-Type": "application/json"},
            json={"model": model, "input": text, "encoding_format": "float"},
            timeout=30,
        )
        res.raise_for_status()
        return res.json()["data"][0]["embedding"]
    except Exception as e:
        logger.warning(f"Embedding failed ({model}): {e}")
        return None


def invalidate_training(reason: str, template_name: str = None):
    """Mark training as stale when templates/companies change.
    If template_name is given, only that template is marked untrained.
    If None, all training status is cleared."""
    try:
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()

        if template_name:
            # Mark specific template as untrained
            cur.execute("UPDATE templates SET ai_trained = FALSE, ai_analyzed = FALSE WHERE name = %s", (template_name,))
        else:
            # Mark ALL templates as untrained
            cur.execute("UPDATE templates SET ai_trained = FALSE, ai_analyzed = FALSE")

        # Clear training status and logs
        cur.execute("UPDATE training_status SET status = 'stale', logs = '[]'::jsonb WHERE training_type = 'templates'")

        conn.commit(); cur.close(); conn.close()
        logger.info(f"[TRAINING] Invalidated: {reason}")
    except Exception as e:
        logger.warning(f"[TRAINING] Failed to invalidate: {e}")


def send_notification_email(to_email: str, subject: str, body: str):
    """Send email notification via SMTP (optional — only if configured)."""
    import smtplib, os
    from email.mime.text import MIMEText

    smtp_host = os.getenv("SMTP_HOST")
    smtp_port = int(os.getenv("SMTP_PORT", "587"))
    smtp_user = os.getenv("SMTP_USER")
    smtp_pass = os.getenv("SMTP_PASS")
    from_email = os.getenv("SMTP_FROM", "noreply@legalscout.com")

    if not smtp_host or not smtp_user:
        return  # Email not configured, skip silently

    try:
        msg = MIMEText(body, "html")
        msg["Subject"] = f"Legal Scout — {subject}"
        msg["From"] = from_email
        msg["To"] = to_email

        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_pass)
            server.send_message(msg)
    except Exception as e:
        print(f"Email send failed: {e}")


# Auto-create admin user on startup
def _init_admin():
    try:
        import os
        from psycopg import connect
        from db.connection import get_db_conn as _get_admin_conn
        conn = _get_admin_conn()
        cur = conn.cursor()
        # Create tables if not exist
        cur.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id SERIAL PRIMARY KEY, email VARCHAR(255) UNIQUE NOT NULL,
                hashed_password VARCHAR(255) NOT NULL, full_name VARCHAR(255) DEFAULT '',
                role VARCHAR(50) DEFAULT 'user', is_active BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP, updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS activity_logs (
                id SERIAL PRIMARY KEY, user_id INTEGER, user_email VARCHAR(255),
                action VARCHAR(100) NOT NULL, details TEXT, ip_address VARCHAR(45),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS document_versions (
                id SERIAL PRIMARY KEY, document_name VARCHAR(255), company_name VARCHAR(255),
                template_name VARCHAR(255), version INTEGER DEFAULT 1, file_name VARCHAR(255),
                file_path TEXT, generated_by INTEGER, generated_by_email VARCHAR(255),
                custom_data JSONB, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS template_versions (
                id SERIAL PRIMARY KEY, template_name VARCHAR(255),
                version INTEGER DEFAULT 1, uploaded_by VARCHAR(255),
                fields JSONB, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        # Create admin if not exists
        admin_email = os.getenv("ADMIN_EMAIL", "admin@legalscout.com")
        # No silent `admin123` fallback: a fresh install that forgot to set
        # ADMIN_PASSWORD used to come up with a guessable admin account and say
        # nothing about it.
        admin_pass = os.getenv("ADMIN_PASSWORD", "")
        cur.execute("SELECT id, hashed_password FROM users WHERE email = %s", (admin_email,))
        row = cur.fetchone()
        if not row:
            if not admin_pass:
                print("Admin NOT created: set ADMIN_PASSWORD to seed the admin account")
            else:
                cur.execute("INSERT INTO users (email, hashed_password, full_name, role) VALUES (%s, %s, %s, %s)",
                    (admin_email, hash_password(admin_pass), "Admin", "admin"))
                print(f"Admin user created: {admin_email}")
        elif admin_pass and not verify_password(admin_pass, row[1]):
            # Rotation. This branch did not exist: the row was only ever
            # INSERTed when absent, so changing ADMIN_PASSWORD in the
            # environment did NOTHING to an existing admin and the old
            # credential kept working — an operator who "rotated" the secret
            # was left believing they had.
            cur.execute("UPDATE users SET hashed_password = %s, updated_at = now() WHERE id = %s",
                (hash_password(admin_pass), row[0]))
            print(f"Admin password rotated from ADMIN_PASSWORD: {admin_email}")
        conn.commit(); cur.close(); conn.close()
    except Exception as e:
        print(f"Admin init warning: {e}")

# Initialize admin on import
try:
    _init_admin()
except Exception as e:
    logging.getLogger("legalscout").warning(f"Admin init failed: {e}")


# ---------------------------------------------------------------------------
# Settings API — admin can configure SMTP, etc. from UI
# ---------------------------------------------------------------------------
@app.get("/api/admin/settings")
async def admin_get_settings(request: Request):
    """Get all app settings (admin only)."""
    require_admin(request)
    try:
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("CREATE TABLE IF NOT EXISTS app_settings (id SERIAL PRIMARY KEY, key VARCHAR(255) UNIQUE NOT NULL, value TEXT, updated_by VARCHAR(255), updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
        cur.execute("SELECT key, value, updated_by, updated_at FROM app_settings ORDER BY key")
        rows = cur.fetchall()
        cur.close(); conn.close()
        settings = {r[0]: {"value": r[1], "updated_by": r[2], "updated_at": r[3].isoformat() if r[3] else None} for r in rows}
        return {"success": True, "settings": settings}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/admin/settings")
async def admin_save_settings(request: Request):
    """Save app settings (admin only)."""
    admin = require_admin(request)
    try:
        body = await request.json()
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("CREATE TABLE IF NOT EXISTS app_settings (id SERIAL PRIMARY KEY, key VARCHAR(255) UNIQUE NOT NULL, value TEXT, updated_by VARCHAR(255), updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
        for key, value in body.items():
            cur.execute("""
                INSERT INTO app_settings (key, value, updated_by, updated_at) VALUES (%s, %s, %s, NOW())
                ON CONFLICT (key) DO UPDATE SET value = %s, updated_by = %s, updated_at = NOW()
            """, (key, str(value), admin.get("email", ""), str(value), admin.get("email", "")))
        conn.commit(); cur.close(); conn.close()
        log_activity(admin.get("user_id"), admin.get("email"), "update_settings", f"Updated {len(body)} settings", "")
        return {"success": True, "message": f"Saved {len(body)} settings"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/admin/test-email")
async def admin_test_email(request: Request):
    """Send a test email to verify SMTP config."""
    admin = require_admin(request)
    try:
        body = await request.json()
        to_email = body.get("to_email", admin.get("email", ""))

        # Get SMTP settings from DB
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT key, value FROM app_settings WHERE key LIKE 'smtp_%'")
        smtp_settings = {r[0]: r[1] for r in cur.fetchall()}
        cur.close(); conn.close()

        smtp_host = smtp_settings.get("smtp_host") or os.getenv("SMTP_HOST", "")
        smtp_port = int(smtp_settings.get("smtp_port") or os.getenv("SMTP_PORT", "587"))
        smtp_user = smtp_settings.get("smtp_user") or os.getenv("SMTP_USER", "")
        smtp_pass = smtp_settings.get("smtp_pass") or os.getenv("SMTP_PASS", "")
        smtp_from = smtp_settings.get("smtp_from") or os.getenv("SMTP_FROM", "noreply@legalscout.com")

        if not smtp_host or not smtp_user:
            return {"success": False, "error": "SMTP not configured. Fill in the settings first."}

        import smtplib
        from email.mime.text import MIMEText
        msg = MIMEText("<h2>Legal Scout — Test Email</h2><p>If you see this, email notifications are working!</p>", "html")
        msg["Subject"] = "Legal Scout — Test Email"
        msg["From"] = smtp_from
        msg["To"] = to_email

        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_pass)
            server.send_message(msg)

        return {"success": True, "message": f"Test email sent to {to_email}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _login_failure(error: str) -> JSONResponse:
    """Failed login → HTTP 401, same JSON body shape as before.

    Every credential path returns 401 with the same status, including a
    disabled account: a distinct 403 there would re-open the account
    enumeration the dummy-bcrypt check below exists to prevent.
    """
    return JSONResponse(status_code=401, content={"success": False, "error": error})


@app.post("/api/auth/login")
async def auth_login(request: Request):
    try:
        body = await request.json()
        email = sanitize_string(body.get("email", ""), 255).lower()
        password = body.get("password", "")
        if not email or not password:
            return _login_failure("Email and password required")
        if not validate_email(email):
            return _login_failure("Invalid email format")
        if len(password) < 3 or len(password) > 200:
            return _login_failure("Invalid password")

        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT id, email, hashed_password, full_name, role, is_active FROM users WHERE email = %s", (email,))
        row = cur.fetchone()
        cur.close(); conn.close()

        if not row:
            # Always hash-check to prevent timing attack (email enumeration)
            bcrypt.checkpw(b"dummy", b"$2b$12$LJ3m4ys3Gn/0FWpfKMNbIeDjQJz2GnnKTjPVTqJjLYKmWFjGEQ3ya")
            return _login_failure("Invalid email or password")
        if not row[5]:
            return _login_failure("Account is disabled")
        if not verify_password(password, row[2]):
            return _login_failure("Invalid email or password")

        token = create_token(row[0], row[1], row[4])
        log_activity(row[0], row[1], "login", "User logged in", request.client.host if request.client else "")

        resp = JSONResponse(content={
            "success": True,
            "token": token,
            "user": {"id": row[0], "email": row[1], "name": row[3], "role": row[4]},
        })
        # Same token, second transport. The frontend keeps reading `token` out
        # of this body into localStorage exactly as before; the cookie exists
        # because the AgentOS routes (/agents, /sessions, ...) are called by
        # code that does not attach an Authorization header. Same-origin fetch
        # sends this automatically, so gating those routes does not black out
        # chat. SameSite=Lax blocks the cross-site CSRF vector.
        resp.set_cookie(
            SESSION_COOKIE_NAME,
            token,
            max_age=JWT_EXPIRY_HOURS * 3600,
            httponly=True,
            samesite="lax",
            secure=(request.url.scheme == "https"),
            path="/",
        )
        return resp
    except Exception as e:
        log_error("auth_login", e)
        return JSONResponse(status_code=500, content={"success": False, "error": "Login failed"})


@app.post("/api/auth/logout")
async def auth_logout():
    """Clears the session cookie.

    The frontend sign-out handlers currently only wipe localStorage, so until
    they call this the cookie outlives sign-out (up to JWT_EXPIRY_HOURS).
    """
    resp = JSONResponse(content={"success": True})
    resp.delete_cookie(SESSION_COOKIE_NAME, path="/")
    return resp


@app.get("/api/auth/me")
async def auth_me(request: Request):
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return {"success": True, "user": user}


@app.get("/api/admin/users")
async def admin_list_users(request: Request):
    require_admin(request)
    try:
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT id, email, full_name, role, is_active, created_at, updated_at FROM users ORDER BY created_at DESC")
        rows = cur.fetchall()
        cur.close(); conn.close()
        users = [{"id": r[0], "email": r[1], "name": r[2], "role": r[3], "is_active": r[4], "created_at": r[5].isoformat() if r[5] else None, "updated_at": r[6].isoformat() if r[6] else None} for r in rows]
        return {"success": True, "users": users}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/admin/users")
async def admin_create_user(request: Request):
    admin = require_admin(request)
    try:
        body = await request.json()
        email = sanitize_string(body.get("email", ""), 255).lower()
        password = body.get("password", "")
        name = sanitize_string(body.get("name", ""), 255)
        role = body.get("role", "user")
        if not email or not password:
            return {"success": False, "error": "Email and password required"}
        if not validate_email(email):
            return {"success": False, "error": "Invalid email format"}
        if len(password) < 10:
            return {"success": False, "error": "Password must be at least 10 characters"}
        if role not in ("user", "editor", "admin"):
            return {"success": False, "error": "Role must be user, editor, or admin"}

        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("INSERT INTO users (email, hashed_password, full_name, role) VALUES (%s, %s, %s, %s) RETURNING id",
            (email, hash_password(password), name, role))
        new_id = cur.fetchone()[0]
        conn.commit(); cur.close(); conn.close()

        log_activity(admin.get("user_id"), admin.get("email"), "create_user", f"Created user: {email}", "")
        return {"success": True, "message": f"User '{email}' created", "id": new_id}
    except Exception as e:
        if "duplicate key" in str(e).lower():
            return {"success": False, "error": "Email already exists"}
        return {"success": False, "error": str(e)}


@app.put("/api/admin/users/{user_id}")
async def admin_update_user(user_id: int, request: Request):
    admin = require_admin(request)
    try:
        body = await request.json()
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()

        updates = []
        params = []
        if "name" in body:
            updates.append("full_name = %s"); params.append(body["name"])
        if "role" in body:
            updates.append("role = %s"); params.append(body["role"])
        if "is_active" in body:
            updates.append("is_active = %s"); params.append(body["is_active"])
        if "password" in body and body["password"]:
            updates.append("hashed_password = %s"); params.append(hash_password(body["password"]))

        if updates:
            updates.append("updated_at = NOW()")
            params.append(user_id)
            cur.execute(f"UPDATE users SET {', '.join(updates)} WHERE id = %s", tuple(params))
            conn.commit()

        cur.close(); conn.close()
        log_activity(admin.get("user_id"), admin.get("email"), "update_user", f"Updated user id={user_id}", "")
        return {"success": True, "message": "User updated"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.delete("/api/admin/users/{user_id}")
async def admin_delete_user(user_id: int, request: Request):
    admin = require_admin(request)
    try:
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("DELETE FROM users WHERE id = %s AND role != 'admin'", (user_id,))
        deleted = cur.rowcount
        conn.commit(); cur.close(); conn.close()
        if deleted:
            log_activity(admin.get("user_id"), admin.get("email"), "delete_user", f"Deleted user id={user_id}", "")
            return {"success": True, "message": "User deleted"}
        return {"success": False, "error": "User not found or cannot delete admin"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/admin/activity-logs")
async def admin_activity_logs(request: Request):
    require_admin(request)
    try:
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT id, user_email, action, details, ip_address, created_at FROM activity_logs ORDER BY created_at DESC LIMIT 200")
        rows = cur.fetchall()
        cur.close(); conn.close()
        logs = [{"id": r[0], "email": r[1], "action": r[2], "details": r[3], "ip": r[4], "created_at": r[5].isoformat() if r[5] else None} for r in rows]
        return {"success": True, "logs": logs}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# People Register (migration 011: people, company_people)
# ---------------------------------------------------------------------------
PEOPLE_ROLES = ("director", "individual_shareholder", "both")
PEOPLE_PAGE_LIMIT = 100

# Column list shared by every people SELECT so row unpacking stays in sync.
_PEOPLE_COLS = (
    "id, full_name, nationality, nrc_passport_no, gender, date_of_birth, "
    "phone, email, residential_address, business_occupation, country_of_residence, "
    "father_name, created_by_email, created_at, updated_at"
)


def _nullable_date(value):
    """DATE columns reject ''. Treat empty/whitespace as SQL NULL."""
    if value is None:
        return None
    if isinstance(value, str):
        value = value.strip()
        if not value:
            return None
    return value


def _nullable_text(value, max_length: int = 500):
    """Sanitize an optional text field; '' becomes NULL so the partial
    unique index on nrc_passport_no behaves predictably."""
    if value is None:
        return None
    cleaned = sanitize_string(value, max_length)
    return cleaned or None


def _person_row(r):
    return {
        "id": r[0],
        "full_name": r[1],
        "nationality": r[2],
        "nrc_passport_no": r[3],
        "gender": r[4],
        "date_of_birth": r[5].isoformat() if r[5] else None,
        "phone": r[6],
        "email": r[7],
        "residential_address": r[8],
        "business_occupation": r[9],
        "country_of_residence": r[10],
        "father_name": r[11],
        "created_by_email": r[12],
        "created_at": r[13].isoformat() if r[13] else None,
        "updated_at": r[14].isoformat() if r[14] else None,
    }


def _link_row(r):
    """company_people row joined with the company's English name."""
    return {
        "id": r[0],
        "company_id": r[1],
        "person_id": r[2],
        "role": r[3],
        "number_of_shares": r[4],
        "capital_amount": r[5],
        "share_class": r[6],
        "appointed_date": r[7].isoformat() if r[7] else None,
        "resigned_date": r[8].isoformat() if r[8] else None,
        "cessation_reason": r[9],
        "cessation_recorded_by": r[10],
        "company_name": r[11] if len(r) > 11 else None,
    }


# ★★★LANDMINE — _LINK_COLS is a PREFIX of three queries (get_person,
# person_companies, list_company_people). The third appends
# p.full_name, p.nationality, p.nrc_passport_no and reads them POSITIONALLY.
# Adding a column here without renumbering that slice puts the person's NAME
# into company_name. No exception is raised — the endpoint just answers wrong.
_LINK_COLS = (
    "cp.id, cp.company_id, cp.person_id, cp.role, cp.number_of_shares, "
    "cp.capital_amount, cp.share_class, cp.appointed_date, cp.resigned_date, "
    "cp.cessation_reason, cp.cessation_recorded_by"
)


def _is_duplicate_nrc(exc: Exception) -> bool:
    """migration 011 partial unique index idx_people_nrc."""
    msg = str(exc).lower()
    return "idx_people_nrc" in msg or ("duplicate key" in msg and "nrc" in msg)


@app.get("/api/people")
async def list_people(request: Request, search: str = "", limit: int = 50, offset: int = 0):
    """List people, optionally filtered by name or NRC/passport number."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")

    limit = max(1, min(int(limit or 50), PEOPLE_PAGE_LIMIT))
    offset = max(0, int(offset or 0))
    search = sanitize_string(search or "", 200)

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        if search:
            pattern = f"%{search}%"
            cur.execute(
                f"SELECT {_PEOPLE_COLS} FROM people "
                "WHERE full_name ILIKE %s OR nrc_passport_no ILIKE %s "
                "ORDER BY full_name ASC LIMIT %s OFFSET %s",
                (pattern, pattern, limit, offset),
            )
        else:
            cur.execute(
                f"SELECT {_PEOPLE_COLS} FROM people "
                "ORDER BY full_name ASC LIMIT %s OFFSET %s",
                (limit, offset),
            )
        rows = cur.fetchall()

        if search:
            pattern = f"%{search}%"
            cur.execute(
                "SELECT COUNT(*) FROM people WHERE full_name ILIKE %s OR nrc_passport_no ILIKE %s",
                (pattern, pattern),
            )
        else:
            cur.execute("SELECT COUNT(*) FROM people")
        total = cur.fetchone()[0]
        cur.close()

        people = [_person_row(r) for r in rows]
        return {"success": True, "people": people, "count": len(people),
                "total": total, "limit": limit, "offset": offset}
    except Exception as e:
        logger.warning(f"[PEOPLE] list failed: {e}")
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.post("/api/people")
async def create_person(request: Request):
    """Create a person. Admin only."""
    admin = require_admin(request)
    conn = None
    try:
        body = await request.json()
        full_name = sanitize_string(body.get("full_name", ""), 500)
        if not full_name:
            return {"success": False, "error": "full_name is required"}

        email = _nullable_text(body.get("email"), 255)
        if email and not validate_email(email):
            return {"success": False, "error": "Invalid email format"}

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO people (full_name, nationality, nrc_passport_no, gender, "
            "date_of_birth, phone, email, residential_address, business_occupation, "
            "country_of_residence, father_name, created_by_email) "
            "VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING id",
            (
                full_name,
                _nullable_text(body.get("nationality"), 100),
                _nullable_text(body.get("nrc_passport_no"), 100),
                _nullable_text(body.get("gender"), 20),
                _nullable_date(body.get("date_of_birth")),
                _nullable_text(body.get("phone"), 50),
                email,
                _nullable_text(body.get("residential_address"), 2000),
                _nullable_text(body.get("business_occupation"), 255),
                _nullable_text(body.get("country_of_residence"), 100),
                _nullable_text(body.get("father_name"), 500),
                admin.get("email"),
            ),
        )
        new_id = cur.fetchone()[0]
        conn.commit()
        cur.close()

        log_activity(admin.get("user_id"), admin.get("email"), "create_person",
                     f"Created person: {full_name} (id={new_id})", "")
        return {"success": True, "message": f"Person '{full_name}' created", "id": new_id}
    except Exception as e:
        if conn:
            conn.rollback()
        if _is_duplicate_nrc(e):
            logger.info(f"[PEOPLE] duplicate NRC rejected: {e}")
            raise HTTPException(status_code=409, detail="A person with this NRC/passport number already exists")
        logger.warning(f"[PEOPLE] create failed: {e}")
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.post("/api/people/sync-from-companies")
async def sync_people_from_companies(request: Request):
    """Backfill the People register from every stored company. Admin only.

    New companies sync themselves on save; this replays the projection over
    records that predate the sync (or that were edited straight in the DB).
    Idempotent — existing person fields are filled, never overwritten.
    Declared above /api/people/{person_id} so the literal path wins the match.
    """
    admin = require_admin(request)
    conn = None
    try:
        from scout.tools.people_sync import sync_all_companies

        conn = get_db_conn()
        totals = sync_all_companies(conn, admin.get("email"))
        conn.commit()

        log_activity(
            admin.get("user_id"), admin.get("email"), "sync_people",
            f"Synced People from {totals['companies']} companies: "
            f"{totals['created']} created, {totals['updated']} updated, {totals['linked']} linked",
            "",
        )
        return {"success": True, **totals}
    except Exception as e:
        if conn:
            conn.rollback()
        logger.warning(f"[PEOPLE] sync from companies failed: {e}")
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.get("/api/people/{person_id}")
async def get_person(person_id: int, request: Request):
    """Single person record including its company links."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(f"SELECT {_PEOPLE_COLS} FROM people WHERE id = %s", (person_id,))
        row = cur.fetchone()
        if not row:
            cur.close()
            return {"success": False, "error": "Person not found"}

        person = _person_row(row)
        cur.execute(
            f"SELECT {_LINK_COLS}, c.company_name_english "
            "FROM company_people cp LEFT JOIN companies c ON c.id = cp.company_id "
            "WHERE cp.person_id = %s ORDER BY cp.company_id ASC, cp.role ASC",
            (person_id,),
        )
        person["companies"] = [_link_row(r) for r in cur.fetchall()]
        cur.close()
        return {"success": True, "person": person}
    except Exception as e:
        logger.warning(f"[PEOPLE] get id={person_id} failed: {e}")
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.put("/api/people/{person_id}")
async def update_person(person_id: int, request: Request):
    """Update a person. Admin only."""
    admin = require_admin(request)
    conn = None
    try:
        body = await request.json()

        updates = []
        params = []
        if "full_name" in body:
            full_name = sanitize_string(body.get("full_name", ""), 500)
            if not full_name:
                return {"success": False, "error": "full_name cannot be empty"}
            updates.append("full_name = %s"); params.append(full_name)
        if "email" in body:
            email = _nullable_text(body.get("email"), 255)
            if email and not validate_email(email):
                return {"success": False, "error": "Invalid email format"}
            updates.append("email = %s"); params.append(email)
        for field, max_len in (("nationality", 100), ("nrc_passport_no", 100),
                               ("business_occupation", 255), ("country_of_residence", 100),
                               ("father_name", 500),
                               ("gender", 20), ("phone", 50),
                               ("residential_address", 2000)):
            if field in body:
                updates.append(f"{field} = %s")
                params.append(_nullable_text(body.get(field), max_len))
        if "date_of_birth" in body:
            updates.append("date_of_birth = %s")
            params.append(_nullable_date(body.get("date_of_birth")))

        if not updates:
            return {"success": False, "error": "No fields to update"}

        conn = get_db_conn()
        cur = conn.cursor()
        updates.append("updated_at = NOW()")
        params.append(person_id)
        cur.execute(f"UPDATE people SET {', '.join(updates)} WHERE id = %s", tuple(params))
        updated = cur.rowcount
        conn.commit()
        cur.close()

        if not updated:
            return {"success": False, "error": "Person not found"}
        log_activity(admin.get("user_id"), admin.get("email"), "update_person",
                     f"Updated person id={person_id}", "")
        return {"success": True, "message": "Person updated"}
    except Exception as e:
        if conn:
            conn.rollback()
        if _is_duplicate_nrc(e):
            logger.info(f"[PEOPLE] duplicate NRC rejected on update: {e}")
            raise HTTPException(status_code=409, detail="A person with this NRC/passport number already exists")
        logger.warning(f"[PEOPLE] update id={person_id} failed: {e}")
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.delete("/api/people/{person_id}")
async def delete_person(person_id: int, request: Request):
    """Delete a person. Admin only. company_people rows cascade."""
    admin = require_admin(request)
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("DELETE FROM people WHERE id = %s", (person_id,))
        deleted = cur.rowcount
        conn.commit()
        cur.close()

        if not deleted:
            return {"success": False, "error": "Person not found"}
        log_activity(admin.get("user_id"), admin.get("email"), "delete_person",
                     f"Deleted person id={person_id}", "")
        return {"success": True, "message": "Person deleted"}
    except Exception as e:
        if conn:
            conn.rollback()
        logger.warning(f"[PEOPLE] delete id={person_id} failed: {e}")
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.get("/api/people/{person_id}/companies")
async def list_person_companies(person_id: int, request: Request):
    """Company links for one person."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(
            f"SELECT {_LINK_COLS}, c.company_name_english "
            "FROM company_people cp LEFT JOIN companies c ON c.id = cp.company_id "
            "WHERE cp.person_id = %s ORDER BY cp.company_id ASC, cp.role ASC",
            (person_id,),
        )
        links = [_link_row(r) for r in cur.fetchall()]
        cur.close()
        return {"success": True, "companies": links, "count": len(links)}
    except Exception as e:
        logger.warning(f"[PEOPLE] companies for id={person_id} failed: {e}")
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.post("/api/companies/{company_id}/people")
async def link_company_person(company_id: int, request: Request):
    """Link a person to a company with a role. Admin only."""
    admin = require_admin(request)
    conn = None
    try:
        body = await request.json()
        person_id = body.get("person_id")
        if person_id in (None, ""):
            return {"success": False, "error": "person_id is required"}
        try:
            person_id = int(person_id)
        except (TypeError, ValueError):
            return {"success": False, "error": "person_id must be an integer"}

        role = sanitize_string(body.get("role", "director"), 30) or "director"
        if role not in PEOPLE_ROLES:
            return {"success": False, "error": f"role must be one of: {', '.join(PEOPLE_ROLES)}"}

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT 1 FROM companies WHERE id = %s", (company_id,))
        if not cur.fetchone():
            cur.close()
            return {"success": False, "error": "Company not found"}
        cur.execute("SELECT 1 FROM people WHERE id = %s", (person_id,))
        if not cur.fetchone():
            cur.close()
            return {"success": False, "error": "Person not found"}

        # ★★★This UPSERT sets EVERY column from EXCLUDED, so a body that OMITS a
        # field used to write NULL over the stored value. That was already firing
        # in production: the "Link company" form never sends share_class, so
        # re-linking someone to correct an appointed date silently cleared their
        # share class, with nothing in the UI or the response to show it.
        # Fields ABSENT from the body now keep what is stored. A field that is
        # PRESENT but empty still clears, so blanking one deliberately still works
        # ('' becomes NULL inside _nullable_text/_nullable_date).
        cur.execute(
            "SELECT number_of_shares, capital_amount, share_class, appointed_date, "
            "resigned_date, cessation_reason, cessation_recorded_by "
            "FROM company_people WHERE company_id = %s AND person_id = %s AND role = %s",
            (company_id, person_id, role),
        )
        stored = cur.fetchone()

        def _keep(key, idx, value):
            """body value when the caller sent the key, otherwise the stored one."""
            if key in body:
                return value
            return stored[idx] if stored else None

        # cessation_recorded_by is an audit field: it names who recorded a
        # cessation, so it is stamped from the SESSION (require_admin above), never
        # read from the body — a recorder the caller chooses is a claim, not an
        # audit trail. It is only re-stamped when this request is actually about a
        # cessation; an unrelated re-link leaves the original recorder in place.
        recording_cessation = "cessation_reason" in body or "resigned_date" in body
        recorded_by = (
            _nullable_text(admin.get("email"), 255)
            if recording_cessation
            else (stored[6] if stored else None)
        )

        cur.execute(
            "INSERT INTO company_people (company_id, person_id, role, number_of_shares, "
            "capital_amount, share_class, appointed_date, resigned_date, "
            "cessation_reason, cessation_recorded_by) "
            "VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s) "
            "ON CONFLICT (company_id, person_id, role) DO UPDATE SET "
            "number_of_shares = EXCLUDED.number_of_shares, "
            "capital_amount = EXCLUDED.capital_amount, "
            "share_class = EXCLUDED.share_class, "
            "appointed_date = EXCLUDED.appointed_date, "
            "resigned_date = EXCLUDED.resigned_date, "
            "cessation_reason = EXCLUDED.cessation_reason, "
            "cessation_recorded_by = EXCLUDED.cessation_recorded_by, "
            "updated_at = NOW() RETURNING id",
            (
                company_id,
                person_id,
                role,
                _keep("number_of_shares", 0, _nullable_text(body.get("number_of_shares"), 100)),
                _keep("capital_amount", 1, _nullable_text(body.get("capital_amount"), 100)),
                _keep("share_class", 2, _nullable_text(body.get("share_class"), 100)),
                _keep("appointed_date", 3, _nullable_date(body.get("appointed_date"))),
                _keep("resigned_date", 4, _nullable_date(body.get("resigned_date"))),
                _keep("cessation_reason", 5, _nullable_text(body.get("cessation_reason"), 500)),
                recorded_by,
            ),
        )
        link_id = cur.fetchone()[0]
        conn.commit()
        cur.close()

        log_activity(admin.get("user_id"), admin.get("email"), "link_company_person",
                     f"Linked person id={person_id} to company id={company_id} as {role}", "")
        return {"success": True, "message": "Person linked to company", "id": link_id}
    except Exception as e:
        if conn:
            conn.rollback()
        logger.warning(f"[PEOPLE] link company={company_id} failed: {e}")
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.delete("/api/companies/{company_id}/people/{person_id}")
async def unlink_company_person(company_id: int, person_id: int, request: Request, role: str = ""):
    """Unlink a person from a company. Admin only.
    Without ?role= every role link for that pair is removed."""
    admin = require_admin(request)
    conn = None
    try:
        role = sanitize_string(role or "", 30)
        if role and role not in PEOPLE_ROLES:
            return {"success": False, "error": f"role must be one of: {', '.join(PEOPLE_ROLES)}"}

        conn = get_db_conn()
        cur = conn.cursor()
        if role:
            cur.execute(
                "DELETE FROM company_people WHERE company_id = %s AND person_id = %s AND role = %s",
                (company_id, person_id, role),
            )
        else:
            cur.execute(
                "DELETE FROM company_people WHERE company_id = %s AND person_id = %s",
                (company_id, person_id),
            )
        deleted = cur.rowcount
        conn.commit()
        cur.close()

        if not deleted:
            return {"success": False, "error": "Link not found"}
        log_activity(admin.get("user_id"), admin.get("email"), "unlink_company_person",
                     f"Unlinked person id={person_id} from company id={company_id}", "")
        return {"success": True, "message": "Person unlinked from company", "deleted": deleted}
    except Exception as e:
        if conn:
            conn.rollback()
        logger.warning(f"[PEOPLE] unlink company={company_id} person={person_id} failed: {e}")
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.get("/api/companies/{company_id}/people")
async def list_company_people(company_id: int, request: Request, role: str = ""):
    """People linked to a company, optionally filtered by role.
    Role 'both' counts as director and as individual_shareholder."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")

    conn = None
    try:
        role = sanitize_string(role or "", 30)
        if role and role not in PEOPLE_ROLES:
            return {"success": False, "error": f"role must be one of: {', '.join(PEOPLE_ROLES)}"}

        conn = get_db_conn()
        cur = conn.cursor()
        sql = (
            f"SELECT {_LINK_COLS}, p.full_name, p.nationality, p.nrc_passport_no "
            "FROM company_people cp JOIN people p ON p.id = cp.person_id "
            "WHERE cp.company_id = %s"
        )
        params = [company_id]
        if role:
            sql += " AND (cp.role = %s OR cp.role = 'both')" if role != "both" else " AND cp.role = %s"
            params.append(role)
        sql += " ORDER BY p.full_name ASC LIMIT %s"
        params.append(PEOPLE_PAGE_LIMIT * 5)

        cur.execute(sql, tuple(params))
        rows = cur.fetchall()
        cur.close()

        people = []
        for r in rows:
            link = _link_row(r[:11])
            link["full_name"] = r[11]
            link["nationality"] = r[12]
            link["nrc_passport_no"] = r[13]
            people.append(link)
        return {"success": True, "people": people, "count": len(people)}
    except Exception as e:
        logger.warning(f"[PEOPLE] company={company_id} people failed: {e}")
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


# ---------------------------------------------------------------------------
# Static Files for Documents (with S3 fallback)
# ---------------------------------------------------------------------------
documents_dir = Path("/documents")


_MIME_BY_EXT = {
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".pdf": "application/pdf",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".txt": "text/plain",
    ".csv": "text/csv",
    ".json": "application/json",
}


def _file_response(path: Path, filename: str | None = None) -> FileResponse:
    """FileResponse with correct MIME and Content-Disposition so browsers don't append .txt."""
    ext = path.suffix.lower()
    media_type = _MIME_BY_EXT.get(ext, "application/octet-stream")
    return FileResponse(
        path,
        media_type=media_type,
        filename=filename or path.name,
    )


@app.get("/documents/legal/{subdir}/{filename}")
async def serve_document_with_s3_fallback(subdir: str, filename: str):
    """Serve document files — local first, S3 fallback if missing."""
    from fastapi.responses import JSONResponse
    base_dir = (documents_dir / "legal" / subdir).resolve()
    local_path = (documents_dir / "legal" / subdir / filename).resolve()
    if not str(local_path).startswith(str(base_dir)):
        return JSONResponse(status_code=400, content={"error": "Invalid filename"})
    if local_path.exists():
        return _file_response(local_path, filename)

    # Try S3 fallback
    if is_s3_enabled():
        s3_key = f"{subdir}/{filename}"
        if s3_download(s3_key, str(local_path)):
            return _file_response(local_path, filename)

    return JSONResponse(status_code=404, content={"error": "File not found"})


if documents_dir.exists():
    app.mount("/documents", StaticFiles(directory=str(documents_dir)), name="documents")

# Mount HTML templates
templates_dir = Path(__file__).parent / "templates"
if templates_dir.exists():
    app.mount("/static", StaticFiles(directory=str(templates_dir)), name="templates")


# ---------------------------------------------------------------------------
# Dashboard API Endpoints
# ---------------------------------------------------------------------------
from scout.tools.document_tracker import (
    get_all_documents,
    get_document_stats,
)
from scout.tools.template_analyzer import list_analyzed_templates
from scout.tools.smart_doc import analyze_template
from scout.tools.template_analyzer import analyze_template as ai_analyze_template, save_template_knowledge


@app.get("/api/dashboard/stats")
async def dashboard_stats():
    """Get dashboard statistics including template, company, and document counts."""
    stats = get_document_stats()
    try:
        from psycopg import connect
        import os
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM templates")
        stats["templates"] = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM companies")
        stats["companies"] = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM documents")
        stats["documents"] = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM knowledge_vec WHERE embedding IS NOT NULL")
        stats["embeddings"] = cur.fetchone()[0]
        cur.close(); conn.close()
    except Exception:
        stats.setdefault("templates", 0)
        stats.setdefault("companies", 0)
        stats.setdefault("documents", 0)
        stats.setdefault("embeddings", 0)
    return stats


@app.get("/api/dashboard/documents")
async def list_documents(limit: int = 50):
    """List all tracked documents."""
    return {"documents": get_all_documents(limit)}


@app.get("/api/dashboard/templates")
async def list_available_templates():
    """List all available templates."""
    return {"templates": list_analyzed_templates()}


def load_field_mapping(template_name: str) -> dict:
    """Load a stored field_mapping, upgraded to the slot contract."""
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT field_mapping FROM templates WHERE name = %s", (template_name,))
        row = cur.fetchone()
        cur.close()
        if not row or not row[0]:
            return {}
        raw = row[0] if isinstance(row[0], dict) else json.loads(row[0])
        return normalise_mapping(raw)
    except Exception as e:
        logging.getLogger("legalscout").warning(f"Failed to load field mapping for '{template_name}': {e}")
        return {}
    finally:
        if conn:
            conn.close()


@app.get("/api/dashboard/templates/{template_name}")
async def get_template_info(template_name: str):
    """Get template details."""
    base_dir = Path("/documents/legal/templates").resolve()
    safe_path = (base_dir / template_name).resolve()
    if not str(safe_path).startswith(str(base_dir)):
        return {"error": "Invalid filename"}
    info = analyze_template(template_name)
    if isinstance(info, dict):
        info["field_mapping"] = load_field_mapping(template_name)
    return info


@app.get("/api/templates/field-mapping/{template_name}")
async def get_template_field_mapping(template_name: str):
    """Get a template's field_mapping in slot-contract form."""
    return {"template": template_name, "field_mapping": load_field_mapping(template_name)}


@app.get("/api/templates/categories")
async def get_template_categories():
    """Get all template categories."""
    cat_file = Path("/documents/legal/data/template_categories.json")
    if cat_file.exists():
        with open(cat_file) as f:
            return json.load(f)
    return {"categories": [], "template_categories": {}}


@app.post("/api/templates/categories")
async def set_template_category(request: dict):
    """Set category for a template."""
    template_name = request.get("template_name")
    category = request.get("category")

    cat_file = Path("/documents/legal/data/template_categories.json")
    data = {"categories": [], "template_categories": {}}

    if cat_file.exists():
        with open(cat_file) as f:
            data = json.load(f)

    data["template_categories"][template_name] = category

    with open(cat_file, "w") as f:
        json.dump(data, f, indent=2)

    return {"success": True}


@app.post("/api/templates/group")
async def set_template_group_endpoint(request: Request, body: dict):
    """Tag a template with a group (e.g. 'new_company_setup') or clear it with null.

    Lets admins add newly uploaded consent forms to the new-company setup set, so
    the agent offers only setup documents when a new company is being set up.
    """
    require_admin(request)
    name = body.get("template_name") or body.get("name")
    group = body.get("template_group") or body.get("group")  # None/"" clears the tag
    if not name:
        return {"success": False, "error": "Template name required"}
    from scout.tools.template_analyzer import set_template_group
    ok = set_template_group(name, group or None)
    return {"success": ok, "template_name": name, "template_group": group or None}


@app.post("/api/templates/upload")
async def upload_template(request: Request, file: UploadFile = File(...)):
    """Upload a new template."""
    require_admin(request)
    templates_dir = Path("/documents/legal/templates")
    templates_dir.mkdir(parents=True, exist_ok=True)

    try:
        # Read in chunks, abort if too large
        MAX_SIZE = 50 * 1024 * 1024
        chunks = []
        size = 0
        while True:
            chunk = await file.read(8192)
            if not chunk:
                break
            size += len(chunk)
            if size > MAX_SIZE:
                return JSONResponse(status_code=413, content={"error": f"File too large (max {MAX_SIZE // 1024 // 1024}MB)"})
            chunks.append(chunk)
        content = b"".join(chunks)
        filename = file.filename or "template.docx"

        if not filename.endswith(".docx"):
            return {"success": False, "error": "Only .docx files are supported"}

        file_path = templates_dir / filename

        with open(file_path, "wb") as f:
            f.write(content)

        result = ai_analyze_template(filename)

        if result.get("success"):
            save_template_knowledge(filename, result.get("fields", []))
            return {
                "success": True,
                "template_name": filename,
                "fields": result.get("fields", []),
                "analyze": False,
            }

        return {"success": True, "template_name": filename, "analyze": True}

    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/templates/analyze")
async def analyze_template_endpoint(request: Request, name: str = Form(...)):
    """Analyze a template and extract fields."""
    require_admin(request)
    result = ai_analyze_template(name)

    if result.get("success"):
        save_template_knowledge(name, result.get("fields", []))

    return result


@app.delete("/api/templates/delete")
async def delete_template(request: Request, name: str = None):
    """Delete a template file."""
    require_admin(request)
    if not name:
        return {"success": False, "error": "Template name required"}

    template_path = Path("/documents/legal/templates") / name
    deleted_file = False

    try:
        if template_path.exists():
            template_path.unlink()
            deleted_file = True

        # Delete from S3
        s3_delete_async(f"templates/{name}")

        # Also delete from database
        try:
            import os
            from psycopg import connect

            conn = get_db_conn()
            conn.autocommit = True
            cur = conn.cursor()
            # Delete template record
            cur.execute("DELETE FROM templates WHERE name = %s", (name,))
            # Clean up ALL knowledge data for this template
            cur.execute("DELETE FROM knowledge_vec WHERE source_file = %s", (f"template:{name}",))
            cur.execute("DELETE FROM knowledge_lookup WHERE source_file = %s", (f"template:{name}",))
            try: cur.execute("DELETE FROM knowledge_raw WHERE source_file = %s", (f"template:{name}",))
            except Exception: pass
            # Clean agent memory/learnings about this template
            for mem_table in ["agno_memories", "agno_learnings", "scout_learnings"]:
                try: cur.execute(f"DELETE FROM {mem_table} WHERE content ILIKE %s OR content ILIKE %s",
                        (f"%{name}%", f"%{name.replace('.docx','')}%"))
                except Exception: pass
            # Clean scout_knowledge entries
            try: cur.execute("DELETE FROM scout_knowledge WHERE content ILIKE %s", (f"%{name}%",))
            except Exception: pass
            try: cur.execute("DELETE FROM scout_knowledge_contents WHERE content ILIKE %s", (f"%{name}%",))
            except Exception: pass
            cur.close()
            conn.close()
        except Exception as db_err:
            print(f"Database delete error: {db_err}")

        # Delete PDF preview
        pdf_path = Path("/documents/legal/previews") / name.replace(".docx", ".pdf")
        if pdf_path.exists():
            pdf_path.unlink()

        invalidate_training(f"Template deleted: {name}", template_name=name)
        clear_cache()  # Clear dashboard cache

        # Refresh agent's system prompt so it forgets this template
        try:
            _refresh_agent_knowledge()
            logger.info(f"Agent instructions refreshed after deleting {name}")
        except Exception as _ref_err:
            logger.warning(f"Agent refresh failed: {_ref_err}")

        if deleted_file:
            return {"success": True, "message": "Template deleted. Re-training required.", "training_invalidated": True}
        else:
            return {"success": True, "message": "Template removed from database. Re-training required.", "training_invalidated": True}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/templates/preview/{template_name}")
async def preview_template(request: Request, template_name: str):
    """Convert Word document to HTML for preview."""
    get_current_user(request)
    from docx import Document

    base_dir = Path("/documents/legal/templates").resolve()
    template_path = (base_dir / template_name).resolve()
    if not str(template_path).startswith(str(base_dir)):
        return {"error": "Invalid filename"}

    if not template_path.exists():
        return {"error": "Template not found"}

    try:
        doc = Document(str(template_path))

        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body { font-family: Arial, sans-serif; padding: 20px; max-width: 800px; margin: 0 auto; }
                p { margin: 10px 0; }
                table { border-collapse: collapse; width: 100%; margin: 10px 0; }
                td, th { border: 1px solid #ddd; padding: 8px; }
                .placeholder { background-color: #e8f5e9; padding: 2px 6px; border-radius: 4px; color: #2e7d32; font-weight: bold; }
            </style>
        </head>
        <body>
        """

        for para in doc.paragraphs:
            text = para.text
            if text.strip():
                import re

                text = re.sub(
                    r"\{\{([^}]+)\}\}|\{([^}]+)\}|\[([^\]]+)\]", r'<span class="placeholder">\1\2\3</span>', text
                )
                html_content += f"<p>{text}</p>"

        for table in doc.tables:
            html_content += "<table>"
            for row in table.rows:
                html_content += "<tr>"
                for cell in row.cells:
                    html_content += f"<td>{cell.text}</td>"
                html_content += "</tr>"
            html_content += "</table>"

        html_content += """
        </body>
        </html>
        """

        return {"html": html_content, "success": True}

    except Exception as e:
        return {"error": str(e)}


# ---------------------------------------------------------------------------
# Knowledge Base API Endpoints
# ---------------------------------------------------------------------------
from scout.tools.knowledge_base import (
    process_file,
    get_knowledge_sources,
    search_knowledge,
    lookup_value,
    get_source_data,
)

knowledge_dir = Path("/documents/legal/knowledge")
knowledge_dir.mkdir(parents=True, exist_ok=True)


@app.post("/api/knowledge/upload")
async def upload_knowledge(request: Request, file: UploadFile = File(...)):
    """Upload knowledge file (Excel, CSV, Word)."""
    require_admin(request)
    try:
        filename = file.filename or "knowledge.xlsx"
        file_path = knowledge_dir / filename

        content = await file.read()
        if len(content) > 50 * 1024 * 1024:
            return {"success": False, "error": "File too large (max 50MB)"}

        allowed_ext = [".xlsx", ".xls", ".csv", ".docx"]
        if not any(filename.lower().endswith(ext) for ext in allowed_ext):
            return {
                "success": False,
                "error": "Only Excel (.xlsx, .xls), CSV (.csv), and Word (.docx) files are supported",
            }

        with open(file_path, "wb") as f:
            f.write(content)

        result = process_file(str(file_path), filename)

        return result

    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/knowledge/sources")
async def list_knowledge_sources(request: Request):
    """List all uploaded knowledge sources."""
    require_admin(request)
    sources = get_knowledge_sources()
    return {"sources": sources}


@app.post("/api/knowledge/sync/companies")
async def sync_companies_to_knowledge(request: Request):
    """Sync companies from DB to knowledge base."""
    require_admin(request)
    try:
        from scout.tools.knowledge_base import store_cleaned_data, get_db_connection
        import os, json
        from psycopg import connect

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT company_name_english, company_registration_number, registered_office_address,
                   directors, status, company_type, principal_activity,
                   total_shares_issued, currency_of_share_capital,
                   ultimate_holding_company_name, members
            FROM companies ORDER BY company_name_english
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()

        if not rows:
            return {"success": False, "error": "No companies in database. Add companies first."}

        # Clean ALL old company data first
        conn2 = get_db_conn()
        cur2 = conn2.cursor()
        cur2.execute("DELETE FROM knowledge_vec WHERE source_file LIKE '%compan%'")
        cur2.execute("DELETE FROM knowledge_raw WHERE source_file LIKE '%compan%'")
        cur2.execute("DELETE FROM knowledge_lookup WHERE source_file LIKE '%compan%'")
        cur2.execute("DELETE FROM knowledge_sources WHERE filename LIKE '%compan%'")
        conn2.commit(); cur2.close(); conn2.close()

        # Build records for knowledge base
        records = []
        for row in rows:
            dirs = row[3] if isinstance(row[3], list) else []
            mems = row[10] if isinstance(row[10], list) else []
            records.append({
                "company_name": row[0] or "",
                "company_registration_number": row[1] or "",
                "registered_office": row[2] or "",
                "directors": ", ".join(d.get("name", "") for d in dirs) if dirs else "",
                "status": row[4] or "",
                "company_type": row[5] or "",
                "principal_activity": row[6] or "",
                "total_shares": row[7] or "",
                "currency": row[8] or "",
                "ultimate_holding_company": row[9] or "",
                "shareholders": ", ".join(m.get("name", "") for m in mems) if mems else "",
            })

        store_cleaned_data("companies_db", records, "database")

        # AI Company Analysis — Claude Haiku generates insights per company
        openrouter_key = os.getenv("OPENROUTER_API_KEY")
        if openrouter_key:
            import httpx
            conn3 = get_db_conn()
            cur3 = conn3.cursor()
            cur3.execute("""
                SELECT id, company_name_english, company_registration_number, registered_office_address,
                       directors, members, status, company_type, principal_activity,
                       total_shares_issued, currency_of_share_capital,
                       ultimate_holding_company_name, date_of_last_annual_return,
                       registration_date, filing_history, foreign_company, small_company,
                       under_corpsec_management, group_company
                FROM companies
            """)
            company_rows = cur3.fetchall()

            for crow in company_rows:
                try:
                    cname = crow[1] or "Unknown"
                    dirs = crow[4] if isinstance(crow[4], list) else []
                    mems = crow[5] if isinstance(crow[5], list) else []
                    filings = crow[14] if isinstance(crow[14], list) else []

                    company_text = f"""Company: {cname}
Reg: {crow[2]}, Status: {crow[6]}, Type: {crow[7]}
Activity: {crow[8]}, Foreign: {crow[15]}, Group: {crow[18]}
Directors: {json.dumps(dirs)}
Members/Shareholders: {json.dumps(mems)}
Shares: {crow[9]} {crow[10]}, Holding: {crow[11]}
Registration: {crow[13]}, Last AR: {crow[12]}
Filings: {len(filings)} records"""

                    ai_prompt = f"""Analyze this Myanmar company and generate insights:

{company_text}

Return JSON:
{{
  "profile_summary": "2-3 sentence company profile",
  "shareholder_structure": {{"majority_holder": "", "ownership_type": "sole/multiple", "notes": ""}},
  "compliance_status": {{"annual_return_status": "up to date/overdue/unknown", "next_filing_due": "", "notes": ""}},
  "key_dates": {{"incorporation": "", "financial_year_end": "", "next_agm_deadline": "", "next_ar_deadline": ""}},
  "missing_information": ["List critical missing fields"],
  "company_relationships": {{"parent": "", "subsidiaries": [], "group": ""}},
  "industry_requirements": "Any industry-specific compliance notes",
  "contact_info": {{"phone": "", "email": "", "source": ""}},
  "risk_flags": ["Any compliance or data quality concerns"]
}}
Return ONLY JSON."""

                    ai_res = httpx.post(
                        f"{OPENROUTER_BASE_URL}/chat/completions",
                        headers={"Authorization": f"Bearer {openrouter_key}", "Content-Type": "application/json"},
                        json={"model": get_model("training"), "messages": [{"role": "user", "content": ai_prompt}], "temperature": 0},
                        timeout=60,
                    )
                    ai_res.raise_for_status()
                    ai_text = ai_res.json()["choices"][0]["message"]["content"].strip()
                    if ai_text.startswith("```"):
                        ai_text = ai_text.split("```")[1]
                        if ai_text.startswith("json"): ai_text = ai_text[4:]
                        ai_text = ai_text.strip()
                    company_analysis = json.loads(ai_text)

                    # Store in knowledge_vec for semantic search
                    profile = company_analysis.get("profile_summary", "")
                    knowledge_text = f"Company: {cname}\n{profile}\nDirectors: {', '.join(d.get('name','') for d in dirs)}\n{company_text}"

                    cur3.execute("DELETE FROM knowledge_vec WHERE source_file = %s", (f"company:{cname}",))
                    cur3.execute("INSERT INTO knowledge_vec (content, source_file, metadata) VALUES (%s, %s, %s)",
                        (knowledge_text, f"company:{cname}", json.dumps(company_analysis)))
                    cur3.execute("DELETE FROM knowledge_raw WHERE source_file = %s", (f"company:{cname}",))
                    cur3.execute("INSERT INTO knowledge_raw (source_file, file_type, data) VALUES (%s, %s, %s)",
                        (f"company:{cname}", "company_analysis", json.dumps(company_analysis)))
                    conn3.commit()

                    # Generate vector embedding for company
                    _c_vec = generate_embedding(knowledge_text)
                    if _c_vec:
                        cur3.execute("UPDATE knowledge_vec SET embedding = %s WHERE source_file = %s",
                            (str(_c_vec), f"company:{cname}"))
                        conn3.commit()

                    print(f"  ✓ Company: {cname} — {profile[:50]}")
                except Exception as ce:
                    print(f"  Company analysis skip {crow[1]}: {ce}")

            cur3.close(); conn3.close()

        # Save training timestamp
        try:
            conn2 = get_db_conn()
            cur2 = conn2.cursor()
            cur2.execute("INSERT INTO training_status (training_type, last_trained, record_count) VALUES ('companies', NOW(), %s) ON CONFLICT (training_type) DO UPDATE SET last_trained = NOW(), record_count = %s", (len(records), len(records)))
            conn2.commit(); cur2.close(); conn2.close()
        except Exception: pass

        return {"success": True, "message": f"Synced {len(records)} companies to knowledge base"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/knowledge/train-companies-stream")
async def train_companies_stream(request: Request):
    """Stream company training progress via SSE — real per-company analysis."""
    require_admin(request)
    from starlette.responses import StreamingResponse
    import httpx, os

    def _sse(step, msg, **kw):
        d = {"step": step, "msg": msg, **kw}
        return f"data: {json.dumps(d)}\n\n"

    def generate():
        conn = None
        try:
            yield _sse("start", "Starting company deep training...")

            conn = get_db_conn()
            cur = conn.cursor()
            cur.execute("""
                SELECT id, company_name_english, company_registration_number, registered_office_address,
                       directors, members, status, company_type, principal_activity,
                       total_shares_issued, currency_of_share_capital,
                       ultimate_holding_company_name, date_of_last_annual_return,
                       registration_date, filing_history, foreign_company
                FROM companies ORDER BY company_name_english
            """)
            rows = cur.fetchall()
            cur.close(); conn.close(); conn = None

            if not rows:
                yield _sse("error", "No companies in database")
                return

            total = len(rows)
            yield _sse("load", f"Loaded {total} companies from database", count=total)

            # Step 1: Sync to knowledge_lookup
            yield _sse("sync_start", "Syncing to knowledge lookup table...")
            try:
                from scout.tools.knowledge_base import store_cleaned_data
                records = []
                for row in rows:
                    dirs = row[4] if isinstance(row[4], list) else []
                    mems = row[5] if isinstance(row[5], list) else []
                    records.append({
                        "company_name": row[1] or "", "company_registration_number": row[2] or "",
                        "registered_office": row[3] or "",
                        "directors": ", ".join(d.get("name","") for d in dirs) if dirs else "",
                        "status": row[6] or "", "company_type": row[7] or "",
                        "principal_activity": row[8] or "",
                        "total_shares": row[9] or "", "currency": row[10] or "",
                        "shareholders": ", ".join(m.get("name","") for m in mems) if mems else "",
                    })
                # Clean old data
                _c = get_db_conn(); _c.autocommit = True; _cc = _c.cursor()
                _cc.execute("DELETE FROM knowledge_vec WHERE source_file LIKE '%compan%'")
                _cc.execute("DELETE FROM knowledge_raw WHERE source_file LIKE '%compan%'")
                _cc.execute("DELETE FROM knowledge_lookup WHERE source_file LIKE '%compan%'")
                _cc.close(); _c.close()
                store_cleaned_data("companies_db", records, "database")
                yield _sse("sync", f"Synced {len(records)} companies to knowledge base")
            except Exception as e:
                yield _sse("sync_warn", f"Sync warning: {e}")

            # Step 2: Per-company AI analysis
            openrouter_key = os.getenv("OPENROUTER_API_KEY", "")
            training_model = get_model("training")
            analyzed = 0

            if openrouter_key:
                yield _sse("ai_start", f"Starting AI analysis for {total} companies...")
                for idx, row in enumerate(rows):
                    cname = row[1] or "Unknown"
                    yield _sse("company_start", f"[{idx+1}/{total}] Analyzing {cname}...", company=cname, index=idx+1)

                    try:
                        dirs = row[4] if isinstance(row[4], list) else []
                        mems = row[5] if isinstance(row[5], list) else []
                        filings = row[14] if isinstance(row[14], list) else []

                        company_text = f"Company: {cname}\nReg: {row[2]}, Status: {row[6]}, Type: {row[7]}\nActivity: {row[8]}\nDirectors: {json.dumps(dirs)}\nMembers: {json.dumps(mems)}\nShares: {row[9]} {row[10]}"

                        ai_prompt = f"""Analyze this Myanmar company and generate insights:

{company_text}

Return JSON:
{{
  "profile_summary": "2-3 sentence company profile",
  "shareholder_structure": {{"majority_holder": "", "ownership_type": "sole/multiple"}},
  "compliance_status": {{"annual_return_status": "up to date/overdue/unknown", "notes": ""}},
  "missing_information": ["List critical missing fields"],
  "risk_flags": ["Any compliance or data quality concerns"]
}}
Return ONLY JSON."""

                        ai_res = httpx.post(
                            f"{OPENROUTER_BASE_URL}/chat/completions",
                            headers={"Authorization": f"Bearer {openrouter_key}", "Content-Type": "application/json"},
                            json={"model": training_model, "messages": [{"role": "user", "content": ai_prompt}], "temperature": 0},
                            timeout=60)
                        ai_res.raise_for_status()
                        ai_text = ai_res.json()["choices"][0]["message"]["content"].strip()
                        if ai_text.startswith("```"):
                            ai_text = ai_text.split("```")[1]
                            if ai_text.startswith("json"): ai_text = ai_text[4:]
                            ai_text = ai_text.strip()
                        analysis = json.loads(ai_text)

                        profile = analysis.get("profile_summary", "")
                        knowledge_text = f"Company: {cname}\n{profile}\nDirectors: {', '.join(d.get('name','') for d in dirs)}"

                        _ac = get_db_conn(); _ac.autocommit = True; _acc = _ac.cursor()
                        _acc.execute("DELETE FROM knowledge_vec WHERE source_file = %s", (f"company:{cname}",))
                        _acc.execute("INSERT INTO knowledge_vec (content, source_file, metadata) VALUES (%s, %s, %s)",
                            (knowledge_text, f"company:{cname}", json.dumps(analysis)))
                        _acc.close(); _ac.close()

                        # Emit details
                        risks = analysis.get("risk_flags", [])
                        missing = analysis.get("missing_information", [])
                        compliance = analysis.get("compliance_status", {}).get("annual_return_status", "unknown")
                        yield _sse("company_done", f"  ✓ {cname} — {profile[:60]}",
                            company=cname, profile=profile[:80], compliance=compliance,
                            risks=len(risks), missing=len(missing))
                        analyzed += 1

                    except Exception as ce:
                        yield _sse("company_warn", f"  ⚠ {cname}: {str(ce)[:50]}", company=cname)

                yield _sse("ai_done", f"AI analysis complete: {analyzed}/{total} companies")
            else:
                yield _sse("ai_skip", "Skipped AI analysis: no API key configured")

            # Step 3: Embeddings
            yield _sse("embed_start", "Generating vector embeddings...")
            try:
                embed_count = 0
                _ec = get_db_conn(); _ecc = _ec.cursor()
                _ecc.execute("SELECT source_file, content FROM knowledge_vec WHERE source_file LIKE 'company:%' AND (embedding IS NULL OR embedding = '')")
                to_embed = _ecc.fetchall()
                for sf, content in to_embed:
                    vec = generate_embedding(content)
                    if vec:
                        _ecc.execute("UPDATE knowledge_vec SET embedding = %s WHERE source_file = %s", (str(vec), sf))
                        embed_count += 1
                _ec.commit(); _ecc.close(); _ec.close()
                yield _sse("embed", f"Generated {embed_count} embeddings")
            except Exception as ee:
                yield _sse("embed_warn", f"Embedding warning: {ee}")

            # Step 4: Save training status
            try:
                _tc = get_db_conn(); _tc.autocommit = True; _tcc = _tc.cursor()
                _tcc.execute("INSERT INTO training_status (training_type, last_trained, record_count) VALUES ('companies', NOW(), %s) ON CONFLICT (training_type) DO UPDATE SET last_trained = NOW(), record_count = %s", (total, total))
                _tcc.close(); _tc.close()
            except Exception as e:
                logging.getLogger("legalscout").warning(f"Training status save failed: {e}")

            # Summary
            yield _sse("summary", f"Training complete: {total} companies, {analyzed} analyzed", total=total, analyzed=analyzed)
            yield _sse("done", "Complete")

        except Exception as e:
            yield _sse("error", str(e))
        finally:
            if conn:
                try: conn.close()
                except: pass

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.get("/api/training/status")
async def get_training_status():
    """Get training status + persisted logs for all types."""
    try:
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT training_type, last_trained, record_count, status, logs FROM training_status ORDER BY training_type")
        rows = cur.fetchall()
        cur.close(); conn.close()
        result = {}
        for row in rows:
            result[row[0]] = {
                "last_trained": row[1].isoformat() if row[1] else None,
                "record_count": row[2],
                "status": row[3],
                "logs": row[4] if row[4] else []
            }
        return {"success": True, "data": result}
    except Exception as e:
        return {"success": True, "data": {}}


@app.post("/api/training/save-logs")
async def save_training_logs(request: Request):
    """Save training logs to DB so they persist across page reloads."""
    try:
        import os, json
        from psycopg import connect
        body = await request.json()
        training_type = body.get("type", "")  # "templates" or "companies"
        logs = body.get("logs", [])
        if not training_type:
            return {"success": False, "error": "type required"}

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO training_status (training_type, logs, last_trained) VALUES (%s, %s, NOW())
            ON CONFLICT (training_type) DO UPDATE SET logs = %s, last_trained = NOW()
        """, (training_type, json.dumps(logs), json.dumps(logs)))
        conn.commit(); cur.close(); conn.close()
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/knowledge/sync/templates")
async def sync_templates_to_knowledge():
    """Sync template placeholders to knowledge base."""
    try:
        from scout.tools.knowledge_base import store_cleaned_data
        from scout.tools.template_analyzer import analyze_template

        templates_dir = Path("/documents/legal/templates")
        if not templates_dir.exists():
            return {"success": False, "error": "Templates directory not found"}

        records = []
        for template_file in templates_dir.glob("*.docx"):
            result = analyze_template(template_file.name)
            fields = result.get("fields", [])

            records.append(
                {"template_name": template_file.name, "fields": ", ".join(fields), "field_count": len(fields)}
            )

        if records:
            store_cleaned_data("templates.xlsx", records, "excel")

        return {"success": True, "message": f"Synced {len(records)} templates to knowledge base"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.delete("/api/dashboard/document/{doc_id}")
async def delete_document(request: Request, doc_id: str):
    """Delete a generated document by ID, synthetic doc_<mtime>, or filename."""
    require_admin(request)
    try:
        import urllib.parse, re

        doc_id_str = urllib.parse.unquote(doc_id)
        conn = get_db_conn()
        conn.autocommit = True
        cur = conn.cursor()

        file_name = None
        output_base = Path("/documents/legal/output").resolve()

        # 1. Numeric DB id
        if doc_id_str.isdigit():
            cur.execute("SELECT file_name FROM documents WHERE id = %s", (int(doc_id_str),))
            row = cur.fetchone()
            if row:
                file_name = row[0]
                cur.execute("DELETE FROM documents WHERE id = %s", (int(doc_id_str),))

        # 2. Synthetic id "doc_<mtime>" — disk-only file. Find by mtime match.
        if not file_name:
            m = re.match(r"^doc_(\d+)$", doc_id_str)
            if m:
                target_mtime = int(m.group(1))
                if output_base.exists():
                    for f in output_base.glob("*.docx"):
                        if int(f.stat().st_mtime) == target_mtime:
                            file_name = f.name
                            break

        # 3. Direct filename match
        if not file_name:
            file_name = doc_id_str
            cur.execute("DELETE FROM documents WHERE file_name = %s", (file_name,))

        cur.close(); conn.close()

        # Delete file from disk
        deleted = False
        if file_name:
            file_path = (output_base / file_name).resolve()
            if str(file_path).startswith(str(output_base)) and file_path.exists():
                file_path.unlink()
                deleted = True

        # Delete from S3 if enabled
        try:
            from app.s3_storage import s3_delete_async
            if file_name:
                s3_delete_async(f"output/{file_name}")
        except Exception: pass

        clear_cache()
        if not deleted and not file_name:
            return {"success": False, "error": "Document not found"}
        if not deleted:
            return {"success": False, "error": f"DB row removed but file not found on disk: {file_name}"}
        return {"success": True, "message": f"Deleted {file_name}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/documents/sync")
async def sync_existing_documents():
    """Sync all generated documents to database for tracking."""
    try:
        from pathlib import Path
        from datetime import datetime
        import re
        import os
        from psycopg import connect

        output_dir = Path("/documents/legal/output")
        if not output_dir.exists():
            return {"success": False, "error": "Output directory not found"}

        conn = get_db_conn()
        cur = conn.cursor()

        count = 0
        for doc_file in output_dir.glob("*.docx"):
            # Parse filename pattern: TemplateName_CompanyName_2026-03-06_08-15-58.docx
            # or: Template_Name_Company_Name_2026-03-06_08-15-58.docx
            try:
                # Extract timestamp from end (always YYYY-MM-DD_HH-MM-SS.docx)
                timestamp_match = re.search(r'(\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2})\.docx$', doc_file.name)
                if not timestamp_match:
                    continue

                timestamp_str = timestamp_match.group(1)
                timestamp = datetime.strptime(timestamp_str, "%Y-%m-%d_%H-%M-%S")

                # Everything before timestamp is template_company
                base_name = doc_file.name[:timestamp_match.start()].rstrip('_')

                # Try to split into template and company (find last underscore)
                parts = base_name.rsplit('_', 1)
                if len(parts) == 2:
                    template_part, company_part = parts
                    template_name = template_part.replace('_', ' ') + ".docx"
                    company_name = company_part.replace('_', ' ')
                else:
                    # Can't split - use whole name
                    template_name = base_name.replace('_', ' ') + ".docx"
                    company_name = "Unknown"

                # Check if already tracked
                cur.execute("SELECT id FROM documents WHERE file_name = %s", (doc_file.name,))
                if cur.fetchone():
                    continue  # Already tracked

                # Insert into documents table
                cur.execute(
                    """
                    INSERT INTO documents (template_name, company_name, file_name, file_path, version, created_at)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """,
                    (template_name, company_name, doc_file.name, str(doc_file), 1, timestamp)
                )
                # Also insert into document_versions for audit trail
                cur.execute(
                    """
                    INSERT INTO document_versions (document_name, company_name, template_name, version, file_name, file_path, generated_by_email, created_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    """,
                    (doc_file.name, company_name, template_name, 1, doc_file.name, str(doc_file), "system@sync", timestamp)
                )
                count += 1

            except Exception as e:
                print(f"Error processing {doc_file.name}: {e}")
                continue

        conn.commit()
        cur.close()
        conn.close()

        return {"success": True, "message": f"Synced {count} documents to database"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/knowledge/train-stream/{template_name}")
async def train_single_template_stream(request: Request, template_name: str):
    """Train a single template with SSE streaming — sends each step as it happens."""
    require_admin(request)
    import urllib.parse, subprocess, json as _sj
    from starlette.responses import StreamingResponse
    template_name = urllib.parse.unquote(template_name)

    # Path traversal protection
    _tmpl_base = Path("/documents/legal/templates").resolve()
    _tmpl_safe = (_tmpl_base / template_name).resolve()
    if not str(_tmpl_safe).startswith(str(_tmpl_base)):
        from fastapi.responses import JSONResponse
        return JSONResponse(status_code=400, content={"error": "Invalid filename"})

    def _sse(step: str, msg: str, **extra):
        data = {"step": step, "msg": msg, **extra}
        return f"data: {_sj.dumps(data)}\n\n"

    def generate():
        conn = None
        try:
            from scout.tools.template_analyzer import analyze_template, classify_template_fields, get_db_connection
            from scout.tools.knowledge_base import get_db_connection as get_kb_conn
            from docx import Document
            import os

            template_path = _tmpl_safe
            if not template_path.exists():
                yield _sse("error", f"Template not found: {template_name}")
                return

            # Step 1: Extract placeholders
            yield _sse("extract_start", "Extracting placeholders...")
            result = analyze_template(template_name)
            fields = result.get("fields", [])
            yield _sse("extract", f"Extracted {len(fields)} placeholders", fields=fields)

            # Emit each placeholder individually
            for _fname in fields:
                yield _sse("field_found", f"  → {_fname}", field=_fname)

            # Step 2: Read document
            yield _sse("read_start", "Reading document content...")
            doc = Document(str(template_path))
            content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        content += "\n" + row_text
            yield _sse("read", f"Read {len(content)} chars of document content")

            # Step 3: AI Analysis
            openrouter_key = os.getenv("OPENROUTER_API_KEY")
            from app.model_config import get_model
            training_model = get_model("training")

            # Fallback defaults
            from app.main import _infer_category, _get_when_to_use, _get_how_to_use, _get_prerequisites, _get_filing_deadline, _get_fees, _get_validity_period, _get_approval_chain, _get_required_attachments, _get_common_mistakes, _get_industry_tags, _get_complexity, _get_estimated_time

            category = _infer_category(template_name)
            purpose = "Legal document template"
            when_to_use = _get_when_to_use(template_name)
            how_to_use = _get_how_to_use(template_name)
            prerequisites = _get_prerequisites(template_name)
            filing_deadline = _get_filing_deadline(template_name)
            fees = _get_fees(template_name)
            common_mistakes = _get_common_mistakes(template_name)
            legal_refs = []
            static_warnings = []
            related_templates = []
            workflow = {}
            field_details = {}
            required_fields = []
            optional_fields = []
            company_specific = []
            language_notes = ""
            agent_summary = ""
            signatures = {}
            regulatory = {}
            copies = {}
            quorum = {}
            stamp = {}

            ai_analysis = None
            if openrouter_key:
                yield _sse("ai_start", f"Sending to {training_model}...")
                try:
                    import httpx
                    ai_prompt = f"""You are a legal document analyst. Analyze this Myanmar legal document template.
TEMPLATE NAME: {template_name}
PLACEHOLDER FIELDS: {', '.join(fields)}
FULL TEMPLATE TEXT:
{content[:6000]}

Return ONLY a JSON object with: purpose, when_to_use, how_to_use (array), category, legal_references (array), legal_context, related_templates (array), workflow_sequence (before/after arrays), field_details (object), required_fields (array), optional_fields (array), static_text_warnings (array), prerequisites (array), filing_deadline, fees, common_mistakes (array), signatures_required (signers array), validity_period, regulatory_body (name), language_notes, summary_for_agent. Return ONLY JSON."""

                    ai_res = httpx.post(
                        f"{OPENROUTER_BASE_URL}/chat/completions",
                        headers={"Authorization": f"Bearer {openrouter_key}", "Content-Type": "application/json"},
                        json={"model": training_model, "messages": [{"role": "user", "content": ai_prompt}], "temperature": 0},
                        timeout=60,
                    )
                    ai_res.raise_for_status()
                    ai_text = ai_res.json()["choices"][0]["message"]["content"].strip()
                    if ai_text.startswith("```"):
                        ai_text = ai_text.split("```")[1]
                        if ai_text.startswith("json"): ai_text = ai_text[4:]
                        ai_text = ai_text.strip()
                    ai_analysis = _sj.loads(ai_text)
                    yield _sse("ai_done", f"AI analyzed successfully")
                except Exception as ai_err:
                    yield _sse("ai_warn", f"AI analysis warning: {ai_err}")

            if ai_analysis:
                category = ai_analysis.get("category", category)
                purpose = ai_analysis.get("purpose", purpose)
                when_to_use = ai_analysis.get("when_to_use", when_to_use)
                how_to_use = ai_analysis.get("how_to_use", how_to_use)
                prerequisites = ai_analysis.get("prerequisites", prerequisites)
                filing_deadline = ai_analysis.get("filing_deadline", filing_deadline)
                fees = ai_analysis.get("fees", fees) if ai_analysis.get("fees") else fees
                common_mistakes = ai_analysis.get("common_mistakes", common_mistakes)
                legal_refs = ai_analysis.get("legal_references", [])
                static_warnings = ai_analysis.get("static_text_warnings", [])
                related_templates = ai_analysis.get("related_templates", [])
                required_fields = ai_analysis.get("required_fields", [])
                optional_fields = ai_analysis.get("optional_fields", [])
                agent_summary = ai_analysis.get("summary_for_agent", "")
                language_notes = ai_analysis.get("language_notes", "")
                signatures = ai_analysis.get("signatures_required", {})

            yield _sse("ai_analysis", f"Category: {category}", purpose=purpose[:80],
                        legal_refs=legal_refs[:3], required=len(required_fields), optional=len(optional_fields))

            # Step 4: Save metadata
            yield _sse("save_start", "Saving metadata to database...")
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("""
                    UPDATE templates SET category=%s, purpose=%s, when_to_use=%s, how_to_use=%s,
                    prerequisites=%s, filing_deadline=%s, fees=%s, common_mistakes=%s,
                    ai_trained=TRUE, ai_analyzed=TRUE, jurisdiction=%s, complexity=%s, estimated_time=%s
                    WHERE name=%s
                """, (category, purpose, when_to_use, _sj.dumps(how_to_use) if isinstance(how_to_use, list) else how_to_use,
                      _sj.dumps(prerequisites) if isinstance(prerequisites, list) else prerequisites,
                      filing_deadline, fees,
                      _sj.dumps(common_mistakes) if isinstance(common_mistakes, list) else common_mistakes,
                      "Myanmar", _get_complexity(template_name), _get_estimated_time(template_name),
                      template_name))
                conn.commit(); cur.close(); conn.close(); conn = None
            except Exception as e:
                yield _sse("save_warn", f"Metadata save warning: {e}")
            yield _sse("metadata", f"Saved: {category}")

            # Step 5: Classify fields
            classification_model = get_model("classification")
            yield _sse("classify_start", f"Classifying fields via {classification_model}...")
            try:
                classification = classify_template_fields(content, fields)
                if classification:
                    conn2 = get_db_connection()
                    cur2 = conn2.cursor()
                    cur2.execute("UPDATE templates SET fields = %s WHERE name = %s",
                        (_sj.dumps(classification), template_name))
                    conn2.commit(); cur2.close(); conn2.close()
                    db_count = len(classification.get("db_fields", []))
                    user_count = len(classification.get("user_input_fields", []))
                    yield _sse("classify", f"DB fields: {db_count}, User input: {user_count}",
                              db_fields=classification.get("db_fields", []),
                              user_input_fields=classification.get("user_input_fields", []))
                else:
                    yield _sse("classify", "Classification returned empty")
            except Exception as e:
                yield _sse("classify_warn", f"Classification warning: {e}")

            # Step 5.5: Generate field_mapping (learned placeholder → DB column mapping)
            yield _sse("mapping_start", "Generating field mapping...")
            try:
                mapping_prompt = f"""You are a database mapping expert for a Myanmar corporate law document system. Map each template placeholder to the correct data source.

TEMPLATE: {template_name}
PLACEHOLDERS: {', '.join(fields)}

AVAILABLE DATABASE COLUMNS on the companies table (these describe the COMPANY, never a person):
- company_name_english (registered English name)
- company_name_myanmar (registered Myanmar name)
- company_registration_number (DICA registration number)
- registration_date (date of incorporation)
- status (active/inactive)
- company_type (private/public)
- foreign_company (yes/no)
- small_company (yes/no)
- principal_activity (business activity)
- registered_office_address (registered office address)
- principal_place_of_business (business address)
- date_of_last_annual_return (date)
- previous_registration_number (prior DICA number)
- total_shares_issued (total shares number)
- currency_of_share_capital (currency code)
- total_capital (total capital amount)
- consideration_amount_paid (consideration paid)
- under_corpsec_management (yes/no)
- group_company (yes/no)
- ultimate_holding_company_name (parent company)
- ultimate_holding_company_jurisdiction (parent jurisdiction)
- ultimate_holding_company_registration_number (parent registration number)
- financial_year_end_date (financial year end, free text)
- next_financial_year_end_date (next financial year end, free text)
- auditor_name (appointed auditor)
- auditor_fee (auditor remuneration)
- custom_fields (JSONB of per-company extras; address a key as custom_fields.some_key)

HARD RULES:
1. NEVER put an array index in db_column. "directors[0].name", "members[1].name" and any other
   "[n]" path are FORBIDDEN — they freeze one person at training time and produce the wrong signatory.
2. "directors", "members" and "shareholder_links" are person lists and are NEVER valid db_column values.
3. NEVER use "today" (or "now"/"current_date") as a default. Dates are chosen by the user, never auto-filled.
4. Any placeholder naming a PERSON or a SIGNING PARTY — signatory, chairperson, attendee, the director
   resigning, the director being appointed, a shareholder, a corporate representative, the auditor as a
   signing party — is a SLOT, not a db column and not plain user_input.

A SLOT means "a person or party the USER must choose when the document is generated". It is resolved
at generation time against the people register, so it must never be frozen to a position.

"slot": {{"kind": <kind>, "of": <of>, "multi": true|false}}

slot.kind is EXACTLY one of:
  signatory           — the person who signs the document
  attendee            — a person present at the meeting
  chairperson         — the person chairing the meeting
  resigning_director  — the director who is resigning
  new_director        — the director being appointed
  representative      — the individual signing on behalf of a corporate entity
  shareholder_list    — the list of shareholders
  auditor             — the auditor as a named party

slot.of is EXACTLY one of:
  document_company      — candidates are the people linked to the company the document is for
  corporate_shareholder — candidates are the directors of a CORPORATE SHAREHOLDER of the document company
                          (use with kind "representative" when a company signs through its own director)
  people_register       — candidates are any person in the people register

slot.multi is true when the placeholder holds several people (a list of attendees or shareholders),
false when it holds exactly one.

For EACH placeholder return a JSON object with:
- "source": "db" | "user_input" | "slot"
- "db_column": exact column name when source is "db", otherwise null. No array indexes, ever.
- "slot": the slot object when source is "slot", otherwise null
- "default": a literal fallback string for source "user_input", otherwise null. Never "today".
- "description": short description of what this field is

WORKED EXAMPLES:
db field:
  "company_name": {{"source":"db","db_column":"company_name_english","slot":null,"default":null,"description":"Company registered name"}}
user_input field:
  "meeting_date": {{"source":"user_input","db_column":null,"slot":null,"default":null,"description":"Date the meeting is held, supplied by the user"}}
single signatory slot:
  "director_name": {{"source":"slot","db_column":null,"slot":{{"kind":"signatory","of":"document_company","multi":false}},"default":null,"description":"Director who signs this document, chosen by the user"}}
multi attendee slot:
  "attendees": {{"source":"slot","db_column":null,"slot":{{"kind":"attendee","of":"document_company","multi":true}},"default":null,"description":"People present at the meeting, chosen by the user"}}

Return ONLY a JSON object where keys are placeholder names."""

                import httpx as _mhx
                _m_res = _mhx.post(
                    f"{OPENROUTER_BASE_URL}/chat/completions",
                    headers={"Authorization": f"Bearer {os.getenv('OPENROUTER_API_KEY','')}", "Content-Type": "application/json"},
                    json={"model": classification_model, "messages": [{"role": "user", "content": mapping_prompt}], "temperature": 0},
                    timeout=60,
                )
                _m_res.raise_for_status()
                _m_text = _m_res.json()["choices"][0]["message"]["content"].strip()
                if _m_text.startswith("```"):
                    _m_text = _m_text.split("```")[1]
                    if _m_text.startswith("json"): _m_text = _m_text[4:]
                    _m_text = _m_text.strip()
                field_mapping, _m_rejected, _m_repaired = sanitise_mapping(_sj.loads(_m_text))
                if _m_repaired:
                    logging.getLogger("legalscout").warning(
                        f"Field mapping for '{template_name}' repaired {len(_m_repaired)} entries: "
                        + "; ".join(_m_repaired))
                    yield _sse("mapping_repair", f"Repaired {len(_m_repaired)} off-contract mapping entries",
                              repaired=_m_repaired[:10])
                if _m_rejected:
                    logging.getLogger("legalscout").warning(
                        f"Field mapping for '{template_name}' dropped {len(_m_rejected)} invalid entries: "
                        + "; ".join(_m_rejected))
                    yield _sse("mapping_reject", f"Rejected {len(_m_rejected)} invalid mapping entries",
                              rejected=_m_rejected[:10])

                # Save to DB
                _mc = get_db_connection()
                _mcur = _mc.cursor()
                _mcur.execute("UPDATE templates SET field_mapping = %s WHERE name = %s",
                    (_sj.dumps(field_mapping), template_name))
                _mc.commit(); _mcur.close(); _mc.close()

                db_mapped = sum(1 for v in field_mapping.values() if v.get("source") == "db")
                user_mapped = sum(1 for v in field_mapping.values() if v.get("source") == "user_input")
                slot_mapped = sum(1 for v in field_mapping.values() if v.get("source") == "slot")
                yield _sse("mapping", f"Mapped {db_mapped} fields to DB, {slot_mapped} slots, {user_mapped} need user input",
                          db_mapped=db_mapped, user_mapped=user_mapped, slot_mapped=slot_mapped)
            except Exception as _me:
                yield _sse("mapping_warn", f"Field mapping warning: {_me}")

            # Step 6: Knowledge base storage
            yield _sse("kb_start", "Storing in knowledge base...")
            try:
                knowledge_text = f"Template: {template_name}\nCategory: {category}\nPurpose: {purpose}\nWhen: {when_to_use}\nFields: {', '.join(fields)}\nContent: {content[:2000]}"
                conn3 = get_kb_conn()
                cur3 = conn3.cursor()
                cur3.execute("DELETE FROM knowledge_vec WHERE source_file = %s", (f"template:{template_name}",))
                cur3.execute("DELETE FROM knowledge_lookup WHERE source_file = %s", (f"template:{template_name}",))
                cur3.execute("INSERT INTO knowledge_vec (content, source_file, metadata) VALUES (%s, %s, %s)",
                    (knowledge_text, f"template:{template_name}", _sj.dumps({"type": "template", "name": template_name, "category": category})))
                for kk, kv in [("template_name", template_name), ("template_category", category),
                               ("template_purpose", purpose), ("template_when_to_use", when_to_use)]:
                    if kv:
                        cur3.execute("INSERT INTO knowledge_lookup (key_name, key_value, source_file) VALUES (%s, %s, %s)",
                            (kk, kv, f"template:{template_name}"))
                conn3.commit(); cur3.close(); conn3.close()
                yield _sse("knowledge", "Stored in knowledge base (vector + lookup)")
            except Exception as e:
                yield _sse("kb_warn", f"Knowledge storage warning: {e}")

            # Step 7: Embedding
            yield _sse("embed_start", "Generating vector embedding...")
            try:
                embedding = generate_embedding(knowledge_text)
                if embedding:
                    conn4 = get_kb_conn()
                    cur4 = conn4.cursor()
                    cur4.execute("UPDATE knowledge_vec SET embedding = %s WHERE source_file = %s",
                        (str(embedding), f"template:{template_name}"))
                    conn4.commit(); cur4.close(); conn4.close()
                    yield _sse("embedding", f"Vector embedding generated ({len(embedding)} dimensions)")
                else:
                    yield _sse("embed_skip", "Embedding skipped: no API key")
            except Exception as e:
                yield _sse("embed_warn", f"Embedding warning: {e}")

            # Step 8: PDF with highlighted placeholders
            yield _sse("pdf_start", "Generating PDF preview with highlighted placeholders...")
            try:
                pdf_dir = Path("/documents/legal/previews")
                pdf_dir.mkdir(parents=True, exist_ok=True)
                # Delete old cached PDF to force regeneration with highlights
                _old_pdf = pdf_dir / (template_path.stem + ".pdf")
                if _old_pdf.exists():
                    _old_pdf.unlink()
                _hl_path = pdf_dir / f"_hl_{template_name}"
                try:
                    _highlight_placeholders_in_docx(template_path, _hl_path)
                    _conv_src = _hl_path
                except Exception:
                    _conv_src = template_path
                subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(_conv_src)],
                    capture_output=True, text=True, timeout=30)
                _hl_pdf = pdf_dir / (_conv_src.stem + ".pdf")
                _expected_pdf = pdf_dir / (template_path.stem + ".pdf")
                if _hl_pdf.exists() and _hl_pdf != _expected_pdf:
                    _hl_pdf.rename(_expected_pdf)
                if _hl_path.exists():
                    _hl_path.unlink(missing_ok=True)
                yield _sse("pdf", "PDF preview generated (placeholders highlighted)")
            except Exception:
                yield _sse("pdf_warn", "PDF conversion skipped")

            # ── Deep Training Steps 9–15 ──────────────────────────

            _openrouter_key = getenv("OPENROUTER_API_KEY", "")
            _training_model = get_model("training") if 'get_model' in dir() else "google/gemini-3.6-flash"

            # Step 9: Field-level deep analysis
            yield _sse("field_deep_start", "Analyzing each field in detail...")
            _field_deep = {}
            try:
                if _openrouter_key and fields:
                    import httpx as _fd_httpx
                    _fd_prompt = f"""You are a Myanmar corporate law document expert. Analyze each placeholder field in this template.

Template: {template_name}
Placeholders: {json.dumps(fields)}
Document content (first 2000 chars): {content[:2000]}

For EACH placeholder, return a JSON object keyed by field name:
{{
  "field_name": {{
    "data_type": "text|date|number|currency|name|address|nrc|phone|email",
    "format": "expected format or null (e.g. DD/MM/YYYY, RM X,XXX.XX)",
    "validation": "validation rule in plain English",
    "required": true/false,
    "default_logic": "how to get default value (e.g. 'From company database', 'Today\\'s date')",
    "description": "what this field represents in 1 sentence"
  }}
}}
Return ONLY the JSON object, no markdown."""
                    _fd_res = _fd_httpx.post(f"{OPENROUTER_BASE_URL}/chat/completions",
                        headers={"Authorization": f"Bearer {_openrouter_key}", "Content-Type": "application/json"},
                        json={"model": _training_model, "messages": [{"role": "user", "content": _fd_prompt}], "temperature": 0},
                        timeout=60)
                    _fd_res.raise_for_status()
                    _fd_text = _fd_res.json().get("choices", [{}])[0].get("message", {}).get("content", "{}")
                    _fd_text = _fd_text.strip().strip("`").strip()
                    if _fd_text.startswith("json"):
                        _fd_text = _fd_text[4:].strip()
                    _field_deep = json.loads(_fd_text)

                    # Save to DB
                    _fd_conn = get_db_conn(); _fd_conn.autocommit = True; _fd_cur = _fd_conn.cursor()
                    _fd_cur.execute("UPDATE templates SET field_deep_analysis = %s WHERE name = %s", (json.dumps(_field_deep), template_name))
                    _fd_cur.close(); _fd_conn.close()

                    # Emit per-field details
                    _db_fields_set = set(classification.get("db_fields", [])) if isinstance(classification, dict) else set()
                    for _fn, _fd in _field_deep.items():
                        _src = "DB auto-fill" if _fn in _db_fields_set else "user input"
                        yield _sse("field_detail", f"  → {_fn} ({_src}) — {_fd.get('data_type','text')}, {'required' if _fd.get('required') else 'optional'}", field=_fn, detail=_fd)
                    yield _sse("field_deep", f"Deep analysis complete for {len(_field_deep)} fields")
                else:
                    yield _sse("field_deep_warn", "Skipped: no API key or no fields")
            except Exception as _fde:
                yield _sse("field_deep_warn", f"Field analysis warning: {_fde}")

            # Step 10: Legal reference extraction
            yield _sse("legal_ref_start", "Extracting Myanmar legal references...")
            try:
                if _openrouter_key:
                    import httpx as _lr_httpx
                    _lr_prompt = f"""You are a Myanmar Companies Law 2017 expert. For this legal template, identify applicable law sections.

Template: {template_name}
Category: {category}
Purpose: {purpose}
Document content (first 1500 chars): {content[:1500]}

Return JSON:
{{
  "sections": [
    {{"section": "Section 166", "title": "Annual General Meeting", "relevance": "Requires company to hold AGM annually"}}
  ],
  "compliance": ["Must be filed within 14 days of AGM"],
  "filing_obligations": ["File with DICA via MyCO portal"]
}}
Return ONLY the JSON object."""
                    _lr_res = _lr_httpx.post(f"{OPENROUTER_BASE_URL}/chat/completions",
                        headers={"Authorization": f"Bearer {_openrouter_key}", "Content-Type": "application/json"},
                        json={"model": _training_model, "messages": [{"role": "user", "content": _lr_prompt}], "temperature": 0},
                        timeout=60)
                    _lr_res.raise_for_status()
                    _lr_text = _lr_res.json().get("choices", [{}])[0].get("message", {}).get("content", "{}")
                    _lr_text = _lr_text.strip().strip("`").strip()
                    if _lr_text.startswith("json"):
                        _lr_text = _lr_text[4:].strip()
                    _legal_data = json.loads(_lr_text)

                    _lr_conn = get_db_conn(); _lr_conn.autocommit = True; _lr_cur = _lr_conn.cursor()
                    _lr_cur.execute("UPDATE templates SET legal_references = %s WHERE name = %s", (json.dumps(_legal_data), template_name))
                    _lr_cur.close(); _lr_conn.close()

                    _sec_count = len(_legal_data.get("sections", []))
                    yield _sse("legal_ref", f"Found {_sec_count} legal sections from Myanmar Companies Law 2017")
                else:
                    yield _sse("legal_ref_warn", "Skipped: no API key")
            except Exception as _lre:
                yield _sse("legal_ref_warn", f"Legal reference warning: {_lre}")

            # Step 11: Sample filled document
            yield _sse("sample_start", "Generating sample filled document...")
            try:
                if _openrouter_key and fields:
                    import httpx as _sf_httpx
                    _sf_prompt = f"""Generate realistic sample values for a Myanmar legal document template.

Template: {template_name}
Category: {category}
Placeholders: {json.dumps(fields)}

Return a JSON object mapping each placeholder to a realistic sample value.
Use Myanmar company names, addresses, director names, dates in DD/MM/YYYY format.
Example: {{"company_name": "City Holdings Limited", "meeting_date": "15/03/2026", "director_name": "U Aung Kyaw"}}
Return ONLY the JSON object."""
                    _sf_res = _sf_httpx.post(f"{OPENROUTER_BASE_URL}/chat/completions",
                        headers={"Authorization": f"Bearer {_openrouter_key}", "Content-Type": "application/json"},
                        json={"model": _training_model, "messages": [{"role": "user", "content": _sf_prompt}], "temperature": 0.3},
                        timeout=60)
                    _sf_res.raise_for_status()
                    _sf_text = _sf_res.json().get("choices", [{}])[0].get("message", {}).get("content", "{}")
                    _sf_text = _sf_text.strip().strip("`").strip()
                    if _sf_text.startswith("json"):
                        _sf_text = _sf_text[4:].strip()
                    _sample = json.loads(_sf_text)

                    _sf_conn = get_db_conn(); _sf_conn.autocommit = True; _sf_cur = _sf_conn.cursor()
                    _sf_cur.execute("UPDATE templates SET sample_filled_document = %s WHERE name = %s", (json.dumps(_sample), template_name))
                    _sf_cur.close(); _sf_conn.close()
                    yield _sse("sample", f"Generated sample values for {len(_sample)} fields")
                else:
                    yield _sse("sample_warn", "Skipped: no API key or no fields")
            except Exception as _sfe:
                yield _sse("sample_warn", f"Sample generation warning: {_sfe}")

            # Step 12: Document workflow
            yield _sse("workflow_start", "Analyzing document workflow...")
            try:
                if _openrouter_key:
                    import httpx as _wf_httpx
                    _wf_prompt = f"""For this Myanmar corporate legal template, determine the document workflow.

Template: {template_name}
Category: {category}
Purpose: {purpose}

Return JSON:
{{
  "trigger": "Business event that triggers this document (e.g. 'Annual General Meeting held')",
  "before": ["Documents that must be prepared BEFORE this one"],
  "after": ["Documents that should be prepared AFTER this one"],
  "notes": "Any important workflow notes"
}}
Return ONLY the JSON object."""
                    _wf_res = _wf_httpx.post(f"{OPENROUTER_BASE_URL}/chat/completions",
                        headers={"Authorization": f"Bearer {_openrouter_key}", "Content-Type": "application/json"},
                        json={"model": _training_model, "messages": [{"role": "user", "content": _wf_prompt}], "temperature": 0},
                        timeout=60)
                    _wf_res.raise_for_status()
                    _wf_text = _wf_res.json().get("choices", [{}])[0].get("message", {}).get("content", "{}")
                    _wf_text = _wf_text.strip().strip("`").strip()
                    if _wf_text.startswith("json"):
                        _wf_text = _wf_text[4:].strip()
                    _workflow = json.loads(_wf_text)

                    _wf_conn = get_db_conn(); _wf_conn.autocommit = True; _wf_cur = _wf_conn.cursor()
                    _wf_cur.execute("UPDATE templates SET document_workflow = %s WHERE name = %s", (json.dumps(_workflow), template_name))
                    _wf_cur.close(); _wf_conn.close()
                    _before = len(_workflow.get("before", []))
                    _after = len(_workflow.get("after", []))
                    yield _sse("workflow", f"Workflow mapped: {_before} prerequisite(s), {_after} follow-up(s)")
                else:
                    yield _sse("workflow_warn", "Skipped: no API key")
            except Exception as _wfe:
                yield _sse("workflow_warn", f"Workflow warning: {_wfe}")

            # Step 13: Q&A pairs generation
            yield _sse("qa_start", "Generating Q&A pairs for agent knowledge...")
            try:
                if _openrouter_key:
                    import httpx as _qa_httpx
                    _qa_prompt = f"""Generate 10 practical Q&A pairs about this Myanmar legal document template.
These will be used by an AI assistant to answer user questions.

Template: {template_name}
Category: {category}
Purpose: {purpose}
Fields: {json.dumps(fields[:15])}

Return a JSON array:
[
  {{"question": "What do I need to prepare for an AGM?", "answer": "You need AGM Minutes template, list of shareholders..."}},
  ...
]
Make questions practical — what users would actually ask. Include questions about requirements, deadlines, common issues.
Return ONLY the JSON array."""
                    _qa_res = _qa_httpx.post(f"{OPENROUTER_BASE_URL}/chat/completions",
                        headers={"Authorization": f"Bearer {_openrouter_key}", "Content-Type": "application/json"},
                        json={"model": _training_model, "messages": [{"role": "user", "content": _qa_prompt}], "temperature": 0.3},
                        timeout=60)
                    _qa_res.raise_for_status()
                    _qa_text = _qa_res.json().get("choices", [{}])[0].get("message", {}).get("content", "[]")
                    _qa_text = _qa_text.strip().strip("`").strip()
                    if _qa_text.startswith("json"):
                        _qa_text = _qa_text[4:].strip()
                    _qa_pairs = json.loads(_qa_text)

                    if isinstance(_qa_pairs, list) and _qa_pairs:
                        _qa_conn = get_db_conn(); _qa_conn.autocommit = True; _qa_cur = _qa_conn.cursor()
                        # Clean old Q&A for this template
                        _qa_cur.execute("DELETE FROM knowledge_vec WHERE source_file = %s", (f"template_qa:{template_name}",))
                        for _qa in _qa_pairs:
                            _qa_content = f"Q: {_qa.get('question','')}\nA: {_qa.get('answer','')}"
                            _qa_cur.execute("INSERT INTO knowledge_vec (content, source_file, metadata) VALUES (%s, %s, %s)",
                                (_qa_content, f"template_qa:{template_name}", json.dumps({"type": "template_qa", "template": template_name})))
                        _qa_cur.close(); _qa_conn.close()
                        yield _sse("qa", f"Generated {len(_qa_pairs)} Q&A pairs for knowledge base")
                    else:
                        yield _sse("qa_warn", "No Q&A pairs generated")
                else:
                    yield _sse("qa_warn", "Skipped: no API key")
            except Exception as _qae:
                yield _sse("qa_warn", f"Q&A generation warning: {_qae}")

            # Step 14: Cross-template relationships
            yield _sse("cross_ref_start", "Mapping cross-template relationships...")
            try:
                if _openrouter_key:
                    import httpx as _cr_httpx
                    _cr_conn = get_db_conn(); _cr_cur = _cr_conn.cursor()
                    _cr_cur.execute("SELECT name, category, purpose FROM templates WHERE name != %s", (template_name,))
                    _other_templates = [{"name": r[0], "category": r[1] or "", "purpose": (r[2] or "")[:80]} for r in _cr_cur.fetchall()]
                    _cr_cur.close(); _cr_conn.close()

                    if _other_templates:
                        _cr_prompt = f"""Identify relationships between this template and others.

This template: {template_name} (Category: {category}, Purpose: {purpose})

Other templates in the system:
{json.dumps(_other_templates, indent=2)}

Return a JSON array of relationships:
[
  {{"template": "Board_Resolution.docx", "relationship": "prerequisite|follow_up|related|alternative", "description": "Board resolution authorizing this action"}}
]
Only include templates that have a real relationship. Return ONLY the JSON array."""
                        _cr_res = _cr_httpx.post(f"{OPENROUTER_BASE_URL}/chat/completions",
                            headers={"Authorization": f"Bearer {_openrouter_key}", "Content-Type": "application/json"},
                            json={"model": _training_model, "messages": [{"role": "user", "content": _cr_prompt}], "temperature": 0},
                            timeout=60)
                        _cr_res.raise_for_status()
                        _cr_text = _cr_res.json().get("choices", [{}])[0].get("message", {}).get("content", "[]")
                        _cr_text = _cr_text.strip().strip("`").strip()
                        if _cr_text.startswith("json"):
                            _cr_text = _cr_text[4:].strip()
                        _cross_refs = json.loads(_cr_text)

                        _cr_conn2 = get_db_conn(); _cr_conn2.autocommit = True; _cr_cur2 = _cr_conn2.cursor()
                        _cr_cur2.execute("UPDATE templates SET cross_template_relationships = %s WHERE name = %s", (json.dumps(_cross_refs), template_name))
                        _cr_cur2.close(); _cr_conn2.close()
                        yield _sse("cross_ref", f"Found {len(_cross_refs)} related template(s)")
                    else:
                        yield _sse("cross_ref_warn", "No other templates to compare")
                else:
                    yield _sse("cross_ref_warn", "Skipped: no API key")
            except Exception as _cre:
                yield _sse("cross_ref_warn", f"Cross-reference warning: {_cre}")

            # Step 15: Confidence scoring
            yield _sse("confidence_start", "Calculating training confidence score...")
            try:
                _score = 0
                _checks = [
                    (bool(purpose and purpose != "Legal document template"), 10, "purpose"),
                    (bool(when_to_use), 5, "when_to_use"),
                    (bool(category and category != "General"), 5, "category"),
                    (bool(fields), 10, "placeholders"),
                    (bool(isinstance(classification, dict) and classification.get("db_fields")), 10, "field_classification"),
                    (bool(_field_deep), 10, "field_deep_analysis"),
                    (bool(_legal_data if '_legal_data' in dir() else None), 10, "legal_references"),
                    (bool(_sample if '_sample' in dir() else None), 10, "sample_document"),
                    (bool(_workflow if '_workflow' in dir() else None), 10, "workflow"),
                    (bool(_qa_pairs if '_qa_pairs' in dir() else None), 10, "qa_pairs"),
                    (bool(_cross_refs if '_cross_refs' in dir() else None), 5, "cross_references"),
                    (True, 5, "base"),  # Always get base points for having the template
                ]
                _score = sum(w for c, w, _ in _checks if c)
                _passed = [n for c, _, n in _checks if c]

                _sc_conn = get_db_conn(); _sc_conn.autocommit = True; _sc_cur = _sc_conn.cursor()
                _sc_cur.execute("UPDATE templates SET training_confidence = %s WHERE name = %s", (_score, template_name))
                _sc_cur.close(); _sc_conn.close()
                yield _sse("confidence", f"Training confidence: {_score}% ({len(_passed)}/{len(_checks)} checks passed)", score=_score)
            except Exception as _sce:
                yield _sse("confidence_warn", f"Confidence scoring warning: {_sce}")

            try:
                _reg_count = _refresh_field_registry_internal()
                yield _sse("registry", f"Field registry refreshed: {_reg_count} fields", count=_reg_count)
            except Exception as _re:
                yield _sse("registry_warn", f"Registry refresh warning: {_re}")

            yield _sse("done", "Complete")

        except Exception as e:
            yield _sse("error", str(e))
        finally:
            if conn:
                try: conn.close()
                except: pass

    return StreamingResponse(generate(), media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})



@app.post("/api/knowledge/deep-train")
async def deep_train_templates(request: Request):
    """Deep AI training using Agent to analyze each document.

    ★★★Body may carry {"templates": [...]} to limit the run to those files.
    Without it this globs EVERY .docx in the templates directory, which is how
    uploading 2 templates produced 15 registered rows and ~11 minutes of AI
    work: the training job called this at the end of its queue, and this
    ignored the queue and re-analysed everything sitting on disk — including
    files that had deliberately been removed from the register.
    """
    require_admin(request)
    try:
        only = None
        try:
            _body = await request.json()
            if isinstance(_body, dict):
                _want = _body.get("templates")
                if isinstance(_want, list) and _want:
                    only = {str(x) for x in _want}
        except Exception:
            only = None  # no body / not JSON — fall back to the whole directory
        from scout.tools.knowledge_base import get_db_connection
        from scout.tools.template_analyzer import analyze_template
        from scout.agent import scout

        templates_dir = Path("/documents/legal/templates")
        if not templates_dir.exists():
            return {"success": False, "error": "Templates directory not found"}

        conn = get_db_connection()
        cur = conn.cursor()

        templates = list(templates_dir.glob("*.docx"))
        if only is not None:
            templates = [t for t in templates if t.name in only]
        results = []

        for template_file in templates:
            template_name = template_file.name

            # First analyze to get fields
            result = analyze_template(template_name)
            fields = result.get("fields", [])

            # Now use AI Agent to deeply analyze the document
            try:
                # Read the template content
                from docx import Document

                doc = Document(str(template_file))
                content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

                # Field classification will run after main UPDATE below

                # Create prompt for AI analysis - enhanced with more details
                analysis_prompt = f"""Analyze this legal document template and provide comprehensive detailed information. This is CRITICAL for legal document automation.

DOCUMENT NAME: {template_name}
PLACEHOLDERS NEEDED: {", ".join(fields)}
DOCUMENT CONTENT PREVIEW: {content[:4000]}

Provide your response in this EXACT JSON format (no other text):
{{
    "category": "Type of document (e.g., AGM Minutes, Director Consent, Shareholder Resolution)",
    "purpose": "What is the main purpose of this document in 1-2 sentences?",
    "when_to_use": "When should someone use this document? Be specific about the scenario",
    "how_to_use": ["Step 1: action", "Step 2: action", "Step 3: action"],
    "keywords": "keyword1, keyword2, keyword3, keyword4, keyword5",
    "usage_instructions": "When and why use this document - specific use cases separated by |",
    "sections": ["Main section 1", "Main section 2", "Main section 3"],
    "signatures": [{{"name": "Who signs", "role": "Their role"}}],
    "deadlines": ["Any deadlines or timelines"],
    "legal_references": ["Myanmar Companies Law 2017 Section X", "DICA directives"],
    "related_documents": ["Other documents typically needed/related"],
    "tips": ["Tip 1", "Tip 2", "Tip 3"],
    "prerequisites": ["What is needed before starting this document"],
    "filing_deadline": "Any filing deadline (e.g., Within 30 days)",
    "fees": "Any fees (e.g., DICA online filing fee)",
    "validity_period": "How long is this document valid?",
    "approval_chain": ["Who approves this in order"],
    "required_attachments": ["Required attachments like NRC copy, proof of address"],
    "common_mistakes": ["Common mistake 1", "Common mistake 2"],
    "jurisdiction": "Country (e.g., Myanmar)",
    "industry_tags": ["Industry 1", "Industry 2"],
    "complexity": "Easy | Medium | Complex",
    "estimated_time": "Time to complete (e.g., 30 minutes)"
}}"""

                # Run the agent to analyze - use async version with timeout
                import asyncio
                import concurrent.futures
                import os

                agent_response = None
                ai_analysis = None

                # Skip AI if no API key or use fallback mode
                skip_ai = os.getenv("SKIP_AI_TRAINING", "false").lower() == "true"

                if not skip_ai:

                    def run_agent():
                        return asyncio.run(scout.arun(analysis_prompt))

                    try:
                        with concurrent.futures.ThreadPoolExecutor() as pool:
                            future = pool.submit(run_agent)
                            agent_response = future.result(timeout=30)  # 30 second timeout
                        ai_analysis = (
                            agent_response.content if hasattr(agent_response, "content") else str(agent_response)
                        )
                    except concurrent.futures.TimeoutError:
                        print(f"AI analysis timeout for {template_name}, using fallback")
                        ai_analysis = None
                    except Exception as agent_err:
                        print(f"Agent error: {agent_err}")
                        ai_analysis = None
                else:
                    print(f"Skipping AI for {template_name}, using fallback")
                    ai_analysis = None

                # Parse the AI response
                import re

                # Initialize all fields
                category = ""
                keywords = ""
                purpose = ""
                usage = ""
                use_cases = json.dumps(["Document generation"])
                sections = json.dumps([])
                signatures = json.dumps([])
                deadlines = json.dumps([])
                legal_references = json.dumps([])
                related_docs = json.dumps([])
                tips = json.dumps([])
                when_to_use = ""
                how_to_use = json.dumps([])
                prerequisites = json.dumps([])
                filing_deadline = ""
                fees = ""
                validity_period = ""
                approval_chain = json.dumps([])
                required_attachments = json.dumps([])
                common_mistakes = json.dumps([])
                jurisdiction = "Myanmar"
                industry_tags = json.dumps([])
                complexity = "Medium"
                estimated_time = "30 minutes"

                # Try to find JSON in the response (handle nested braces)
                # Fix: Check if ai_analysis is not None before regex search
                if ai_analysis and isinstance(ai_analysis, str):
                    json_match = re.search(r"\{[\s\S]*\}", ai_analysis)
                else:
                    json_match = None

                if json_match:
                    try:
                        ai_data = json.loads(json_match.group())
                        category = ai_data.get("category", "")
                        keywords = ai_data.get("keywords", "")
                        purpose = ai_data.get("purpose", "")
                        usage = ai_data.get("usage_instructions", ai_data.get("usage", ""))
                        use_cases = json.dumps(ai_data.get("use_cases", ["Document generation"]))
                        sections = json.dumps(ai_data.get("sections", []))
                        signatures = json.dumps(ai_data.get("signatures", []))
                        deadlines = json.dumps(ai_data.get("deadlines", []))
                        legal_references = json.dumps(ai_data.get("legal_references", []))
                        related_docs = json.dumps(ai_data.get("related_documents", []))
                        tips = json.dumps(ai_data.get("tips", []))

                        # NEW: Extract enhanced fields from AI
                        when_to_use = ai_data.get("when_to_use", "")
                        how_to_use = json.dumps(ai_data.get("how_to_use", []))
                        prerequisites = json.dumps(ai_data.get("prerequisites", []))
                        filing_deadline = ai_data.get("filing_deadline", "")
                        fees = ai_data.get("fees", "")
                        validity_period = ai_data.get("validity_period", "")
                        approval_chain = json.dumps(ai_data.get("approval_chain", []))
                        required_attachments = json.dumps(ai_data.get("required_attachments", []))
                        common_mistakes = json.dumps(ai_data.get("common_mistakes", []))
                        jurisdiction = ai_data.get("jurisdiction", "Myanmar")
                        industry_tags = json.dumps(ai_data.get("industry_tags", []))
                        complexity = ai_data.get("complexity", "Medium")
                        estimated_time = ai_data.get("estimated_time", "30 minutes")

                        # If fields are empty, use template name to infer
                        if not category or category == "General Document":
                            category = _infer_category(template_name)
                    except Exception as parse_err:
                        # Not cosmetic: every field below is now a guess from the
                        # template NAME rather than its content. All 15 templates
                        # shipped that way without anyone noticing, because this
                        # was a bare print to stdout.
                        logging.getLogger("legalscout").error(
                            "TEMPLATE ANALYSIS FELL BACK for %r — the AI response "
                            "did not parse (%s). Every analysed field below is "
                            "inferred from the filename.", template_name, parse_err
                        )
                        category = _infer_category(template_name)
                        keywords = _generate_keywords(template_name, fields)
                        purpose = "Legal document template"
                        usage = "Document generation"
                        use_cases = json.dumps(["Document generation"])
                        sections = json.dumps(_get_sections_from_name(template_name))
                        signatures = json.dumps(_get_signatures_from_name(template_name))
                        deadlines = json.dumps([])
                        legal_references = json.dumps(_get_legal_refs_from_name(template_name))
                        related_docs = json.dumps(_get_related_from_name(template_name))
                        tips = json.dumps(
                            ["Fill in all required fields", "Ensure signatures are obtained", "Review for accuracy"]
                        )
                        # Use fallback for new fields
                        when_to_use = _get_when_to_use(template_name)
                        how_to_use = json.dumps(_get_how_to_use(template_name))
                        prerequisites = json.dumps(_get_prerequisites(template_name))
                        filing_deadline = _get_filing_deadline(template_name)
                        fees = _get_fees(template_name)
                        validity_period = _get_validity_period(template_name)
                        approval_chain = json.dumps(_get_approval_chain(template_name))
                        required_attachments = json.dumps(_get_required_attachments(template_name))
                        common_mistakes = json.dumps(_get_common_mistakes(template_name))
                        jurisdiction = "Myanmar"
                        industry_tags = json.dumps(_get_industry_tags(template_name))
                        complexity = _get_complexity(len(fields))
                        estimated_time = _get_estimated_time(len(fields))
                else:
                    # No JSON found, use template name inference
                    logging.getLogger("legalscout").error(
                        "TEMPLATE ANALYSIS FELL BACK for %r — no JSON object in the "
                        "AI response. Every analysed field below is inferred from "
                        "the filename.", template_name
                    )
                    category = _infer_category(template_name)
                    keywords = ", ".join(fields[:10]) if fields else ""
                    purpose = "Legal document template"
                    usage = "Document generation"
                    use_cases = json.dumps(["Document generation"])
                    sections = json.dumps(_get_sections_from_name(template_name))
                    signatures = json.dumps(_get_signatures_from_name(template_name))
                    deadlines = json.dumps([])
                    legal_references = json.dumps(_get_legal_refs_from_name(template_name))
                    related_docs = json.dumps(_get_related_from_name(template_name))
                    tips = json.dumps(
                        ["Fill in all required fields", "Ensure signatures are obtained", "Review for accuracy"]
                    )
                    # Use fallback for new fields
                    when_to_use = _get_when_to_use(template_name)
                    how_to_use = json.dumps(_get_how_to_use(template_name))
                    prerequisites = json.dumps(_get_prerequisites(template_name))
                    filing_deadline = _get_filing_deadline(template_name)
                    fees = _get_fees(template_name)
                    validity_period = _get_validity_period(template_name)
                    approval_chain = json.dumps(_get_approval_chain(template_name))
                    required_attachments = json.dumps(_get_required_attachments(template_name))
                    common_mistakes = json.dumps(_get_common_mistakes(template_name))
                    jurisdiction = "Myanmar"
                    industry_tags = json.dumps(_get_industry_tags(template_name))
                    complexity = _get_complexity(len(fields))
                    estimated_time = _get_estimated_time(len(fields))
                    print(f"  FALLBACK SET: category={category}, purpose={purpose[:30]}")

            except Exception as e:
                logging.getLogger("legalscout").error(
                    "TEMPLATE ANALYSIS FAILED for %r (%s: %s) — falling back to "
                    "filename inference for every analysed field. Fields seen: %s",
                    template_name, type(e).__name__, e, fields
                )
                category = _infer_category(template_name)
                keywords = ", ".join(fields[:10]) if fields else ""
                purpose = "Legal document template"
                usage = "Document generation"
                use_cases = json.dumps(["Document generation"])
                sections = json.dumps(_get_sections_from_name(template_name))
                signatures = json.dumps(_get_signatures_from_name(template_name))
                deadlines = json.dumps([])
                legal_references = json.dumps(_get_legal_refs_from_name(template_name))
                related_docs = json.dumps(_get_related_from_name(template_name))
                tips = json.dumps(
                    ["Fill in all required fields", "Ensure signatures are obtained", "Review for accuracy"]
                )
                # Use fallback for new fields
                when_to_use = _get_when_to_use(template_name)
                how_to_use = json.dumps(_get_how_to_use(template_name))
                prerequisites = json.dumps(_get_prerequisites(template_name))
                filing_deadline = _get_filing_deadline(template_name)
                fees = _get_fees(template_name)
                validity_period = _get_validity_period(template_name)
                approval_chain = json.dumps(_get_approval_chain(template_name))
                required_attachments = json.dumps(_get_required_attachments(template_name))
                common_mistakes = json.dumps(_get_common_mistakes(template_name))
                jurisdiction = "Myanmar"
                industry_tags = json.dumps(_get_industry_tags(template_name))
                complexity = _get_complexity(len(fields))
                estimated_time = _get_estimated_time(len(fields))

            # Save to database
            description = purpose
            cur.execute(
                """
                UPDATE templates 
                SET category = %s, keywords = %s, description = %s, 
                    usage_instructions = %s, ai_trained = TRUE, ai_analyzed = TRUE,
                    fields = %s, total_fields = %s,
                    sections = %s, signatures = %s, deadlines = %s,
                    legal_references = %s, related_documents = %s,
                    use_cases = %s, tips = %s,
                    purpose = %s, when_to_use = %s, how_to_use = %s,
                    prerequisites = %s, filing_deadline = %s, fees = %s,
                    validity_period = %s, approval_chain = %s,
                    required_attachments = %s, common_mistakes = %s,
                    jurisdiction = %s, industry_tags = %s,
                    complexity = %s, estimated_time = %s
                WHERE name = %s
            """,
                (
                    category,
                    keywords,
                    description,
                    usage,
                    json.dumps({f: {} for f in fields}),
                    len(fields),
                    sections,
                    signatures,
                    deadlines,
                    legal_references,
                    related_docs,
                    use_cases,
                    tips,
                    purpose,
                    when_to_use,
                    how_to_use,
                    prerequisites,
                    filing_deadline,
                    fees,
                    validity_period,
                    approval_chain,
                    required_attachments,
                    common_mistakes,
                    jurisdiction,
                    industry_tags,
                    complexity,
                    estimated_time,
                    template_name,
                ),
            )

            conn.commit()

            # ── Knowledge Base Storage (knowledge_vec + knowledge_lookup) ──
            try:
                from scout.tools.knowledge_base import get_db_connection as _get_kb_conn
                _kb_conn = _get_kb_conn()
                _kb_cur = _kb_conn.cursor()

                # Build knowledge text for this template
                _kb_text = f"Template: {template_name}\nCategory: {category}\nPurpose: {purpose}\nWhen to use: {when_to_use}\nFields: {', '.join(fields)}\nContent: {content[:2000]}"

                # Clear old entries for this template
                _kb_cur.execute("DELETE FROM knowledge_vec WHERE source_file = %s", (f"template:{template_name}",))
                _kb_cur.execute("DELETE FROM knowledge_lookup WHERE source_file = %s", (f"template:{template_name}",))

                # Insert into knowledge_vec (semantic search)
                _kb_cur.execute(
                    "INSERT INTO knowledge_vec (content, source_file, metadata) VALUES (%s, %s, %s)",
                    (_kb_text, f"template:{template_name}", json.dumps({
                        "type": "template", "name": template_name, "category": category,
                        "purpose": purpose, "fields": fields,
                    }))
                )

                # Insert into knowledge_lookup (key-value search)
                for _kk, _kv in [("template_name", template_name), ("template_category", category),
                                  ("template_purpose", purpose), ("template_when_to_use", when_to_use)]:
                    if _kv:
                        _kb_cur.execute(
                            "INSERT INTO knowledge_lookup (key_name, key_value, source_file) VALUES (%s, %s, %s)",
                            (_kk, _kv, f"template:{template_name}"))

                _kb_conn.commit()

                # Generate vector embedding via OpenRouter
                _embedding = generate_embedding(_kb_text)
                if _embedding:
                    _kb_cur2 = _kb_conn.cursor()
                    _kb_cur2.execute(
                        "UPDATE knowledge_vec SET embedding = %s WHERE source_file = %s",
                        (str(_embedding), f"template:{template_name}"))
                    _kb_conn.commit()
                    _kb_cur2.close()
                    print(f"  ✓ KB + embedding: {template_name}")

                _kb_cur.close(); _kb_conn.close()
            except Exception as _kb_err:
                print(f"  ⚠ Knowledge storage error for {template_name}: {_kb_err}")

            results.append(
                {
                    "name": template_name,
                    "category": category,
                    "keywords": keywords[:80] + "..." if len(keywords) > 80 else keywords,
                    "description": description[:80] + "..." if len(description) > 80 else description,
                }
            )

        cur.close()
        conn.close()

        # AI field classification — run AFTER all templates are saved
        try:
            from scout.tools.template_analyzer import classify_template_fields, get_db_connection as _get_class_db
            from docx import Document as _ClassDoc
            _cconn = _get_class_db()
            _ccur = _cconn.cursor()
            for template_file in templates:
                try:
                    _cdoc = _ClassDoc(str(template_file))
                    _ccontent = "\n".join([p.text for p in _cdoc.paragraphs if p.text.strip()])
                    # Get fields from DB (don't call analyze_template again — it overwrites metadata)
                    _ccur.execute("SELECT fields FROM templates WHERE name = %s", (template_file.name,))
                    _crow = _ccur.fetchone()
                    _cfields_raw = _crow[0] if _crow else {}
                    # The raw placeholder list to classify. When `fields` is
                    # already a classification dict (a re-train), its KEYS are
                    # bucket names, NOT placeholders — feeding list(dict.keys())
                    # to classify_template_fields corrupts the result. Recover the
                    # real placeholders from the prior classification's own lists.
                    if isinstance(_cfields_raw, list):
                        _cfields = _cfields_raw
                    elif isinstance(_cfields_raw, dict):
                        if _cfields_raw.get("all_fields"):
                            _cfields = _cfields_raw["all_fields"]
                        elif _cfields_raw.get("db_fields") is not None or _cfields_raw.get("user_input_fields") is not None:
                            _cfields = list(_cfields_raw.get("db_fields") or []) + list(_cfields_raw.get("user_input_fields") or [])
                        else:
                            _cfields = list(_cfields_raw.keys())
                    else:
                        _cfields = []
                    _classification = classify_template_fields(_ccontent, _cfields)
                    if _classification:
                        _ccur.execute("UPDATE templates SET fields = %s WHERE name = %s",
                            (json.dumps(_classification), template_file.name))
                        _cconn.commit()
                        print(f"  Classified {template_file.name}: {len(_classification.get('db_fields',[]))} DB, {len(_classification.get('user_input_fields',[]))} user")
                except Exception as _ce:
                    print(f"  Classification skip {template_file.name}: {_ce}")
            _ccur.close()
            _cconn.close()
        except Exception as _ce2:
            print(f"  Classification phase error: {_ce2}")

        # Convert all templates to PDF for preview
        try:
            import subprocess
            pdf_dir = Path("/documents/legal/previews")
            pdf_dir.mkdir(parents=True, exist_ok=True)
            for template_file in templates:
                try:
                    subprocess.run(
                        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(template_file)],
                        capture_output=True, text=True, timeout=30
                    )
                    print(f"  PDF: {template_file.name}")
                except Exception:
                    pass
        except Exception as _pdf_err:
            print(f"  PDF conversion error: {_pdf_err}")

        # Refresh agent's template knowledge after training
        try:
            _refresh_agent_knowledge()
        except Exception as _refresh_err:
            print(f"Agent refresh warning: {_refresh_err}")

        # Save training timestamp
        try:
            from psycopg import connect as _tc2
            import os as _to2
            _tconn2 = get_db_conn()
            _tcur2 = _tconn2.cursor()
            _tcur2.execute("INSERT INTO training_status (training_type, last_trained, record_count) VALUES ('templates', NOW(), %s) ON CONFLICT (training_type) DO UPDATE SET last_trained = NOW(), record_count = %s", (len(results), len(results)))
            _tconn2.commit(); _tcur2.close(); _tconn2.close()
        except Exception: pass
        return {"success": True, "message": f"Deep trained {len(results)} templates with AI Agent", "results": results}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# Server-side background template training (survives tab/browser close)
# ---------------------------------------------------------------------------
@app.post("/api/knowledge/train-start")
async def train_start(request: Request):
    """Start (or attach to) a server-side training job.

    Body may contain `{"templates": [...]}` (explicit names) or
    `{"retrain_all": bool}`. With neither, only UNTRAINED templates are queued.
    Returns the job (with id + queue), or the already-active job, or
    `{"status": "nothing_to_train"}` when the queue would be empty.
    """
    require_admin(request)
    try:
        body = await request.json()
    except Exception:
        body = {}
    body = body or {}

    names = body.get("templates")
    if not names:
        retrain_all = bool(body.get("retrain_all"))
        from scout.tools.template_analyzer import get_db_connection as _gdb
        conn = _gdb()
        try:
            cur = conn.cursor()
            cur.execute("SELECT name, ai_trained FROM templates ORDER BY name")
            rows = cur.fetchall()
        finally:
            conn.close()
        if retrain_all:
            names = [r[0] for r in rows]
        else:
            names = [r[0] for r in rows if not r[1]]

    names = [n for n in (names or []) if n]
    if not names:
        return {"status": "nothing_to_train"}

    return training_jobs.start_job(names)


@app.get("/api/knowledge/train-job")
async def train_job(request: Request):
    """Return the most recent training job (logs capped to last 200)."""
    require_admin(request)
    job = training_jobs.get_latest_job()
    if job is None:
        return {"status": "idle"}
    return job


@app.post("/api/knowledge/train-cancel")
async def train_cancel(request: Request):
    """Flag the active training job cancelled (worker stops between templates)."""
    require_admin(request)
    return {"cancelled": training_jobs.cancel_active()}


def _infer_category(template_name: str) -> str:
    """Group a template by the TASK it serves, from its name.

    Grouped by task, not document form, because that is how the question
    arrives — "appoint a director", "hold the AGM" — not "give me a resolution".

    Two bugs lived here. `"group" in name` also matched "Non-Group", filing the
    non-group consent under Group Company; and a broad "shareholder" test above
    the specific ones swept unrelated templates into "Shareholder Resolution".
    Specific tests now come first, and "non-group" is checked before "group".
    """
    n = template_name.lower()

    appointing = "appointment" in n or "appoint" in n
    resigning = "resignation" in n or "resign" in n

    # Consents are named by who gives them, so test those before the generic
    # appointment/resignation words that also appear in their titles.
    if "corporate shareholder consent" in n:
        return "Shareholder Consent - Corporate"
    if "individual shareholder consent" in n:
        return "Shareholder Consent - Individual"
    if "director consent" in n:
        # "Non-Group" contains "group" — always test the negative first.
        if "non-group" in n:
            return "Director Appointment - Non-Group"
        if "group" in n:
            return "Director Appointment - Group Company"
        return "Director Appointment"

    if resigning and appointing:
        return "Director Resignation & Appointment"
    if resigning:
        return "Director Resignation"
    if appointing:
        return "Director Appointment"

    if "annual general meeting" in n or "agm" in n:
        return "Annual General Meeting"
    if "extraordinary" in n or "egm" in n:
        return "Extraordinary General Meeting"
    if "minutes" in n:
        return "Meeting Minutes"
    if "resolution" in n:
        return "Resolution"
    return "Other"


def _get_sections_from_name(template_name: str) -> list:
    """Extract sections based on template name."""
    name_lower = template_name.lower()
    sections = []

    if "agm" in name_lower or "annual general meeting" in name_lower:
        sections = [
            "Call of Meeting",
            "Notice",
            "Quorum",
            "Chairman Appointment",
            "Minutes",
            "Resolutions",
            "Vote Results",
        ]
    elif "egm" in name_lower or "extraordinary" in name_lower:
        sections = ["Notice of Meeting", "Agenda", "Quorum", "Resolutions", "Voting"]
    elif "director consent" in name_lower or "director appointment" in name_lower:
        sections = ["Director Details", "Appointment Terms", "Consent Declaration", "Signature"]
    elif "shareholder" in name_lower:
        sections = ["Shareholder Details", "Resolution", "Approval", "Signatures"]
    elif "minutes" in name_lower:
        sections = ["Meeting Details", "Attendance", "Agenda Items", "Discussion", "Decisions", "Vote Count"]
    else:
        sections = ["Document Header", "Main Content", "Signatures", "Date"]

    return sections


def _get_signatures_from_name(template_name: str) -> list:
    """Extract required signatures based on template name."""
    name_lower = template_name.lower()
    signatures = []

    if "director consent" in name_lower:
        signatures = [{"name": "Appointed Director", "role": "Director"}, {"name": "Witness", "role": "Witness"}]
    elif "agm" in name_lower or "minutes" in name_lower:
        signatures = [
            {"name": "Chairman", "role": "Meeting Chairman"},
            {"name": "Secretary", "role": "Company Secretary"},
            {"name": "Scrutineer", "role": "Vote Scrutineer"},
        ]
    elif "shareholder" in name_lower:
        signatures = [
            {"name": "Shareholder", "role": "Shareholder"},
            {"name": "Director", "role": "Director"},
            {"name": "Witness", "role": "Witness"},
        ]
    else:
        signatures = [{"name": "Authorized Signatory", "role": "Director/Secretary"}]

    return signatures


def _get_legal_refs_from_name(template_name: str) -> list:
    """Return NO legal references. A citation may never be guessed from a filename.

    This function used to return hardcoded citations picked by substring match on
    the template name:

        "agm"/"egm"   -> Companies Act 2013 - Section 100-104, SEBI Regulations
        "director"    -> Companies Act 2013 - Section 152, SEBI (LODR) 2015
        "shareholder" -> Companies Act 2013 - Section 189, Companies
                         (Management and Administration) Rules 2014
        otherwise     -> Companies Act 2013

    Every one of those is INDIAN law. The Companies Act 2013 is India's, and SEBI
    is India's securities regulator, which has no jurisdiction in Myanmar. This
    product drafts under the Myanmar Companies Law 2017 and files with DICA — the
    same code block that produced these strings set `jurisdiction = "Myanmar"`
    three lines later.

    It was not a stray default. Measured 2026-08-06, all 15 templates in the
    database held exactly these strings, which means the AI extraction step never
    once succeeded and this fallback supplied the citation every single time,
    silently. A lawyer reading the template register was told their Myanmar AGM
    minutes were governed by SEBI (LODR) Regulations 2015.

    The correct references come from the deep-training legal step, which already
    prompts for Myanmar Companies Law sections and DICA filing obligations. When
    that step fails the honest answer is nothing at all: an empty citation list
    is visibly incomplete, while a plausible wrong one is not. Guessing Myanmar
    section numbers here to replace the Indian ones would repeat the bug in a
    harder-to-spot form.
    """
    logging.getLogger("legalscout").error(
        "LEGAL REFERENCE EXTRACTION FAILED for %r — storing no citations. "
        "The deep-training legal step (Myanmar Companies Law 2017 / DICA) did not "
        "produce a result, and citations are never inferred from a template name. "
        "Re-run training for this template; until then its legal_references stays "
        "empty by design.",
        template_name,
    )
    return []


def _get_related_from_name(template_name: str) -> list:
    """Suggest related documents from the template name.

    These are workflow hints, not citations, so a wrong one is a bad suggestion
    rather than bad law — which is why this fallback survives while the legal
    reference one does not. One entry did carry the same Indian assumption:
    "DIN Application", a Director Identification Number filing under Indian
    company law. Myanmar directors are recorded with DICA, so the DICA director
    particulars filing takes its place.
    """
    name_lower = template_name.lower()
    docs = []

    if "agm" in name_lower:
        docs = ["Notice of AGM", "Proxy Form", "Financial Statements", "Board Resolution"]
    elif "egm" in name_lower:
        docs = ["Notice of EGM", "Proxy Form", "Board Resolution"]
    elif "director consent" in name_lower:
        docs = ["Board Resolution", "Appointment Letter", "DICA Director Particulars Filing"]
    elif "shareholder" in name_lower:
        docs = ["Share Certificate", "Register of Members", "Board Resolution"]
    else:
        docs = ["Company Documents", "Board Resolution"]

    return docs


def _generate_keywords(template_name: str, fields: list) -> str:
    """Generate keywords for the template."""
    name_lower = template_name.lower()
    keywords = []

    if "agm" in name_lower:
        keywords.extend(["AGM", "annual", "general meeting", "shareholders", "financial report", "auditor"])
    if "egm" in name_lower:
        keywords.extend(["EGM", "extraordinary", "special resolution", "urgent"])
    if "director" in name_lower:
        keywords.extend(["director", "appointment", "consent", "board"])
    if "shareholder" in name_lower:
        keywords.extend(["shareholder", "resolution", "shares", "voting"])
    if "minutes" in name_lower:
        keywords.extend(["minutes", "meeting", "proceedings", "record"])

    field_keywords = [f.replace("_", " ").replace("-", " ") for f in fields[:5]]
    keywords.extend(field_keywords)

    return ", ".join(keywords[:15])


def _generate_description(template_name: str, fields: list) -> str:
    """Generate AI description for the template."""
    category = _infer_category(template_name)

    description = f"This is a {category} template. "
    description += f"It contains {len(fields)} data fields including: "
    description += ", ".join(fields[:7])
    if len(fields) > 7:
        description += f" and {len(fields) - 7} more fields."

    return description


def _generate_usage_instructions(template_name: str, category: str) -> str:
    """Generate usage instructions for when to use this template."""
    name_lower = template_name.lower()
    instructions = []

    if "agm" in name_lower:
        instructions = [
            "Use for annual general meetings required by company law",
            "Use when presenting annual financial reports to shareholders",
            "Required for auditor appointment/re-appointment",
            "Use for election of directors at yearly meetings",
            "Use for approval of dividend distribution",
        ]
    elif "egm" in name_lower:
        instructions = [
            "Use for extraordinary/special general meetings",
            "Use when urgent shareholder decisions are needed",
            "Required for major corporate changes (mergers, acquisitions)",
            "Use for removal or appointment of directors",
            "Use for changes to company constitution",
        ]
    elif "director consent" in name_lower:
        if "group" in name_lower:
            instructions = [
                "Use when appointing directors from group companies",
                "Required for intra-group director transfers",
                "Use for corporate director appointments",
            ]
        else:
            instructions = [
                "Use when appointing new independent directors",
                "Required for individual director appointments",
                "Use for director consent to act",
            ]
    elif "shareholder" in name_lower:
        instructions = [
            "Use for shareholder resolutions",
            "Required for decisions requiring shareholder approval",
            "Use for new company setup with multiple shareholders",
        ]
    else:
        instructions = ["Use for general corporate documentation", "Review fields and fill accordingly"]

    return " | ".join(instructions)


@app.delete("/api/knowledge/sources/{filename}")
async def delete_knowledge_source(filename: str):
    """Delete a knowledge source."""
    try:
        from scout.tools.knowledge_base import delete_knowledge_source as do_delete

        do_delete(filename)
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/knowledge/search")
async def search_knowledge_base(request: Request, q: str = "", limit: int = 20):
    """Search knowledge base. Admin only — this is the one caller allowed to see
    the per-client rows, because firm staff legitimately search their whole
    register. The AGENT path deliberately cannot: see `_CLIENT_NAMESPACE` in
    scout/tools/knowledge_base.py."""
    require_admin(request)
    if not q:
        return {"results": []}

    results = search_knowledge(q, limit, include_clients=True)
    return {"results": results}


@app.get("/api/knowledge/lookup")
async def lookup_knowledge(key: str = None, value: str = None):
    """Exact lookup in knowledge base."""
    if not key or not value:
        return {"results": []}

    results = lookup_value(key, value)
    return {"results": results}


@app.get("/api/knowledge/table/{table_name}")
async def get_table_data(table_name: str, limit: int = 50):
    """Get data from a database table."""
    try:
        from scout.tools.knowledge_base import get_db_connection

        allowed_tables = [
            "templates",
            "knowledge_sources",
            "knowledge_vec",
            "knowledge_lookup",
            "knowledge_raw",
            "knowledge_sources",
        ]
        if table_name not in allowed_tables:
            return {"error": "Invalid table name", "data": [], "count": 0}

        conn = get_db_connection()
        cur = conn.cursor()

        # First get the count
        cur.execute(f"SELECT COUNT(*) FROM {table_name}")
        count = cur.fetchone()[0]

        # Then get the data
        cur.execute(f"SELECT * FROM {table_name} LIMIT %s", (int(limit),))
        rows = cur.fetchall()

        columns = [desc[0] for desc in cur.description] if cur.description else []

        data = []
        for row in rows:
            data.append(dict(zip(columns, row)))

        cur.close()
        conn.close()

        return {"data": data, "count": count}
    except Exception as e:
        return {"error": str(e), "data": [], "count": 0}


@app.get("/api/knowledge/data/{filename}")
async def get_knowledge_data(filename: str, limit: int = 50):
    """Get data from specific knowledge source."""
    data = get_source_data(filename, limit)
    return {"filename": filename, "data": data}


# ---------------------------------------------------------------------------
# Chat File Upload Endpoint
# ---------------------------------------------------------------------------
import base64
from db.connection import get_db_conn



# ---------------------------------------------------------------------------
@app.get("/api/dashboard/data")
async def get_dashboard_data():
    """Get all data for dashboard visualization.

    NOT cached. `_cache` lives in process memory and we run `--workers 2`, so a
    mutation handled by one worker cannot invalidate the other worker's copy.
    That produced a deleted company that kept reappearing in the admin list for
    up to CACHE_TTL_SECONDS — and a second delete attempt on the phantom row
    then failed with "Company not found". An admin list of a few dozen rows is
    cheap to read; correctness wins over the 30-second saving.
    """
    try:
        from scout.tools.template_analyzer import list_analyzed_templates

        # Get companies from DB
        companies = []
        try:
            import os as _os
            from psycopg import connect as _connect

            _conn = get_db_conn()
            _cur = _conn.cursor()
            _cur.execute("""
                SELECT company_name_english, company_registration_number, registered_office_address,
                       directors, principal_place_of_business, status, company_type, registration_date,
                       principal_activity, total_shares_issued, currency_of_share_capital,
                       ultimate_holding_company_name, date_of_last_annual_return, members, id
                FROM companies ORDER BY company_name_english
            """)
            for row in _cur.fetchall():
                dirs = row[3] if isinstance(row[3], list) else []
                mems = row[13] if isinstance(row[13], list) else []
                director_names = ", ".join(d.get("name", "") for d in dirs) if dirs else ""
                shareholder_names = ", ".join(m.get("name", "") for m in mems) if mems else ""
                total_shares = row[9] or ""
                currency = row[10] or ""
                shares_display = f"{total_shares} {currency}".strip() if total_shares else ""
                companies.append({
                    "id": row[14],
                    "company_name": row[0] or "",
                    "company_registration_number": row[1] or "",
                    "registered_office": row[2] or "",
                    "directors": director_names,
                    "shareholders": shareholder_names,
                    "total_shares": shares_display,
                    "principal_place_of_business": row[4] or "",
                    "status": row[5] or "",
                    "company_type": row[6] or "",
                    "registration_date": row[7].isoformat() if row[7] else "",
                    "principal_activity": row[8] or "",
                    "total_shares_issued": row[9] or "",
                    "currency_of_share_capital": row[10] or "",
                    "ultimate_holding_company": row[11] or "",
                    "date_of_last_annual_return": row[12].isoformat() if row[12] else "",
                    "members": mems,
                })
            _cur.close()
            _conn.close()
        except Exception as _e:
            print(f"Companies DB read failed: {_e}")
            companies = []

        # Get templates with timestamps from database
        templates = list_analyzed_templates()

        # Get documents from PostgreSQL
        documents = []
        try:
            import os
            from psycopg import connect

            conn = get_db_conn()
            cur = conn.cursor()
            cur.execute("""
                SELECT id, template_name, company_name, file_name, file_path, 
                       download_url, created_at 
                FROM documents 
                ORDER BY created_at DESC LIMIT 50
            """)
            for row in cur.fetchall():
                documents.append(
                    {
                        "id": row[0],
                        "template_name": row[1],
                        "company_name": row[2],
                        "file_name": row[3],
                        "file_path": row[4],
                        "download_url": row[5],
                        "created_at": row[6].isoformat() if row[6] else None,
                    }
                )
            cur.close()
            conn.close()
        except Exception as e:
            print(f"Error fetching documents: {e}")

        # Also scan output directory for any files not in DB
        output_dir = Path("/documents/legal/output")
        if output_dir.exists():
            existing_files = set(d.get("file_name", "") for d in documents)
            for f in sorted(output_dir.glob("*.docx"), key=lambda x: x.stat().st_mtime, reverse=True):
                if f.name not in existing_files:
                    stat = f.stat()
                    stem = f.stem
                    import re

                    match = re.search(r"_(\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2})$", stem)
                    if match:
                        parts = stem[: match.start()].rsplit("_", 1)
                        template = parts[0] if len(parts) > 0 else stem
                        company = parts[1] if len(parts) > 1 else "Unknown"
                    else:
                        template = stem
                        company = "Unknown"
                    documents.append(
                        {
                            "template_name": template,
                            "company_name": company,
                            "file_name": f.name,
                            "file_path": str(f),
                            "download_url": f"{API_HOST}/documents/legal/output/{f.name}",
                            "created_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                            "id": f"doc_{int(stat.st_mtime)}",
                        }
                    )

        # Sort by date and limit to 50
        documents = sorted(documents, key=lambda x: x.get("created_at", ""), reverse=True)[:50]

        return {"companies": companies, "templates": templates, "documents": documents}
    except Exception as e:
        return {"error": str(e)}


@app.get("/api/dashboard/search-documents")
async def search_documents(request: Request, q: str = ""):
    """Search across generated documents by name, company, template."""
    if not q or len(q) < 2:
        return {"success": False, "error": "Search query must be at least 2 characters"}
    try:
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, template_name, company_name, file_name, created_at
            FROM documents
            WHERE file_name ILIKE %s OR company_name ILIKE %s OR template_name ILIKE %s
            ORDER BY created_at DESC LIMIT 50
        """, (f"%{q}%", f"%{q}%", f"%{q}%"))
        rows = cur.fetchall()
        cur.close(); conn.close()
        results = [{"id": r[0], "template": r[1], "company": r[2], "file": r[3],
                   "created": r[4].isoformat() if r[4] else None} for r in rows]
        return {"success": True, "results": results, "count": len(results)}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/dashboard/analytics")
async def dashboard_analytics(request: Request):
    """Get analytics data for dashboard charts."""
    try:
        import os
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()

        # Documents per day (last 30 days)
        cur.execute("""
            SELECT DATE(created_at) as day, COUNT(*) as count
            FROM documents WHERE created_at > NOW() - INTERVAL '30 days'
            GROUP BY DATE(created_at) ORDER BY day
        """)
        docs_per_day = [{"date": r[0].isoformat() if r[0] else "", "count": r[1]} for r in cur.fetchall()]

        # Documents by template
        cur.execute("""
            SELECT template_name, COUNT(*) as count FROM documents
            GROUP BY template_name ORDER BY count DESC LIMIT 10
        """)
        by_template = [{"name": r[0] or "Unknown", "count": r[1]} for r in cur.fetchall()]

        # Documents by company
        cur.execute("""
            SELECT company_name, COUNT(*) as count FROM documents
            GROUP BY company_name ORDER BY count DESC LIMIT 10
        """)
        by_company = [{"name": (r[0] or "Unknown").split('\n')[0], "count": r[1]} for r in cur.fetchall()]

        # User activity (last 30 days)
        cur.execute("""
            SELECT user_email, action, COUNT(*) as count
            FROM activity_logs WHERE created_at > NOW() - INTERVAL '30 days'
            GROUP BY user_email, action ORDER BY count DESC LIMIT 20
        """)
        user_activity = [{"email": r[0], "action": r[1], "count": r[2]} for r in cur.fetchall()]

        # Total counts
        cur.execute("SELECT COUNT(*) FROM documents")
        total_docs = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM templates")
        total_templates = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM companies")
        total_companies = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM users")
        total_users = cur.fetchone()[0]

        cur.close(); conn.close()

        return {
            "success": True,
            "docs_per_day": docs_per_day,
            "by_template": by_template,
            "by_company": by_company,
            "user_activity": user_activity,
            "totals": {"documents": total_docs, "templates": total_templates,
                      "companies": total_companies, "users": total_users}
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/dashboard/key-dates")
async def dashboard_key_dates(request: Request):
    """Get key dates and upcoming deadlines for all companies."""
    try:
        import os
        from psycopg import connect
        from datetime import timedelta
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT company_name_english, registration_date, date_of_last_annual_return,
                   company_type, status
            FROM companies WHERE company_name_english IS NOT NULL
        """)
        rows = cur.fetchall()
        cur.close(); conn.close()

        dates = []
        today = datetime.now().date()
        for r in rows:
            name = r[0]
            reg_date = r[1]
            last_ar = r[2]

            # Incorporation anniversary
            if reg_date:
                next_anniv = reg_date.replace(year=today.year)
                if next_anniv < today:
                    next_anniv = next_anniv.replace(year=today.year + 1)
                days_until = (next_anniv - today).days
                dates.append({"company": name, "event": "Incorporation Anniversary",
                    "date": next_anniv.isoformat(), "days_until": days_until,
                    "urgency": "info" if days_until > 30 else "warning" if days_until > 7 else "urgent"})

            # Annual Return deadline (within 30 days of AR anniversary)
            if last_ar:
                next_ar = last_ar.replace(year=today.year)
                if next_ar < today:
                    next_ar = next_ar.replace(year=today.year + 1)
                ar_deadline = next_ar + timedelta(days=30)
                days_until = (ar_deadline - today).days
                dates.append({"company": name, "event": "Annual Return Deadline",
                    "date": ar_deadline.isoformat(), "days_until": days_until,
                    "urgency": "info" if days_until > 30 else "warning" if days_until > 14 else "urgent"})

            # AGM deadline (18 months from incorporation or last AGM)
            if reg_date:
                agm_deadline = reg_date.replace(year=today.year) + timedelta(days=547)  # ~18 months
                if agm_deadline < today:
                    agm_deadline = agm_deadline.replace(year=today.year + 1)
                days_until = (agm_deadline - today).days
                dates.append({"company": name, "event": "AGM Deadline",
                    "date": agm_deadline.isoformat(), "days_until": days_until,
                    "urgency": "info" if days_until > 60 else "warning" if days_until > 30 else "urgent"})

        # Sort by urgency
        dates.sort(key=lambda x: x["days_until"])
        return {"success": True, "dates": dates}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/dashboard/export/excel")
async def export_companies_excel(request: Request):
    """Export companies as Excel from database."""
    require_admin(request)
    try:
        from fastapi.responses import StreamingResponse
        from openpyxl import Workbook

        # Read companies from DB
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT company_name_english, company_registration_number, registered_office_address, company_type, status, directors, members, total_shares_issued, currency_of_share_capital, principal_activity FROM companies ORDER BY company_name_english LIMIT 10000")
        rows = cur.fetchall()
        cols = [desc[0] for desc in cur.description]
        companies = []
        for row in rows:
            company = {}
            for i, col in enumerate(cols):
                val = row[i]
                if isinstance(val, list):
                    val = ", ".join(d.get("name", "") for d in val if isinstance(d, dict))
                company[col] = str(val) if val else ""
            companies.append(company)
        cur.close(); conn.close()

        if not companies:
            return {"error": "No companies to export"}

        # Map internal columns back to user format
        INTERNAL_TO_USER = {
            "company_name": "Company Name",
            "company_name_original": "Company Name (Original)",
            "company_registration_number": "Registration Number",
            "date_of_incorporation": "Date of Incorporation",
            "registered_office": "Registered Office Address",
            "principal_place_of_business": "Principal Place of Business",
            "company_type": "Company Type (DICA)",
            "foreign_company": "Foreign Company",
            "under_corpset_management": "Under CorpSec Management",
            "group_company": "Group Company",
            "directors": "Directors",
            "total_shares": "Total Shares",
            "total_capital": "Total Capital",
            "shareholders": "Shareholders",
            "currency": "Currency",
            "meeting_date": "Meeting Date",
            "meeting_location": "Meeting Location",
            "agenda": "Agenda",
            "director_name": "Director Name",
            "individual_shareholder_1_name": "Individual Shareholder 1 Name",
            "individual_shareholder_2_name": "Individual Shareholder 2 Name",
            "individual_shareholder_3_name": "Individual Shareholder 3 Name",
            "individual_shareholder_4_name": "Individual Shareholder 4 Name",
            "corporate_shareholder_1_name": "Corporate Shareholder 1 Name",
            "financial_year_end_date": "Financial Year End Date",
            "pronoun": "Pronoun",
            "auditor_name": "Auditor Name",
            "auditor_fee": "Auditor Fee",
            "next_financial_year_end_date": "Next Financial Year End Date",
        }

        column_order = [
            "Number",
            "Company Name (Original)",
            "Registration Number",
            "Company Name",
            "Date of Incorporation",
            "Registered Office Address",
            "Principal Place of Business",
            "Company Type (DICA)",
            "Foreign Company",
            "Under CorpSec Management",
            "Group Company",
            "Directors",
            "Total Shares",
            "Total Capital",
            "Shareholders",
            "Currency",
            "Meeting Date",
            "Meeting Location",
            "Agenda",
            "Director Name",
            "Individual Shareholder 1 Name",
            "Individual Shareholder 2 Name",
            "Individual Shareholder 3 Name",
            "Individual Shareholder 4 Name",
            "Corporate Shareholder 1 Name",
            "Financial Year End Date",
            "Pronoun",
            "Auditor Name",
            "Auditor Fee",
            "Next Financial Year End Date",
        ]

        wb = Workbook()
        ws = wb.active
        ws.title = "Companies"

        # Write headers
        for col_idx, header in enumerate(column_order, 1):
            ws.cell(1, col_idx, header)

        # Write data
        for row_idx, company in enumerate(companies, 2):
            ws.cell(row_idx, 1, row_idx - 1)  # Number
            for col_idx, header in enumerate(column_order[1:], 2):
                # Find internal key
                internal_key = None
                for ik, uk in INTERNAL_TO_USER.items():
                    if uk == header:
                        internal_key = ik
                        break
                value = company.get(internal_key, "") if internal_key else ""
                ws.cell(row_idx, col_idx, str(value) if value else "")

        from io import BytesIO

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=companies_export.xlsx"},
        )
    except Exception as e:
        import traceback

        traceback.print_exc()
        return {"error": str(e)}


# ---------------------------------------------------------------------------
# Backup / Export endpoint (Task 2)
# ---------------------------------------------------------------------------
@app.get("/api/admin/backup")
async def admin_backup(request: Request):
    """Export all data as JSON backup."""
    require_admin(request)
    try:
        import os, json
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()

        backup = {"exported_at": datetime.now().isoformat(), "version": "1.0"}

        # Export each table
        for table in ["templates", "companies", "documents", "users", "knowledge_lookup", "knowledge_raw",
                       "training_status", "document_versions", "app_settings", "activity_logs",
                       "agno_sessions", "agno_memories", "agno_learnings",
                       "scout_knowledge", "scout_learnings", "scout_knowledge_contents"]:
            try:
                cur.execute(f"SELECT * FROM {table}")
                cols = [desc[0] for desc in cur.description]
                rows = []
                for row in cur.fetchall():
                    r = {}
                    for i, col in enumerate(cols):
                        val = row[i]
                        if hasattr(val, 'isoformat'): val = val.isoformat()
                        elif isinstance(val, (dict, list)): pass  # JSONB is fine
                        else: val = str(val) if val is not None else None
                        r[col] = val
                    rows.append(r)
                backup[table] = rows
            except Exception:
                backup[table] = []

        cur.close(); conn.close()

        from starlette.responses import Response
        return Response(
            content=json.dumps(backup, indent=2, default=str),
            media_type="application/json",
            headers={"Content-Disposition": f"attachment; filename=legal_scout_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"}
        )
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# Field Registry — dynamic per-template user_input fields
# ---------------------------------------------------------------------------
@app.get("/api/dashboard/field-registry")
async def get_field_registry(request: Request):
    """Return all registered company-level user_input fields with metadata."""
    get_current_user(request)
    try:
        conn = get_db_conn(); cur = conn.cursor()
        cur.execute("""
            SELECT field_key, label, description, field_type, used_by_templates
            FROM company_field_registry
            ORDER BY field_key
        """)
        rows = cur.fetchall()
        cur.close(); conn.close()
        fields = [{"field_key": r[0], "label": r[1] or r[0].replace("_", " ").title(),
                   "description": r[2] or "", "field_type": r[3] or "text",
                   "used_by_templates": r[4] or []} for r in rows]
        return {"success": True, "fields": fields}
    except Exception as e:
        return {"success": False, "error": str(e), "fields": []}


def _refresh_field_registry_internal() -> int:
    """Rebuild company_field_registry from templates table. Returns row count."""
    SKIP = {"company", "company_name", "company_name_english", "company_registration_number",
            "directors", "members", "shareholders", "registered_office", "registered_office_address",
            "principal_activity", "status", "company_type",
            "individual_shareholder_1_name", "individual_shareholder_2_name",
            "corporate_shareholder_3_name", "shareholder_1_name", "shareholder_2_name",
            "authorized_director_name", "director_name", "pronoun", "date", "meeting_location",
            "financial_year_end_date", "next_financial_year_end_date",
            "auditor_name", "auditor_fee"}
    conn = get_db_conn(); conn.autocommit = True; cur = conn.cursor()
    cur.execute("SELECT name, fields FROM templates WHERE fields IS NOT NULL")
    rows = cur.fetchall()
    registry: dict = {}
    for tname, fields in rows:
        if not isinstance(fields, dict): continue
        user_inputs = fields.get("user_input_fields") or []
        descs = fields.get("field_descriptions") or {}
        for f in user_inputs:
            fk = str(f).lower().replace(" ", "_").replace("-", "_")
            if fk in SKIP: continue
            desc = descs.get(f, "")
            ft = "date" if "date" in fk else ("number" if any(x in fk for x in ("fee","amount","number","capital","share")) else "text")
            if fk not in registry:
                registry[fk] = {"label": f.replace("_", " ").title(), "desc": desc, "type": ft, "tpls": []}
            if tname not in registry[fk]["tpls"]:
                registry[fk]["tpls"].append(tname)
    for fk, meta in registry.items():
        cur.execute("""
            INSERT INTO company_field_registry (field_key, label, description, field_type, used_by_templates, updated_at)
            VALUES (%s, %s, %s, %s, %s, NOW())
            ON CONFLICT (field_key) DO UPDATE SET
                label = EXCLUDED.label, description = EXCLUDED.description,
                field_type = EXCLUDED.field_type, used_by_templates = EXCLUDED.used_by_templates,
                updated_at = NOW()
        """, (fk, meta["label"], meta["desc"], meta["type"], meta["tpls"]))
    cur.close(); conn.close()
    return len(registry)


@app.post("/api/dashboard/field-registry/refresh")
async def refresh_field_registry(request: Request):
    """Rebuild registry from all templates' user_input_fields."""
    require_admin(request)
    try:
        count = _refresh_field_registry_internal()
        return {"success": True, "count": count}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# Document Detail endpoint — full row with custom_data + validation
# ---------------------------------------------------------------------------
@app.get("/api/dashboard/document-detail/{doc_id}")
async def get_document_detail(doc_id: str, request: Request):
    """Return full document record incl custom_data + validation_result."""
    user = get_current_user(request)
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, template_name, company_name, file_name, file_path, download_url,
                   preview_url, validation_result, custom_data, version, created_at
            FROM documents WHERE id = %s
        """, (int(doc_id),))
        r = cur.fetchone()
        cur.close(); conn.close()
        if not r:
            return {"success": False, "error": "Not found"}
        return {"success": True, "data": {
            "id": r[0], "template_name": r[1], "company_name": r[2],
            "file_name": r[3], "file_path": r[4], "download_url": r[5],
            "preview_url": r[6], "validation_result": r[7], "custom_data": r[8],
            "version": r[9], "created_at": r[10].isoformat() if r[10] else None,
        }}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# Document History endpoint (Task 3)
# ---------------------------------------------------------------------------
@app.get("/api/dashboard/document-history/{doc_name}")
async def get_document_history(doc_name: str, request: Request):
    """Get version history for a document."""
    import urllib.parse, os
    from psycopg import connect
    doc_name = urllib.parse.unquote(doc_name)
    user = get_current_user(request)
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT version, file_name, generated_by_email, custom_data, created_at
            FROM document_versions
            WHERE document_name ILIKE %s OR template_name ILIKE %s
            ORDER BY created_at DESC LIMIT 20
        """, (f"%{doc_name}%", f"%{doc_name}%"))
        rows = cur.fetchall()
        cur.close(); conn.close()
        versions = [{"version": r[0], "file_name": r[1], "generated_by": r[2],
                     "custom_data": r[3], "created_at": r[4].isoformat() if r[4] else None} for r in rows]
        return {"success": True, "versions": versions}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# Bulk Operations endpoints (Task 5)
# ---------------------------------------------------------------------------
@app.post("/api/dashboard/bulk/generate")
async def bulk_generate_documents(request: Request):
    """Generate a document for ALL companies using a specific template."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        body = await request.json()
        template_name = body.get("template_name", "")
        if not template_name:
            return {"success": False, "error": "template_name required"}

        import os
        from psycopg import connect
        from scout.tools.smart_doc import create_smart_document_tool

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT company_name_english FROM companies WHERE company_name_english IS NOT NULL")
        companies = [row[0] for row in cur.fetchall()]
        cur.close(); conn.close()

        if not companies:
            return {"success": False, "error": "No companies in database"}

        tools = create_smart_document_tool("/documents")
        generate = tools["generate_document"]

        results = []
        for company in companies:
            try:
                result = generate(template_name=template_name, company_name=company, custom_data={})
                results.append({"company": company, "success": result.get("success", False),
                               "file": result.get("file_name", ""), "error": result.get("error", "")})
                log_activity(user.get("user_id"), user.get("email"), "bulk_generate",
                           f"Generated {template_name} for {company}", "")
            except Exception as e:
                results.append({"company": company, "success": False, "error": str(e)})

        success_count = sum(1 for r in results if r["success"])
        return {"success": True, "message": f"Generated {success_count}/{len(companies)} documents", "results": results}
    except Exception as e:
        return {"success": False, "error": str(e)}





@app.get("/api/documents/fill-view")
async def documents_fill_view(request: Request, template: str, company: str):
    """Render a template as fill-in blocks: the whole document with each blank
    shown in place and click-to-pick candidates from the registers."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        from scout.tools.fill_view import build_fill_view
        return build_fill_view(template, company, "/documents")
    except Exception as e:
        logging.getLogger("legalscout").warning(f"fill-view failed: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/documents/fill-generate")
async def documents_fill_generate(request: Request):
    """Generate a document straight from the fill-in view's chosen values."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        body = await request.json()
        template_name = (body.get("template") or body.get("template_name") or "").strip()
        company_name = (body.get("company") or body.get("company_name") or "").strip()
        custom_data = body.get("custom_data") or {}
        if not template_name or not company_name:
            return {"success": False, "error": "template and company are required"}
        if not isinstance(custom_data, dict):
            return {"success": False, "error": "custom_data must be an object"}
        from scout.tools.smart_doc import create_smart_document_tool
        tools = create_smart_document_tool("/documents", host=API_HOST)
        result = tools["generate_document"](
            template_name=template_name, company_name=company_name, custom_data=custom_data
        )
        if result.get("success"):
            log_activity(user.get("user_id"), user.get("email"), "fill_generate",
                         f"Generated {template_name} for {company_name} from fill-view", "")
        return result
    except Exception as e:
        logging.getLogger("legalscout").warning(f"fill-generate failed: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/dashboard/add/company")
async def add_dashboard_company(request: Request):
    """Add a single company manually."""
    try:
        from scout.tools.knowledge_base import add_company
        user = get_current_user(request)
        body = await request.json()

        result = add_company(body)

        if result.get("success"):
            # Which filing this record came from — resolved server-side against
            # the uploads directory, never taken on the client's word.
            _persist_company_source_pdf(body)
            company_name = body.get("company_name_english", "Unknown")
            creator_email = user.get("email", "unknown") if user else "unknown"
            # Track creator
            try:
                from psycopg import connect as _ac
                import os as _ao
                _aconn = get_db_conn()
                _acur = _aconn.cursor()
                _acur.close(); _aconn.close()
            except Exception: pass
            log_activity(user.get("user_id") if user else None, creator_email, "add_company", f"Added: {company_name}", "")
            clear_cache()  # Clear dashboard cache so list refreshes
            return {"success": True, "message": "Company added successfully"}
        return {"success": False, "error": result.get("error", "Failed to add company")}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.delete("/api/dashboard/company/{company_name}")
async def delete_dashboard_company(request: Request, company_name: str):
    """Delete a company by name from the companies DB table."""
    require_admin(request)
    try:
        import urllib.parse
        import os
        from psycopg import connect

        company_name = urllib.parse.unquote(company_name)

        conn = get_db_conn()
        conn.autocommit = True
        cur = conn.cursor()

        # ★★★Was `ILIKE %s` with the raw unquoted path parameter and no LIMIT.
        # ILIKE reads as case-insensitive equality but LIKE metacharacters in the
        # parameter are LIVE, and `conn.autocommit = True` two lines above means
        # there is no transaction to roll back. Measured on throwaway rows:
        #   exact name                     4 seeded -> 1 deleted   (intended)
        #   a name containing '_'          4 seeded -> 2 deleted   (took a sibling)
        #   DELETE /api/dashboard/company/%25  ->  '%'  ->  4 deleted  (ALL of them)
        # The '_' case needs no attacker: it fires on any registered name with an
        # underscore. Equality on lower() removes both.
        cur.execute(
            "DELETE FROM companies WHERE lower(company_name_english) = lower(%s)",
            (company_name,),
        )
        deleted = cur.rowcount > 0

        # Clean up ALL knowledge data for this company
        if deleted:
            # ★ Still a LIKE, and deliberately so — this one is a genuine substring
            # sweep over `source_file`. But the metacharacters are escaped now, or
            # deleting a company named '%' would empty the whole table.
            _like_safe = (company_name.replace("\\", "\\\\")
                                      .replace("%", "\\%")
                                      .replace("_", "\\_"))
            cur.execute(
                "DELETE FROM knowledge_lookup WHERE source_file ILIKE %s ESCAPE '\\'",
                (f"%{_like_safe}%",),
            )
            cur.execute("DELETE FROM knowledge_vec WHERE source_file = %s", (f"company:{company_name}",))
            try: cur.execute("DELETE FROM knowledge_raw WHERE source_file = %s", (f"company:{company_name}",))
            except Exception: pass
            for mem_table in ["agno_memories", "agno_learnings", "scout_learnings"]:
                try: cur.execute(f"DELETE FROM {mem_table} WHERE content ILIKE %s", (f"%{company_name}%",))
                except Exception: pass
            try: cur.execute("DELETE FROM scout_knowledge WHERE content ILIKE %s", (f"%{company_name}%",))
            except Exception: pass
            try: cur.execute("DELETE FROM scout_knowledge_contents WHERE content ILIKE %s", (f"%{company_name}%",))
            except Exception: pass

        # The company_people links went with the row (ON DELETE CASCADE), which
        # can leave sync-created people on the register belonging to nobody.
        pruned = 0
        if deleted:
            try:
                from scout.tools.people_sync import prune_orphan_people

                pruned = prune_orphan_people(conn)["removed"]
            except Exception as e:
                logger.warning(f"[COMPANY DELETE] orphan prune skipped: {e}")

        cur.close()
        conn.close()

        if deleted:
            clear_cache()  # Clear dashboard cache so list refreshes immediately
            invalidate_training(f"Company deleted: {company_name}")
            # Refresh agent's knowledge so it forgets this company
            try:
                from scout.agent import _build_template_knowledge
                _refresh_agent_knowledge()
            except Exception: pass
            msg = f"Company '{company_name}' deleted. Re-training recommended."
            if pruned:
                msg = (f"Company '{company_name}' deleted, plus {pruned} "
                       f"{'person' if pruned == 1 else 'people'} left with no company. "
                       "Re-training recommended.")
            return {"success": True, "message": msg, "training_invalidated": True,
                    "people_removed": pruned}

        # Already gone — treat as done rather than an error. Deleting a row that
        # does not exist has the outcome the caller asked for, and reporting it
        # as a failure only confuses an admin whose list was a moment stale.
        return {
            "success": True,
            "message": f"Company '{company_name}' is no longer on file.",
            "already_absent": True,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/api/dashboard/company/{company_id}")
async def get_company_by_id(request: Request, company_id: int):
    """Get full company data by ID for editing."""
    get_current_user(request)
    try:
        import os
        from psycopg import connect

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT * FROM companies WHERE id = %s", (company_id,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return {"success": False, "error": "Company not found"}

        cols = [desc[0] for desc in cur.description]
        company = {}
        for i, col in enumerate(cols):
            val = row[i]
            if hasattr(val, 'isoformat'):
                val = val.isoformat()
            company[col] = val

        cur.close(); conn.close()
        return {"success": True, "data": company}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.put("/api/dashboard/company/{company_name}")
async def update_dashboard_company(request: Request, company_name: str, body: dict):
    """Update a company in the DB by name (upsert via add_company)."""
    require_admin(request)
    try:
        import urllib.parse
        from scout.tools.knowledge_base import add_company

        company_name = urllib.parse.unquote(company_name)

        # The company this PUT names has to be resolved and its identity forced
        # onto the body. add_company UPSERTs on company_registration_number, so a
        # body that omits it (or carries a different one) INSERTS instead of
        # updating: a PUT naming an existing company created a second, NAMELESS
        # row holding the payload and still answered "Company updated". The path
        # parameter was unquoted into this local and then never used again.
        conn = get_db_conn()
        try:
            cur = conn.cursor()
            cur.execute(
                """SELECT * FROM companies
                   WHERE lower(company_name_english) = lower(%s)
                   ORDER BY id LIMIT 1""",
                (company_name,),
            )
            row = cur.fetchone()
            columns = [d[0] for d in cur.description] if cur.description else []
        finally:
            conn.close()

        if not row:
            # Refuse rather than insert. An update to a name that is not on file
            # is a mistake, and answering it by creating a company hides that.
            return {"success": False, "error": f"Company '{company_name}' not found"}

        stored = dict(zip(columns, row))
        stored_name = stored.get("company_name_english")
        stored_reg = stored.get("company_registration_number")

        # MERGE, do not replace. add_company's ON CONFLICT sets every column from
        # EXCLUDED, so a PUT carrying one field nulls all the others — a partial
        # update that changed only `principal_activity` emptied a company's
        # `directors` array (measured: 3 directors -> 0) and, because the people
        # sync runs off that array, took the register link with it. Stored values
        # are the base; only keys the caller actually sent override them.
        merged = {k: v for k, v in stored.items() if k not in ("id", "created_at", "updated_at")}
        merged.update(body or {})
        body = merged

        body["company_name_english"] = stored_name
        if stored_reg:
            body["company_registration_number"] = stored_reg
        elif not body.get("company_registration_number"):
            # No registration number on either side: the upsert would insert a
            # duplicate keyed on an empty string. Nothing safe to update against.
            return {
                "success": False,
                "error": f"Company '{stored_name}' has no registration number; cannot update by name",
            }

        result = add_company(body)
        if result.get("success"):
            _persist_company_source_pdf(body)
            return {"success": True, "message": f"Company updated"}
        return {"success": False, "error": result.get("error", "Update failed")}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/dashboard/upload/template")
async def upload_dashboard_template(request: Request, file: UploadFile = File(...)):
    """Upload template Word file."""
    try:
        user = get_current_user(request)
        uploader_email = user.get("email", "unknown") if user else "unknown"
        content = await file.read()
        filename = file.filename or ""

        if not filename.endswith(".docx"):
            return {"success": False, "error": "Only .docx files allowed"}
        if not validate_filename(filename):
            return {"success": False, "error": "Invalid filename"}

        # Check if template already exists
        from scout.tools.template_analyzer import get_db_connection
        _check_conn = get_db_connection()
        _check_cur = _check_conn.cursor()
        _check_cur.execute("SELECT id, total_fields, category, ai_trained, created_at FROM templates WHERE name = %s", (filename,))
        existing = _check_cur.fetchone()
        if existing:
            _check_cur.close(); _check_conn.close()
            return {
                "success": False,
                "error": f"Template '{filename}' already exists. Delete it first before re-uploading.",
                "exists": True,
                "existing_template": {
                    "name": filename,
                    "fields": existing[1] or 0,
                    "category": existing[2] or "Unknown",
                    "trained": bool(existing[3]),
                    "uploaded_at": existing[4].isoformat() if existing[4] else None,
                }
            }
        _check_cur.close(); _check_conn.close()

        if len(content) > 50 * 1024 * 1024:  # 50MB limit
            return {"success": False, "error": "File too large (max 50MB)"}
        if len(content) < 100:
            return {"success": False, "error": "File too small — appears corrupted"}

        # Save to templates directory — verify path stays within allowed directory
        templates_dir = Path("/documents/legal/templates")
        save_path = (templates_dir / filename).resolve()
        if not str(save_path).startswith(str(templates_dir.resolve())):
            return {"success": False, "error": "Invalid file path"}
        save_path.parent.mkdir(parents=True, exist_ok=True)
        with open(save_path, "wb") as f:
            f.write(content)

        # Upload to S3 (background, non-blocking)
        s3_upload_async(str(save_path))

        # Quick: just extract placeholders and save to DB (no AI, no PDF)
        from scout.tools.template_analyzer import save_template_knowledge
        try:
            from docx import Document as _UpDoc
            from scout.tools.smart_doc import extract_placeholders_from_template
            analysis = extract_placeholders_from_template(save_path)
            field_list = analysis.get("fields", [])
        except Exception:
            field_list = []

        save_template_knowledge(filename, field_list, "/documents")
        invalidate_training(f"New template uploaded: {filename}", template_name=filename)

        # Track who uploaded
        try:
            from psycopg import connect as _uc
            import os as _uo
            _uconn = get_db_conn()
            _ucur = _uconn.cursor()
            _ucur.close(); _uconn.close()
        except Exception: pass
        log_activity(user.get("user_id") if user else None, uploader_email, "upload_template", f"Uploaded: {filename}", "")
        clear_cache()  # Clear dashboard cache

        return {
            "success": True,
            "message": f"Template uploaded: {filename}. Click 'Start Training' for AI analysis.",
            "fields": field_list,
            "training_required": True,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# DOCX → PDF conversion for template preview
# ---------------------------------------------------------------------------
@app.get("/api/templates/preview-pdf/{template_name}")
async def get_template_pdf(request: Request, template_name: str, token: str = ""):
    """Serve pre-generated PDF preview. Falls back to on-the-fly conversion if needed."""
    if not token:
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            token = auth[7:]
    if not token:
        return JSONResponse(status_code=401, content={"error": "Token required"})
    try:
        jwt.decode(token, getenv("JWT_SECRET_KEY", ""), algorithms=["HS256"])
    except Exception:
        return JSONResponse(status_code=401, content={"error": "Invalid token"})

    import urllib.parse
    import subprocess

    template_name = urllib.parse.unquote(template_name)
    base_dir = Path("/documents/legal/templates").resolve()
    docx_path = (base_dir / template_name).resolve()
    if not str(docx_path).startswith(str(base_dir)):
        return {"error": "Invalid filename"}

    if not docx_path.exists():
        return {"error": "Template not found"}

    pdf_dir = Path("/documents/legal/previews")
    pdf_dir.mkdir(parents=True, exist_ok=True)
    pdf_name = docx_path.stem + ".pdf"
    pdf_path = pdf_dir / pdf_name

    # If PDF doesn't exist or docx is newer, convert now (with yellow-highlighted placeholders)
    if not pdf_path.exists() or docx_path.stat().st_mtime > pdf_path.stat().st_mtime:
        import shutil
        if not shutil.which("libreoffice"):
            return {"error": "PDF conversion not available — LibreOffice not installed"}
        try:
            # Create highlighted copy, convert that to PDF
            highlighted_docx = pdf_dir / f"_hl_{template_name}"
            try:
                _highlight_placeholders_in_docx(docx_path, highlighted_docx)
                convert_source = highlighted_docx
            except Exception:
                convert_source = docx_path  # fallback to original

            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(convert_source)],
                capture_output=True, text=True, timeout=30
            )
            # Rename highlighted PDF to expected name
            hl_pdf = pdf_dir / (convert_source.stem + ".pdf")
            if hl_pdf.exists() and hl_pdf != pdf_path:
                hl_pdf.rename(pdf_path)
            # Clean up highlighted docx
            if highlighted_docx.exists():
                highlighted_docx.unlink(missing_ok=True)
            if result.returncode != 0:
                logger.warning(f"LibreOffice PDF conversion failed: {result.stderr[:200]}")
        except subprocess.TimeoutExpired:
            return {"error": "PDF conversion timed out"}
        except Exception as e:
            return {"error": f"PDF conversion failed: {e}"}

    if pdf_path.exists():
        from starlette.responses import Response
        pdf_bytes = pdf_path.read_bytes()
        return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": "inline"})
    return {"error": "PDF not available — conversion may have failed"}


@app.get("/api/documents/preview-pdf/{doc_name}")
async def get_document_pdf(request: Request, doc_name: str, token: str = ""):
    """Convert a generated document (output) to PDF for preview."""
    if not token:
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            token = auth[7:]
    if not token:
        return JSONResponse(status_code=401, content={"error": "Token required"})
    try:
        jwt.decode(token, getenv("JWT_SECRET_KEY", ""), algorithms=["HS256"])
    except Exception:
        return JSONResponse(status_code=401, content={"error": "Invalid token"})

    import urllib.parse
    import subprocess

    doc_name = urllib.parse.unquote(doc_name)
    output_base = Path("/documents/legal/output").resolve()
    docx_path = (output_base / doc_name).resolve()
    # Error responses MUST carry the right HTTP status, not a 200 with a JSON
    # body. The FE PdfViewer hands the response bytes straight to pdf.js; a
    # 200+JSON error was parsed as a PDF and surfaced the cryptic "Invalid PDF
    # structure" instead of a real reason. 4xx/5xx let the viewer show the cause.
    if not str(docx_path).startswith(str(output_base)):
        return JSONResponse(status_code=400, content={"error": "Invalid filename"})

    if not docx_path.exists():
        return JSONResponse(status_code=404, content={"error": "Source document file not found"})

    pdf_dir = Path("/documents/legal/previews")
    pdf_dir.mkdir(parents=True, exist_ok=True)
    pdf_name = docx_path.stem + ".pdf"
    pdf_path = pdf_dir / pdf_name

    if not pdf_path.exists() or docx_path.stat().st_mtime > pdf_path.stat().st_mtime:
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(docx_path)],
                capture_output=True, text=True, timeout=30
            )
        except Exception as e:
            logging.getLogger("legalscout").warning(f"preview-pdf convert failed for '{doc_name}': {e}")

    if pdf_path.exists():
        from starlette.responses import Response
        pdf_bytes = pdf_path.read_bytes()
        return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": "inline"})
    return JSONResponse(status_code=502, content={"error": "Preview could not be generated for this document"})


# ---------------------------------------------------------------------------
# Uploaded DICA filings — storage and provenance
#
# The uploaded PDF is the document every extracted field was copied out of, so
# the company row has to be able to point back at it. Two rules make that link
# trustworthy: the server, not the client, decides where the file lands, and the
# path is re-resolved inside the uploads directory before it is ever written to
# the database.
# ---------------------------------------------------------------------------
UPLOADS_DIR = Path("/documents/legal/uploads")


def _store_uploaded_pdf(content: bytes, client_filename: str | None) -> tuple[Path, str]:
    """Write an uploaded PDF into the uploads directory under a server-minted name.

    The client's filename is never used as a path. It is reduced to its basename,
    stripped to safe characters and prefixed with a UTC timestamp — two firms both
    filing "extract.pdf" must not overwrite each other's provenance, and a name
    like "../../app/main.py" must not escape the directory at all.

    Returns (absolute path on disk, stored filename).
    """
    raw = Path(client_filename or "upload.pdf").name  # drops any directory part
    cleaned = _re.sub(r"[^\w\-.]+", "_", raw).strip("._")
    if not cleaned.lower().endswith(".pdf"):
        cleaned = (cleaned or "upload") + ".pdf"
    cleaned = cleaned[-120:].lstrip("._") or "upload.pdf"

    stored_name = f"{datetime.utcnow().strftime('%Y%m%dT%H%M%S%f')}_{cleaned}"
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    base_dir = UPLOADS_DIR.resolve()
    dest = (base_dir / stored_name).resolve()
    if not str(dest).startswith(str(base_dir)):
        raise ValueError("Refusing to write outside the uploads directory")
    dest.write_bytes(content)
    return dest, stored_name


def _resolve_stored_upload(value: str | None) -> str | None:
    """Turn a client-supplied reference to an uploaded filing into a stored path.

    Accepts either the bare stored name or any of the URL forms the frontend
    carries around ("/documents/legal/uploads/x.pdf"). Only the basename is
    trusted; it is resolved inside the uploads directory and the file must
    actually be there. Anything else — traversal, a name for a file that was
    never uploaded, an absolute path — returns None rather than being stored.

    Returns the path relative to the documents root, e.g. "legal/uploads/x.pdf".
    """
    if not value or not isinstance(value, str):
        return None
    name = Path(value.strip()).name
    if not name or not name.lower().endswith(".pdf"):
        return None
    base_dir = UPLOADS_DIR.resolve()
    candidate = (base_dir / name).resolve()
    if not str(candidate).startswith(str(base_dir)):
        return None
    if not candidate.is_file():
        return None
    return f"legal/uploads/{name}"


def _persist_company_source_pdf(body: dict) -> None:
    """Record which uploaded filing a company row was extracted from.

    Called after add_company() has upserted the row. The reference comes from the
    client, so it is re-resolved against the uploads directory here and dropped
    unless it names a file that is really on disk.

    A save that carries no usable reference leaves whatever is already stored
    alone: an edit through the manual form must not silently erase the link to
    the filing the record came from.
    """
    stored = _resolve_stored_upload(body.get("source_pdf_path") or body.get("pdf_url"))
    if not stored:
        return

    reg_no = (body.get("company_registration_number") or "").strip()
    name = (body.get("company_name_english") or body.get("company_name") or "").strip()
    if not reg_no and not name:
        return

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        if reg_no:
            cur.execute(
                "UPDATE companies SET source_pdf_path = %s, updated_at = NOW() "
                "WHERE company_registration_number = %s",
                (stored, reg_no),
            )
        else:
            cur.execute(
                "UPDATE companies SET source_pdf_path = %s, updated_at = NOW() "
                "WHERE company_name_english = %s",
                (stored, name),
            )
        conn.commit()
        cur.close()
    except Exception as e:
        logging.getLogger("legalscout").warning(
            f"source PDF not recorded for '{name or reg_no}': {e}"
        )
    finally:
        if conn is not None:
            conn.close()


# ---------------------------------------------------------------------------
# PDF Upload (preview only, no AI extraction)
# ---------------------------------------------------------------------------
@app.post("/api/company/upload-pdf")
async def upload_company_pdf(file: UploadFile = File(...)):
    """Upload a PDF for preview. Returns the PDF URL only — no AI extraction."""
    try:
        content = await file.read()
        if len(content) > 50 * 1024 * 1024:
            return {"success": False, "error": "File too large (max 50MB)"}
        if not (file.filename or "").lower().endswith(".pdf"):
            return {"success": False, "error": "Only PDF files allowed"}
        _, stored_name = _store_uploaded_pdf(content, file.filename)
        return {
            "success": True,
            "pdf_url": f"/documents/legal/uploads/{stored_name}",
            "source_pdf_path": f"legal/uploads/{stored_name}",
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# PDF Company Extract - AI-powered extraction
# ---------------------------------------------------------------------------

@app.post("/api/company/extract-pdf-stream")
async def extract_company_from_pdf_stream(file: UploadFile = File(...)):
    """Streaming version of extract-pdf. Yields ND-JSON progress events."""
    from fastapi.responses import StreamingResponse
    import asyncio
    import tempfile
    import os
    import httpx

    content = await file.read()

    async def gen():
        def evt(stage: str, msg: str, **extra):
            return json.dumps({"stage": stage, "msg": msg, **extra}) + "\n"

        ai_text = ""
        try:
            yield evt("upload", f"Received {file.filename} ({len(content)} bytes)")
            await asyncio.sleep(0)

            if len(content) > 50 * 1024 * 1024:
                yield evt("error", "File too large (max 50MB)"); return

            try:
                import pdfplumber
            except ImportError:
                yield evt("error", "pdfplumber not installed"); return

            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(content); tmp_path = tmp.name

            # The saved file is the provenance for every field extracted below,
            # so the server names it — a client filename is not a path here.
            pdf_save_path, stored_name = _store_uploaded_pdf(content, file.filename)
            source_pdf_path = f"legal/uploads/{stored_name}"
            yield evt("save", f"Saved to {pdf_save_path}")

            try:
                from app.s3_storage import s3_upload_async
                s3_upload_async(str(pdf_save_path))
            except Exception:
                pass

            yield evt("parse", "Reading PDF text…")
            full_text = ""
            with pdfplumber.open(tmp_path) as pdf:
                num_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
                    yield evt("parse", f"Page {i}/{num_pages} parsed ({len(text or '')} chars)")
                    await asyncio.sleep(0)
            Path(tmp_path).unlink(missing_ok=True)

            if not full_text.strip():
                yield evt("error", "Could not extract text from PDF"); return

            openrouter_key = os.getenv("OPENROUTER_API_KEY")
            if not openrouter_key:
                yield evt("error", "OPENROUTER_API_KEY not set"); return

            chat_model = get_model("training")
            yield evt("ai", f"Calling {chat_model} for structured extraction…")

            extraction_prompt = (
                "Extract ALL company information from this Myanmar DICA Company Extract PDF. "
                "Return ONLY a JSON object with fields: company_name_english, company_name_myanmar, "
                "company_registration_number, registration_date (YYYY-MM-DD), status, company_type, "
                "foreign_company, small_company, under_corpsec_management, group_company, "
                "principal_activity, date_of_last_annual_return, previous_registration_number, "
                "registered_office_address, principal_place_of_business, directors[], "
                "ultimate_holding_company_name, ultimate_holding_company_jurisdiction, "
                "ultimate_holding_company_registration_number, total_shares_issued, total_capital, "
                "consideration_amount_paid, currency_of_share_capital, members[], filing_history[]. "
                "directors[] items: {name, type, date_of_appointment, date_of_birth, nationality, "
                "nrc_passport, gender, business_occupation}. members[] items: {type, name, "
                "registration_number, jurisdiction, share_quantity, amount_paid, amount_unpaid, "
                "share_class, nrc_passport, date_of_birth, gender, nationality}. "
                # The member table lists individuals with their full identity, exactly
                # like the officer table. Omitting those fields threw that identity
                # away and made the same human look like two people on the register.
                "For members[]: set type to \"Individual\" for a natural person or "
                "\"Company\" for a body corporate — never leave type null. A member "
                "with a registration_number is a Company; a member with an "
                "N.R.C./Passport, date of birth or gender is an Individual, and you "
                "MUST copy those identity fields across exactly as printed. "
                "Copy every N.R.C./Passport character-for-character; do not correct "
                "or normalise the township code even if it looks misspelt. "
                "filing_history[]: {form_type, effective_date}. "
                "Use null when unknown. Set company_name_myanmar to null. No markdown.\n\n"
                f"DOCUMENT:\n{full_text[:8000]}"
            )

            def _call_llm(model: str, prompt: str) -> str:
                r = httpx.post(
                    f"{OPENROUTER_BASE_URL}/chat/completions",
                    headers={"Authorization": f"Bearer {openrouter_key}", "Content-Type": "application/json"},
                    json={"model": model, "messages": [{"role": "user", "content": prompt}], "temperature": 0},
                    timeout=60,
                )
                r.raise_for_status()
                t = r.json()["choices"][0]["message"]["content"].strip()
                if t.startswith("```"):
                    t = t.split("```")[1]
                    if t.startswith("json"): t = t[4:]
                    t = t.strip()
                return t

            ai_text = await asyncio.to_thread(_call_llm, chat_model, extraction_prompt)
            yield evt("ai", f"AI returned {len(ai_text)} chars, parsing JSON…")
            extracted = json.loads(ai_text)
            yield evt("ai", f"Extracted: {extracted.get('company_name_english') or 'unknown'} · "
                            f"{len(extracted.get('directors') or [])} directors · "
                            f"{len(extracted.get('members') or [])} members")

            has_cid = "(cid:" in full_text
            if has_cid and not extracted.get("company_name_myanmar"):
                yield evt("ai", "CID codes detected — extracting Myanmar name…")
                try:
                    cls_model = get_model("classification")
                    myanmar_prompt = (
                        f"This PDF text contains Myanmar/Burmese company name encoded as CID codes. "
                        f"The English company name is: {extracted.get('company_name_english','Unknown')}. "
                        f"Return ONLY the Myanmar name (Burmese script) or 'null'.\n\nTEXT:\n{full_text[:3000]}"
                    )
                    myanmar_text = (await asyncio.to_thread(_call_llm, cls_model, myanmar_prompt)).strip().strip('"')
                    if myanmar_text and myanmar_text != "null" and "(cid:" not in myanmar_text:
                        extracted["company_name_myanmar"] = myanmar_text
                        yield evt("ai", f"Myanmar name: {myanmar_text}")
                except Exception as e:
                    yield evt("warn", f"Myanmar name extraction failed: {e}")

            # Carried inside the extracted record as well as on the event: the
            # record is what gets posted back to /api/dashboard/add/company, so
            # the link to the filing travels with the data it was read from and
            # cannot be dropped on the way.
            extracted["source_pdf_path"] = source_pdf_path
            extracted["pdf_url"] = f"/documents/legal/uploads/{stored_name}"

            yield evt("done", "Extraction complete", success=True, data=extracted,
                      pdf_url=f"/documents/legal/uploads/{stored_name}",
                      source_pdf_path=source_pdf_path)
        except json.JSONDecodeError as e:
            yield evt("error", f"AI returned invalid JSON: {e}", raw=(ai_text or "")[:500])
        except Exception as e:
            yield evt("error", str(e))

    return StreamingResponse(gen(), media_type="application/x-ndjson")


# ---------------------------------------------------------------------------
# System Management — reset/delete endpoints (admin only)
# ---------------------------------------------------------------------------
@app.post("/api/admin/reset/documents")
async def reset_documents(request: Request):
    """Delete all generated documents."""
    user = require_admin(request)
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("DELETE FROM documents")
        cur.execute("DELETE FROM document_versions")
        count = cur.rowcount
        conn.commit(); cur.close(); conn.close()

        import shutil
        output_dir = Path("/documents/legal/output")
        if output_dir.exists():
            for f in output_dir.glob("*.docx"):
                f.unlink()

        log_activity(user.get("user_id"), user.get("email"), "reset_documents", f"Deleted all documents", "")
        return {"success": True, "message": f"All documents deleted"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/admin/reset/chat")
async def reset_chat(request: Request):
    """Delete all chat sessions, AI memory, and learnings."""
    user = require_admin(request)
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        for t in ["agno_sessions", "agno_memories", "agno_learnings",
                   "scout_knowledge", "scout_learnings", "scout_knowledge_contents"]:
            try:
                cur.execute(f"DELETE FROM {t}")
            except Exception:
                conn.rollback()
        conn.commit(); cur.close(); conn.close()

        log_activity(user.get("user_id"), user.get("email"), "reset_chat", "Deleted all chat sessions and AI memory", "")
        return {"success": True, "message": "All chat sessions, AI memory, and learnings deleted"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/admin/reset/companies")
async def reset_companies(request: Request):
    """Delete all companies and knowledge data."""
    user = require_admin(request)
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("DELETE FROM knowledge_lookup")
        cur.execute("DELETE FROM knowledge_raw")
        cur.execute("DELETE FROM companies")
        conn.commit(); cur.close(); conn.close()

        invalidate_training("All companies deleted")
        log_activity(user.get("user_id"), user.get("email"), "reset_companies", "Deleted all companies", "")
        return {"success": True, "message": "All companies and knowledge data deleted. Re-training required."}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# AI Model Configuration — admin can change models from Settings
# ---------------------------------------------------------------------------
@app.get("/api/admin/timezone")
async def get_tz(request: Request):
    """Get current timezone setting."""
    require_admin(request)
    return {"success": True, "timezone": get_timezone(), "current_datetime": get_current_datetime()}


@app.post("/api/admin/timezone")
async def set_tz(request: Request):
    """Set timezone."""
    user = require_admin(request)
    try:
        body = await request.json()
        tz = body.get("timezone", "Asia/Yangon")
        # Validate timezone
        import zoneinfo
        zoneinfo.ZoneInfo(tz)  # throws if invalid
        save_timezone(tz)
        log_activity(user.get("user_id"), user.get("email"), "set_timezone", f"Timezone: {tz}", "")
        return {"success": True, "timezone": tz, "current_datetime": get_current_datetime()}
    except Exception as e:
        return {"success": False, "error": f"Invalid timezone: {e}"}


@app.post("/api/admin/test-model")
async def test_model(request: Request):
    """Test if a model works by sending a simple prompt."""
    require_admin(request)
    try:
        import httpx, time as _t
        body = await request.json()
        model = body.get("model", "")
        purpose = body.get("purpose", "chat")

        if not model:
            return {"success": False, "error": "No model specified"}

        openrouter_key = getenv("OPENROUTER_API_KEY")
        if not openrouter_key:
            return {"success": False, "error": "OPENROUTER_API_KEY not set"}

        start = _t.time()

        if purpose == "embedding":
            res = httpx.post(
                f"{OPENROUTER_BASE_URL}/embeddings",
                headers={"Authorization": f"Bearer {openrouter_key}", "Content-Type": "application/json"},
                json={"model": model, "input": "test embedding", "encoding_format": "float"},
                timeout=15,
            )
            res.raise_for_status()
            data = res.json()
            dims = len(data["data"][0]["embedding"])
            elapsed = round((_t.time() - start) * 1000)
            return {"success": True, "message": f"Working — {dims} dimensions, {elapsed}ms", "time_ms": elapsed}
        else:
            res = httpx.post(
                f"{OPENROUTER_BASE_URL}/chat/completions",
                headers={"Authorization": f"Bearer {openrouter_key}", "Content-Type": "application/json"},
                # 20 tokens cannot cover a reasoning model's thinking, so the
                # probe reported "(empty response)" for a perfectly healthy model.
                json={"model": model, "messages": [{"role": "user", "content": "Say OK"}], "max_tokens": 400, "temperature": 0},
                timeout=15,
            )
            res.raise_for_status()
            data = res.json()
            reply = (data["choices"][0]["message"].get("content") or "").strip() or "(empty response)"
            elapsed = round((_t.time() - start) * 1000)
            result = {"success": True, "message": f"Working — replied \"{reply}\", {elapsed}ms", "time_ms": elapsed}

        # Save test result to DB
        try:
            _tconn = get_db_conn()
            _tcur = _tconn.cursor()
            _tcur.execute(
                "INSERT INTO app_settings (key, value) VALUES (%s, %s) ON CONFLICT (key) DO UPDATE SET value = %s, updated_at = CURRENT_TIMESTAMP",
                (f"model_test_{purpose}", json.dumps({"model": model, "ok": result["success"], "msg": result.get("message", result.get("error","")), "time": datetime.now().isoformat()}),
                 json.dumps({"model": model, "ok": result["success"], "msg": result.get("message", result.get("error","")), "time": datetime.now().isoformat()})))
            _tconn.commit(); _tcur.close(); _tconn.close()
        except Exception as e:
            logging.getLogger("legalscout").warning(f"Model test log failed: {e}")

        return result
    except httpx.HTTPStatusError as e:
        return {"success": False, "error": f"API error {e.response.status_code}: {e.response.text[:100]}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# S3 Cloud Storage — config + sync
# ---------------------------------------------------------------------------
@app.get("/api/admin/s3")
async def get_s3_config(request: Request):
    """Get S3 configuration."""
    require_admin(request)
    config = _get_s3_config()
    # Mask secret key
    if config.get("secret_key"):
        config["secret_key"] = f"...{config['secret_key'][-4:]}"
    return {"success": True, "config": config}


@app.post("/api/admin/s3")
async def set_s3_config(request: Request):
    """Save S3 configuration."""
    user = require_admin(request)
    try:
        body = await request.json()
        config = {
            "enabled": body.get("enabled", False),
            "bucket": body.get("bucket", ""),
            "region": body.get("region", "us-east-1"),
            "access_key": body.get("access_key", ""),
            "secret_key": body.get("secret_key", ""),
            "endpoint_url": body.get("endpoint_url", ""),
        }
        # Don't overwrite secret if masked
        if config["secret_key"].startswith("..."):
            old = _get_s3_config()
            config["secret_key"] = old.get("secret_key", "")

        save_s3_config(config)
        log_activity(user.get("user_id"), user.get("email"), "update_s3", f"S3: {config['bucket']}", "")
        return {"success": True, "message": "S3 configuration saved"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/admin/s3/test")
async def test_s3(request: Request):
    """Test S3 connection."""
    require_admin(request)
    return s3_test()


@app.post("/api/admin/s3/sync")
async def sync_to_s3(request: Request):
    """Sync all local files to S3."""
    user = require_admin(request)
    result = s3_sync_all()
    if result.get("success"):
        log_activity(user.get("user_id"), user.get("email"), "s3_sync",
                     f"Synced {result.get('synced', 0)} files", "")
    return result


@app.get("/api/admin/s3/files")
async def list_s3_files(request: Request, prefix: str = ""):
    """List files in S3 bucket."""
    require_admin(request)
    files = s3_list(prefix)
    return {"success": True, "files": files, "total": len(files)}


@app.get("/api/admin/model-tests")
async def get_model_tests(request: Request):
    """Get saved model test results."""
    require_admin(request)
    try:
        from psycopg import connect
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT key, value FROM app_settings WHERE key LIKE 'model_test_%'")
        results = {}
        for row in cur.fetchall():
            purpose = row[0].replace("model_test_", "")
            data = json.loads(row[1]) if isinstance(row[1], str) else row[1]
            results[purpose] = data
        cur.close(); conn.close()
        return {"success": True, "tests": results}
    except Exception:
        return {"success": True, "tests": {}}


@app.get("/api/admin/api-keys")
async def get_api_key_status(request: Request):
    """Show API key status (set/not set, last 4 chars). Keys are read-only — edit .env to change."""
    require_admin(request)
    keys = {}
    for name in ["OPENROUTER_API_KEY"]:
        val = getenv(name, "")
        if val:
            keys[name] = {"set": True, "hint": f"...{val[-4:]}"}
        else:
            keys[name] = {"set": False, "hint": "Not set"}
    return {"success": True, "keys": keys}


@app.get("/api/changelog")
async def get_changelog():
    """Parsed CHANGELOG.md for the rail's version panel.

    ★ Registered HERE, beside /api/version and well above the frontend
    catch-all. A GET added below `@app.get("/{full_path:path}")` answers 200
    with the frontend's HTML and never runs — see the landmine in CLAUDE.md.

    Parsed server-side so no markdown library ships in the bundle. The format is
    fixed and ours: `## [x.y.z] — date`, `### Section`, `- item`.
    """
    from pathlib import Path
    import re as _re

    cf = Path("/app/CHANGELOG.md")
    if not cf.exists():
        return {"success": False, "error": "CHANGELOG.md not found", "versions": []}

    running = "unknown"
    vf = Path("/app/VERSION")
    if vf.exists():
        running = vf.read_text().strip()

    def _inline(text: str) -> str:
        """Strip the markdown we use; the panel renders plain text."""
        text = _re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = _re.sub(r"`([^`]+)`", r"\1", text)
        text = _re.sub(r"~~(.+?)~~", r"\1", text)
        return text.strip()

    versions, cur, sec = [], None, None
    for raw in cf.read_text().split("\n"):
        line = raw.rstrip()
        m = _re.match(r"^##\s+\[([^\]]+)\]\s*[—-]\s*(.+)$", line)
        if m:
            cur = {"version": m.group(1).strip(), "date": m.group(2).strip(),
                   "running": m.group(1).strip() == running, "sections": []}
            versions.append(cur)
            sec = None
            continue
        if cur is None:
            continue
        m = _re.match(r"^###\s+(.+)$", line)
        if m:
            sec = {"title": m.group(1).strip(), "items": []}
            cur["sections"].append(sec)
            continue
        m = _re.match(r"^[-*]\s+(.+)$", line)
        if m:
            if sec is None:
                sec = {"title": "", "items": []}
                cur["sections"].append(sec)
            sec["items"].append(_inline(m.group(1)))
            continue
        # a wrapped continuation of the previous bullet
        if sec and sec["items"] and line.startswith(("  ", "\t")) and line.strip():
            if not line.strip().startswith("|"):
                sec["items"][-1] = (sec["items"][-1] + " " + _inline(line)).strip()

    return {"success": True, "running": running, "versions": versions}


@app.get("/api/version")
async def get_version():
    """Get app version."""
    from pathlib import Path
    version = "unknown"
    vf = Path("/app/VERSION")
    if vf.exists():
        version = vf.read_text().strip()
    return {"version": version, "app": "Legal Scout"}


@app.post("/api/admin/log-level")
async def set_log_level(request: Request):
    """Change log level at runtime without restart."""
    user = require_admin(request)
    import logging
    body = await request.json()
    level = body.get("level", "INFO").upper()
    if level not in ("DEBUG", "INFO", "WARNING", "ERROR"):
        return {"success": False, "error": f"Invalid level: {level}. Use DEBUG, INFO, WARNING, ERROR"}
    logging.getLogger().setLevel(getattr(logging, level))
    log_activity(user.get("user_id"), user.get("email"), "set_log_level", f"Log level: {level}", "")
    return {"success": True, "message": f"Log level set to {level}", "level": level}


@app.get("/api/admin/models")
async def get_models(request: Request):
    """Get current AI model configuration."""
    require_admin(request)
    return {"success": True, "models": get_all_models()}


@app.post("/api/admin/models")
async def update_models(request: Request):
    """Save AI model configuration."""
    user = require_admin(request)
    try:
        body = await request.json()
        models = body.get("models", {})

        # Validate — only allow known purposes
        valid_keys = {"chat", "training", "classification", "embedding"}
        cleaned = {k: v for k, v in models.items() if k in valid_keys and v}

        if not cleaned:
            return {"success": False, "error": "No valid model configuration provided"}

        save_models(cleaned)
        clear_model_cache()

        log_activity(user.get("user_id"), user.get("email"), "update_models",
                     f"Models updated: {cleaned}", "")

        # Hot-reload agent model
        try:
            if "chat" in cleaned:
                from agno.models.openai import OpenAIChat as _OAC
                new_model = cleaned["chat"]
                scout.model = _OAC(
                    id=new_model,
                    api_key=getenv("OPENROUTER_API_KEY"),
                    base_url=OPENROUTER_BASE_URL,
                )
                logger.info(f"Agent model reloaded: {new_model}")
        except Exception as e:
            logger.warning(f"Model reload failed: {e}")

        return {
            "success": True,
            "message": "Model configuration saved and applied.",
            "models": get_all_models(),
            "restart_required": False,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# Activity Dashboard — tracks all system events
# ---------------------------------------------------------------------------
@app.get("/api/admin/activity")
async def get_activity(request: Request, days: int = 30, limit: int = 200):
    """Get system activity for dashboard — logins, uploads, training, generation, etc."""
    require_admin(request)
    try:
        from psycopg import connect
        import os
        conn = get_db_conn()
        cur = conn.cursor()

        # Recent activity logs
        cur.execute("""
            SELECT id, user_email, action, details, ip_address, created_at
            FROM activity_logs
            WHERE created_at > NOW() - INTERVAL '%s days'
            ORDER BY created_at DESC LIMIT %s
        """, (days, limit))
        logs = [{"id": r[0], "user": r[1] or "system", "action": r[2], "details": r[3] or "",
                 "ip": r[4] or "", "time": r[5].isoformat() if r[5] else None} for r in cur.fetchall()]

        # Activity by action type (for chart)
        cur.execute("""
            SELECT action, COUNT(*) FROM activity_logs
            WHERE created_at > NOW() - INTERVAL '%s days'
            GROUP BY action ORDER BY count DESC
        """, (days,))
        by_action = {r[0]: r[1] for r in cur.fetchall()}

        # Activity by day (for timeline chart)
        cur.execute("""
            SELECT DATE(created_at) as day, COUNT(*) FROM activity_logs
            WHERE created_at > NOW() - INTERVAL '%s days'
            GROUP BY day ORDER BY day
        """, (days,))
        by_day = [{"date": r[0].isoformat(), "count": r[1]} for r in cur.fetchall()]

        # Activity by user (for user chart)
        cur.execute("""
            SELECT COALESCE(user_email, 'system'), COUNT(*) FROM activity_logs
            WHERE created_at > NOW() - INTERVAL '%s days'
            GROUP BY user_email ORDER BY count DESC LIMIT 10
        """, (days,))
        by_user = {r[0]: r[1] for r in cur.fetchall()}

        # Recent documents
        cur.execute("""
            SELECT template_name, company_name, file_name, created_at
            FROM documents ORDER BY created_at DESC LIMIT 10
        """)
        recent_docs = [{"template": r[0], "company": r[1], "file": r[2],
                        "time": r[3].isoformat() if r[3] else None}
                       for r in cur.fetchall()]

        # Recent templates
        cur.execute("""
            SELECT name, created_at FROM templates ORDER BY created_at DESC LIMIT 10
        """)
        recent_templates = [{"name": r[0],
                            "time": r[1].isoformat() if r[1] else None} for r in cur.fetchall()]

        # Recent companies
        cur.execute("""
            SELECT company_name_english, created_at FROM companies ORDER BY created_at DESC LIMIT 10
        """)
        recent_companies = [{"name": r[0],
                            "time": r[1].isoformat() if r[1] else None} for r in cur.fetchall()]

        cur.close(); conn.close()

        return {
            "success": True,
            "logs": logs,
            "summary": {
                "total_events": len(logs),
                "by_action": by_action,
                "by_day": by_day,
                "by_user": by_user,
            },
            "recent": {
                "documents": recent_docs,
                "templates": recent_templates,
                "companies": recent_companies,
            }
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# Email Logs — view all sent emails
# ---------------------------------------------------------------------------
@app.get("/api/admin/emails")
async def get_email_logs(request: Request, limit: int = 100):
    """Get all sent email logs."""
    require_admin(request)
    try:
        from psycopg import connect
        import os
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, to_email, subject, body, attachment_name, attachment_path,
                   sent_by_email, status, error_message, created_at
            FROM email_logs ORDER BY created_at DESC LIMIT %s
        """, (limit,))
        emails = [{
            "id": r[0], "to": r[1], "subject": r[2], "body": r[3],
            "attachment": r[4], "attachment_path": r[5],
            "sent_by": r[6] or "unknown", "status": r[7],
            "error": r[8], "time": r[9].isoformat() if r[9] else None,
        } for r in cur.fetchall()]
        cur.close(); conn.close()
        return {"success": True, "emails": emails, "total": len(emails)}
    except Exception as e:
        return {"success": True, "emails": [], "total": 0}


@app.post("/api/admin/reset/templates")
async def reset_templates(request: Request):
    """Delete all templates from DB and filesystem."""
    user = require_admin(request)
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("DELETE FROM templates")
        cur.execute("DELETE FROM template_versions")
        conn.commit(); cur.close(); conn.close()

        templates_dir = Path("/documents/legal/templates")
        if templates_dir.exists():
            for f in templates_dir.glob("*.docx"):
                f.unlink()
            for f in templates_dir.glob("*.pdf"):
                f.unlink()

        log_activity(user.get("user_id"), user.get("email"), "reset_templates", "Deleted all templates", "")
        return {"success": True, "message": "All templates deleted"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/admin/reset/all")
async def reset_all(request: Request):
    """Delete ALL data — companies, documents, knowledge. Templates are preserved."""
    user = require_admin(request)
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        # App data
        cur.execute("DELETE FROM documents")
        cur.execute("DELETE FROM document_versions")
        cur.execute("DELETE FROM knowledge_lookup")
        cur.execute("DELETE FROM knowledge_raw")
        cur.execute("DELETE FROM knowledge_vec")
        cur.execute("DELETE FROM companies")
        cur.execute("DELETE FROM activity_logs")
        cur.execute("DELETE FROM training_status")
        # Chat sessions & AI memory
        for t in ["agno_sessions", "agno_memories", "agno_learnings",
                   "scout_knowledge", "scout_learnings", "scout_knowledge_contents"]:
            try:
                cur.execute(f"DELETE FROM {t}")
            except Exception:
                conn.rollback()
        conn.commit(); cur.close(); conn.close()

        output_dir = Path("/documents/legal/output")
        if output_dir.exists():
            for f in output_dir.glob("*.docx"):
                f.unlink()

        return {"success": True, "message": "All data deleted (documents, companies, knowledge, chat sessions, AI memory). Templates preserved."}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/admin/restore")
async def restore_backup(request: Request, file: UploadFile = File(...)):
    """Restore data from a JSON backup file."""
    user = require_admin(request)
    try:
        content = await file.read()
        backup = json.loads(content)

        conn = get_db_conn()
        cur = conn.cursor()

        restored = {}

        # Restore order matters — tables with foreign keys last
        restore_tables = [
            ("companies", ["company_name_english", "company_registration_number", "company_type",
                          "registered_office_address", "principal_place_of_business", "directors",
                          "total_shares_issued", "currency_of_share_capital", "members",
                          "total_capital", "source"]),
            ("knowledge_lookup", ["key_name", "key_value", "source_file"]),
            ("knowledge_raw", ["source_file", "file_type", "sheet_name", "row_number", "data"]),
            ("training_status", ["training_type", "last_trained", "record_count", "status", "logs"]),
            ("documents", ["template_name", "company_name", "file_name", "file_path", "download_url",
                          "validation_result", "custom_data", "version"]),
            ("document_versions", ["document_name", "company_name", "template_name", "version",
                                  "file_name", "file_path", "generated_by_email", "custom_data"]),
            ("activity_logs", ["user_email", "action", "details", "ip_address"]),
            ("app_settings", ["key", "value", "updated_by"]),
        ]

        # Agno tables (simple key-value style, restore if present)
        agno_tables = ["agno_sessions", "agno_memories", "agno_learnings",
                       "scout_knowledge", "scout_learnings", "scout_knowledge_contents"]

        for table_name, columns in restore_tables:
            rows = backup.get(table_name, [])
            if not rows:
                continue

            count = 0
            for row in rows:
                # Filter to only columns that exist in the row
                available_cols = [c for c in columns if c in row and row[c] is not None]
                if not available_cols:
                    continue

                values = []
                for c in available_cols:
                    val = row[c]
                    # Convert JSONB fields
                    if isinstance(val, (dict, list)):
                        val = json.dumps(val)
                    values.append(val)

                placeholders = ", ".join(["%s"] * len(available_cols))
                col_names = ", ".join(available_cols)

                try:
                    cur.execute(
                        f"INSERT INTO {table_name} ({col_names}) VALUES ({placeholders}) ON CONFLICT DO NOTHING",
                        values
                    )
                    count += 1
                except Exception:
                    conn.rollback()

            restored[table_name] = count

        # Restore agno tables with column whitelist to prevent SQL injection
        AGNO_SAFE_COLUMNS = {
            "agno_sessions": {"session_id", "agent_id", "user_id", "memory", "agent_data", "session_data", "created_at", "updated_at"},
            "agno_memories": {"id", "user_id", "memory", "topics", "created_at", "updated_at"},
            "agno_learnings": {"id", "user_id", "learning", "topics", "created_at", "updated_at"},
            "scout_learnings": {"id", "user_id", "learning", "topics", "created_at", "updated_at"},
        }
        _safe_col_re = _re.compile(r"^[a-zA-Z_][a-zA-Z0-9_]*$")

        for table_name in agno_tables:
            rows = backup.get(table_name, [])
            if not rows:
                continue

            safe_cols_set = AGNO_SAFE_COLUMNS.get(table_name, set())
            count = 0
            for row in rows:
                cols = [k for k in row.keys() if k != "id" and k in safe_cols_set and _safe_col_re.match(k)]
                if not cols:
                    continue
                values = []
                for c in cols:
                    val = row[c]
                    if isinstance(val, (dict, list)):
                        val = json.dumps(val)
                    values.append(val)
                placeholders = ", ".join(["%s"] * len(cols))
                col_names = ", ".join(cols)
                try:
                    cur.execute(f"INSERT INTO {table_name} ({col_names}) VALUES ({placeholders}) ON CONFLICT DO NOTHING", values)
                    count += 1
                except Exception:
                    conn.rollback()
            restored[table_name] = count

        conn.commit(); cur.close(); conn.close()

        log_activity(user.get("user_id"), user.get("email"), "restore_backup",
                     f"Restored: {restored}", "")

        total = sum(restored.values())
        return {"success": True, "message": f"Restored {total} records across {len(restored)} tables", "details": restored}
    except json.JSONDecodeError:
        return {"success": False, "error": "Invalid JSON file"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# Legal Skills — Agent-Skills engine (L1 metadata in prompt, L2 body on demand)
# Registered before the frontend catch-all so the GET routes aren't shadowed.
# ---------------------------------------------------------------------------
import re as _skills_re

_SKILL_NAME_RE = _skills_re.compile(r"^[a-z0-9]+(?:-[a-z0-9]+)*$")


def _valid_skill_name(name: str) -> bool:
    return bool(name) and len(name) <= 64 and bool(_SKILL_NAME_RE.match(name))


# Tool references inside a skill BODY, checked against the live registry.
#
# Skill bodies are L2 playbooks — markdown the firm's lawyers edit in this admin
# UI — and they instruct the agent to call tools by name. A body that names a
# tool which is not registered fails SILENTLY and in the worst possible way: the
# model follows the instruction, finds no such tool, and simply skips that step.
# No exception, no log line, no missing document — just a step that quietly did
# not happen. Seven enabled skills shipped naming `preview_document` (the export
# key; @wraps made the real name `preview_doc`) and `list_tracked_documents`.
#
# The startup contract audit cannot catch this: _build_legal_skills_block()
# selects only name + description for the L1 prompt block, so bodies never pass
# through _audit_prompt_tool_contract(). This is the only gate they get.
#
# WHY A PREFIX HEURISTIC AND NOT A WORD LIST:
# a skill body is prose. Any rule that flags every unknown snake_case token
# would reject legitimate writing — placeholder names, DB columns, file paths,
# ordinary phrases in backticks. So a candidate is flagged only when it BOTH
# looks like a tool call (backticked, or written with call parens) AND shares a
# 6-character prefix with a real tool in either direction. That is precisely the
# near-miss shape of the actual failure mode — a typo or a stale name for a tool
# that does exist — and it ignores everything else. It is deliberately
# incomplete: a wholly invented name sharing no prefix with anything registered
# passes. Better a gap than a validator lawyers learn to work around.
_SKILL_TOOL_BACKTICKED_RE = _skills_re.compile(r"`([a-z_][a-z0-9_]{3,})\(?\)?`")
_SKILL_TOOL_CALL_RE = _skills_re.compile(r"\b([a-z_][a-z0-9_]{3,})\(\)")


def _unknown_skill_tools(skill_body: str) -> list[str]:
    """Names in `skill_body` that read as tool calls but are not registered.

    Returns a list of ready-to-display messages, each naming the closest
    registered match. Empty list means the body is acceptable.
    """
    if not skill_body:
        return []
    import difflib

    # Same in-function import the other agent-touching helpers use
    # (_refresh_agent_knowledge, _refresh_legal_skills, health_check) —
    # app/main.py and scout.agent have a circular relationship, and this reads
    # the registry LIVE rather than a copy taken at import time.
    import scout.agent as _am

    registry = _am._registered_tool_names(_am.scout.tools or [])
    if not registry:
        return []

    candidates: list[str] = []
    seen: set[str] = set()
    for match in _SKILL_TOOL_BACKTICKED_RE.findall(skill_body) + _SKILL_TOOL_CALL_RE.findall(skill_body):
        if match not in seen:
            seen.add(match)
            candidates.append(match)

    problems: list[str] = []
    for cand in candidates:
        if cand in registry:
            continue
        near = [t for t in registry if t.startswith(cand[:6]) or cand.startswith(t[:6])]
        if not near:
            continue
        # The 6-char prefix decides IF this is flagged; difflib only picks which
        # of the near names to name in the message, so several tools sharing a
        # prefix ('list_*') don't get suggested in alphabetical order.
        best = difflib.get_close_matches(cand, sorted(near), n=1, cutoff=0.0)[0]
        problems.append(f"Unknown tool '{cand}' (did you mean '{best}'?)")
    return problems


def _reject_unknown_skill_tools(skill_body: str, skill_name: str) -> None:
    """Raise 400 if the body names a near-miss tool. WRITES ONLY.

    Never call this on a read path or from _refresh_legal_skills(): the seeded
    rows still carry the bad names until a migration fixes them, and validating
    on read would make the Skills tab fail to load.
    """
    problems = _unknown_skill_tools(skill_body)
    if not problems:
        return
    logging.getLogger("legalscout").warning(
        f"[SKILLS] Rejected write to '{skill_name}' — {len(problems)} unregistered tool "
        f"reference(s): {'; '.join(problems)}"
    )
    raise HTTPException(
        status_code=400,
        detail=(
            "This skill body names tools the agent cannot call, so those steps would be "
            "skipped silently. " + " ".join(problems)
        ),
    )


@app.get("/api/skills")
async def list_legal_skills(request: Request):
    """List all legal skills (any authenticated role). Bodies excluded — body_chars only."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, name, description, version, enabled, source, updated_at,
                   char_length(body)
            FROM legal_skills ORDER BY name
        """)
        rows = cur.fetchall()
        cur.close()
        skills = [{
            "id": r[0], "name": r[1], "description": r[2], "version": r[3],
            "enabled": r[4], "source": r[5],
            "updated_at": r[6].isoformat() if r[6] else None,
            "body_chars": r[7] or 0,
        } for r in rows]
        return {"success": True, "skills": skills, "count": len(skills)}
    except Exception as e:
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.get("/api/skills/{name}")
async def get_legal_skill(name: str, request: Request):
    """Get a single legal skill including its full body (any authenticated role)."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT id, name, description, body, version, enabled, source, updated_at
            FROM legal_skills WHERE name = %s
        """, (name,))
        row = cur.fetchone()
        cur.close()
        if not row:
            raise HTTPException(status_code=404, detail="Skill not found")
        return {"success": True, "skill": {
            "id": row[0], "name": row[1], "description": row[2], "body": row[3],
            "version": row[4], "enabled": row[5], "source": row[6],
            "updated_at": row[7].isoformat() if row[7] else None,
        }}
    except HTTPException:
        raise
    except Exception as e:
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.post("/api/skills")
async def create_legal_skill(request: Request):
    """Create a new manual legal skill (admin only)."""
    user = require_admin(request)
    conn = None
    try:
        body = await request.json()
        name = (body.get("name") or "").strip()
        description = (body.get("description") or "").strip()
        skill_body = body.get("body") or ""
        if not _valid_skill_name(name):
            return {"success": False, "error": "Invalid name — use kebab-case, max 64 chars (a-z, 0-9, hyphens)."}
        if not description:
            return {"success": False, "error": "Description is required."}
        if not skill_body.strip():
            return {"success": False, "error": "Body is required."}
        _reject_unknown_skill_tools(skill_body, name)
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT 1 FROM legal_skills WHERE name = %s", (name,))
        if cur.fetchone():
            cur.close()
            return {"success": False, "error": f"Skill '{name}' already exists."}
        cur.execute("""
            INSERT INTO legal_skills (name, description, body, source)
            VALUES (%s, %s, %s, 'manual')
            RETURNING id
        """, (name, description, skill_body))
        new_id = cur.fetchone()[0]
        conn.commit(); cur.close()
        _refresh_legal_skills()
        log_activity(user.get("user_id"), user.get("email"), "create_skill", f"Created skill {name}", "")
        return {"success": True, "id": new_id, "name": name}
    except HTTPException:
        # Without this the validator's 400 is caught below and returned as a
        # 200 with success:false — the status code the UI checks would be a lie.
        if conn:
            conn.rollback()
        raise
    except Exception as e:
        if conn:
            conn.rollback()
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.put("/api/skills/{name}")
async def update_legal_skill(name: str, request: Request):
    """Update a legal skill's description/body/version (admin only)."""
    user = require_admin(request)
    conn = None
    try:
        body = await request.json()
        fields = []
        params = []
        if "description" in body:
            desc = (body.get("description") or "").strip()
            if not desc:
                return {"success": False, "error": "Description cannot be empty."}
            fields.append("description = %s"); params.append(desc)
        if "body" in body:
            new_body = body.get("body") or ""
            if not new_body.strip():
                return {"success": False, "error": "Body cannot be empty."}
            _reject_unknown_skill_tools(new_body, name)
            fields.append("body = %s"); params.append(new_body)
        if "version" in body:
            ver = (body.get("version") or "").strip()[:16]
            fields.append("version = %s"); params.append(ver)
        if not fields:
            return {"success": False, "error": "Nothing to update."}
        fields.append("updated_at = now()")
        conn = get_db_conn()
        cur = conn.cursor()
        params.append(name)
        cur.execute(
            f"UPDATE legal_skills SET {', '.join(fields)} WHERE name = %s RETURNING id",
            tuple(params),
        )
        row = cur.fetchone()
        if not row:
            cur.close()
            raise HTTPException(status_code=404, detail="Skill not found")
        conn.commit(); cur.close()
        _refresh_legal_skills()
        log_activity(user.get("user_id"), user.get("email"), "update_skill", f"Updated skill {name}", "")
        return {"success": True, "name": name}
    except HTTPException:
        raise
    except Exception as e:
        if conn:
            conn.rollback()
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.post("/api/skills/{name}/toggle")
async def toggle_legal_skill(name: str, request: Request):
    """Enable/disable a legal skill (admin only)."""
    user = require_admin(request)
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(
            "UPDATE legal_skills SET enabled = NOT enabled, updated_at = now() WHERE name = %s RETURNING enabled",
            (name,),
        )
        row = cur.fetchone()
        if not row:
            cur.close()
            raise HTTPException(status_code=404, detail="Skill not found")
        conn.commit(); cur.close()
        _refresh_legal_skills()
        log_activity(user.get("user_id"), user.get("email"), "toggle_skill", f"Skill {name} enabled={row[0]}", "")
        return {"success": True, "name": name, "enabled": row[0]}
    except HTTPException:
        raise
    except Exception as e:
        if conn:
            conn.rollback()
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


@app.delete("/api/skills/{name}")
async def delete_legal_skill(name: str, request: Request):
    """Delete a legal skill (admin only). Only source='manual' is deletable;
    'adapted'/'template' skills are toggle-only."""
    user = require_admin(request)
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT source FROM legal_skills WHERE name = %s", (name,))
        row = cur.fetchone()
        if not row:
            cur.close()
            raise HTTPException(status_code=404, detail="Skill not found")
        if row[0] != "manual":
            cur.close()
            return {"success": False, "error": f"Skill '{name}' (source={row[0]}) cannot be deleted — disable it instead."}
        cur.execute("DELETE FROM legal_skills WHERE name = %s", (name,))
        conn.commit(); cur.close()
        _refresh_legal_skills()
        log_activity(user.get("user_id"), user.get("email"), "delete_skill", f"Deleted skill {name}", "")
        return {"success": True, "name": name}
    except HTTPException:
        raise
    except Exception as e:
        if conn:
            conn.rollback()
        return {"success": False, "error": str(e)}
    finally:
        if conn:
            conn.close()


# ---------------------------------------------------------------------------
# Frontend Static Files — serve Next.js export from /app/static-frontend/
# ---------------------------------------------------------------------------
# Outbound email approval gate
# ---------------------------------------------------------------------------
# The agent can QUEUE an email (scout/agent.py:send_email_tool). It cannot send
# one. Delivery happens here, and only here, behind a request carrying the
# user's own JWT — which the agent has no way to obtain. That asymmetry is the
# entire control: an approval the agent could grant itself would not be an
# approval, so there is deliberately no "confirmed" flag the model can set.
#
# ★ These must stay ABOVE the `@app.get("/{full_path:path}")` catch-all below.
# Defined after it, GET /api/email/queued matched the catch-all instead and
# returned the frontend's HTML with a 200 — so the queue looked permanently
# empty and no approval could ever be given. The POSTs still worked, because
# the catch-all only claims GET, which is exactly the kind of half-working that
# survives a quick test.


@app.get("/api/email/queued")
async def list_queued_emails(request: Request):
    """List emails the agent has queued and nobody has approved yet."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, to_email, subject, body, attachment_name, created_at
            FROM email_logs
            WHERE status = 'queued'
            ORDER BY created_at DESC
            LIMIT 50
            """
        )
        rows = cur.fetchall()
        cur.close()
        return {
            "success": True,
            "queued": [
                {
                    "id": r[0],
                    "to_email": r[1],
                    "subject": r[2],
                    "body": r[3],
                    "attachment_name": r[4],
                    "created_at": r[5].isoformat() if r[5] else None,
                }
                for r in rows
            ],
        }
    except Exception as e:
        logging.getLogger("legalscout").error(f"Listing queued emails failed: {e}")
        raise HTTPException(status_code=500, detail="Could not list queued emails")
    finally:
        if conn:
            conn.close()


@app.post("/api/email/queued/{email_id}/discard")
async def discard_queued_email(email_id: int, request: Request):
    """Reject a queued email. It is never delivered."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(
            """
            UPDATE email_logs
            SET status = 'discarded', decided_at = NOW(), decided_by_email = %s
            WHERE id = %s AND status = 'queued'
            RETURNING to_email
            """,
            (user.get("email", "unknown"), email_id),
        )
        row = cur.fetchone()
        conn.commit()
        cur.close()
        if not row:
            raise HTTPException(status_code=404, detail="No such queued email")
        log_activity(user.get("user_id"), user.get("email", "unknown"),
                     "email_discarded", f"Discarded queued email #{email_id} to {row[0]}", "")
        return {"success": True, "discarded": email_id}
    except HTTPException:
        raise
    except Exception as e:
        logging.getLogger("legalscout").error(f"Discarding email {email_id} failed: {e}")
        raise HTTPException(status_code=500, detail="Could not discard the email")
    finally:
        if conn:
            conn.close()


@app.post("/api/email/queued/{email_id}/send")
async def send_queued_email(email_id: int, request: Request):
    """Deliver a queued email. This is the ONLY path to SMTP for agent-composed mail."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")

    approver = user.get("email", "unknown")
    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        # Claim the row before doing any work. Two clicks, or a click racing a
        # second tab, must not send the same email twice — and an email cannot
        # be recalled, so the claim has to happen before SMTP, not after.
        cur.execute(
            """
            UPDATE email_logs
            SET status = 'sending', decided_at = NOW(), decided_by_email = %s
            WHERE id = %s AND status = 'queued'
            RETURNING to_email, subject, body, attachment_name, attachment_path
            """,
            (approver, email_id),
        )
        row = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()
        conn = None
        if not row:
            raise HTTPException(
                status_code=409,
                detail="That email is not waiting for approval — it may already have been sent or discarded.",
            )
        to_email, subject, body, attachment_name, attachment_path = row
    except HTTPException:
        raise
    except Exception as e:
        logging.getLogger("legalscout").error(f"Claiming email {email_id} failed: {e}")
        if conn:
            conn.close()
        raise HTTPException(status_code=500, detail="Could not send the email")

    def _finish(status: str, error: str = ""):
        try:
            c = get_db_conn()
            k = c.cursor()
            k.execute(
                "UPDATE email_logs SET status = %s, error_message = %s, sent_by_email = %s WHERE id = %s",
                (status, error or None, approver, email_id),
            )
            c.commit()
            k.close()
            c.close()
        except Exception as e:  # noqa: BLE001
            logging.getLogger("legalscout").error(f"Recording email {email_id} outcome failed: {e}")

    try:
        import os
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        from email.mime.base import MIMEBase
        from email import encoders

        full_path = None
        if attachment_path:
            # The path was resolved when queued, but re-check it here: this is
            # the boundary that actually opens the file, and the row in between
            # was written by a tool the model drives.
            output_base = Path("/documents/legal/output").resolve()
            templates_base = Path("/documents/legal/templates").resolve()
            candidate = Path(attachment_path).resolve()
            inside = str(candidate).startswith(str(output_base)) or str(
                candidate
            ).startswith(str(templates_base))
            if not inside or not candidate.exists():
                _finish("failed", f"Attachment outside the documents directory or missing: {attachment_path}")
                return {"success": False, "error": "The attachment is no longer available."}
            full_path = candidate

        sconn = get_db_conn()
        scur = sconn.cursor()
        scur.execute("SELECT key, value FROM app_settings WHERE key LIKE 'smtp_%'")
        smtp = {r[0]: r[1] for r in scur.fetchall()}
        scur.close()
        sconn.close()

        smtp_host = smtp.get("smtp_host") or os.getenv("SMTP_HOST", "")
        smtp_port = int(smtp.get("smtp_port") or os.getenv("SMTP_PORT", "587"))
        smtp_user = smtp.get("smtp_user") or os.getenv("SMTP_USER", "")
        smtp_pass = smtp.get("smtp_pass") or os.getenv("SMTP_PASS", "")
        smtp_from = smtp.get("smtp_from") or os.getenv("SMTP_FROM", "noreply@legalscout.com")

        if not smtp_host or not smtp_user:
            # Put it back so the approval is not lost to a config problem.
            _finish("queued")
            return {"success": False, "error": "Email is not configured. Go to Settings to configure SMTP."}

        msg = MIMEMultipart()
        msg["Subject"] = f"Legal Scout — {subject}" if subject else "Legal Scout"
        msg["From"] = smtp_from
        msg["To"] = to_email
        msg.attach(MIMEText(body or "", "plain"))

        if full_path is not None:
            with open(full_path, "rb") as f:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header("Content-Disposition", f"attachment; filename={full_path.name}")
                msg.attach(part)

        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_pass)
            server.send_message(msg)

        _finish("sent")
        log_activity(user.get("user_id"), approver, "email_approved_and_sent",
                     f"Sent queued email #{email_id} to {to_email}"
                     + (f" with {attachment_name}" if attachment_name else ""), "")
        return {
            "success": True,
            "message": f"Sent to {to_email}"
            + (f" with {attachment_name} attached" if attachment_name else ""),
        }
    except Exception as e:
        # 'failed', not back to 'queued'. SMTP may have accepted the message
        # before the error, and re-offering Send on something that might already
        # be delivered is how a client receives the same filing twice.
        _finish("failed", str(e))
        logging.getLogger("legalscout").error(f"Sending queued email {email_id} failed: {e}")
        return {"success": False, "error": f"Delivery failed: {e}"}


# ---------------------------------------------------------------------------
# Message feedback — thumbs up / down on one assistant reply
# ---------------------------------------------------------------------------
# STORAGE ONLY. These two handlers write and read message_feedback
# (db/migration_027_message_feedback.sql) and nothing else reads that table:
# no retraining, no prompt change, no retrieval weighting. Wiring the votes
# into anything is a separate decision that has not been taken.
#
# ★ Like /api/email/queued above, these must stay ABOVE the
# `@app.get("/{full_path:path}")` catch-all. A GET registered below it is
# answered by the frontend's HTML with a 200, which is indistinguishable from
# working until someone parses the body.
#
# ★ WHICH MESSAGE A VOTE IS ABOUT. Agno's own message ids never reach the
# client — the streaming parser drops them — so MessageFeedback.tsx DERIVES one
# from the fields ChatMessage does have: "<role>-<run_id>-<index>", falling back
# to "<role>-<created_at>-<index>" before a run_id exists. That string is the
# identity here and it is OPAQUE to this file: never parsed, never split, never
# assumed to be a UUID, validated only for presence and length. Its format
# belongs to the client and has already changed twice.
# message_index and run_id are accepted and stored when sent (ordering and
# provenance) but are never part of the identity.

# Distinguishes "vote key absent" from "vote: null". None cannot do that job
# here, because null is a meaningful value in this contract (clear the vote).
_MISSING = object()


@app.post("/api/feedback")
async def record_message_feedback(request: Request):
    """Record, change, or CLEAR this user's vote on one message.

    `vote` of "up"/"down" records; `vote: null` clears (the user un-toggled).
    Switching up->down is one call with vote "down", not a clear plus a set.

    Voting again on the same message UPDATEs; it never appends. The identity is
    (session_id, user_id, message_id) and it is enforced by
    uq_message_feedback_vote in the database, not by a read-then-write here —
    two clicks racing would both find no row and both insert.
    """
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")

    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Body must be JSON")

    # ★ Required, though the client types it `string|null`. A vote with no
    # session can never be read back — GET is keyed on session_id — so storing
    # one would be writing a row nothing can ever render. Better a 400 the
    # client can gate on than a silent write-only row.
    session_id = str(body.get("session_id") or "").strip()
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")

    # Opaque on purpose: two client-side forms exist ("<role>-<run_id>", and
    # "<role>-<created_at>-<index>" before a run_id is assigned), so this is
    # checked for presence and length ONLY — never parsed, never assumed to be
    # a UUID. The length cap mirrors the CHECK constraint and exists because
    # message_id is indexed: an oversized value would otherwise surface as a
    # bare 500 from btree, and only for values that fail to compress.
    message_id = str(body.get("message_id") or "").strip()
    if not message_id:
        raise HTTPException(status_code=400, detail="message_id is required")
    if len(message_id) > 255:
        raise HTTPException(status_code=400, detail="message_id must be 255 characters or fewer")

    # ★ null means CLEAR, and it must be distinguished from absent-or-garbage.
    # `body.get("vote") or ""` would collapse null and "" and "meh" into one
    # branch, turning a typo into a silent delete of a real vote.
    raw_vote = body.get("vote", _MISSING)
    if raw_vote is _MISSING:
        raise HTTPException(status_code=400, detail="vote is required ('up', 'down', or null to clear)")
    vote = None if raw_vote is None else str(raw_vote).strip().lower()
    if vote is not None and vote not in ("up", "down"):
        raise HTTPException(status_code=400, detail="vote must be 'up', 'down', or null to clear")

    # Optional. bool is a subclass of int in Python, so `isinstance(True, int)`
    # is True and {"message_index": true} would otherwise be stored as index 1.
    raw_index = body.get("message_index")
    message_index = None
    if raw_index is not None and not isinstance(raw_index, bool):
        try:
            message_index = int(raw_index)
        except (TypeError, ValueError):
            raise HTTPException(status_code=400, detail="message_index must be an integer")
        if message_index < 0:
            raise HTTPException(status_code=400, detail="message_index must not be negative")

    comment = body.get("comment")
    comment = str(comment).strip()[:4000] if comment else None
    run_id = str(body.get("run_id")).strip() if body.get("run_id") else None

    # ★ The author of the vote comes from the token, never from the body. A
    # user_id the caller supplies is a user_id the caller chose, which makes the
    # whole table unattributable. Any user_id in the body is ignored.
    user_id = user.get("user_id")
    if user_id is None:
        raise HTTPException(status_code=401, detail="Token carries no user id")

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()

        if vote is None:
            # Un-toggle. DELETE rather than a 'none' vote value: the CHECK
            # constraint is what keeps every row countable, and a third state
            # would occupy the unique slot while counting for nothing.
            cur.execute(
                """
                DELETE FROM message_feedback
                WHERE session_id = %s AND user_id = %s AND message_id = %s
                """,
                (session_id, user_id, message_id),
            )
            cleared = cur.rowcount
            conn.commit()
            cur.close()
            # Clearing a vote that was never cast is the same end state, so it
            # is a success, not a 404 — the UI would have nothing to do with one.
            return {"success": True, "vote": None, "cleared": bool(cleared)}

        cur.execute(
            """
            INSERT INTO message_feedback
                (session_id, message_id, message_index, run_id, user_id, vote, comment)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT ON CONSTRAINT uq_message_feedback_vote DO UPDATE
                SET vote = EXCLUDED.vote,
                    comment = EXCLUDED.comment,
                    -- Only overwrite these when the caller actually has them; a
                    -- client that sends run_id must not have it wiped by an
                    -- older tab that does not.
                    message_index = COALESCE(EXCLUDED.message_index, message_feedback.message_index),
                    run_id = COALESCE(EXCLUDED.run_id, message_feedback.run_id),
                    updated_at = CURRENT_TIMESTAMP
            RETURNING id, vote, (created_at <> updated_at) AS changed
            """,
            (session_id, message_id, message_index, run_id, user_id, vote, comment),
        )
        row = cur.fetchone()
        conn.commit()
        cur.close()
        return {"success": True, "id": row[0], "vote": row[1], "updated": bool(row[2])}
    except HTTPException:
        raise
    except Exception as e:
        logging.getLogger("legalscout").error(f"Recording message feedback failed: {e}")
        raise HTTPException(status_code=500, detail="Could not record the feedback")
    finally:
        if conn:
            conn.close()


@app.get("/api/feedback")
async def list_message_feedback(request: Request):
    """This user's votes for one session, so the UI can restore them on reload.

    Scoped to the caller's own user_id — one reader's opinion is not another
    reader's to see, and the UI only ever renders its own thumbs.
    """
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")

    session_id = (request.query_params.get("session_id") or "").strip()
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")

    user_id = user.get("user_id")
    if user_id is None:
        raise HTTPException(status_code=401, detail="Token carries no user id")

    conn = None
    try:
        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(
            """
            SELECT message_id, vote, comment, message_index, run_id, updated_at
            FROM message_feedback
            WHERE session_id = %s AND user_id = %s
            ORDER BY message_index NULLS LAST, id
            """,
            (session_id, user_id),
        )
        rows = cur.fetchall()
        cur.close()
        return {
            "success": True,
            "session_id": session_id,
            # `votes` is the map the client renders from; `feedback` is the same
            # data with the comment and provenance attached. Both are sent
            # because the client normalises either, and the map alone would
            # drop the comment.
            "votes": {r[0]: r[1] for r in rows},
            "feedback": [
                {
                    "message_id": r[0],
                    "vote": r[1],
                    "comment": r[2],
                    "message_index": r[3],
                    "run_id": r[4],
                    "updated_at": r[5].isoformat() if r[5] else None,
                }
                for r in rows
            ],
        }
    except HTTPException:
        raise
    except Exception as e:
        logging.getLogger("legalscout").error(f"Listing message feedback failed: {e}")
        raise HTTPException(status_code=500, detail="Could not list the feedback")
    finally:
        if conn:
            conn.close()


# ---------------------------------------------------------------------------
# ★★★MOVED 2026-08-24 — this route was defined BELOW the frontend catch-all.
# `@app.get("/{full_path:path}")` inside `if _frontend_dir.exists():` claims
# every GET registered after it, so this endpoint answered with the frontend's
# HTML at HTTP 200 — byte-identical to a path that does not exist. Verified
# live before the move: GET /api/company/generate-extract/1 with a valid admin
# token returned text/html 200 with the same Next.js build id as
# /definitely-not-a-route, while /api/email/queued (registered above the
# catch-all) returned JSON. A 200 is not proof a route runs.
# ---------------------------------------------------------------------------
@app.get("/api/company/generate-extract/{company_id}")
async def generate_company_extract(company_id: int, request: Request):
    """Generate a DICA-style Company Extract PDF from company data."""
    try:
        import os, json
        from psycopg import connect
        from starlette.responses import Response

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT company_name_english, company_name_myanmar, company_registration_number,
                   registration_date, status, company_type, foreign_company, small_company,
                   principal_activity, date_of_last_annual_return, previous_registration_number,
                   registered_office_address, principal_place_of_business,
                   directors, ultimate_holding_company_name,
                   ultimate_holding_company_jurisdiction, ultimate_holding_company_registration_number,
                   total_shares_issued, currency_of_share_capital, members, filing_history
            FROM companies WHERE id = %s
        """, (company_id,))
        row = cur.fetchone()
        cur.close(); conn.close()

        if not row:
            return {"success": False, "error": "Company not found"}

        # Build HTML report
        name = row[0] or "Unknown"
        myanmar_name = row[1] or ""
        reg_no = row[2] or ""
        reg_date = row[3].strftime("%d %B, %Y") if row[3] else ""
        status = row[4] or ""
        company_type = row[5] or ""
        foreign = row[6] or "No"
        small = row[7] or "No"
        activity = row[8] or ""
        last_ar = row[9].strftime("%d %B, %Y") if row[9] else ""
        prev_reg = row[10] or ""
        office = row[11] or ""
        business = row[12] or ""
        dirs = row[13] if isinstance(row[13], list) else []
        uhc_name = row[14] or ""
        uhc_juris = row[15] or ""
        uhc_reg = row[16] or ""
        shares = row[17] or ""
        currency = row[18] or ""
        members = row[19] if isinstance(row[19], list) else []
        filings = row[20] if isinstance(row[20], list) else []

        # Build directors HTML
        directors_html = ""
        for d in dirs:
            directors_html += f"""
            <tr><td>Name:</td><td><b>{d.get('name','')}</b></td><td>Type:</td><td>{d.get('type','Director')}</td></tr>
            <tr><td>Date of Appointment:</td><td>{d.get('date_of_appointment','')}</td><td>Date of Birth:</td><td>{d.get('date_of_birth','')}</td></tr>
            <tr><td>Nationality:</td><td>{d.get('nationality','')}</td><td>N.R.C./Passport:</td><td>{d.get('nrc_passport','')}</td></tr>
            <tr><td>Gender:</td><td>{d.get('gender','')}</td><td>Business Occupation:</td><td>{d.get('business_occupation','-')}</td></tr>
            <tr><td colspan="4" style="border-bottom:1px solid #ddd;height:10px;"></td></tr>
            """

        # Build members HTML
        members_html = ""
        for m in members:
            members_html += f"<tr><td>{m.get('name','')}</td><td>{m.get('type','')}</td><td>{m.get('share_quantity','')}</td><td>{m.get('amount_paid','')}</td><td>{m.get('share_class','')}</td></tr>"

        # Build filing history HTML
        filings_html = ""
        for f in filings[:15]:
            filings_html += f"<tr><td>{f.get('form_type','')}</td><td>{f.get('effective_date','')}</td></tr>"

        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
    body {{ font-family: Arial, sans-serif; font-size: 11px; margin: 40px; color: #333; }}
    h1 {{ text-align: center; color: #2d5016; font-size: 16px; }}
    h2 {{ background: #f0f0f0; padding: 8px 12px; font-size: 12px; border-left: 4px solid #2d5016; margin-top: 20px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
    td {{ padding: 4px 8px; vertical-align: top; }}
    .header {{ text-align: center; margin-bottom: 30px; }}
    .header img {{ height: 50px; }}
    .label {{ color: #666; font-weight: bold; width: 200px; }}
    .value {{ color: #000; }}
    .footer {{ text-align: right; color: #999; font-size: 9px; margin-top: 30px; }}
    .confidential {{ text-align: right; color: red; font-weight: bold; font-size: 10px; }}
</style></head><body>
<div class="confidential">STRICTLY CONFIDENTIAL</div>
<div class="header">
    <h1>Myanmar Companies Online Registry - Company Extract</h1>
    <p><b>Company Name (English):</b> {name}</p>
    {f'<p><b>Company Name (Myanmar):</b> {myanmar_name}</p>' if myanmar_name else ''}
</div>

<h2>Company Information</h2>
<table>
    <tr><td class="label">Registration Number</td><td class="value">{reg_no}</td><td class="label">Registration Date</td><td class="value">{reg_date}</td></tr>
    <tr><td class="label">Status</td><td class="value">{status}</td><td class="label">Company Type</td><td class="value">{company_type}</td></tr>
    <tr><td class="label">Foreign Company</td><td class="value">{foreign}</td><td class="label">Small Company</td><td class="value">{small}</td></tr>
    <tr><td class="label">Principal Activity</td><td class="value" colspan="3">{activity}</td></tr>
    <tr><td class="label">Date of Last Annual Return</td><td class="value">{last_ar}</td><td class="label">Previous Registration Number</td><td class="value">{prev_reg}</td></tr>
</table>

<h2>Addresses</h2>
<table>
    <tr><td class="label">Registered Office</td><td class="value">{office}</td></tr>
    <tr><td class="label">Principal Place of Business</td><td class="value">{business}</td></tr>
</table>

<h2>Officers</h2>
<table>{directors_html}</table>

{'<h2>Ultimate Holding Company</h2><table><tr><td class="label">Name</td><td class="value">' + uhc_name + '</td><td class="label">Jurisdiction</td><td class="value">' + uhc_juris + '</td><td class="label">Reg No</td><td class="value">' + uhc_reg + '</td></tr></table>' if uhc_name else ''}

<h2>Share Capital Structure</h2>
<table>
    <tr><td class="label">Total Shares Issued</td><td class="value">{shares}</td><td class="label">Currency</td><td class="value">{currency}</td></tr>
</table>

{f'<h2>Members</h2><table><tr><th>Name</th><th>Type</th><th>Shares</th><th>Amount Paid</th><th>Class</th></tr>{members_html}</table>' if members_html else ''}

{f'<h2>Filing History</h2><table><tr><th>Form / Filing Type</th><th>Effective Date</th></tr>{filings_html}</table>' if filings_html else ''}

<div class="footer">
    EXTRACT GENERATED ON {datetime.now().strftime("%d/%m/%Y")} AT {datetime.now().strftime("%H:%M")}<br>
    Generated by Legal Scout
</div>
</body></html>"""

        # Convert HTML to PDF using LibreOffice
        import subprocess, tempfile
        with tempfile.NamedTemporaryFile(suffix=".html", delete=False, mode='w') as tmp:
            tmp.write(html)
            tmp_path = tmp.name

        pdf_dir = Path("/documents/legal/extracts")
        pdf_dir.mkdir(parents=True, exist_ok=True)

        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_dir), tmp_path],
            capture_output=True, text=True, timeout=30
        )

        pdf_name = Path(tmp_path).stem + ".pdf"
        pdf_path = pdf_dir / pdf_name

        if pdf_path.exists():
            pdf_bytes = pdf_path.read_bytes()
            # Also save with company name
            final_name = f"DICA_Extract_{name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
            final_path = pdf_dir / final_name
            final_path.write_bytes(pdf_bytes)

            Path(tmp_path).unlink(missing_ok=True)
            pdf_path.unlink(missing_ok=True)

            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={"Content-Disposition": f"inline; filename={final_name}"}
            )

        Path(tmp_path).unlink(missing_ok=True)
        return {"success": False, "error": "PDF generation failed"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# The frontend is built as a static export (HTML/JS/CSS) and served directly
# by FastAPI. This eliminates the need for a separate Node.js container.
#
# Route priority: /api/* and /agents/* are handled by FastAPI routes above.
# Everything else falls through to the static frontend.
# ---------------------------------------------------------------------------
_frontend_dir = Path("/app/static-frontend")

if _frontend_dir.exists():
    @app.get("/dashboard")
    async def dashboard_redirect():
        """Redirect /dashboard to /admin/dashboard/."""
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/admin/dashboard/")

    @app.get("/{full_path:path}")
    async def serve_frontend(full_path: str):
        """Catch-all: serve static frontend files for client-side routing."""
        from fastapi.responses import FileResponse, HTMLResponse

        # ★ Nothing here sent Cache-Control at all, so every asset was left to
        # the browser's heuristic caching. That is fine for content-hashed files
        # and actively harmful for stable paths: `/pdf.worker.min.mjs` never
        # changes name, and this port hosts a different app on this machine every
        # few weeks, so a browser could reuse a worker cached from an unrelated
        # app — producing
        #   The API version "4.0.379" does not match the Worker version "5.2.133"
        # while the server was serving 4.0.379 throughout.
        #
        # `_next/static/*` filenames contain a content hash, so a changed file is
        # a changed URL and it is safe to cache forever. Everything else must
        # revalidate; the ETag above makes that a cheap 304, not a re-download.
        def _cached(path):
            immutable = "/_next/static/" in f"/{full_path}"
            return FileResponse(
                path,
                headers={
                    "Cache-Control": "public, max-age=31536000, immutable"
                    if immutable
                    else "no-cache"
                },
            )

        # Try exact file match (e.g., /favicon.ico, /robots.txt)
        file_path = _frontend_dir / full_path
        if file_path.is_file():
            return _cached(file_path)

        # Try directory with index.html (e.g., /admin/dashboard/ → admin/dashboard/index.html)
        index_path = _frontend_dir / full_path / "index.html"
        if index_path.is_file():
            return FileResponse(index_path, headers={"Cache-Control": "no-cache"})

        # Try with trailing slash stripped
        stripped = full_path.rstrip("/")
        if stripped:
            index_path = _frontend_dir / stripped / "index.html"
            if index_path.is_file():
                return FileResponse(index_path, headers={"Cache-Control": "no-cache"})

        # Fallback to root index.html (SPA client-side routing)
        root_index = _frontend_dir / "index.html"
        if root_index.is_file():
            return FileResponse(root_index, headers={"Cache-Control": "no-cache"})

        return HTMLResponse("<h1>Not Found</h1>", status_code=404)
else:
    @app.get("/dashboard")
    async def dashboard():
        """Redirect to frontend (when running separately in dev mode)."""
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url=f"{FRONTEND_HOST}/dashboard")

    @app.get("/")
    async def root_redirect():
        """Redirect to frontend in dev mode."""
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url=FRONTEND_HOST)


if __name__ == "__main__":
    agent_os.serve(
        app="main:app",
        reload=False,
    )


# ========== NEW FALLBACK FUNCTIONS FOR ENHANCED FIELDS ==========


def _get_when_to_use(template_name: str) -> str:
    """Describe when a template applies, from its name.

    Order matters and used to be wrong: a broad "shareholder" test sat above the
    specific ones, so nine differently-purposed templates all collapsed onto
    "When shareholders need to pass a resolution on specific matters" — useless
    for choosing between them. The distinguishing words (resignation,
    appointment, consent, notice) are now tested first, broad words last.
    """
    n = template_name.lower()

    appointing = "appointment" in n or "appoint" in n
    resigning = "resignation" in n or "resign" in n

    if "corporate shareholder consent" in n:
        return ("When a corporate shareholder must consent — signed by that shareholder "
                "company's OWN directors, never the new company's board")
    if "individual shareholder consent" in n:
        return "When an individual shareholder must give written consent in their own name"
    if "director consent" in n:
        if "non-group" in n:
            return "When appointing a director who is NOT from a group company"
        if "group" in n:
            return "When appointing a director from a group company to the board"
        return "When a proposed director consents to act"
    if "resignation letter" in n:
        return "When a director resigns — their own letter of resignation to the company"
    if "notice of calling" in n:
        return "To call an annual general meeting — the internal notice convening it"
    if "notice of annual general meeting" in n:
        return "To give shareholders formal notice that an AGM will be held"
    if resigning and appointing:
        return "When a director resigns and a replacement is appointed in the same decision"
    if resigning:
        return "When a director resigns and the company records the departure"
    if appointing:
        return "When a new director is appointed to the board"
    if "annual general meeting" in n or "agm" in n:
        return "For the annual general meeting — recording its proceedings and resolutions"
    if "egm" in n:
        return "For an extraordinary general meeting on urgent matters needing shareholder approval"
    if "minutes" in n:
        return "To record the proceedings and decisions of a meeting"
    if "resolution" in n:
        return "When a formal decision must be documented and approved"
    return ""


def _get_how_to_use(template_name: str) -> list:
    """Generate step-by-step instructions based on template name."""
    name_lower = template_name.lower()

    if "director consent" in name_lower:
        return [
            "Fill in director's personal details (name, NRC / passport, address)",
            "Enter appointment terms and position",
            "Obtain director's signature",
            "Get witness signature (Company Secretary)",
            "File the change of director with DICA via the MyCO online registry within 21 days",
        ]
    elif "agm" in name_lower or "egm" in name_lower:
        return [
            "Prepare notice of meeting with agenda",
            "Send notice to all shareholders (14-21 days before)",
            "Conduct meeting with quorum present",
            "Record minutes of proceedings",
            "Document resolutions passed",
            "File any required forms with DICA via MyCO",
        ]
    elif "shareholder" in name_lower:
        return [
            "Identify the resolution topic",
            "Prepare resolution draft",
            "Circulate to shareholders for approval",
            "Collect votes/signatures",
            "File with company records",
        ]
    elif "minutes" in name_lower:
        return [
            "Record meeting date, time, and venue",
            "List attendees and quorum",
            "Document each agenda item discussion",
            "Record decisions and resolutions",
            "Note voting results if applicable",
            "Obtain chairman and secretary signatures",
        ]
    else:
        return [
            "Review template requirements",
            "Fill in all required fields",
            "Obtain necessary signatures",
            "Review for accuracy",
            "File or submit as required",
        ]


def _get_prerequisites(template_name: str) -> list:
    """Generate prerequisites based on template name."""
    name_lower = template_name.lower()

    if "director consent" in name_lower:
        return [
            "Director's NRC / passport (certified true copy)",
            "Address proof (utility bill < 3 months)",
            "Passport (for foreign director)",
            "Board resolution appointing director",
        ]
    elif "agm" in name_lower:
        return [
            "Financial statements",
            "Auditor's report",
            "Notice of meeting",
            "Proxy forms (if applicable)",
            "Attendance register",
        ]
    elif "shareholder" in name_lower:
        return ["Shareholder details", "Share certificate copies", "Board resolution (if required)"]
    else:
        return ["Company letterhead", "Previous documents (if any)", "Authorized signatures"]


def _get_filing_deadline(template_name: str) -> str:
    """Generate filing deadline based on template name."""
    name_lower = template_name.lower()

    if "director consent" in name_lower:
        return "Within 21 days of appointment (Myanmar Companies Law 2017, filed with DICA)"
    elif "agm" in name_lower:
        return "Within 30 days of meeting date"
    elif "form 24" in name_lower or "notice of appointment" in name_lower:
        return "Within 30 days of appointment"
    elif "form 8" in name_lower or "annual return" in name_lower:
        return "Annual return filed with DICA within 2 months of the AGM"
    else:
        return "As required by law or company policy"


def _get_fees(template_name: str) -> str:
    """Generate fees information based on template name."""
    name_lower = template_name.lower()

    if "director consent" in name_lower:
        return "DICA online filing fee; late filing attracts a penalty"
    elif "agm" in name_lower:
        return "No fee for meeting, may have costs for venue/catering"
    elif "form" in name_lower:
        return "DICA online filing fee per form"
    else:
        return "Varies by document type and filing authority"


def _get_validity_period(template_name: str) -> str:
    """Generate validity period based on template name."""
    name_lower = template_name.lower()

    if "director consent" in name_lower:
        return "Valid for 6 months from signing date"
    elif "form" in name_lower:
        return "Valid until next update required"
    elif "resolution" in name_lower:
        return "Permanent record, effective until revoked"
    elif "certificate" in name_lower:
        return "Valid for 12 months, renewable"
    else:
        return "As specified in document terms"


def _get_approval_chain(template_name: str) -> list:
    """Generate approval chain based on template name."""
    name_lower = template_name.lower()

    if "director consent" in name_lower:
        return ["Board of Directors (approval)", "Company Secretary (witness)", "DICA (filing)"]
    elif "agm" in name_lower or "shareholder" in name_lower:
        return [
            "Board of Directors (convene meeting)",
            "Shareholders (vote/approve)",
            "Chairman (preside meeting)",
            "Company Secretary (certify)",
        ]
    elif "minutes" in name_lower:
        return ["Chairman (review and sign)", "Company Secretary (certify)"]
    else:
        return ["Authorized Director", "Company Secretary"]


def _get_required_attachments(template_name: str) -> list:
    """Generate required attachments based on template name."""
    name_lower = template_name.lower()

    if "director consent" in name_lower:
        return ["NRC copy (certified true copy)", "Address proof", "Passport (foreign directors)", "Photo"]
    elif "agm" in name_lower:
        return ["Financial Statements", "Auditor's Report", "Notice of AGM", "Proxy Forms"]
    elif "shareholder" in name_lower:
        return ["Share Certificate", "Board Resolution"]
    else:
        return ["Supporting documents as applicable"]


def _get_common_mistakes(template_name: str) -> list:
    """Generate common mistakes based on template name."""
    name_lower = template_name.lower()

    if "director consent" in name_lower:
        return [
            "Missing certified NRC / passport copy",
            "Address doesn't match the NRC",
            "Signature not witnessed by valid witness",
            "Filed after 30-day deadline",
            "Missing board resolution",
        ]
    elif "agm" in name_lower:
        return [
            "Insufficient notice period",
            "No quorum present",
            "Voting procedures not followed",
            "Minutes not signed by chairman",
            "Missing required agenda items",
        ]
    elif "minutes" in name_lower:
        return [
            "Incomplete attendance record",
            "Decisions not clearly documented",
            "Missing vote counts",
            "Not signed by required parties",
            "No record of disclosures of interest",
        ]
    else:
        return ["Missing required fields", "Incorrect data format", "Missing signatures", "Not reviewed for accuracy"]


def _get_industry_tags(template_name: str) -> list:
    """Generate industry tags based on template name."""
    name_lower = template_name.lower()

    tags = ["General Business"]

    if "bank" in name_lower or "financial" in name_lower:
        tags = ["Banking", "Finance", "Financial Services"]
    elif "manufacturing" in name_lower or "factory" in name_lower:
        tags = ["Manufacturing", "Industrial"]
    elif "technology" in name_lower or "tech" in name_lower:
        tags = ["Technology", "IT", "Software"]
    elif "construction" in name_lower or "building" in name_lower:
        tags = ["Construction", "Real Estate"]
    elif "agm" in name_lower or "director" in name_lower or "shareholder" in name_lower:
        tags = ["Corporate", "Compliance", "All Industries"]

    return tags


def _get_complexity(field_count: int) -> str:
    """Determine complexity based on number of fields."""
    if field_count <= 5:
        return "Easy"
    elif field_count <= 15:
        return "Medium"
    else:
        return "Complex"


def _get_estimated_time(field_count: int) -> str:
    """Estimate time to complete based on field count."""
    if field_count <= 5:
        return "10 minutes"
    elif field_count <= 10:
        return "20 minutes"
    elif field_count <= 20:
        return "30 minutes"
    elif field_count <= 30:
        return "45 minutes"
    else:
        return "1 hour or more"


# ---------------------------------------------------------------------------
# DICA Company Extract Report Generation
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Send Document via Email
# ---------------------------------------------------------------------------
@app.post("/api/documents/send-email")
async def send_document_email(request: Request):
    """Send a document via email with optional message."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        body = await request.json()
        to_email = body.get("to_email", "")
        subject = body.get("subject", "Document from Legal Scout")
        message = body.get("message", "")
        file_path = body.get("file_path", "")  # e.g. /documents/legal/output/file.docx

        if not to_email or not file_path:
            return {"success": False, "error": "to_email and file_path required"}

        # Resolve file path with traversal protection
        output_base = Path("/documents/legal/output").resolve()
        templates_base = Path("/documents/legal/templates").resolve()

        full_path = Path(file_path).resolve()
        # Ensure file is within allowed directories
        if not (str(full_path).startswith(str(output_base)) or str(full_path).startswith(str(templates_base))):
            # Try with /documents prefix using just the filename
            fname = Path(file_path).name
            full_path = (output_base / fname).resolve()
            if not str(full_path).startswith(str(output_base)):
                return {"success": False, "error": "Invalid filename"}
            if not full_path.exists():
                full_path = (templates_base / fname).resolve()
                if not str(full_path).startswith(str(templates_base)):
                    return {"success": False, "error": "Invalid filename"}
                if not full_path.exists():
                    return {"success": False, "error": "File not found"}
        elif not full_path.exists():
            return {"success": False, "error": "File not found"}

        # Get SMTP settings
        import os, smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        from email.mime.base import MIMEBase
        from email import encoders
        from psycopg import connect

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT key, value FROM app_settings WHERE key LIKE 'smtp_%'")
        smtp = {r[0]: r[1] for r in cur.fetchall()}
        cur.close(); conn.close()

        smtp_host = smtp.get("smtp_host") or os.getenv("SMTP_HOST", "")
        smtp_port = int(smtp.get("smtp_port") or os.getenv("SMTP_PORT", "587"))
        smtp_user = smtp.get("smtp_user") or os.getenv("SMTP_USER", "")
        smtp_pass = smtp.get("smtp_pass") or os.getenv("SMTP_PASS", "")
        smtp_from = smtp.get("smtp_from") or os.getenv("SMTP_FROM", "noreply@legalscout.com")

        if not smtp_host or not smtp_user:
            return {"success": False, "error": "Email not configured. Go to Settings to configure SMTP."}

        # Build email with attachment
        msg = MIMEMultipart()
        msg["Subject"] = subject
        msg["From"] = smtp_from
        msg["To"] = to_email

        html_body = f"""<div style="font-family:Arial,sans-serif;max-width:600px;">
            <div style="background:linear-gradient(135deg,#FF6B35,#FF4017);padding:20px;border-radius:8px 8px 0 0;">
                <h2 style="color:white;margin:0;">Legal Scout</h2>
            </div>
            <div style="padding:20px;border:1px solid #eee;border-top:0;">
                <p>{message or "Please find the attached document."}</p>
                <p style="color:#666;font-size:12px;">Sent by {user.get('email', 'Legal Scout')}</p>
            </div>
        </div>"""
        msg.attach(MIMEText(html_body, "html"))

        # Attach file
        with open(full_path, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={full_path.name}")
            msg.attach(part)

        # Send
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_pass)
            server.send_message(msg)

        sender_email = user.get("email", "unknown") if user else "unknown"
        log_activity(user.get("user_id") if user else None, sender_email, "send_email",
            f"Sent {full_path.name} to {to_email}", "")

        # Log to email_logs table
        try:
            _econn = get_db_conn()
            _ecur = _econn.cursor()
            _ecur.execute(
                "INSERT INTO email_logs (to_email, subject, body, attachment_name, attachment_path, sent_by_email, status) VALUES (%s,%s,%s,%s,%s,%s,%s)",
                (to_email, subject, message, full_path.name, str(full_path), sender_email, "sent"))
            _econn.commit(); _ecur.close(); _econn.close()
        except Exception as e:
            logging.getLogger("legalscout").warning(f"Email log failed: {e}")

        return {"success": True, "message": f"Email sent to {to_email} with {full_path.name}"}
    except Exception as e:
        # Log failed email
        try:
            _econn2 = get_db_conn()
            _ecur2 = _econn2.cursor()
            _ecur2.execute(
                "INSERT INTO email_logs (to_email, subject, body, sent_by_email, status, error_message) VALUES (%s,%s,%s,%s,%s,%s)",
                (to_email, subject, message, "unknown", "failed", str(e)))
            _econn2.commit(); _ecur2.close(); _econn2.close()
        except Exception as e2:
            logging.getLogger("legalscout").warning(f"Email error log failed: {e2}")
        return {"success": False, "error": str(e)}


# ---------------------------------------------------------------------------
# Import-time startup
# ---------------------------------------------------------------------------
# Runs here rather than via @app.on_event("startup"), which Starlette ignores
# because AgentOS installs a lifespan. See startup_sync's docstring.
try:
    startup_sync()
except Exception as _startup_err:  # noqa: BLE001 — never take the app down on boot
    logger.error(f"[STARTUP] startup_sync failed: {_startup_err}")
