"""
Fill-path suite — slot_resolver / repeat_regions / smart_doc, offline.

WHY THIS EXISTS
---------------
The three tracker layers all stop short of the fill path. Layer 1 reads
`/api/documents/fill-view`; Layers 2 and 3 (now scripted) gate the agent
harness. None of them executes `collect_slot_requests`, `expand_repeat_regions`
or `fill_template_with_validation`, so the code that decides WHICH PARTIES GET
ASKED ABOUT and HOW MANY LINES A DOCUMENT GROWS had no offline regression gate
at all.

This drives that code directly, in-process, against the 15 real `.docx`
templates on disk. No DB, no HTTP, no model, no container.

THE CASE THIS WAS BUILT FOR
---------------------------
CLAUDE.md, "Known open":

    ⚠️ Still unreconciled: `slot_resolver.collect_slot_requests` asks about the
    NUMBERED slots the template declares (1, 2, 3). A 7-party document renders
    correctly and may still only be asked about 3.

A template's numbered placeholders are a LAYOUT, not a party count.
`repeat_regions` grows or shrinks that block to the real party list at fill
time, so a seven-director appointment renders seven lines out of a three-slot
template — and the ask step, walking `mapping` alone, could never see past the
third position. The document came out right and the user was asked about three
parties.

The fix has to make the ask count come from the same source the expander uses.
This suite asserts that in BOTH directions, because there are two ways to get it
wrong and only one of them is the original bug:

    too few   — 7 real parties, still only 3 questions (the original bug)
    too many  — 2 real parties but 3 questions, or a count that is genuinely
                UNKNOWN and the ask step invents one question per numbered slot
                instead of one per role (the over-correction)

Every case asserts A NUMBER, never that a call "worked".

WHAT IS AND IS NOT COVERED
--------------------------
Covered: the ask-step party arithmetic (`collect_slot_requests`), the expander's
growth and shrink (`expand_repeat_regions`), and that `smart_doc` actually
reaches the expander (proved by effect — row counts move with the party count —
because it imports it lazily inside the function and there is no module
attribute to check).

Not covered: `slot_resolver._member_position_covered`. It has no reachable
path — `_correct_member_slot` rewrites every member-family placeholder to
`shareholder_list`/`attendee`, and the kind-level rule returns before the
position check; with an empty member list the function returns False on its
first line. Measured across all eight slot kinds with a 7-member list: zero
calls. A case for it was written, could not be made to fail, and was removed
rather than left in as a test that asserts nothing. Reported, not patched.

Not covered: anything needing the `templates` table. `smart_doc` loads
`field_mapping` from the DB, so with no database its slot pass runs on an empty
mapping. The mapping fixtures here are therefore constructed, not loaded — see
`TEMPLATE_SLOTS` for how they are derived from the real files.

Run:  python3.12 tests/tracker_fill.py            (needs Python >= 3.10)
"""

import ast
import hashlib
import importlib
import importlib.util
import sys
import types
from pathlib import Path

# NO VERSION FLOOR — deliberately.
#
# This file used to hard-exit below python 3.10: `scout/tools/slot_resolver.py`
# annotates with PEP 604 unions (`str | None`) and had no
# `from __future__ import annotations`, so the annotation was evaluated at def
# time and the module raised TypeError on IMPORT under the 3.9.6 system python.
# That import is now at `slot_resolver.py:28` and the floor is gone — verified
# by importing the module and calling `collect_slot_requests` on 3.9.6.
#
# The gate was replaced rather than re-pointed at a lower number: a hard-coded
# version is a guess about WHY an import might fail, and it goes stale silently
# the moment somebody fixes the real cause — as happened here, where the number
# would have kept refusing to run on an interpreter that works. The import
# itself is the only honest check, so failures surface with the real reason
# below.

REPO = Path(__file__).resolve().parent.parent
TEMPLATE_DIR = REPO / "documents" / "legal" / "templates"


# ── offline harness ─────────────────────────────────────────────────


class DBTripwire:
    """Stands in for `db.connection.get_db_conn`.

    Never returns a connection. Counting the attempts is the point: a case that
    claims to be offline can be checked rather than trusted, and a code path
    that starts reaching for the database shows up as a number instead of a
    hang. `armed` is False only for the smart_doc case, which legitimately tries
    to load `field_mapping` and logs the failure.
    """

    def __init__(self):
        self.attempts = 0

    def __call__(self, *args, **kwargs):
        self.attempts += 1
        raise RuntimeError("DB access attempted in an offline test")


TRIPWIRE = DBTripwire()


def _install_stubs():
    """Make the product modules importable without importing the product.

    Three separate walls, each for a different reason:

    `scout` / `scout.tools` — `scout/__init__.py` does `from scout.agent import
    scout`, and `scout/agent.py` imports `agno.tools.mcp.MCPTools`, which raises
    unless `mcp` is installed. `scout/tools/__init__.py` is worse: it imports
    eight tool factories, most of the product. Registering both names as bare
    packages with a `__path__` lets `scout.tools.slot_resolver` resolve while
    neither `__init__` ever runs.

    `db.connection` — the tripwire above.

    `agno.run` — `smart_doc` imports `RunContext` for a type annotation and
    nothing else. Stubbing an import-only dependency does not change the code
    under test; installing agno to satisfy an annotation would.
    """
    sys.path.insert(0, str(REPO))

    for name, rel in (("scout", "scout"), ("scout.tools", "scout/tools")):
        mod = types.ModuleType(name)
        mod.__path__ = [str(REPO / rel)]
        sys.modules[name] = mod

    db = types.ModuleType("db")
    db.__path__ = [str(REPO / "db")]
    sys.modules["db"] = db
    conn = types.ModuleType("db.connection")
    conn.get_db_conn = TRIPWIRE
    sys.modules["db.connection"] = conn

    agno = types.ModuleType("agno")
    agno.__path__ = []
    sys.modules["agno"] = agno
    run = types.ModuleType("agno.run")
    run.RunContext = type("RunContext", (), {})
    sys.modules["agno.run"] = run


def _install_slot_contract(enabled):
    """Wire (or unwire) `app.slot_contract` before `slot_resolver` is imported.

    ★ This is a real divergence, not a test convenience. `slot_resolver` does:

        try:
            from app import slot_contract as _slot_contract
        except Exception:
            _slot_contract = None

    and `app/__init__.py` does `from app.main import agent_os, app` — the whole
    FastAPI application. In the container that import succeeds and the CONTRACT
    normaliser is used. On a dev laptop it dies on a missing `jwt` and the
    except swallows it, so the module silently falls back to its own
    `_fallback_normalise`. Two different normalisers, chosen by which machine
    you are on, with nothing in the logs to say which one ran.

    So the suite runs every case under BOTH and asserts they agree. Loading the
    file by path is what makes the contract reachable here at all: it bypasses
    `app/__init__.py` and therefore the entire app import.
    """
    for name in ("app", "app.slot_contract"):
        sys.modules.pop(name, None)
    if not enabled:
        return False

    app_pkg = types.ModuleType("app")
    app_pkg.__path__ = [str(REPO / "app")]
    sys.modules["app"] = app_pkg

    spec = importlib.util.spec_from_file_location("app.slot_contract", str(REPO / "app" / "slot_contract.py"))
    module = importlib.util.module_from_spec(spec)
    sys.modules["app.slot_contract"] = module
    spec.loader.exec_module(module)
    app_pkg.slot_contract = module
    return True


def load_product(contract):
    """Fresh copies of the modules under test, for one normalisation mode.

    The contract is read into module-level globals at import time
    (`_contract_normalise`, `SLOT_KINDS`, ...), so switching modes means
    re-importing rather than re-assigning.
    """
    for name in list(sys.modules):
        if name.startswith("scout.tools."):
            del sys.modules[name]
    wired = _install_slot_contract(contract)
    slot_resolver = importlib.import_module("scout.tools.slot_resolver")
    repeat_regions = importlib.import_module("scout.tools.repeat_regions")
    try:
        smart_doc = importlib.import_module("scout.tools.smart_doc")
    except TypeError as exc:
        # The PEP 604 floor again, one module along. `slot_resolver.py` was
        # fixed on 2026-08-24 by adding `from __future__ import annotations`;
        # `smart_doc.py` has the same annotations and NOT that import, so on
        # 3.9 it still dies at def time — `smart_doc.py:105`,
        # `def _lookup_recent_generation(key: str) -> dict[str, Any] | None`.
        #
        # Fixing slot_resolver did not remove the floor, it MOVED it. Say which
        # module and which line, or the next person re-diagnoses it from a
        # traceback that points into `importlib`.
        raise SystemExit(
            f"scout/tools/smart_doc.py cannot be imported on Python "
            f"{sys.version.split()[0]}: {exc}\n"
            "  Cause: PEP 604 unions (`dict[str, Any] | None`) evaluated at def "
            "time — see smart_doc.py:105.\n"
            "  Fix:   add `from __future__ import annotations` to smart_doc.py, "
            "as slot_resolver.py:28 and repeat_regions.py already have.\n"
            "  Until then run this suite on 3.10+ (the container is 3.12.8)."
        ) from exc
    return slot_resolver, repeat_regions, smart_doc, wired


# ── fixtures — FICTIONAL ONLY ───────────────────────────────────────
# No client name appears in this file. The people register and company tables
# are empty anyway, but the rule stands regardless of what is loaded: these
# names are invented, and they follow the project's own convention for fixture
# companies (see LoginShowcase.tsx, "FICTIONAL data only").

COMPANY = "GOLDEN LOTUS HOLDINGS LIMITED"
CORPORATE_MEMBER = "EMERALD HOLDINGS LIMITED"


def parties(n, prefix="FICTIONAL PARTY", party_type="individual"):
    """`n` distinct fictional parties, in the shape a picker hands back."""
    return [{"name": f"{prefix} {i}", "party_type": party_type, "identifier": f"NRC-{i:03d}"} for i in range(1, n + 1)]


def slot_entry(kind, of="document_company", multi=False):
    """One `field_mapping` entry whose source is a slot.

    Written out longhand rather than fetched from the `templates` table: the
    table is empty on this install, and a fixture that silently degrades to
    "no slots" would make every count below zero and every case pass.
    """
    return {
        "source": "slot",
        "db_column": None,
        "slot": {"kind": kind, "of": of, "multi": multi},
        "default": None,
        "description": "",
    }


# ── what the real templates actually declare ────────────────────────


def _docx_texts(doc):
    for paragraph in doc.paragraphs:
        yield paragraph.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text


def declared_slots(placeholders_mod, family_of, position_of):
    """Per template, the numbered repeat-family placeholders it declares.

    Derived from the real `.docx` files rather than hard-coded, so a template
    that is re-uploaded with a different number of slots re-derives its own
    expectations and the invariant still holds. The FAMILY and POSITION rules
    come from the product (`repeat_regions._family`, `_slot_position`) because
    those are the definitions the fill path itself uses — re-deriving them here
    is how the ask step and the fill step drifted apart in the first place.
    """
    from docx import Document

    out = {}
    for path in sorted(TEMPLATE_DIR.glob("*.docx")):
        counter = placeholders_mod.new_empty_counter()
        families = {}
        for text in _docx_texts(Document(str(path))):
            for match in placeholders_mod.PLACEHOLDER_PATTERN.finditer(text or ""):
                name = placeholders_mod.placeholder_name(match.groups(), counter)
                if not name or placeholders_mod.is_empty_placeholder(name):
                    continue
                family, position = family_of(name), position_of(name)
                if family and position:
                    families.setdefault(family, {}).setdefault(position, set()).add(name)
        if families:
            out[path.name] = families
    return out


# ── cases ───────────────────────────────────────────────────────────
#
# Each returns (label, got, want, note). `got` and `want` are NUMBERS.

APPOINTMENT_TEMPLATE = (
    "Corporate Shareholder Consent - Directors Resolution for New Company Setup and Director Appointment.docx"
)
MEMBER_TEMPLATE = "Annual General Meeting Minutes.docx"


def appointed_mapping(declared, template=APPOINTMENT_TEMPLATE, family="appointed_director"):
    """The template's own appointed-director placeholders, as slot entries.

    `new_director` is asserted rather than assumed: `_KIND_PATTERNS` maps
    `appointed[ _-]?director` to `new_director`, and that kind matters — it is
    in `_NEVER_SUPPRESS_KINDS`, so if the classification drifted, the
    suppression rules under test would silently change meaning.
    """
    spellings = declared[template][family]
    return {name: slot_entry("new_director") for names in spellings.values() for name in names}


def positions_asked(requests, position_of):
    """The distinct numbered positions a request list covers."""
    return {p for p in (position_of(r["placeholder"]) for r in requests) if p}


def case_grow(sr, declared):
    """★ THE CASE. 7 real parties against a template that declares 3.

    The template's numbered slots are the layout; the party list is the count.
    Seven appointees must produce seven questions, and the four the template
    never spelled out must be synthesised with the template's own spelling.
    """
    mapping = appointed_mapping(declared)
    declared_max = max(declared[APPOINTMENT_TEMPLATE]["appointed_director"])
    data = {"members": parties(1), "appointed_directors": parties(7)}
    requests = sr.collect_slot_requests(mapping, list(mapping), data, company_name=COMPANY)
    got = len(positions_asked(requests, sr._slot_position))
    return (
        f"grow: 7 parties, template declares {declared_max}",
        got,
        7,
        f"asked positions {sorted(positions_asked(requests, sr._slot_position))}",
    )


def case_grow_spelling(sr, declared):
    """The synthesised positions must keep the template's spelling.

    `appointed_director_4_nrc`, not `appointed_director_4_name` and not a
    re-spelling. These placeholders carry U+00A0 and mixed separators in five
    templates, so anything that rebuilds the name instead of substituting the
    digits produces a placeholder that matches nothing at fill time.
    """
    mapping = appointed_mapping(declared)
    data = {"members": parties(1), "appointed_directors": parties(7)}
    requests = sr.collect_slot_requests(mapping, list(mapping), data, company_name=COMPANY)
    spelt = {r["placeholder"] for r in requests}
    pattern = sr._position_pattern(sorted(mapping)[0])
    expected = {sr._placeholder_at(pattern, i) for i in range(1, 8)}
    return (
        "grow: synthesised spelling matches the template's",
        len(spelt & expected),
        7,
        f"unexpected: {sorted(spelt - expected)}",
    )


def case_shrink(sr, declared):
    """The other direction: 2 real parties, template declares 3.

    The expander DELETES the third unit, so the finished document has no third
    line and a question about it is a question about nothing. An over-correction
    that only ever grows would answer 3 here.
    """
    mapping = appointed_mapping(declared)
    data = {"members": parties(1), "appointed_directors": parties(2)}
    requests = sr.collect_slot_requests(mapping, list(mapping), data, company_name=COMPANY)
    got = len(positions_asked(requests, sr._slot_position))
    return ("shrink: 2 parties, template declares 3", got, 2, "")


def case_exact(sr, declared):
    """3 real parties against 3 declared — neither grow nor shrink."""
    mapping = appointed_mapping(declared)
    data = {"members": parties(1), "appointed_directors": parties(3)}
    requests = sr.collect_slot_requests(mapping, list(mapping), data, company_name=COMPANY)
    got = len(positions_asked(requests, sr._slot_position))
    return ("exact: 3 parties, template declares 3", got, 3, "")


def case_unknown_count_collapses(sr, declared):
    """★ The over-correction guard. No party data at all.

    When the size of a repeating region cannot be established, the ask step must
    fall back to ONE question per ROLE — not one per numbered slot. A fix that
    started asking three times whenever it could not count would be a
    regression in the opposite direction from the original bug, and it would
    look like "more thorough" behaviour rather than a defect.

    Pinned on `signing_director`, NOT `appointed_director`. This case has to use
    a family whose size is establishable IN PRINCIPLE and merely unknown here —
    signing directors come off the register, so an unregistered company leaves a
    genuine "could not count". `appointed_director` can never be counted from any
    store (the appointees are the answer being collected), so for that family
    "unknown" is the normal state rather than a failure to look, and collapsing
    there is the original under-asking bug — see
    `case_appointed_asks_per_declared_position`.
    """
    mapping = appointed_mapping(declared, family="signing_director")
    data = {"members": parties(1)}  # nothing that names the signatories
    requests = sr.collect_slot_requests(mapping, list(mapping), data, company_name=COMPANY)
    return ("unknown count: collapses to one ask per role", len(requests), 1, "")


def case_appointed_asks_per_declared_position(sr, declared):
    """★ The under-asking guard for appointees. No party data at all.

    An APPOINTED director is a person being placed on a NEW company's board, so
    no register can enumerate them and the party count is permanently 0. Under a
    rule that asks per position only when a count is known, the three numbered
    slots this template declares collapsed to a SINGLE question and positions 2
    and 3 were never put to anyone — measured live on THAZIN VALLEY HOLDINGS,
    where 6 numbered placeholders produced 2 requests.

    For this family the template's declared positions are the only evidence of
    how many distinct people are wanted, so each declared position is asked
    about. The cost of being wrong in this direction is a question the user
    skips; the cost in the other direction is a legal document that silently
    omits two directors.
    """
    mapping = appointed_mapping(declared)
    data = {"members": parties(1)}  # nothing that names the appointees
    requests = sr.collect_slot_requests(mapping, list(mapping), data, company_name=COMPANY)
    want = len({p for p in (sr._slot_position(n) for n in mapping) if p})
    return ("appointed: one ask per declared position", len(positions_asked(requests, sr._slot_position)), want, "")


def case_member_family_suppressed(sr, declared):
    """A known member list must SUPPRESS the numbered member asks.

    7 members, 0 questions. That zero is CORRECT and it is not new — read on
    before filing it as a regression.

    The rule that produces it is a blanket one in `collect_slot_requests`:

        if kind in ("shareholder_list", "attendee") and member_parties:
            continue

    added 2026-08-06 for finding F1, on the grounds that `repeat_regions` fills
    those slots from the very same member list, so asking about them is noise.
    It fires the moment a member list is resolvable at all, and it returns
    BEFORE any of the party-count arithmetic — so for this family the ask count
    is 0 whether or not the count logic is present, and it reads the same in the
    old code and the new.

    ★ The growth cases therefore use `appointed_director`, not `member`. That
    family has no register fallback and no blanket rule, so its count is
    observable: 7 parties -> 7 asks, 2 -> 2, unknown -> 1. Expecting 7 member
    questions here would be encoding a misreading of this rule, and the case
    would sit at 0 forever looking like a bug in somebody's fix.

    Retiring the blanket rule is a semantics decision about whether a member the
    register already knows should still be confirmed. Deliberately not touched.
    """
    spellings = declared[MEMBER_TEMPLATE]["member"]
    mapping = {name: slot_entry("shareholder_list", multi=True) for names in spellings.values() for name in names}
    # `corporate_shareholder_name` is supplied because the corporate spelling
    # (`corporate shareholder_3_name`) routes through
    # `_corporate_shareholder_name`, which otherwise SELECTs shareholder_links
    # to find the member company. The agent supplies it in a real run too.
    data = {"members": parties(7), "corporate_shareholder_name": CORPORATE_MEMBER}
    requests = sr.collect_slot_requests(mapping, list(mapping), data, company_name=COMPANY)
    return ("member family: 7 known members suppress the numbered asks", len(requests), 0, "")


def case_member_family_unknown_asks(sr, declared):
    """…and with NO member list, it must still ask.

    The mirror of the case above. An empty register proves nothing, so silence
    here would mean the suppression fires on absence of evidence — which is how
    a document gets signed by whoever happened to be first.
    """
    spellings = declared[MEMBER_TEMPLATE]["member"]
    mapping = {name: slot_entry("shareholder_list", multi=True) for names in spellings.values() for name in names}
    requests = sr.collect_slot_requests(mapping, list(mapping), {}, company_name=None)
    return ("member family: nothing on file, still asks", 1 if requests else 0, 1, f"{len(requests)} request(s)")


def case_render_grows(sr, rr):
    """The expander's side of the same number, on the real template.

    7 members must render 7 signature blocks. Counted as `__rr_N__` tokens,
    which is what the expander returns and the highlighter then fills.
    """
    from docx import Document

    doc = Document(str(TEMPLATE_DIR / MEMBER_TEMPLATE))
    synth = rr.expand_repeat_regions(doc, {"members": parties(7)}, template_name=MEMBER_TEMPLATE, company_name=COMPANY)
    return ("render: 7 members expand the document", len(synth), 14, "two tokens per member (name + identifier)")


def case_render_shrinks(sr, rr):
    """…and 1 member must shrink it, from a template that ships 3 slots."""
    from docx import Document

    doc = Document(str(TEMPLATE_DIR / MEMBER_TEMPLATE))
    synth = rr.expand_repeat_regions(doc, {"members": parties(1)}, template_name=MEMBER_TEMPLATE, company_name=COMPANY)
    return ("render: 1 member shrinks the document", len(synth), 2, "")


def case_smart_doc_reaches_expander(sr, sd):
    """`smart_doc` must actually reach the expander — proved by EFFECT.

    `fill_template_with_validation` imports `expand_repeat_regions` lazily
    inside the function body, so there is no module attribute to compare and a
    source grep would only prove the line exists, not that it runs. Driving the
    real entry point and counting table rows proves the call happened.

    Row counts on this template: 3 slots ship as 4 rows; 7 parties render 22.
    """

    path = TEMPLATE_DIR / MEMBER_TEMPLATE
    before = hashlib.md5(path.read_bytes()).hexdigest()
    doc = sd.fill_template_with_validation(
        path,
        {"members": parties(7), "company_name": COMPANY},
        template_name=MEMBER_TEMPLATE,
        company_name=COMPANY,
    )
    rows = sum(len(t.rows) for t in doc.tables)
    after = hashlib.md5(path.read_bytes()).hexdigest()
    if before != after:
        return ("smart_doc: MUTATED THE TEMPLATE ON DISK", 1, 0, "template md5 changed")
    return ("smart_doc: reaches the expander (7 parties -> table rows)", rows, 22, "template unchanged on disk")


SLOT_CASES = [
    case_grow,
    case_grow_spelling,
    case_shrink,
    case_exact,
    case_unknown_count_collapses,
    case_appointed_asks_per_declared_position,
    case_member_family_suppressed,
    case_member_family_unknown_asks,
]


# ── mutants ─────────────────────────────────────────────────────────
#
# `scout/tools/slot_resolver.py` belongs to another agent and is being written
# to concurrently, so no mutant edits a file. Each one monkeypatches a SEAM at
# runtime, reproduces a specific defect, and is restored in a finally. What is
# asserted is that the golden NUMBER MOVES: a mutant that leaves the count
# unchanged means the case was not measuring what it claims to measure.


def mutate_layout_only(sr, rr):
    """THE ORIGINAL BUG: the count comes from the template's layout.

    Blanking the party list is what the ask step effectively saw before it
    learned to consult the expander's source — it could only ever see the
    positions `mapping` declared.
    """
    original = rr._parties_for_family

    def layout_only(family, tail, data, company_name):
        return []

    rr._parties_for_family = layout_only
    return lambda: setattr(rr, "_parties_for_family", original)


def mutate_never_shrink(sr, rr):
    """OVER-CORRECTION: a family's real count is reported as the declared max.

    Reproduces a fix that grows the ask list but never lets it shrink — every
    template would ask about all of its numbered slots regardless of how many
    parties exist.
    """
    original = sr._family_parties

    def padded(family, tail, data, company_name):
        real = original(family, tail, data, company_name)
        return real if len(real) >= 3 else real + [{"name": "PHANTOM"}] * (3 - len(real))

    sr._family_parties = padded
    return lambda: setattr(sr, "_family_parties", original)


def mutate_member_list_unresolvable(sr, rr):
    """The member list stops being resolvable at all.

    This is the seam that actually governs `case_member_family_suppressed`. The
    first version of this mutant patched `_member_position_covered` and the
    number did NOT move — because the kind-level rule
    (`kind in ("shareholder_list", "attendee") and member_parties`) returns
    first and the position check is never reached. The suite caught that: a
    mutant that leaves the golden number alone is reported as a failure, which
    is exactly what it did.
    """
    original = sr._member_family_parties
    sr._member_family_parties = lambda data, company_name: []
    return lambda: setattr(sr, "_member_family_parties", original)


def mutate_collapse_distinct_positions(sr, rr):
    """Appointed-director positions collapse back to one question per role.

    Reproduces the defect measured on THAZIN VALLEY HOLDINGS: the dedup key used
    the position ONLY when a real party count was known, and `appointed_director`
    can never be counted, so three declared positions became one request. Emptying
    `_DISTINCT_POSITION_FAMILIES` restores exactly that rule without touching the
    key's construction, so what moves the number is the family exemption itself.
    """
    original = sr._DISTINCT_POSITION_FAMILIES
    sr._DISTINCT_POSITION_FAMILIES = frozenset()
    return lambda: setattr(sr, "_DISTINCT_POSITION_FAMILIES", original)


MUTANTS = [
    # (name, apply, the cases whose number MUST move)
    (
        "layout-only (the original bug)",
        mutate_layout_only,
        [
            "grow: 7 parties, template declares 3",
            "grow: synthesised spelling matches the template's",
            "shrink: 2 parties, template declares 3",
        ],
    ),
    ("never-shrink (over-correction)", mutate_never_shrink, ["shrink: 2 parties, template declares 3"]),
    (
        "member list unresolvable",
        mutate_member_list_unresolvable,
        ["member family: 7 known members suppress the numbered asks"],
    ),
    (
        "collapse distinct positions (the under-asking bug)",
        mutate_collapse_distinct_positions,
        ["appointed: one ask per declared position"],
    ),
]


# ── runner ──────────────────────────────────────────────────────────


def _literal_set(source: str, name: str) -> set:
    """The literal set assigned to `name` anywhere in `source`, via AST.

    Parsed rather than imported because `KNOWN_USER_INPUT` is a local inside
    `find_replacement` and cannot be reached any other way. AST rather than a
    text scan so a name merely MENTIONED in a comment cannot satisfy the check —
    only a real assignment of a real set literal counts.
    """
    tree = ast.parse(source)
    for node in ast.walk(tree):
        if not isinstance(node, ast.Assign):
            continue
        for target in node.targets:
            if isinstance(target, ast.Name) and target.id == name:
                try:
                    return set(ast.literal_eval(node.value))
                except (ValueError, SyntaxError):
                    return set()
    return set()


def case_defaults_have_real_defaults(smart_doc_source, mutant=""):
    """★ `DEFAULT_FIELDS` must not name a field the resolver leaves unfilled.

    That set exempts its members from the value check, so a field listed there
    is reported "available", never put to the user, and reaches the document
    unfilled. Twelve of its fourteen original members had no default at all:
    five resolved to the literal "TBD" (they were ALSO in the resolver's own
    `KNOWN_USER_INPUT`, which says outright that they are user input), and seven
    resolved to None, which leaves the raw placeholder in the output with no
    highlight.

    Pinned as a contradiction between two sets in one file, so it holds without
    a database: nothing may be exempt from being asked AND declared user input,
    and nothing may be exempt while matching the resolver's TBD patterns.
    """
    defaults = _literal_set(smart_doc_source, "DEFAULT_FIELDS")
    user_input = _literal_set(smart_doc_source, "KNOWN_USER_INPUT")
    if mutant == "default_fields_drift":
        defaults = defaults | {"auditor_name", "financial_year_end_date"}
    overlap = defaults & user_input
    tbd_shaped = {f for f in defaults if "auditor" in f or "financial_year" in f}
    return (
        "DEFAULT_FIELDS names no unfillable field",
        len(overlap | tbd_shaped),
        0,
        f"offenders={sorted(overlap | tbd_shaped)}" if (overlap | tbd_shaped) else "",
    )


def run_mode(contract, rows):
    label = "contract" if contract else "fallback"
    sr, rr, sd, wired = load_product(contract)
    if contract and not wired:
        rows.append((f"[{label}]", "app.slot_contract could not be wired", 0, 1, False, ""))
        return
    if contract is False and sr._contract_normalise is not None:
        rows.append((f"[{label}]", "expected NO contract, got one", 1, 0, False, ""))
        return

    placeholders = importlib.import_module("scout.tools.placeholders")
    declared = declared_slots(placeholders, rr._family, sr._slot_position)

    print(f"\n{'=' * 82}\nmode={label}  (slot_contract wired: {sr._contract_normalise is not None})")
    print(
        f"templates on disk: {len(list(TEMPLATE_DIR.glob('*.docx')))}  "
        f"declaring numbered repeat families: {len(declared)}"
    )
    for name, families in declared.items():
        summary = "  ".join(f"{fam}={sorted(pos)}" for fam, pos in sorted(families.items()))
        print(f"  · {name[:62]:<62} {summary}")

    golden = {}
    before = TRIPWIRE.attempts
    for case in SLOT_CASES:
        name, got, want, note = case(sr, declared)
        golden[name] = got
        ok = got == want
        rows.append((f"[{label}]", name, got, want, ok, note))
        print(f"  {'OK ' if ok else 'BAD'} {name:<58} got={got:<3} want={want}" + (f"   {note}" if note else ""))
    sd_source = (REPO / "scout" / "tools" / "smart_doc.py").read_text()
    for mut in ("", "default_fields_drift"):
        name, got, want, note = case_defaults_have_real_defaults(sd_source, mut)
        if mut:
            # the mutant must MOVE the number, otherwise the case is inert
            name, want = f"MUTANT {mut} moves DEFAULT_FIELDS check", 2
        ok = got == want
        rows.append((f"[{label}]", name, got, want, ok, note))
        print(f"  {'OK ' if ok else 'BAD'} {name:<58} got={got:<3} want={want}" + (f"   {note}" if note else ""))

    slot_db = TRIPWIRE.attempts - before
    rows.append((f"[{label}]", "slot cases attempted no DB access", slot_db, 0, slot_db == 0, ""))
    print(f"  {'OK ' if slot_db == 0 else 'BAD'} {'slot cases attempted no DB access':<58} got={slot_db:<3} want=0")

    for case in (case_render_grows, case_render_shrinks):
        name, got, want, note = case(sr, rr)
        golden[name] = got
        ok = got == want
        rows.append((f"[{label}]", name, got, want, ok, note))
        print(f"  {'OK ' if ok else 'BAD'} {name:<58} got={got:<3} want={want}" + (f"   {note}" if note else ""))

    # smart_doc legitimately tries the DB for `field_mapping`; the tripwire
    # refuses every attempt, which is why its slot pass is out of scope here.
    before = TRIPWIRE.attempts
    name, got, want, note = case_smart_doc_reaches_expander(sr, sd)
    ok = got == want
    rows.append((f"[{label}]", name, got, want, ok, note))
    print(f"  {'OK ' if ok else 'BAD'} {name:<58} got={got:<3} want={want}   {note}")
    refused = TRIPWIRE.attempts - before
    print(
        f"      (smart_doc made {refused} field_mapping DB attempt(s); all refused, "
        f"so its slot pass ran on an empty mapping — see the module docstring)"
    )

    # ── mutants ──
    print("\n  mutants (fault injected at a seam, restored after):")
    for mutant_name, apply, must_move in MUTANTS:
        restore = apply(sr, rr)
        try:
            for case in SLOT_CASES:
                name, got, _want, _note = case(sr, declared)
                if name not in must_move:
                    continue
                moved = got != golden[name]
                rows.append(
                    (f"[{label}]", f"MUTANT {mutant_name} moves {name[:28]}", got, golden[name], moved, "must differ")
                )
                print(
                    f"    {'OK ' if moved else 'BAD'} {mutant_name:<34} "
                    f"{name[:30]:<30} golden={golden[name]} mutant={got}"
                    + ("" if moved else "   <<< NUMBER DID NOT MOVE")
                )
        finally:
            restore()

    # the seam really is restored
    name, got, want, _ = case_grow(sr, declared)
    rows.append((f"[{label}]", "seams restored after mutants", got, want, got == want, ""))
    print(f"  {'OK ' if got == want else 'BAD'} {'seams restored after mutants':<58} got={got:<3} want={want}")
    return golden


def main():
    _install_stubs()
    if not TEMPLATE_DIR.is_dir() or not list(TEMPLATE_DIR.glob("*.docx")):
        sys.exit(f"No templates at {TEMPLATE_DIR} — this suite reads the real .docx files.")

    rows = []
    golden_fallback = run_mode(False, rows)
    golden_contract = run_mode(True, rows)

    # ★ The two normalisers must agree. They are selected by whether
    # `app.slot_contract` imports, which depends on the machine — see
    # `_install_slot_contract`. A number that differs between them is a bug that
    # only appears in the container, or only on a laptop.
    print(f"\n{'=' * 82}\nfallback vs contract")
    if golden_fallback and golden_contract:
        for name in sorted(golden_fallback):
            a, b = golden_fallback[name], golden_contract.get(name)
            ok = a == b
            rows.append(("[agree]", f"fallback == contract: {name[:44]}", a, b, ok, ""))
            print(f"  {'OK ' if ok else 'BAD'} {name[:62]:<62} fallback={a} contract={b}")

    failures = [r for r in rows if not r[4]]
    print(f"\n\n{'MODE':<12} {'CHECK':<58} {'GOT':>5} {'WANT':>5}  ")
    print("-" * 92)
    for mode, name, got, want, ok, _note in rows:
        print(f"{mode:<12} {name[:58]:<58} {got!s:>5} {want!s:>5}  {'ok' if ok else 'BAD'}")

    print(f"\nSUMMARY: {len(rows)} checks · {len(failures)} failed · {TRIPWIRE.attempts} DB attempt(s), all refused")
    if failures:
        print(
            "A mutant whose number did NOT move means the case is not measuring "
            "what it claims to measure — treat it as a broken gate, not a flake."
        )
    return 1 if failures else 0


if __name__ == "__main__":
    sys.exit(main())
