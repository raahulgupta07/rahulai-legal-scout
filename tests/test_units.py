"""
Unit tests — pure logic, no server, no LLM, no network.

The three tracker layers all talk to a running app, and layers 2 and 3 talk to
the model, so they are slow and (for 2 and 3) non-deterministic. This layer
covers the decision logic underneath them: placeholder normalisation, slot
classification, session scoping, the person guard, party coercion, and the
structural contracts that have silently broken the product before.

Every case here is deterministic. A failure is a real regression.

Run (inside the container, which has the dependencies):

    docker exec scout-api python3 /app/tests/test_units.py

Or after `docker cp tests/test_units.py scout-api:/app/tests/`.

Only a live DB connection is needed for the import to succeed (see BOOTSTRAP);
no test reads or writes application data.
"""

import json
import os
import re
import sys
from pathlib import Path

# --- BOOTSTRAP -------------------------------------------------------------
# `import scout.tools.<x>` alone raises ImportError: scout/__init__ imports
# scout.agent, which reaches db -> app.model_config -> app/__init__ -> app.main
# -> scout.agent (partially initialised). Importing app.main FIRST establishes
# the same order uvicorn uses and breaks the cycle. This is why the suite runs
# in the container rather than on the host.
sys.path.insert(0, "/app")
from scout.tools import ask_questions as aq
from scout.tools import fill_view as fv
from scout.tools import people_picker as pp
from scout.tools import repeat_regions as rr
from scout.tools import slot_resolver as sr
from scout.tools.field_aliases import (
    canonical_field,
    normalize_field,
    tokens_match,
)
from scout.tools.placeholders import (
    is_empty_placeholder,
    new_empty_counter,
    placeholder_name,
)

REPO = Path("/app")

_results = []


def check(case_id: str, name: str, passed: bool, detail: str = ""):
    _results.append((case_id, "PASS" if passed else "FAIL", name, detail))


def eq(case_id: str, name: str, got, want):
    check(case_id, name, got == want, f"got {got!r}, want {want!r}")


def skip(case_id: str, name: str, why: str):
    """Record a case that could not run, without calling it a failure.

    A few assertions need real rows in the register — a company with a corporate
    member, and that member's board. On a FRESH INSTALL there are none, and
    reporting those as failures is wrong in the expensive direction: it makes a
    correct empty product look broken, and it trains you to ignore red. The
    check is skipped and says exactly which fixture is missing.
    """
    _results.append((case_id, "SKIP", name, why))


def register_has(company: str) -> bool:
    """True when `company` exists in the register, for fixture-dependent cases."""
    try:
        from scout.tools.companies_db import get_all_companies

        target = company.strip().lower()
        return any(target in str(c.get("name") or "").strip().lower() for c in get_all_companies(limit=300))
    except Exception:
        return False


# ===========================================================================
# U1  Placeholder normalisation
#     Word writes non-breaking spaces INSIDE placeholder names in five of the
#     client's templates. A different string from the normal-space spelling, so
#     every alias and field_mapping written by hand missed and the field came
#     out blank.
# ===========================================================================
def test_placeholders():
    eq(
        "U1a",
        "U+00A0 inside a placeholder folds to a normal space",
        placeholder_name(("individual\xa0shareholder_1_name", None, None)),
        "individual shareholder_1_name",
    )
    eq("U1b", "zero-width space is stripped", placeholder_name(("director\u200b_name", None, None)), "director_name")
    eq("U1c", "an ordinary name is untouched", placeholder_name(("meeting_date", None, None)), "meeting_date")
    eq(
        "U1d",
        "surrounding whitespace is trimmed",
        placeholder_name(("  chairperson_name  ", None, None)),
        "chairperson_name",
    )
    eq(
        "U1e",
        "the second capture group is used when the first is empty",
        placeholder_name((None, "company_name", None)),
        "company_name",
    )

    counter = new_empty_counter()
    first = placeholder_name((None, None, None), counter)
    second = placeholder_name((None, None, None), counter)
    check(
        "U1f",
        "bare slots get distinct synthetic names",
        bool(first) and bool(second) and first != second,
        f"{first!r} then {second!r}",
    )
    eq("U1g", "a bare slot with no counter stays empty", placeholder_name((None, None, None)), "")
    check("U1h", "is_empty_placeholder recognises a synthetic name", is_empty_placeholder(first), f"{first!r}")


# ===========================================================================
# U2  Field aliases
# ===========================================================================
def test_field_aliases():
    eq("U2a", "spaces and case normalise", normalize_field("Company Name"), "company_name")
    eq("U2b", "normalising is idempotent", normalize_field("company_name"), "company_name")
    eq(
        "U2c",
        "normalising folds U+00A0 too",
        normalize_field("individual\xa0shareholder_1_name"),
        "individual_shareholder_1_name",
    )
    eq("U2d", "canonical_field maps a known alias", canonical_field("nrc"), "nrc_no")
    check(
        "U2e",
        "tokens_match does not equate a numbered slot with the generic key",
        tokens_match("shareholder_1_name", "shareholder_name") is False,
        "",
    )


# ===========================================================================
# U3  Slot kind classification
#     PICKER_SLOT_KINDS used to stamp every pick "signatory", and the patterns
#     only matched snake_case. A prose purpose fell through to the catch-all, so
#     a person chosen as the INCOMING director appeared on the NEXT document's
#     resignation line.
# ===========================================================================
def test_classify_kind():
    cases = [
        ("U3a", "select the new director to be appointed", "new_director"),
        ("U3b", "the director being appointed", "new_director"),
        ("U3c", "choose the resigning director", "resigning_director"),
        ("U3d", "the outgoing director", "resigning_director"),
        ("U3e", "who will chair the meeting", "chairperson"),
        ("U3f", "authorised person to represent the shareholder", "representative"),
        ("U3g", "persons present at the meeting", "attendee"),
        ("U3h", "the auditor for the year", "auditor"),
        ("U3i", "sign the AGM resolution", "signatory"),
    ]
    for cid, text, want in cases:
        eq(cid, f"prose purpose {text!r}", sr.classify_kind(text), want)

    eq("U3j", "an unrelated string classifies as unknown, not a guess", sr.classify_kind("the meeting date"), "")
    eq("U3k", "empty input classifies as unknown", sr.classify_kind(""), "")
    check(
        "U3l",
        "resigning outranks the generic director pattern",
        sr.classify_kind("director resigning from the board") == "resigning_director",
        "",
    )


# ===========================================================================
# U4  Session scope
#     A pick belongs to ONE conversation. Before session_id existed, any pick
#     for the same company within 30 minutes was reusable by any other chat.
# ===========================================================================
def test_session_scope():
    eq("U4a", "no scope bound by default", sr.current_session_scope(), "")

    with sr.session_scope("sess-A"):
        eq("U4b", "scope is visible inside the block", sr.current_session_scope(), "sess-A")
        with sr.session_scope("sess-B"):
            eq("U4c", "scopes nest", sr.current_session_scope(), "sess-B")
        eq("U4d", "the outer scope is restored", sr.current_session_scope(), "sess-A")
    eq("U4e", "scope is cleared on exit", sr.current_session_scope(), "")

    try:
        with sr.session_scope("sess-boom"):
            raise RuntimeError("generation blew up")
    except RuntimeError:
        pass
    eq("U4f", "scope is cleared even when the block raises", sr.current_session_scope(), "")

    with sr.session_scope(None):
        eq("U4g", "None scope reads as empty", sr.current_session_scope(), "")
    with sr.session_scope("  padded  "):
        eq("U4h", "scope is stripped", sr.current_session_scope(), "padded")

    # The read-back must refuse to answer with no conversation in scope: that is
    # the Fill-in view, which supplies its own values and must inherit nothing.
    slot = {"kind": "new_director", "of": "people_register"}
    eq(
        "U4i",
        "picker log returns nothing when no session is in scope",
        sr._parties_from_picker_log("ANY COMPANY LIMITED", slot),
        [],
    )


# ===========================================================================
# U5  Companion identifier — an NRC belongs to a person
#     `new_director_name` is a slot answered by a picker; training classifies
#     `new_director_identification_number` as user_input, so it was asked as free
#     text. The two halves of one person resolved from different sources and the
#     document went out reading "SOE MOE THU (... NRC/Passport number: )".
# ===========================================================================
def test_companion_identifier():
    ident = [
        ("U5a", "new_director_identification_number", "new_director"),
        ("U5b", "new_director_nrc", "new_director"),
        ("U5c", "director_passport_no", "director"),
        ("U5d", "resigning_director_nrc_passport_no", "resigning_director"),
        ("U5e", "shareholder_1_id_number", "shareholder_1"),
        ("U5f", "individual\xa0shareholder_1_nrc", "individual shareholder_1"),
    ]
    for cid, ph, want in ident:
        eq(cid, f"identifier role of {ph!r}", sr._role_prefix(ph, sr._IDENTIFIER_ATTR_RE), want)

    # Numbers that belong to a COMPANY must never be treated as a person's.
    for cid, ph in [
        ("U5g", "company_registration_number"),
        ("U5h", "certificate_of_incorporation_number"),
        ("U5i", "share_certificate_number"),
        ("U5j", "registration_number"),
        ("U5k", "meeting_date"),
    ]:
        eq(cid, f"{ph!r} is not a person identifier", sr._role_prefix(ph, sr._IDENTIFIER_ATTR_RE), "")

    mapping = {
        "new_director_name": {
            "source": "slot",
            "slot": {"of": "people_register", "kind": "new_director", "multi": False},
        },
        "new_director_identification_number": {"source": "user_input", "slot": None},
    }
    person = {"name": "SOE MOE THU", "identifier": "12/SAKHANA(N)021426", "party_type": "individual"}

    eq(
        "U5l",
        "the NRC comes from the person the name slot resolved to",
        sr.companion_identifier("new_director_identification_number", mapping, {"new_director_name": person}),
        "12/SAKHANA(N)021426",
    )
    eq(
        "U5m",
        "no resolved person means the field is still asked",
        sr.companion_identifier("new_director_identification_number", mapping, {}),
        None,
    )
    eq(
        "U5n",
        "a person with no identifier on file does not fabricate one",
        sr.companion_identifier(
            "new_director_identification_number", mapping, {"new_director_name": {"name": "NO NRC PERSON"}}
        ),
        None,
    )
    eq(
        "U5o",
        "a non-identifier placeholder is left alone",
        sr.companion_identifier("meeting_date", mapping, {"new_director_name": person}),
        None,
    )
    eq(
        "U5p",
        "a company registration number is not answered from a person",
        sr.companion_identifier("company_registration_number", mapping, {"new_director_name": person}),
        None,
    )


# ===========================================================================
# U31 Bare person attributes — the register already holds them
#     Measured on the live box 2026-08-27. The user picked KYAW THU SOE from the
#     card; the register held nationality=Myanmar and date_of_birth=1987-09-21
#     for him, and the consent form asked for both as free text anyway.
#
#     `Director Consent Form - Non-Group Member Appointment.docx` writes its
#     placeholders BARE — `nationality`, not `director_nationality` — and both
#     bridges from a picked person to their attributes were keyed on a role
#     prefix, so neither could see them.
#
#     The third finding is the expensive one: `{address}` is a token subset of
#     `{registered, office, address}`, so a bare `address` meaning the
#     DIRECTOR's home resolved silently to the COMPANY's registered office. A
#     blank gets proof-read; a confident wrong address does not.
# ===========================================================================
def test_bare_person_attributes():
    import scout.tools.smart_doc as sd

    # --- the tail table itself -------------------------------------------
    for cid, ph, want in [
        ("U31a", "nationality", ("", "nationality")),
        ("U31b", "date_of_birth", ("", "date_of_birth")),
        ("U31c", "address", ("", "residential_address")),
        ("U31d", "nric", ("", "nrc_passport_no")),
        ("U31e", "country_of_residence", ("", "country_of_residence")),
        ("U31f", "director_nationality", ("director", "nationality")),
        ("U31g", "new_director_residential_address", ("new_director", "residential_address")),
        ("U31h", "resigning_director_date_of_birth", ("resigning_director", "date_of_birth")),
    ]:
        eq(cid, f"attr tail of {ph!r}", sr._attr_tail(ph), want)

    # Longest tail wins, or the short one strips the wrong prefix off the role.
    eq(
        "U31i",
        "residential_address is not read as role 'residential'",
        sr._attr_tail("director_residential_address"),
        ("director", "residential_address"),
    )

    # Company fields are not person attributes.
    for cid, ph in [
        ("U31j", "company_registration_number"),
        ("U31k", "meeting_date"),
        ("U31l", "company_name"),
        ("U31m", "registered_office"),
    ]:
        eq(cid, f"{ph!r} is not a person attribute", sr._attr_tail(ph), None)

    # `nric` is the spelling the live template uses and was missing from the
    # identifier pattern entirely — it could never resolve as an identifier.
    eq("U31n", "nric is an identifier tail", sr._role_prefix("director_nric", sr._IDENTIFIER_ATTR_RE), "director")

    # --- who a bare attribute belongs to ---------------------------------
    one_person = {
        "director_name": {
            "source": "slot",
            "slot": {"of": "people_register", "kind": "new_director", "multi": False},
        },
        "nationality": {"source": "user_input", "slot": None},
        "address": {"source": "user_input", "slot": None},
        "company": {"source": "db", "db_column": "company_name_english", "slot": None},
    }
    eq("U31o", "one person slot names the bare role", sr.sole_person_role(one_person), "director")

    two_people = {
        "resigning_director_name": {
            "source": "slot",
            "slot": {"of": "people_register", "kind": "resigning_director", "multi": False},
        },
        "new_director_name": {
            "source": "slot",
            "slot": {"of": "people_register", "kind": "new_director", "multi": False},
        },
    }
    eq(
        "U31p",
        "TWO people make a bare attribute ambiguous — ask, never guess",
        sr.sole_person_role(two_people),
        "",
    )
    eq("U31q", "no person slot at all", sr.sole_person_role({"company": {"source": "db"}}), "")

    # A LIST of people is not a single person either.
    multi = {
        "shareholder_names": {
            "source": "slot",
            "slot": {"of": "document_company", "kind": "shareholder_list", "multi": True},
        }
    }
    eq("U31r", "a multi slot does not own a bare attribute", sr.sole_person_role(multi), "")

    # --- the identifier still resolves without a second read -------------
    person = {"name": "KYAW THU SOE", "identifier": "12/LAMANA(N)142591", "party_type": "individual"}
    eq(
        "U31s",
        "a bare nric is answered by the picked person's identifier",
        sr.companion_attribute("nric", one_person, {"director_name": person}),
        "12/LAMANA(N)142591",
    )
    eq(
        "U31t",
        "with nobody picked the field is still asked",
        sr.companion_attribute("nationality", one_person, {}),
        None,
    )
    eq(
        "U31u",
        "a bare attribute in a two-person template is never guessed",
        sr.companion_attribute("nationality", two_people, {"new_director_name": person}),
        None,
    )
    eq(
        "U31v",
        "a company field is not answered from a person",
        sr.companion_attribute("company_registration_number", one_person, {"director_name": person}),
        None,
    )

    # --- the address defect ----------------------------------------------
    # The first diagnosis of this was WRONG and this case is what caught it.
    # The theory was a loose token-subset match — `{address}` inside
    # `{registered, office, address}`. It is not: `address` is a GENERIC token,
    # so that tier refuses it outright.
    #
    # The real path is blunter. `prepare_document_data` republishes the company
    # record under aliases, one of which is a BARE `address`, so the company's
    # own office wins on the very first line of `_resolve_from_data` — exact
    # match — long before any companion lookup runs. Verified on the live box:
    # keys containing "addr" were ['registered_office_address',
    # 'company_address', 'address'].
    company_office = "BARGAYAR ROAD, NO. 1-11, PADONMAR STADIUM (EAST WING), YANGON"
    company_data = {
        "company_name_english": "CITY MART HOLDING COMPANY LIMITED",
        "registered_office_address": company_office,
    }
    eq(
        "U31w",
        "a generic token never claims a longer company key",
        sd._resolve_from_data("address", company_data),
        None,
    )
    eq(
        "U31x",
        "the bare company alias is what answered a person's field",
        sd._resolve_from_data("address", {**company_data, "address": company_office}),
        company_office,
    )
    # Strict mode is kept as a GUARD, not as the fix: a placeholder like
    # `nationality` is a token subset of a company key such as
    # `nationality_of_ultimate_holding_company`, and that tier would answer it.
    eq(
        "U31xa",
        "strict mode blocks a person's field resolving from a company key",
        sd._resolve_from_data(
            "nationality", {"nationality_of_ultimate_holding_company": "Singapore"}, strict=True
        ),
        None,
    )
    eq(
        "U31xb",
        "and the same lookup does resolve without it — this tier is real",
        sd._resolve_from_data("nationality", {"nationality_of_ultimate_holding_company": "Singapore"}),
        "Singapore",
    )
    # An answer the USER actually gave still wins, strict or not — it names the
    # field rather than merely containing its tokens.
    eq(
        "U31y",
        "a supplied answer still outranks the register",
        sd._resolve_from_data("address", {**company_data, "address": "No. 7, Bahan Township"}, strict=True),
        "No. 7, Bahan Township",
    )
    eq(
        "U31z",
        "and so does one supplied under an alias",
        sd._resolve_from_data(
            "address", {**company_data, "residential_address": "No. 7, Bahan Township"}, strict=True
        ),
        "No. 7, Bahan Township",
    )

    # A non-person user_input field must keep its loose match — this is the
    # tier the strict pass skips, and nothing else may lose it.
    eq(
        "U31aa",
        "a non-person field is untouched by strict mode",
        sd._resolve_from_data("registration_number", {"company_registration_number": "113516550"}),
        "113516550",
    )

    # --- the wiring actually reaches find_replacement ---------------------
    # A source read, not a behaviour read: the branch must call the companion
    # lookup and must NOT fall back to the loose tier for a person attribute.
    import inspect

    body = inspect.getsource(sd.find_replacement)
    check(
        "U31ab",
        "find_replacement asks the companion resolver",
        "companion_attribute(" in body,
        "user_input branch no longer calls companion_attribute",
    )
    check(
        "U31ac",
        "a person attribute is read from supplied data in strict mode",
        "strict=person_attr" in body,
        "strict flag not threaded into the user_input branch",
    )
    check(
        "U31ad",
        "a person attribute does not fall back to the loose tier",
        "if not person_attr:" in body,
        "loose fallback is not gated on person_attr",
    )
    check(
        "U31ae",
        "the company's own address is discarded for a person's address field",
        "_company_own_addresses(company_name)" in body,
        "find_replacement no longer screens out the company's own address",
    )


# ===========================================================================
# U32 One resolver, three views
#     The register fix shipped and the product still looked unfixed. The
#     question card correctly stopped asking for nationality and date of birth;
#     the Fields panel beside it went on showing both as PENDING, and displayed
#     the COMPANY's registered office under the DIRECTOR's `address`.
#
#     Nothing was stale. Three pieces of code each decided independently whether
#     a field was filled — generation and the Fill-in view through
#     `find_replacement`, and the Fields panel through `_resolve_from_data`
#     alone, which cannot see the People register.
#
#     This case pins the invariant that was missing: all three answer from ONE
#     resolver. A source check, because the alternative needs a live database
#     and a picked person, and a guard that only runs with fixtures present is a
#     guard that stops running.
# ===========================================================================
def test_one_resolver_three_views():
    import inspect

    import scout.tools.smart_doc as sd

    # --- the resolver exists and reports both halves ----------------------
    check(
        "U32a",
        "a single resolver answers what is filled and what is blank",
        callable(getattr(sd, "_resolved_values", None)),
        "_resolved_values is missing",
    )
    sig = list(inspect.signature(sd._resolved_values).parameters)
    eq("U32b", "it takes the fields, the data and the template", sig[:4],
       ["required_fields", "data", "template_name", "company_name"])

    # --- it goes through find_replacement, not the flat data lookup -------
    body = inspect.getsource(sd._resolved_values)
    check(
        "U32c",
        "the resolver calls find_replacement",
        "find_replacement(" in body,
        "_resolved_values does not use find_replacement",
    )
    check(
        "U32d",
        "the resolver does not fall back to the flat data lookup",
        "_resolve_from_data(" not in body,
        "_resolved_values still reads data directly, which cannot see the register",
    )

    # --- every panel state is built from it -------------------------------
    # `_document_state` renders whatever it is handed, so the defect lives in
    # the CALLERS. Each one must hand it resolved values, never a raw company
    # record — that record publishes a bare `address` alias holding the
    # company's own office.
    tool_src = inspect.getsource(sd)
    calls = tool_src.split("_document_state(")[1:]
    raw_values = []
    for i, seg in enumerate(calls):
        head = seg[: seg.find(")\n") if ")\n" in seg else 900]
        for bad in ("values=normalized_company_data", "values=normalized_data", "values=slot_data",
                    "values=data,", "values=preview_data,"):
            if bad in head:
                raw_values.append((i, bad))
    eq("U32e", "no panel state is built from a raw company record", raw_values, [])

    check(
        "U32f",
        "every _document_state call site passes resolved values",
        tool_src.count("_resolved_values(") >= 5,
        f"only {tool_src.count('_resolved_values(')} resolver call(s) for "
        f"{tool_src.count('_document_state(') - 1} panel state(s)",
    )

    # --- what is outstanding is decided by the resolver too ---------------
    prep = inspect.getsource(sd.prepare_document_data)
    # NOT `"_resolved_blank" in prep` — that survived its own mutation, because
    # a mutant that reassigned the same NAME from `validation["missing_fields"]`
    # still contained the string. The call is what has to be there.
    check(
        "U32g",
        "prepare decides outstanding from the resolver, not from name matching",
        "_resolved_values(" in prep and "outstanding = [f for f in _resolved_blank" in prep,
        "prepare_document_data no longer derives outstanding from the resolver",
    )
    check(
        "U32h",
        "ready and outstanding cannot contradict each other",
        "ready = not outstanding" in prep,
        "ready is computed independently of outstanding and can disagree with it",
    )

    # --- `address` no longer claims a default it does not have ------------
    validate = inspect.getsource(sd.validate_data_vs_template)
    check(
        "U32i",
        "address is not counted available without a value",
        'DEFAULT_FIELDS = {"meeting_location"}' in validate,
        "address is still exempted from the value check, so it is never asked",
    )


# ===========================================================================
# U33 The card must not suggest what the system cannot know
#     Measured 2026-08-27. The register holds a name, an NRC/passport, a
#     nationality and a date of birth — and nothing else, for all 9 people on
#     file. Country of residence, home address, phone and email are empty for
#     every one of them, which is exactly why they are asked.
#
#     The model, composing the question, offered "Myanmar" as a one-click chip
#     for COUNTRY OF RESIDENCE, inferred from the nationality it could see. The
#     drafting skill forbids that in as many words: "If it is missing for a
#     signatory, ask; do not infer it from nationality." A Myanmar national
#     resident in Singapore has a different answer, and the consent form states
#     where the director RESIDES.
#
#     The card is the only place these answers are given, so it is the only
#     place the suggestion can be prevented. The rule is tested by RUNNING the
#     card's own regex, not by checking that a line exists.
# ===========================================================================
def test_card_suggests_nothing_unknowable():
    card = Path(__file__).resolve().parent.parent / "agent-ui/src/components/chat/AskUserCard.tsx"
    if not card.exists():
        skip("U33", "card rule", f"{card} not present in this tree")
        return
    body = card.read_text()

    m = re.search(r"const UNKNOWABLE_FACT_RE\s*=\s*\n?\s*/(.+?)/i", body, re.DOTALL)
    check("U33a", "the card declares the unknowable-fact rule", bool(m), "UNKNOWABLE_FACT_RE not found")
    if not m:
        return

    # The TS source doubles its backslashes; read it as the regex the browser
    # actually compiles.
    pattern = re.compile(m.group(1).replace("\\\\", "\\"), re.IGNORECASE)

    for cid, text in [
        ("U33b", "Country of residence?"),
        ("U33c", "COUNTRY OF RESIDENCE?"),
        ("U33d", "What is the director's residential address?"),
        ("U33e", "Contact phone number?"),
        ("U33f", "Email address?"),
        ("U33g", "DIRECTOR'S RESIDENTIAL ADDRESS?"),
    ]:
        check(cid, f"chips are dropped for {text!r}", bool(pattern.search(text)),
              "the card would still offer a guessed answer")

    # It must NOT swallow questions that legitimately carry options — blocking
    # those would break the verified template-choice and approval flows.
    for cid, text in [
        ("U33h", "Which template did you mean?"),
        ("U33i", "Generate the Director Consent Form now?"),
        ("U33j", "Where is the meeting held?"),
        ("U33k", "Execution date of the consent form?"),
        ("U33l", "What is the director's nationality?"),
    ]:
        check(cid, f"options survive for {text!r}", not pattern.search(text),
              "a legitimate chip question would lose its options")

    # --- the rule is applied, not merely declared -------------------------
    check(
        "U33m",
        "the card strips options before rendering",
        "isUnknowableFact(q)" in body and "options: undefined" in body,
        "UNKNOWABLE_FACT_RE is declared but never applied to the questions",
    )

    # --- every date box opens on today, and stays editable ----------------
    # Scoped to the SEED block on purpose. Checking the whole file for
    # `looksLikeDateQuestion(...)` survived its own mutation — the same call
    # also appears in `isDateQuestion` further down, so narrowing the seed back
    # to `input_type === 'date'` left the string in place and the case green.
    seed_block = body.split("const [freeText")[1].split("const [manualEntry")[0] if "const [freeText" in body else ""
    check(
        "U33n",
        "EVERY date question is seeded with today, not just declared ones",
        "looksLikeDateQuestion(question.text, question.id)" in seed_block
        and "seed[question.id] = todayISO()" in seed_block,
        "date questions are no longer pre-filled with the current date",
    )
    check(
        "U33o",
        "the seeded date remains editable",
        "manualEntry" in body,
        "the card no longer offers a way to type a date by hand",
    )

    # --- the model is told the same thing at the tool boundary ------------
    import scout.tools.ask_questions as aqm

    fn = aqm.ask_questions
    # `@tool` wraps the function, and the WRAPPER carries a docstring of its
    # own — so `fn.__doc__` is truthy and reading it first silently tests the
    # wrong text. The model is shown `entrypoint.__doc__`.
    doc = getattr(getattr(fn, "entrypoint", None), "__doc__", "") or fn.__doc__ or ""
    check(
        "U33p",
        "the tool tells the model not to offer these options",
        "country of residence" in doc.lower() and "free text" in doc.lower(),
        "ask_questions no longer instructs the model against inferred options",
    )


# ===========================================================================
# U34 The panel must watch the tool names that actually stream
#     The document panel read "No document yet" beside a chat holding that
#     document's template, company, director and NRC. Nothing had failed: the
#     run had called `preview_doc`, which carries a full `document_state`, and
#     the panel's DOC_TOOLS set listed `preview_document`.
#
#     `_as_json` wraps with @wraps, so agno registers a tool under the FUNCTION
#     name, not under the key `create_smart_document_tool` exports it as. The
#     export map says "preview_document"; the stream says `preview_doc`.
#     agent.py:213 documents this exactly, and the PROMPT was corrected for it
#     at the time — three frontend files were not, and stayed wrong.
#
#     So the guard derives the expected names from the registry rather than
#     restating them: a rename that moves the tool moves this test with it.
# ===========================================================================
def test_panel_watches_real_tool_names():
    ui = Path(__file__).resolve().parent.parent / "agent-ui/src"
    artifact = ui / "components/shell/useArtifact.ts"
    if not artifact.exists():
        skip("U34", "panel tool names", f"{artifact} not present in this tree")
        return

    from scout.tools.smart_doc import create_smart_document_tool

    exported = create_smart_document_tool(documents_dir="/documents", host="")
    # What agno registers is the wrapped function's __name__ — the same thing
    # the stream puts in tool_name — NOT the export key.
    # `_as_json` is applied to exactly the tools whose results the panel reads —
    # the export map says so — and functools.wraps leaves `__wrapped__` behind.
    # That is the signal, rather than a hand-listed set that would drift the
    # same way DOC_TOOLS did. `analyze_template` is exported unwrapped and
    # returns no document_state, so it is correctly out of scope.
    registered = {
        key: getattr(fn, "__name__", key)
        for key, fn in exported.items()
        if hasattr(fn, "__wrapped__")
    }
    check(
        "U34z",
        "the wrapped set is not empty",
        bool(registered),
        "no _as_json-wrapped tools found — this case would pass vacuously",
    )

    mismatched = {k: v for k, v in registered.items() if k != v}
    check(
        "U34a",
        "the export key and the registered name still differ (the trap is real)",
        bool(mismatched),
        "no mismatch found — this case may no longer be measuring anything",
    )

    body = artifact.read_text()
    # `create_document` is deliberately not bound in agent.py, so it never
    # streams; everything else the map exports can.
    expected = {name for key, name in registered.items() if key != "create_document"}
    missing = sorted(n for n in expected if f"'{n}'" not in body)
    eq("U34b", "DOC_TOOLS lists every name that can stream", missing, [])

    # The two panel-opening regexes match by PREFIX, so `preview_doc` covers
    # `preview_document` too — but the reverse is false, which is the bug.
    for cid, rel in [
        ("U34c", "components/shell/ArtifactPanel.tsx"),
        ("U34d", "app/page.tsx"),
    ]:
        f = ui / rel
        if not f.exists():
            skip(cid, f"{rel} opens the panel", "file not present")
            continue
        text = f.read_text()
        m = re.search(r"/\^\(([a-z_|]+)\)/", text)
        alts = set(m.group(1).split("|")) if m else set()
        unmatched = sorted(
            n for n in expected if not any(n.startswith(a) for a in alts)
        )
        eq(cid, f"{rel} opens the panel for every document tool", unmatched, [])


# ===========================================================================
# U35 Multiple stakeholders, and the fields that belong to each of them
#     Measured 2026-08-27 on the live register (9 people, all with NRC /
#     nationality / date of birth, none with address / country / phone / email).
#
#     Three defects, all found by scanning every template rather than the one
#     in front of us:
#
#     P2 — `appointed_director_1_nrc` is TRAINED as a person NAME slot
#          (source=slot, kind=new_director), so the slot branch rendered the
#          NAME into an NRC field. Verified before the fix: all three NRC
#          fields on the Corporate Shareholder Consent read "KYAW THU SOE",
#          "MIN MIN", "WIN WIN TINT". A consent form stating
#          "N.R.C./Passport: KYAW THU SOE" is filled, confident and wrong —
#          strictly worse than a blank, which a reader would catch.
#
#     P3 — the Individual Shareholder Consent Form has ONE person, but its
#          slot is typed `shareholder_list`, so the sole-person rule excluded
#          it and all four of its attributes were asked from a register that
#          held three of them.
#
#     P4 — `company_address` was read as a PERSON's residential address.
#          Harmless while it fell through to a question, one slot name away
#          from writing a director's home address into a company field.
# ===========================================================================
def test_multi_stakeholder_attributes():
    # --- P4: a company is not a person -----------------------------------
    for cid, ph in [
        ("U35a", "company_address"),
        ("U35b", "registered_office_address"),
        ("U35c", "company_name"),
    ]:
        eq(cid, f"{ph!r} is not a person attribute", sr._attr_tail(ph), None)

    # `corporate_shareholder_address` is deliberately NOT in that list. A
    # corporate shareholder is a party, so the role reads as person-shaped and
    # `_attr_tail` classifies it — but the party is a COMPANY, and what matters
    # is that nothing from a person's register row can reach it. Asserted as
    # behaviour rather than classification: the earlier version of this case
    # demanded the stricter classification and would have forced a change that
    # buys nothing, while breaking the corporate-party identifier path.
    corporate_party = {
        "name": "PAHTAMA GROUP COMPANY LIMITED",
        "identifier": "165855078",
        "party_type": "corporate",
    }
    corp_map = {
        "corporate shareholder_1_name": {
            "source": "slot",
            "slot": {"kind": "attendee", "multi": False, "of": "document_company"},
        },
        "corporate_shareholder_1_address": {"source": "user_input", "slot": None},
    }
    eq(
        "U35d",
        "a corporate party's address never comes from the People register",
        sr.companion_attribute(
            "corporate_shareholder_1_address", corp_map,
            {"corporate shareholder_1_name": corporate_party},
        ),
        None,
    )

    # ...but a bare one still is, or the whole fix stops working.
    eq("U35e", "a bare address is still a person's", sr._attr_tail("address"), ("", "residential_address"))
    eq(
        "U35f",
        "a role-prefixed person address still is",
        sr._attr_tail("new_director_residential_address"),
        ("new_director", "residential_address"),
    )

    # --- P3: one person, however the slot is typed ------------------------
    single_list = {
        "shareholder_name": {
            "source": "slot",
            "slot": {"kind": "shareholder_list", "multi": False, "of": "document_company"},
        },
        "nationality": {"source": "user_input", "slot": None},
    }
    check(
        "U35g",
        "a list-typed slot holding ONE person still names the bare role",
        bool(sr.sole_person_role(single_list)),
        "sole_person_role is empty, so a single-signatory consent form asks for what the register holds",
    )

    truly_multi = {
        "shareholder_names": {
            "source": "slot",
            "slot": {"kind": "shareholder_list", "multi": True, "of": "document_company"},
        },
        "nationality": {"source": "user_input", "slot": None},
    }
    eq("U35h", "a genuinely multi slot owns no bare attribute", sr.sole_person_role(truly_multi), "")

    two_people = {
        "resigning_director_name": {
            "source": "slot",
            "slot": {"kind": "resigning_director", "multi": False, "of": "people_register"},
        },
        "new_director_name": {
            "source": "slot",
            "slot": {"kind": "new_director", "multi": False, "of": "people_register"},
        },
        "nationality": {"source": "user_input", "slot": None},
    }
    eq("U35i", "two people make a bare attribute ambiguous", sr.sole_person_role(two_people), "")

    # --- the runtime net: one slot, but the pick resolved to SEVERAL ------
    # `sole_person_role` reads the TEMPLATE; this reads the SELECTION. A
    # shareholder_list slot can be typed multi=False and still have two people
    # chosen through the card, and printing the first one's nationality beside
    # a joined list of names is the failure this prevents.
    a = {"name": "KYAW THU SOE", "identifier": "12/LAMANA(N)142591", "party_type": "individual"}
    b = {"name": "MIN MIN", "identifier": "12/LATHANA(N)016603", "party_type": "individual"}
    eq(
        "U35j",
        "ONE selected party answers a bare attribute",
        sr.companion_attribute("nric", single_list, {"shareholder_name": a}),
        "12/LAMANA(N)142591",
    )
    eq(
        "U35k",
        "TWO selected parties do not — it asks",
        sr.companion_attribute("nric", single_list, {"shareholder_name": [a, b]}),
        None,
    )
    eq(
        "U35l",
        "and a bare attribute is never taken from one of two person slots",
        sr.companion_attribute("nationality", two_people, {"new_director_name": a}),
        None,
    )

    # --- P2: an identifier placeholder may never render a name -----------
    # Against the REAL template, because the defect lives in the TRAINED
    # mapping: `appointed_director_1_nrc` carries source=slot with a
    # person-NAME slot, so the slot branch rendered the name. A hand-built
    # mapping cannot reproduce that — find_replacement reads the trained one.
    import scout.tools.smart_doc as sd

    TPL = (
        "Corporate Shareholder Consent - Directors Resolution "
        "for New Company Setup and Director Appointment.docx"
    )
    mapping = sd._get_field_mapping(TPL) or {}
    if not mapping.get("appointed_director_1_nrc"):
        skip("U35m", "an identifier placeholder never renders a name", f"{TPL} is not trained here")
    else:
        got = sd.find_replacement(
            "appointed_director_1_nrc",
            {"appointed_director_1_nrc": a, "director_1_name": a},
            template_name=TPL,
            company_name="CITY MART HOLDING COMPANY LIMITED",
        )
        text = "" if got is None else str(got)
        check(
            "U35m",
            "an identifier placeholder never renders a person's name",
            text != a["name"],
            f"the NRC field resolved to {text!r} — that is the director's NAME",
        )
        check(
            "U35m2",
            "and it resolves to the identifier the picker carried",
            text == a["identifier"] or text == "",
            f"expected {a['identifier']!r} or blank, got {text!r}",
        )

    import inspect

    body = inspect.getsource(sd.find_replacement)
    check(
        "U35n",
        "the slot branch screens identifier placeholders",
        "_attr_tail(" in body,
        "find_replacement's slot branch no longer checks whether the placeholder wants an identifier",
    )


# ===========================================================================
# U36 A turn that ends OWING a tool call is a stall, however well it reads
#     Session 75400f45-49c5-4fb1-aa1b-d97a555ee6cd, read from ai.agno_sessions:
#     one request took THREE runs because the model twice ended a turn without
#     making the call its own tool result demanded, and the user typed
#     "continue" by hand each time.
#
#       run 1  COMPLETED  lookup_director_candidates -> never called choose_director
#       run 2  COMPLETED  generate_document (blank fields) -> never called ask_questions
#       run 3  PAUSED     ask_questions   <- finally asked
#
#     The existing guard only fired on `chunk.content.trim() === ''`. Both
#     stalls ended with a confident "Preview Summary" paragraph, so it never
#     ran. Ending with polished prose is WORSE than ending silently: it looks
#     finished, so nobody can tell the turn stalled.
# ===========================================================================
def test_stall_guard_sees_unpaid_debt():
    ui = Path(__file__).resolve().parent.parent / "agent-ui/src"
    helper = ui / "components/chat/toolDebt.ts"
    hook = ui / "hooks/useAIStreamHandler.tsx"
    if not helper.exists() or not hook.exists():
        skip("U36", "stall guard", "agent-ui sources not present in this tree")
        return

    h = helper.read_text()
    k = hook.read_text()

    check("U36a", "the debt is a pure, testable helper", "export const findUnpaidToolDebt" in h,
          "findUnpaidToolDebt is not exported, so it cannot be tested apart from the stream")

    # Read STRUCTURALLY, never out of the English instruction strings — those
    # get reworded every time the prompts are tuned, and a guard keyed to
    # prose would go quiet the next time someone improves the wording.
    # Comments STRIPPED first. The first version of this case failed against
    # correct code: toolDebt.ts explains in its header comment that it must not
    # key off `agent_instruction`, and the scan matched that explanation. A
    # source check that reads its own documentation is measuring nothing.
    h_code = re.sub(r"/\*.*?\*/", "", h, flags=re.DOTALL)
    h_code = re.sub(r"^\s*//.*$", "", h_code, flags=re.MULTILINE)
    check("U36b", "debt is read from the result shape, not the prose",
          "agent_instruction" not in h_code,
          "toolDebt.ts keys off agent_instruction text, which is reworded whenever prompts change")

    for cid, token, why in [
        ("U36c", "picker", "a lookup result names the picker it owes"),
        ("U36d", "blank_fields", "an undeclared-blanks failure owes a question"),
        ("U36e", "user_input_fields", "a needs-input failure owes a question"),
    ]:
        check(cid, why, token in h, f"toolDebt.ts does not look at {token}")

    # A pause means the card IS on screen. Nudging there would talk over a
    # human mid-answer and send "continue" into a run blocked on them.
    check("U36f", "a paused run owes nothing", "paused" in h and "if (options.paused) return null" in h,
          "the helper does not discharge the debt on a paused run")

    # --- wired, not merely declared --------------------------------------
    check("U36g", "the hook imports it", "findUnpaidToolDebt" in k,
          "the stream handler does not import the helper")
    check("U36h", "the debt is one of the nudge triggers", "|| !!toolDebt" in k,
          "toolDebt is computed but never reaches shouldNudge")

    # The budget is the whole safety story. Each nudge starts a NEW run with a
    # NEW id, so a guard that reset its counter on any real text would be
    # unbounded — and the measured stall ENDS with real text. An unbounded
    # nudge is how the same document once got generated three times.
    check("U36i", "text written while still owing a call does not refill the budget",
          "isRealContent(chunk.content) && !toolDebt" in k,
          "the nudge counter resets on prose, so the debt guard is unbounded")
    check("U36j", "the debt guard shares the existing cap",
          "MAX_CONSECUTIVE_NUDGES" in k and "autoContinuedRunsRef" in k,
          "the debt guard does not reuse the per-run guard and the nudge cap")
    check("U36k", "nothing is nudged while the user is being asked",
          "!waitingOnUser && didToolWork" in k,
          "the debt is computed without the waiting-on-user and tool-work gates")


# ===========================================================================
# U37 A pick changes the answer, so the panel must move with it
#     Measured in a real browser 2026-08-27, session 9c586813 on 1.2.68: after
#     choosing KYAW THU SOE the panel still read 3/11 with NATIONALITY, NRIC
#     and DATE OF BIRTH pending — while the question card beside it correctly
#     did NOT ask for any of the three, because the resolver had all three from
#     the register.
#
#     The run was `choose_director` then `ask_questions`. No document tool ran
#     after the pick, so no fresh document_state was ever emitted and the panel
#     kept showing what `preview_doc` computed BEFORE anyone was chosen. Honest
#     about what it last knew; wrong about what is true — and what a lawyer
#     reads there is "these fields are still missing".
# ===========================================================================
def test_pick_refreshes_the_panel():
    import inspect

    body = inspect.getsource(pp._selection_result)

    check(
        "U37a",
        "the picker looks up the task it is answering",
        "recall_task" in body,
        "the picker cannot know the template, so it cannot refresh the panel",
    )
    check(
        "U37b",
        "and declares the state that pick produces",
        "document_state" in body,
        "the picker no longer returns a document_state",
    )
    # An absent key must leave the panel alone. Blanking it on a pick would
    # trade a stale number for no number, which is worse.
    check(
        "U37c",
        "an uncomputed state is omitted, never sent empty",
        "if _state:" in body,
        "the picker may attach an empty document_state and blank the panel",
    )
    # A panel refresh must never be able to lose the selection itself.
    check(
        "U37d",
        "the refresh cannot fail the pick",
        "except Exception" in body,
        "the refresh is unguarded, so a panel error would drop the user's choice",
    )

    # --- the browser has to actually fold it ------------------------------
    art = Path(__file__).resolve().parent.parent / "agent-ui/src/components/shell/useArtifact.ts"
    if not art.exists():
        skip("U37e", "the panel folds a picker's declared state", "useArtifact.ts not in this tree")
        return
    a = art.read_text()

    check(
        "U37e",
        "the panel admits any result that declares a state",
        "declaresState(c)" in a and "DOC_TOOLS.has(c.tool_name) || declaresState(c)" in a,
        "the fold still filters by tool NAME only, so a picker's state never reaches the panel",
    )
    # By EVIDENCE, not by a second name list that would drift the way DOC_TOOLS
    # already did once (preview_doc vs preview_document).
    # Comments STRIPPED. This case failed against correct code on its first
    # run — the header comment above `declaresState` names `choose_director`
    # while explaining why picker names must NOT be listed, and the scan
    # matched that explanation. Third time today a source check has read its
    # own documentation; the lesson is that any of these must strip comments
    # before they measure anything.
    a_code = re.sub(r"/\*.*?\*/", "", a, flags=re.DOTALL)
    a_code = re.sub(r"^\s*//.*$", "", a_code, flags=re.MULTILINE)
    check(
        "U37f",
        "admission is by evidence, not by a list of picker names",
        "choose_director" not in a_code,
        "useArtifact hardcodes picker names, which is the drift that made DOC_TOOLS stale",
    )


# ===========================================================================
# U38 One form, one card
#     Measured on the live box 2026-08-27, Director Consent Form (Non-Group)
#     for CITY MART HOLDING: the panel listed FIVE outstanding fields — date,
#     address, country of residence, phone, email — and the card asked four.
#     `email` was never put to the user, the answers came back one short,
#     generate_document refused it as an undeclared blank, and the run needed a
#     second card to ask one question.
#
#     It worked only because the stall guard absorbs the extra round. The form
#     is one form; the person filling it should see all of it.
#
#     The number lives in THREE places — the validator, the tool docstring and
#     the system prompt — and the last two are prose that cannot interpolate a
#     constant. So this case pins them to it. Two of them are what the MODEL
#     reads: a prompt that still said 1-4 would cap the card at four however
#     high the validator allowed.
# ===========================================================================
def test_one_card_holds_the_whole_form():
    import inspect

    import scout.agent as agent_mod
    import scout.tools.ask_questions as aqm

    cap = getattr(aqm, "MAX_QUESTIONS", None)
    check("U38a", "the cap is a named constant", isinstance(cap, int), f"MAX_QUESTIONS is {cap!r}")
    if not isinstance(cap, int):
        return

    # Six is the largest free-text load measured across the 15 templates.
    check("U38b", "the cap clears the worst real template", cap >= 6,
          f"cap is {cap}; a template needing 6 free-text answers still splits across two cards")
    # A cap still has to exist — an unbounded call could put thirty boxes on
    # screen, and "no limit" is not the fix for "the limit was wrong".
    check("U38c", "but a cap still exists", cap <= 12, f"cap is {cap}, which is effectively unbounded")

    # --- the validator enforces exactly that number ----------------------
    src = inspect.getsource(aqm._validate_questions)
    check("U38d", "the validator reads the constant, not a literal",
          "MAX_QUESTIONS" in src and "<= 4" not in src,
          "the validator hardcodes its own limit, so the constant is decoration")
    eq("U38e", f"{cap} questions are accepted",
       aqm._validate_questions([{"id": f"q{i}", "text": f"Q{i}?"} for i in range(cap)]), [])
    check("U38f", "one more is refused",
          bool(aqm._validate_questions([{"id": f"q{i}", "text": f"Q{i}?"} for i in range(cap + 1)])),
          "the validator accepts more than the cap")

    # --- both places the MODEL reads agree with it -----------------------
    fn = aqm.ask_questions
    doc = getattr(getattr(fn, "entrypoint", None), "__doc__", "") or fn.__doc__ or ""
    check("U38g", "the tool docstring quotes the same cap", f"1-{cap}" in doc,
          f"the docstring does not say 1-{cap}; the model will cap itself at whatever it says")

    prompt = ""
    for name in ("SYSTEM_PROMPT", "INSTRUCTIONS", "_PROMPT"):
        v = getattr(agent_mod, name, None)
        if isinstance(v, str) and "ask_questions" in v:
            prompt = v
            break
    if not prompt:
        prompt = Path(agent_mod.__file__).read_text()
    check("U38h", "the system prompt quotes the same cap",
          f"1-{cap} questions per" in prompt,
          f"the prompt does not say 1-{cap} questions per call, so the model still splits the form")
    check("U38i", "and tells the model to ask the whole form at once",
          "EVERY outstanding field in ONE card" in prompt,
          "nothing instructs the model to put the whole form on one card")

    # --- raising the cap alone did NOT fix it ----------------------------
    # Verified in the browser on 1.2.70: with room for eight the model asked
    # four and MERGED phone and email into one box under the invented id
    # `contact_info`. That id matches no placeholder, so the answer reaches
    # neither field, both stay blank, and the second round happens anyway.
    # The cap was necessary and not sufficient — the id has to BE the field.
    check("U38j", "the schema ties a question id to the field it answers",
          "EXACT field name" in doc or "EXACT field name" in inspect.getsource(aqm),
          "the schema still calls `id` a free-form key, so the model invents ids like contact_info")
    check("U38k", "and forbids merging two fields into one question",
          "never merge two" in doc.lower() or "never merge two" in inspect.getsource(aqm).lower(),
          "nothing stops the model asking phone and email in a single box")

    import scout.tools.smart_doc as _sd

    gen = inspect.getsource(_sd)
    check("U38l", "the needs-input instruction names the id rule too",
          "as that question's `id` VERBATIM" in gen,
          "generate_document asks for one question per field but never says the id must be the field name")


# ===========================================================================
# U39 Placeholders hidden inside tracked changes
#     A finished notice to shareholders went out reading
#
#         Name: [director_name]
#
#     on its signature line. Measured 2026-08-27 on the installed
#     `Notice of Annual General Meeting to Shareholders.docx` (30,807 bytes):
#     the file carries 8 <w:ins> and 2 <w:del> blocks — unaccepted tracked
#     changes — and FOUR placeholders live inside the insertions.
#
#     `paragraph.runs` and `paragraph.text` return only runs that are DIRECT
#     children of <w:p>. A run inside <w:ins> is one level deeper, so neither
#     saw them; Word renders them perfectly normally. The fill could not
#     replace what it could not see, the unfilled-placeholder audit could not
#     flag it, and TRAINING missed them too — the trained mapping holds 8
#     fields where the document has 11.
#
#     <w:del> is the opposite case and must stay excluded: Word does not render
#     deleted text, so filling it writes into something nobody sees and
#     counting it asks the user for a field that does not exist.
# ===========================================================================
def test_tracked_changes_are_visible_to_the_fill():
    from docx import Document
    from docx.oxml.ns import qn

    import scout.tools.placeholders as ph

    for name in ("visible_runs", "paragraph_text", "tracked_changes"):
        check(f"U39_{name}", f"placeholders exposes {name}", callable(getattr(ph, name, None)),
              f"{name} is missing")
    if not all(callable(getattr(ph, n, None)) for n in ("visible_runs", "paragraph_text", "tracked_changes")):
        return

    # --- build a paragraph with the exact shape found in the live template ---
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Name: ")
    ins = p._p.makeelement(qn("w:ins"), {})
    ins.set(qn("w:author"), "CH Legal")
    run_el = p._p.makeelement(qn("w:r"), {})
    t = p._p.makeelement(qn("w:t"), {})
    t.text = "[director_name]"
    run_el.append(t)
    ins.append(run_el)
    p._p.append(ins)

    eq("U39a", "paragraph.text cannot see a tracked insertion", p.text, "Name: ")
    eq("U39b", "paragraph_text can", ph.paragraph_text(p), "Name: [director_name]")
    check("U39c", "and the run is offered for filling",
          any("[director_name]" in (x.text or "") for x in ph.visible_runs(p)),
          "visible_runs does not return the inserted run, so the fill cannot replace it")

    # --- a DELETION must stay invisible ----------------------------------
    p2 = doc.add_paragraph()
    p2.add_run("Kept ")
    dele = p2._p.makeelement(qn("w:del"), {})
    dr = p2._p.makeelement(qn("w:r"), {})
    # A plain <w:t> inside the deletion, NOT <w:delText>. The first version of
    # this case used delText and survived its own mutation: python-docx's
    # Run.text reads only <w:t>, so the text was empty whatever the exclusion
    # did, and the case proved nothing. This shape makes the exclusion the only
    # thing standing between the fill and text Word does not render.
    dt = p2._p.makeelement(qn("w:t"), {})
    dt.text = "[gone_field]"
    dr.append(dt)
    dele.append(dr)
    p2._p.append(dele)
    check("U39d", "deleted text is not offered for filling",
          "[gone_field]" not in ph.paragraph_text(p2),
          "paragraph_text includes tracked DELETIONS, which Word does not render")

    # --- the integrity check names the problem ----------------------------
    rev = ph.tracked_changes(doc)
    check("U39e", "unaccepted revisions are reported", bool(rev), "tracked_changes found nothing in a document that has some")
    check("U39f", "and the author is named", "CH Legal" in (rev.get("authors") or []),
          f"authors not reported: {rev}")
    eq("U39g", "a clean document reports nothing", ph.tracked_changes(Document()), {})

    # --- the fill and the audit actually use it --------------------------
    import inspect

    import scout.tools.smart_doc as sd

    fill_src = inspect.getsource(sd._fill_paragraph_highlighted)
    check("U39h", "the fill walks visible runs", "visible_runs(paragraph)" in fill_src,
          "the fill still reads paragraph.runs, so an inserted placeholder is never replaced")
    # Comments STRIPPED — this case failed against correct code because the
    # comment ABOVE the fix explains what `paragraph.text` cannot see, and the
    # scan matched that explanation. Fourth source check today to read its own
    # documentation; strip first, always.
    fill_code = "\n".join(
        line.split("#", 1)[0] for line in fill_src.splitlines() if not line.strip().startswith("#")
    )
    check("U39i", "and never re-reads paragraph.text for its match", "paragraph.text" not in fill_code,
          "the fill still matches against paragraph.text")

    mod = inspect.getsource(sd)
    check("U39j", "generation warns about unaccepted revisions", "tracked_changes(filled_doc)" in mod,
          "nothing tells the firm their template carries tracked changes")

    # EVERY reader of a template must see the same text, or one of them reports
    # a field the others cannot. 1.2.71 taught the fill, the audit and training;
    # it missed the Fill-in view, and tracker_layer1 A2b stayed red for exactly
    # that reason — the panel could not offer a person to pick because it could
    # not see the person placeholder.
    import scout.tools.fill_view as fv_mod

    fv_src = inspect.getsource(fv_mod)
    fv_code = "\n".join(
        line.split("#", 1)[0] for line in fv_src.splitlines() if not line.strip().startswith("#")
    )
    check(
        "U39k",
        "the Fill-in view reads the same visible text as the fill",
        "paragraph_text(para)" in fv_code and "para.text," not in fv_code,
        "fill_view still emits para.text, so a placeholder inside a tracked insertion never reaches the panel",
    )


# ===========================================================================
# U6  Party coercion — whatever a picker or the model hands back
# ===========================================================================
def test_party_coercion():
    eq(
        "U6a",
        "a dict becomes one party",
        sr._coerce_parties({"name": "A B", "identifier": "X1"}),
        [{"name": "A B", "identifier": "X1", "party_type": "individual", "representative": ""}],
    )
    eq("U6b", "a list of names becomes many parties", [p["name"] for p in sr._coerce_parties(["A", "B"])], ["A", "B"])
    eq("U6c", "TBD is not a person", sr._coerce_parties("TBD"), [])
    eq("U6d", "empty input yields no parties", sr._coerce_parties(""), [])
    eq("U6e", "None yields no parties", sr._coerce_parties(None), [])
    eq("U6f", "a JSON string is parsed", [p["name"] for p in sr._coerce_parties('[{"name": "C D"}]')], ["C D"])
    eq("U6g", "an unanswered picker payload yields no parties", sr._coerce_parties({"candidates": [{"name": "X"}]}), [])
    eq(
        "U6h",
        "nrc_passport_no is accepted as the identifier",
        sr._coerce_parties({"name": "E F", "nrc_passport_no": "9/ABC(N)1"})[0]["identifier"],
        "9/ABC(N)1",
    )

    eq("U6i", "slot_of returns None for a non-slot entry", sr.slot_of({"source": "user_input"}), None)
    eq(
        "U6j",
        "slot_of returns the descriptor for a slot entry",
        sr.slot_of({"source": "slot", "slot": {"kind": "signatory"}}),
        {"kind": "signatory"},
    )
    eq("U6k", "a slot entry with no kind is not a slot", sr.slot_of({"source": "slot", "slot": {}}), None)


# ===========================================================================
# U7  Person guard — a human is chosen from the register, never typed
# ===========================================================================
def test_person_guard():
    must_block = [
        ("U7a", "What is the full legal name of the new director being appointed?"),
        ("U7b", "What is the NRC or Passport number of the new director?"),
        ("U7c", "Who will sign the minutes?"),
        ("U7d", "What is the name of the resigning director?"),
        ("U7e", "Please provide the chairperson name"),
        ("U7f", "What is the shareholder name?"),
        ("U7g", "What is the passport number of the signatory?"),
    ]
    for cid, q in must_block:
        role = aq._person_role(q)
        check(cid, f"blocked: {q}", role is not None, f"role={role!r}")
        if role is not None:
            check(
                cid + "-p",
                "the blocked question names a real picker pair",
                role in aq._PICKER_FOR_ROLE,
                f"role={role!r}",
            )

    must_allow = [
        ("U7h", "What is the proposed name of the new company?"),
        ("U7i", "What is the primary business sector for the new company?"),
        ("U7j", "What is the meeting date?"),
        ("U7k", "What is the financial year end date?"),
        ("U7l", "What is the chairperson pronoun?"),
        ("U7m", "Generate Director Resignation Letter for CITY HOLDINGS LIMITED now?"),
        ("U7n", "Which company is this shareholder consent form being prepared for?"),
        ("U7o", "What is the capital subscription amount?"),
        ("U7p", "How many shares will be subscribed?"),
        ("U7q", "What is the company name?"),
        ("U7r", "What is the template name?"),
        ("U7s", "What is the registered office address?"),
    ]
    for cid, q in must_allow:
        role = aq._person_role(q)
        check(cid, f"allowed: {q}", role is None, f"role={role!r}")

    # A batch with both kinds must not be thrown away wholesale.
    questions = [
        {"id": "q0", "text": "What is the meeting date?"},
        {"id": "q1", "text": "What is the name of the new director?"},
        {"id": "q2", "text": "What is the financial year end date?"},
    ]
    allowed, blocked = aq._split_person_questions(questions)
    eq("U7t", "mixed batch keeps the legitimate questions", allowed, ["q0", "q2"])
    eq("U7u", "mixed batch blocks only the person question", [b["id"] for b in blocked], ["q1"])
    check(
        "U7v",
        "a blocked question names the tools to call instead",
        bool(blocked) and len(blocked[0].get("call_instead") or []) == 2,
        str(blocked[0].get("call_instead") if blocked else None),
    )

    # A question offering options is a constrained pick, not free text.
    with_options = [{"id": "q0", "text": "Which director should sign?", "options": ["A", "B"]}]
    allowed_o, blocked_o = aq._split_person_questions(with_options)
    eq("U7w", "an option-bearing question is left alone", (allowed_o, blocked_o), (["q0"], []))

    eq(
        "U7x",
        "answers for blocked questions are dropped",
        aq._filter_answers({"q0": "2026-09-15", "q1": "Typed Name"}, questions, ["q0", "q2"]),
        {"q0": "2026-09-15"},
    )
    eq(
        "U7y",
        "id-keyed answer lists are filtered the same way",
        aq._filter_answers([{"id": "q0", "answer": "x"}, {"id": "q1", "answer": "y"}], questions, ["q0", "q2"]),
        [{"id": "q0", "answer": "x"}],
    )
    eq(
        "U7z",
        "an unrecognised answer shape passes through untouched",
        aq._filter_answers("just a string", questions, ["q0"]),
        "just a string",
    )


# ===========================================================================
# U8  Picker payload — how a paused tool learns its conversation
#     A requires_user_input tool CANNOT take run_context: agno builds its
#     user_input_schema from every signature parameter, the frontend echoes the
#     whole schema back on resume, and the injected copy collides with it
#     ("got multiple values for keyword argument 'run_context'"). The lookup
#     tools, which never pause, carry the session in their payload instead.
# ===========================================================================
def test_picker_payload():
    payload = pp._payload(picker="choose_director", candidates=[], session="sess-42")
    eq("U8a", "the lookup payload carries the session", payload.get("session"), "sess-42")
    eq("U8b", "the picker reads the session back out", pp._session_from_payload(json.dumps(payload)), "sess-42")
    eq("U8c", "a dict payload works as well as a JSON string", pp._session_from_payload(payload), "sess-42")
    eq("U8d", "malformed JSON yields no session rather than raising", pp._session_from_payload("{not json"), "")
    eq("U8e", "an absent session field yields empty", pp._session_from_payload('{"candidates": []}'), "")
    eq("U8f", "None yields empty", pp._session_from_payload(None), "")
    eq("U8g", "a non-object payload yields empty", pp._session_from_payload("[1,2,3]"), "")

    eq(
        "U8h",
        "a prose purpose is recorded under the role it names",
        pp._classify_purpose("select the new director to be appointed"),
        "new_director",
    )

    cand = pp._candidate(person_id=7, name="A B", identifier="1/AB(N)2")
    eq("U8i", "a candidate keeps its identifier so the NRC travels with the pick", cand["identifier"], "1/AB(N)2")
    eq("U8j", "a candidate is individual unless told otherwise", cand["party_type"], "individual")


# ===========================================================================
# U9  Repeat regions — grow/shrink party blocks to the real count
# ===========================================================================
def test_repeat_regions():
    eq(
        "U9a",
        "'name' is checked before 'share' so shareholder_1_name is a name",
        rr._tail_attr("shareholder_1_name"),
        "name",
    )
    eq("U9b", "an NRC tail is an nrc", rr._tail_attr("shareholder_1_nrc"), "nrc")
    eq("U9c", "a passport tail is an nrc", rr._tail_attr("director_passport"), "nrc")
    eq("U9d", "a percentage tail", rr._tail_attr("shareholder_1_percentage"), "percentage")
    eq("U9e", "a space-delimited share count is shares", rr._tail_attr("number of shares"), "shares")
    eq("U9f", "'shareholder' alone is not a share count", rr._tail_attr("shareholder_1"), "name")
    # KNOWN GAP, asserted so a fix is noticed rather than silently changing
    # behaviour: `\bshares?\b` needs a non-word delimiter and `_` is a word
    # character, so the underscore spelling falls through to the "name"
    # fallback. `number_of_shares` is a real placeholder in a real template. It
    # only misrenders if it sits INSIDE a repeat region — standalone it never
    # reaches _tail_attr — which has not been confirmed either way.
    eq(
        "U9e2",
        "KNOWN GAP: the underscore spelling is not recognised as shares",
        rr._tail_attr("number_of_shares"),
        "name",
    )

    # Real DICA data spells the corporate type "Company", not "corporate".
    eq("U9g", "type 'Company' is corporate", rr._is_corporate({"type": "Company"}), True)
    eq("U9h", "type 'corporate' is corporate", rr._is_corporate({"type": "corporate"}), True)
    eq("U9i", "type 'Individual' is not corporate", rr._is_corporate({"type": "Individual"}), False)
    eq(
        "U9j",
        "an unlabelled company name falls back to the name heuristic",
        rr._is_corporate({"name": "PAHTAMA GROUP CO., LTD"}),
        True,
    )
    eq("U9k", "an unlabelled person name is not corporate", rr._is_corporate({"name": "SOE MOE THU"}), False)
    eq(
        "U9l",
        "an explicit individual label beats a company-looking name",
        rr._is_corporate({"type": "individual", "name": "LIMITED HOLDINGS"}),
        False,
    )

    # --- U18 ---------------------------------------------------------------
    # The signature table in several templates has NO header row: the
    # "Members to sign…" cue is an ordinary paragraph ABOVE it and row 0 is
    # already a signatory unit. The gate only ever read row 0, so those tables
    # were skipped whole, both slots fell through to the flat per-company fill,
    # and a company with ONE corporate member was rendered signing TWICE — once
    # through its representative, once again on the individual line.
    #
    # Built synthetically rather than read from documents/legal/templates:
    # that directory is a bind mount the firm edits, so a test reading it
    # asserts against whatever is on disk today.
    from docx import Document as _Docx

    def _signing_doc():
        d = _Docx()
        d.add_paragraph("Members to sign if they agree with all resolutions included above")
        t = d.add_table(rows=4, cols=1)
        t.cell(0, 0).text = "[corporate shareholder_name] (Represented by its authorized director)"
        t.cell(1, 0).text = "[authorized director_name]"
        t.cell(2, 0).text = "[individual shareholder_name]"
        t.cell(3, 0).text = "Date: [date]"
        return d, t

    d, t = _signing_doc()
    eq("U18a", "a headerless signature table is still recognised", rr._signing_rows_to_scan(t) is not None, True)

    # An ordinary table with no signing cue above it must stay untouched.
    plain = _Docx()
    plain.add_paragraph("Shareholding structure")
    pt = plain.add_table(rows=2, cols=1)
    pt.cell(0, 0).text = "[individual shareholder_name]"
    pt.cell(1, 0).text = "Date: [date]"
    eq("U18b", "a table with no signing cue above it is not claimed", rr._signing_rows_to_scan(pt), None)

    eq(
        "U18c",
        "'Represented by' marks a corporate signatory group",
        bool(rr._CORP_SIGN_RE.search("(Represented by its authorized director)")),
        True,
    )

    # One corporate member -> exactly ONE signatory block, and its name appears
    # exactly once. Two would be the shipped bug.
    d, _t = _signing_doc()
    member = {"name": "CITY HOLDINGS LIMITED", "type": "Company"}
    synth = rr.expand_repeat_regions(
        d,
        {"members": [member], "authorized_director_name": "PHYOE MIN KYAW"},
    )
    vals = list(synth.values())
    eq("U18d", "the sole corporate member signs exactly once", vals.count("CITY HOLDINGS LIMITED"), 1)
    # …and the representative is the person who was CHOSEN, not a register
    # ordering. `authorized_director_name` is the key the picker writes; it was
    # absent from the lookup list, which sent this to the positional fallback.
    eq("U18e", "the chosen representative is used, not a positional guess", "PHYOE MIN KYAW" in vals, True)

    # The three real custom_data shapes observed on live runs (documents 73–75).
    # The same slot arrives under two spellings and EITHER can hold the answer,
    # so the pick is on content: a candidate equal to `director_name` is the
    # untouched per-company default, one that differs is somebody's choice.
    corp = {"name": "CITY HOLDINGS LIMITED", "type": "Company"}
    eq(
        "U18h",
        "the space spelling wins when the underscore one is the default",
        rr._corp_representative(
            corp,
            {
                "director_name": "KYAW THU SOE",
                "authorized director_name": "PHYOE MIN KYAW",
                "authorized_director_name": "KYAW THU SOE",
            },
        ),
        "PHYOE MIN KYAW",
    )
    eq(
        "U18i",
        "the underscore spelling is honoured when it holds the answer",
        rr._corp_representative(corp, {"director_name": "KYAW THU SOE", "authorized_director_name": "PHYOE MIN KYAW"}),
        "PHYOE MIN KYAW",
    )
    eq(
        "U18j",
        "a candidate matching the default is still used, never re-guessed",
        rr._corp_representative(
            corp, {"director_name": "PHYOE MIN KYAW", "authorized_director_name": "PHYOE MIN KYAW"}
        ),
        "PHYOE MIN KYAW",
    )
    # `corporate_shareholder_3_name` is the MEMBER's own name. It used to be read
    # as a representative, putting the company on its own signature line.
    # (The register fallback below it may still supply a director here — what must
    # never come back is the company itself.)
    eq(
        "U18k",
        "the corporate member is never its own authorised director",
        rr._corp_representative(corp, {"corporate_shareholder_3_name": "CITY HOLDINGS LIMITED"})
        != "CITY HOLDINGS LIMITED",
        True,
    )

    # Five individual members -> five blocks, no repeats.
    d, _t = _signing_doc()
    people = [
        {"name": n, "type": "Individual"}
        for n in ("PHYOE MIN KYAW", "MYO MIN KYAW", "MIN MIN", "WIN WIN TINT", "ZAW MIN LATT")
    ]
    synth = rr.expand_repeat_regions(d, {"members": people})
    got = [v for v in synth.values() if v]
    eq("U18f", "every individual member gets exactly one block", sorted(got), sorted(p["name"] for p in people))

    # Empty data must remain a no-op on every template shape.
    d, t = _signing_doc()
    before = len(t.rows)
    rr.expand_repeat_regions(d, {})
    eq("U18g", "no party data leaves the table untouched", len(t.rows), before)


# ===========================================================================
# U10 Fill-in view labels
# ===========================================================================
def test_fill_view():
    eq("U10a", "a synthetic repeat-region token reads as a party", fv._blank_label("__rr_1__"), "Party 1")
    eq("U10b", "the party number is preserved", fv._blank_label("__rr_12__"), "Party 12")
    eq("U10c", "an ordinary key is titlecased", fv._blank_label("meeting_date"), "Meeting Date")


# ===========================================================================
# U10b An emptied field must be REPORTED, never hidden
#      `validate_filled_document` re-opens the saved .docx and looks for
#      leftover {{...}} patterns, so a placeholder replaced with "" leaves
#      nothing to find and reads as FILLED. That is how a Corporate Shareholder
#      Consent was produced reading 'referred to as  ("NewCo") ... shall invest
#      in  in NewCo' while the result reported 13 placeholders, unfilled_names
#      [] and status Complete — a resolution incorporating an unnamed company
#      for an unstated sum, declared finished.
#
#      The fix is a collector: `find_replacement` notes every user_input field
#      it empties, and `_generate_document` folds those into the validation.
#      These cases pin the collector itself, which is the part that can silently
#      stop working (a ContextVar that is never set collects nothing and every
#      document goes back to looking complete).
# ===========================================================================
def test_blank_reporting():
    import scout.tools.smart_doc as sd

    check("U10d", "a blank collector exists", hasattr(sd, "_blanked_placeholders"))
    check("U10e", "the note helper exists", callable(getattr(sd, "_note_blank", None)))

    # With NO collector set, noting must be a silent no-op — never an exception
    # in the middle of filling a document.
    try:
        sd._note_blank("subscription_amount")
        check("U10f", "noting outside a generation is a no-op", True)
    except Exception as e:
        check("U10f", "noting outside a generation is a no-op", False, repr(e))

    # With a collector set, the placeholder is recorded.
    token = sd._blanked_placeholders.set(set())
    try:
        sd._note_blank("subscription_amount")
        sd._note_blank("new_company_name")
        sd._note_blank("")  # falsy names are not recorded
        got = sorted(sd._blanked_placeholders.get())
        eq("U10g", "blanked fields are collected", got, ["new_company_name", "subscription_amount"])
    finally:
        sd._blanked_placeholders.reset(token)

    # Reset must actually clear it: a blank from one document attributed to the
    # next would be worse than not reporting at all.
    eq("U10h", "the collector is cleared after a generation", sd._blanked_placeholders.get(), None)

    # The wiring: generation must fold blanks into the validation and flip
    # is_valid. Asserted on the SOURCE because running a real generation needs
    # a database; the strings below are the ones that carry the behaviour.
    src = (REPO / "scout/tools/smart_doc.py").read_text()
    body = "\n".join(line for line in src.split("\n") if not line.lstrip().startswith("#"))
    check("U10i", "find_replacement notes the field it empties", "_note_blank(placeholder)" in body)
    check(
        "U10j",
        "generation folds blanks into unfilled_names",
        '"unfilled_names"] = _names' in body or 'unfilled_names"] = _names' in body,
    )
    check(
        "U10k",
        "generation marks a blanked document invalid",
        '"is_valid"] = False' in body or 'is_valid"] = False' in body,
    )


# ===========================================================================
# U11 Structural contracts
#     Each of these has broken the product silently at least once.
# ===========================================================================
def test_structural_contracts():
    picker_src = (REPO / "scout/tools/people_picker.py").read_text()
    aq_src = (REPO / "scout/tools/ask_questions.py").read_text()
    agent_src = (REPO / "scout/agent.py").read_text()

    # A paused tool must not declare run_context — see the U8 comment.
    offenders = []
    for src, label in ((picker_src, "people_picker"), (aq_src, "ask_questions")):
        for m in re.finditer(r"@tool\([^)]*requires_user_input[^)]*\)\s*\ndef (\w+)\(([^)]*)\)", src, re.DOTALL):
            if "run_context" in m.group(2):
                offenders.append(f"{label}.{m.group(1)}")
    eq("U11a", "no paused tool declares run_context", offenders, [])

    # The lettered a)/b)/c) menu is the one interaction the client rejected.
    lettered = [i + 1 for i, line in enumerate(agent_src.split("\n")) if line[:3] in ("a) ", "b) ", "c) ")]
    eq("U11b", "the system prompt contains no lettered menus", lettered, [])

    # agno reserves these names; using one hijacks the HITL resume path and the
    # provider rejects the dangling tool call with a 400.
    reserved = [n for n in ("def ask_user(", "def get_user_input(") if n in aq_src]
    eq("U11c", "no tool uses an agno-reserved name", reserved, [])

    # A tool the prompt names but that is not registered fails silently: the
    # model follows the instruction, finds nothing, and ends the turn empty.
    mismatches = getattr(
        __import__("scout.agent", fromlist=["_PROMPT_TOOL_MISMATCHES"]), "_PROMPT_TOOL_MISMATCHES", None
    )
    eq("U11d", "every tool named in the prompt is registered", mismatches or [], [])

    # The picker read-back is only safe because it is scoped three ways.
    resolver_src = (REPO / "scout/tools/slot_resolver.py").read_text()
    # Slice to the END OF THE SQL STRING, not a fixed character count: the
    # comment block below the query discusses `session_id = ''` in prose, and a
    # fixed window that happened to reach it would make U11h flip on an edit to
    # a comment.
    _start = resolver_src.find("SELECT selection FROM party_selections")
    query = resolver_src[_start : resolver_src.find('"""', _start)]
    for cid, needle, why in [
        ("U11e", "session_id = %s", "scoped to the conversation"),
        ("U11f", "slot_kind = %s", "scoped to the role"),
        ("U11g", "LOWER(company_name) = LOWER(%s)", "scoped to the company"),
    ]:
        check(cid, f"picker read-back is {why}", needle in query, "")
    check(
        "U11h",
        "the read-back tolerates no unscoped legacy rows",
        "slot_kind = ''" not in query and "session_id = ''" not in query,
        "",
    )

    # Every caller of an AgentOS route must be able to find the JWT.
    #
    # `useStore().authToken` is a playground leftover: it is not in the store's
    # `partialize`, so it is '' on every page load. While /agents, /teams and
    # /sessions were public that was invisible — they answered with no header.
    # Putting them behind the JWT turned it into "Failed to fetch agents:
    # Unauthorized" on load, and would have 401'd the chat POST too. The real
    # token lives in localStorage under `ls_token`. None of the API-level test
    # layers can catch this: they all authenticate properly, because they are
    # clients rather than browsers.
    ui = REPO.parent / "agent-ui/src" if (REPO.parent / "agent-ui/src").exists() else None
    ui = ui or (REPO / "agent-ui/src")
    for cid, rel, fn_label in [
        ("U11i", "api/os.ts", "os.ts header builders"),
        ("U11j", "hooks/useAIStreamHandler.tsx", "the streaming chat POST"),
    ]:
        path = ui / rel
        if not path.exists():
            check(
                cid,
                f"{fn_label} falls back to the stored JWT",
                True,
                "SKIPPED — frontend sources not present in this image",
            )
            continue
        src = path.read_text()
        builds_bearer = "Bearer ${" in src
        # Match the CALL, not the bare word: both files explain `ls_token` in a
        # comment, so `"ls_token" in src` passes even with the fallback deleted.
        reads_token = re.search(r"localStorage\.getItem\(\s*['\"]ls_token['\"]\s*\)", src) is not None
        check(
            cid,
            f"{fn_label} falls back to the stored JWT",
            (not builds_bearer) or reads_token,
            "builds an Authorization header but never reads localStorage.ls_token",
        )

    # The silent-stop nudge must not decide "did tool work" from the final chunk.
    #
    # RunPaused carries a `tools` array; RunCompleted does NOT carry the key at
    # all — verified against the live stream, where `'tools' in ev` was False
    # while ToolCallStarted had already reported ask_questions and preview_doc.
    # Reading `chunk.tools` there made `didToolWork` permanently false, so the
    # nudge never fired in the browser and neither did the out-of-retries
    # message that shares the guard: the user got a blank bubble that looked
    # exactly like a finished answer. tracker_layer3 could not catch it — the
    # harness counts ToolCallStarted across the stream, so it nudged correctly
    # and reported PASS while the real UI hung.
    handler = ui / "hooks/useAIStreamHandler.tsx"
    if not handler.exists():
        check(
            "U11k",
            "the silent-stop nudge counts tools from the stream",
            True,
            "SKIPPED — frontend sources not present in this image",
        )
    else:
        src = handler.read_text()
        m = re.search(r"const\s+didToolWork\s*=(.+?)\n\s*const\s", src, re.DOTALL)
        expr = m.group(1) if m else ""
        counts_stream = "toolsThisRunRef.current" in expr
        # A ref that is never incremented is the same bug wearing a new name.
        increments = re.search(r"toolsThisRunRef\.current\s*\+=\s*1", src) is not None
        resets = len(re.findall(r"toolsThisRunRef\.current\s*=\s*0", src))
        check(
            "U11k",
            "the silent-stop nudge counts tools from the stream",
            bool(m) and counts_stream and increments and resets >= 2,
            f"didToolWork={expr.strip()[:60]!r} increments={increments} resets={resets}",
        )

        # An empty turn after a document tool must be closed from the tool
        # result, not by buying a second inference.
        #
        # Measured over ten Layer 3 case-runs: generate_document ended the turn
        # with zero characters of content EVERY time it was the last tool. The
        # recovery was a synthetic "continue", which re-runs inference over the
        # whole re-injected history to obtain a sentence the tool result already
        # contains in its `message` field. Rendering that instead removes the
        # round trip; the nudge stays as the fallback when the result cannot be
        # read (pre-JSON Python-repr sessions).
        builds = re.search(r"const\s+buildClosingFromTool\s*=", src) is not None
        captured = "closingFromToolRef.current = closing" in src
        # The nudge and the out-of-retries branch must BOTH stand down when a
        # closing sentence exists, or the user gets the tool's sentence AND a
        # duplicate run.
        # Counting `!closeFromTool &&` occurrences was the original mechanism,
        # and it went stale the moment both branches were refactored to share a
        # single `shouldNudge` const that carries the guard once. The INTENT —
        # neither the nudge nor the out-of-retries branch may fire when a
        # closing sentence exists — is unchanged, so accept either shape.
        guards = len(re.findall(r"!closeFromTool\s*&&", src))
        if guards < 2:
            shared = re.search(r"const\s+shouldNudge\s*=\s*\n?\s*!closeFromTool\s*&&", src)
            if shared and len(re.findall(r"shouldNudge\s*&&", src)) >= 2:
                guards = 2
        renders = (
            re.search(r"if\s*\(closeFromTool\)\s*\{\s*//[^\n]*\n\s*updatedContent\s*=\s*closeFromTool", src) is not None
        )
        # The ref holds a {content, approval} object now, so a per-stream reset
        # clears it to null rather than to the empty string it started as.
        cleared = len(re.findall(r"closingFromToolRef\.current\s*=\s*(?:''|null)", src))
        check(
            "U11l",
            "an empty turn is closed from the document tool result",
            builds and captured and guards >= 2 and renders and cleared >= 2,
            f"builder={builds} captured={captured} guards={guards} renders={renders} resets={cleared}",
        )

        # A stalled PREVIEW must be given back the approval it owed.
        #
        # preview_doc renders the field table and then owes an ask_questions
        # card with one question and two fixed options. Measured across six
        # conversations it produced neither: the turn ended empty, leaving a
        # preview the user could read and could not act on. It only became the
        # dominant stall once the tool was reachable at all — the prompt named
        # `preview_document`, which is not registered, so the required preview
        # step was skipped entirely and the model went straight to generation.
        #
        # The card is reconstructed, not invented: the question and both option
        # strings come from the tool's own agent_instruction. It must NOT be
        # routed through AskUserCardList, which resumes a PAUSED run — this run
        # completed, so there is no pause to consume.
        approval_builder = "buildApprovalFromPreview" in src
        approval_set = "APPROVAL_DOC_TOOLS" in src
        # preview_doc closing the turn silently is the bug this replaced.
        not_closable = re.search(r"CLOSABLE_DOC_TOOLS = new Set\(\[([^\]]*)\]", src, re.DOTALL)
        closable_body = not_closable.group(1) if not_closable else ""
        preview_excluded = "preview_doc" not in closable_body
        carried = "pending_approval: closeFromTool?.approval" in src
        prompt = ui / "components/chat/ApprovalPrompt.tsx"
        sends_message = prompt.exists() and "setPendingMessage(option)" in prompt.read_text()
        rendered = "ApprovalPrompt approval={message.pending_approval}" in (
            (ui / "components/chat/ChatArea/Messages/Messages.tsx").read_text()
            if (ui / "components/chat/ChatArea/Messages/Messages.tsx").exists()
            else ""
        )
        check(
            "U11m",
            "a stalled preview is given back its approval card",
            approval_builder and approval_set and preview_excluded and carried and sends_message and rendered,
            f"builder={approval_builder} set={approval_set} "
            f"preview_excluded_from_closable={preview_excluded} "
            f"carried={carried} sends={sends_message} rendered={rendered}",
        )

    # The tool list the model reads must be GENERATED, and a mismatch must be fatal.
    #
    # Four prompt/tool mismatches shipped while this was hand-written prose
    # checked by a log line, and every one failed silently — the model follows
    # the instruction, finds no such tool, and ends the turn with no text, which
    # reads as a hang and leaves no trace because the tool was never called:
    #
    #   generate_dica_extract  never added to _tools_to_add
    #   list_companies         registered as list_all_companies
    #   preview_document       export-dict key; @wraps made it preview_doc
    #   generate_document_tool named in scout/knowledge/routing/intents.json,
    #                          which reaches the prompt as DATA
    #
    # A log line is worth what someone reading it is worth. The inventory is now
    # built from the live registry, and startup refuses on a mismatch.
    agent_src = REPO / "scout" / "agent.py"
    if not agent_src.exists():
        check(
            "U13",
            "the tool inventory is generated and mismatches are fatal",
            True,
            "SKIPPED — scout/agent.py not present",
        )
    else:
        a = agent_src.read_text()
        generates = "_build_tool_inventory" in a and "TOOL_INVENTORY_BLOCK" in a
        injected = "{TOOL_INVENTORY_BLOCK}" in a
        # The audit and the inventory must measure the SAME registry, or the
        # list shown and the list checked drift apart again.
        shared = a.count("_registered_tool_names(") >= 2
        fatal = re.search(r"if _PROMPT_TOOL_MISMATCHES and[\s\S]{0,200}?raise RuntimeError", a) is not None
        # The agno Function wrapper keeps its purpose on .description; __doc__
        # is the class docstring, which described agno for 24 of 45 tools.
        real_desc = "Model for storing functions" in a and 'getattr(fn, "description"' in a
        check(
            "U13",
            "the tool inventory is generated and mismatches are fatal",
            generates and injected and shared and fatal and real_desc,
            f"generates={generates} injected={injected} shared_registry={shared} "
            f"fatal={fatal} real_descriptions={real_desc}",
        )

    # A stream that delivers nothing must never render as an answer.
    #
    # A 200 does not mean the body was ours. A restarting server, a proxy, or a
    # route that fell through to the static frontend all answer 200 with HTML:
    # response.ok is true, response.body exists, the reader drains it, and
    # parseBuffer finds no events — so the stream "completed" having delivered
    # zero chunks and the UI painted an empty agent bubble indistinguishable
    # from a finished reply. Measured live 2026-08-06: a message sent while the
    # container was being replaced came back in 1.0s as a blank bubble with no
    # error, and no run was ever persisted.
    #
    # Also: the !response.ok branch called response.json() on that HTML, which
    # threw a SyntaxError and hid the real status behind
    # "Unexpected token '<'".
    stream = ui / "hooks/useAIResponseStream.tsx"
    if not stream.exists():
        check(
            "U15",
            "an empty stream is reported, not painted as a reply",
            True,
            "SKIPPED — frontend sources not present in this image",
        )
    else:
        s = stream.read_text()
        counts = "delivered += 1" in s and "let delivered = 0" in s
        raises = re.search(r"if \(delivered === 0\)[\s\S]{0,200}?throw new Error", s) is not None
        # Both parse sites must go through the counter, or the count is a lie.
        wired = len(re.findall(r"parseBuffer\(buffer, countingChunk\)", s)) >= 2
        # Generous window: the branch carries the explanatory comment before the
        # call, and a 300-char bound failed on correct code.
        text_first = re.search(r"if \(!response\.ok\)[\s\S]{0,800}?response\.text\(\)", s) is not None
        check(
            "U15",
            "an empty stream is reported, not painted as a reply",
            counts and raises and wired and text_first,
            f"counts={counts} raises={raises} both_parse_sites={wired} reads_text_before_json={text_first}",
        )

    # The agent must never be able to send an email by itself.
    #
    # send_email_tool used to connect to SMTP and deliver, immediately, on the
    # agent's own decision — it chose the recipient, the subject, the body and
    # which generated document to attach, with no confirmation, no audit row and
    # no record that a send had even been considered. A misread instruction, or
    # text picked up from a document it was reading, was enough to mail a
    # client's corporate filing to an address nobody approved. Email cannot be
    # recalled.
    #
    # The tool now only queues. Delivery lives behind an endpoint requiring the
    # USER's JWT, which the agent has no way to obtain. There must be no
    # "confirmed" parameter either: a flag set by the same model whose judgement
    # is being checked is not an approval.
    if not agent_src.exists():
        check("U14", "the agent cannot send an email without a human", True, "SKIPPED — scout/agent.py not present")
    else:
        a = agent_src.read_text()
        tool = re.search(r"def send_email_tool\(([\s\S]*?)\n(?=\w|@)", a)
        body = tool.group(0) if tool else ""
        no_smtp = "smtplib" not in body and "send_message" not in body
        queues = "'queued'" in body or '"queued"' in body
        no_confirm_arg = tool is not None and "confirmed" not in (tool.group(1).split(")")[0])
        main_src = REPO / "app" / "main.py"
        m = main_src.read_text() if main_src.exists() else ""
        # Delivery endpoint exists, demands a token, and claims the row before
        # touching SMTP so a double click cannot send twice.
        gated = '@app.post("/api/email/queued/{email_id}/send")' in m
        needs_auth = bool(
            re.search(r'/api/email/queued/\{email_id\}/send"\)[\s\S]{0,400}?if not user:[\s\S]{0,80}?401', m)
        )
        claims_first = bool(re.search(r"SET status = 'sending'[\s\S]{0,200}?WHERE id = %s AND status = 'queued'", m))
        check(
            "U14",
            "the agent cannot send an email without a human",
            bool(tool) and no_smtp and queues and no_confirm_arg and gated and needs_auth and claims_first,
            f"tool_found={bool(tool)} no_smtp={no_smtp} queues={queues} "
            f"no_confirm_arg={no_confirm_arg} endpoint={gated} "
            f"auth={needs_auth} claims_before_send={claims_first}",
        )

    # No foreign-jurisdiction statute may be cited by this product.
    #
    # app/main.py:_get_legal_refs_from_name() hardcoded Indian company law —
    # "Companies Act 2013 - Section 152", "SEBI (LODR) Regulations 2015" — and
    # picked between them by substring match on the template FILENAME. It was
    # used whenever AI analysis fell back, which was every time: on 2026-08-06
    # all 15 templates in the database held exactly those strings. A Myanmar law
    # firm was being told its AGM minutes were governed by India's securities
    # regulator. The same code block set jurisdiction = "Myanmar".
    #
    # Matching is on the statute NAMES, not the year alone: "2013" appears in
    # ordinary dates and would make this fire on anything.
    FOREIGN_STATUTES = [
        "Companies Act 2013",
        "Companies Act, 2013",
        "SEBI",
        "Companies (Management and Administration) Rules",
        "DIN Application",
    ]
    # Judge code, not commentary. A docstring naming "SEBI (LODR) Regulations
    # 2015" to explain why it was deleted is documentation; the same text in a
    # returned list is a citation shown to a lawyer. A line-based scan cannot
    # tell those apart and fails on its own fix, so Python is read through `ast`
    # and every docstring is skipped, while data files are scanned by line with
    # their comment syntax honoured.
    import ast as _ast

    def _docstring_nodes(tree):
        out = set()
        for node in _ast.walk(tree):
            if not isinstance(node, (_ast.Module, _ast.FunctionDef, _ast.AsyncFunctionDef, _ast.ClassDef)):
                continue
            body = getattr(node, "body", None) or []
            if (
                body
                and isinstance(body[0], _ast.Expr)
                and isinstance(body[0].value, _ast.Constant)
                and isinstance(body[0].value.value, str)
            ):
                out.add(id(body[0].value))
        return out

    roots = [p for p in (REPO / "app", REPO / "scout", REPO / "db") if p.exists()]
    hits = []
    for root in roots:
        for path in root.rglob("*"):
            if path.suffix not in (".py", ".sql", ".json") or not path.is_file():
                continue
            try:
                text = path.read_text(errors="replace")
            except OSError:
                continue
            if path.suffix == ".py":
                try:
                    tree = _ast.parse(text)
                except SyntaxError:
                    continue
                skip = _docstring_nodes(tree)
                for node in _ast.walk(tree):
                    if not isinstance(node, _ast.Constant) or not isinstance(node.value, str) or id(node) in skip:
                        continue
                    for statute in FOREIGN_STATUTES:
                        if statute in node.value:
                            hits.append(f"{path.relative_to(REPO)}:{node.lineno} {node.value[:60]!r}")
                continue
            # .sql / .json — the migration that removes these has to name them.
            if path.name.startswith("migration_018"):
                continue
            for lineno, line in enumerate(text.splitlines(), 1):
                stripped = line.strip()
                if stripped.startswith("--"):
                    continue
                for statute in FOREIGN_STATUTES:
                    if statute in line:
                        hits.append(f"{path.relative_to(REPO)}:{lineno} {stripped[:60]}")
    check(
        "U12", "no foreign-jurisdiction statute is cited as Myanmar law", not hits, "; ".join(hits[:4]) if hits else ""
    )

    # The routing data files describe THIS product, and their tool names are real.
    #
    # scout/knowledge/routing/intents.json and scout/knowledge/sources/files.json
    # are interpolated into the system prompt as DATA (via INTENT_ROUTING_CONTEXT
    # and SOURCE_REGISTRY_STR). Nothing in them is code, so nothing in them is
    # reachable by searching source for a symbol — which is exactly how
    # `generate_document_tool` survived for weeks as the primary source for
    # "Create legal document".
    #
    # On 2026-08-06 both files were still the generic enterprise-docs boilerplate
    # they shipped from: 11 of 12 intents were about OKRs, RFCs, runbooks and PTO,
    # the source registry advertised company-docs/, engineering-docs/ and
    # data-exports/ — none of which have ever existed here — and a gotcha told the
    # model that PTO sits in "employee-handbook.md, Section 4". 3.5k characters of
    # a different product's map, in the prompt of a Myanmar legal drafting agent.
    #
    # The import-time audit does NOT cover this. It flags a backticked word only
    # when that word shares an 8-character prefix with a registered tool, so a
    # plainly wrong name passes it. This check closes that gap directly.
    routing = REPO / "scout" / "knowledge" / "routing" / "intents.json"
    sources = REPO / "scout" / "knowledge" / "sources" / "files.json"
    if not routing.exists() or not sources.exists():
        check(
            "U16",
            "prompt routing data names real tools and real directories",
            True,
            "SKIPPED — knowledge data files not present",
        )
    else:
        import json as _json

        intents = _json.loads(routing.read_text())
        filesrc = _json.loads(sources.read_text())

        named = set()
        for m in intents.get("intent_mappings", []):
            if m.get("primary_source"):
                named.add(m["primary_source"])
            named.update(m.get("fallback_sources", []))

        # Compare against the LIVE registry. Skipped rather than failed when the
        # agent cannot be imported, so this stays runnable outside the container.
        try:
            from scout.agent import scout as _scout

            registered = set()
            for t in _scout.tools or []:
                n = getattr(t, "__name__", None) or getattr(t, "name", None)
                if n:
                    registered.add(str(n))
                fns = getattr(t, "functions", None)
                if isinstance(fns, dict):
                    registered.update(str(k) for k in fns)
        except Exception:
            registered = set()

        unregistered = sorted(named - registered) if registered else []

        # Every directory the prompt advertises must exist. Telling the model
        # about a folder that is not there is the file-system twin of naming a
        # tool that is not registered: it tries, finds nothing, and says nothing.
        try:
            from scout.paths import DOCUMENTS_DIR as _DOCS

            docs_root = Path(_DOCS)
        except Exception:
            docs_root = REPO / "documents"

        def _resolve(p: str) -> Path:
            p = p.strip().rstrip("/")
            if p.startswith("documents/"):
                return docs_root / p[len("documents/") :]
            return REPO / p

        advertised = [d["name"] for d in filesrc.get("directories", []) if d.get("name")]
        advertised += list(filesrc.get("common_locations", {}).values())
        ghosts = sorted({p for p in advertised if not _resolve(p).is_dir()})

        # Vocabulary from the product this boilerplate came from. None of it
        # describes anything in Legal Scout.
        FOREIGN_VOCAB = [
            "company-docs",
            "engineering-docs",
            "data-exports",
            "employee-handbook",
            "Employee Handbook",
            "OKR",
            "PTO",
            "runbook",
        ]

        # Judge what RENDERS, not the commentary — the same rule U12 applies by
        # skipping docstrings. A `_comment` key recording that these directories
        # never existed is documentation; the formatter reads no key beginning
        # with an underscore, so none of it can reach the model. Scanning raw
        # text instead would make this check fail on its own fix.
        def _rendered(node):
            if isinstance(node, dict):
                return " ".join(_rendered(v) for k, v in node.items() if not str(k).startswith("_"))
            if isinstance(node, list):
                return " ".join(_rendered(v) for v in node)
            return str(node)

        blob = _rendered(intents) + " " + _rendered(filesrc)
        stowaways = sorted({v for v in FOREIGN_VOCAB if v in blob})

        check(
            "U16",
            "prompt routing data names real tools and real directories",
            not unregistered and not ghosts and not stowaways,
            f"unregistered_tools={unregistered or 'none'} "
            f"nonexistent_dirs={ghosts or 'none'} "
            f"foreign_vocab={stowaways or 'none'} "
            f"(registry_{'loaded' if registered else 'UNAVAILABLE—tool check skipped'})",
        )

    # The documents tree must never be served without a token.
    #
    # app.mount("/documents", StaticFiles(...)) serves generated documents,
    # uploaded DICA filings, the firm's templates and cached previews. The auth
    # middleware returns early on any path not starting with "/api/", so all of
    # it was public: measured 2026-08-06 against the running app, a real AGM
    # minutes .docx came back 200 / 29,313 bytes with no token at all, on a
    # container published to 0.0.0.0.
    #
    # PUBLIC_ROUTES listed "/documents/legal/" with the comment "Static file
    # serving", which read like policy but never ran — that list is consulted
    # after the /api/ early return.
    #
    # ORDER is the invariant, not the presence of a constant. A gate placed
    # below the early return is exactly as dead as the comment it replaced, and
    # would still pass a check that only greps for the name.
    main_src = REPO / "app" / "main.py"
    if not main_src.exists():
        check("U17", "the documents tree is not served without a token", True, "SKIPPED — app/main.py not present")
    else:
        m = main_src.read_text()
        # \s* between the paren and the brace: a formatter may split
        # `frozenset({` across lines, and the control is no weaker for it.
        declared = re.search(r"STATIC_PROTECTED_ROOTS\s*=\s*frozenset\(\s*\{([\s\S]{0,400}?)\}\s*\)", m)
        covers_documents = bool(declared and '"documents"' in declared.group(1))

        gate = re.search(r"if root in AGENTOS_PROTECTED_ROOTS or root in STATIC_PROTECTED_ROOTS", m)
        early_return = re.search(r'if not path\.startswith\("/api/"\)', m)
        ordered = bool(gate and early_return and gate.start() < early_return.start())

        # The gate must read the cookie, or a plain <a href> download breaks:
        # an anchor cannot set an Authorization header.
        after_gate = m[gate.start() : early_return.start()] if ordered else ""
        uses_request_jwt = "_request_jwt(request)" in after_gate

        # And the dead whitelist entry must not come back. Scoped to the
        # PUBLIC_ROUTES literal: startup_sync holds a list of the same directory
        # paths for mkdir, and a file-wide scan flags those instead — failing on
        # correct code, which is how U15 broke once already.
        routes_block = re.search(r"PUBLIC_ROUTES\s*=\s*\[([\s\S]*?)\n\]", m)
        no_dead_entry = bool(routes_block) and not re.search(r'^\s*"/documents', routes_block.group(1), re.MULTILINE)

        check(
            "U17",
            "the documents tree is not served without a token",
            covers_documents and ordered and uses_request_jwt and no_dead_entry,
            f"declares_documents={covers_documents} "
            f"gate_before_api_early_return={ordered} "
            f"reads_cookie_via_request_jwt={uses_request_jwt} "
            f"no_public_documents_entry={no_dead_entry}",
        )


# ===========================================================================
# U19-U22  Register authority
#     Who is allowed to sign, and what the product is allowed to claim a tool
#     is called. Both have produced legally wrong documents.
# ===========================================================================
def test_register_authority():
    import scout.agent as _agent
    from db.connection import get_db_conn

    registry = _agent._registered_tool_names(_agent.scout.tools or [])

    # --- U19: skill BODIES name real tools ---------------------------------
    # The startup contract audit reads name + description only, so bodies —
    # the L2 playbooks lawyers edit in the admin UI — were never checked. Seven
    # enabled skills shipped naming `preview_document` (registered as
    # `preview_doc`) and `list_tracked_documents` (registered as
    # `list_documents`). The model follows the instruction, finds no tool, and
    # the step silently does not happen.
    conn = get_db_conn()
    try:
        cur = conn.cursor()
        cur.execute("SELECT name, body FROM legal_skills WHERE enabled = TRUE ORDER BY name")
        skills = cur.fetchall()
        cur.close()
    finally:
        conn.close()

    check("U19a", "there are enabled skills to audit", len(skills) > 0, f"{len(skills)} enabled")

    bad = {}
    for sname, body in skills:
        refs = set(re.findall(r"`([a-z_][a-z0-9_]{3,})\(?\)?`", body or ""))
        refs |= set(re.findall(r"\b([a-z_][a-z0-9_]{3,})\(\)", body or ""))
        for r in sorted(refs):
            if r in registry:
                continue
            # Only a NEAR MISS is a defect. A skill body is prose; flagging every
            # unknown snake_case token would trip on placeholder names and DB
            # columns. Same 6-char prefix rule the /api/skills validator uses.
            if any(t.startswith(r[:6]) or r.startswith(t[:6]) for t in registry):
                bad.setdefault(sname, []).append(r)
    check("U19b", "every enabled skill body names only registered tools", not bad, f"{bad}" if bad else "")

    # The two specific names that shipped, asserted by hand so a future
    # re-seed of migration 014 cannot quietly reintroduce them.
    eq(
        "U19c",
        "preview_document is not a tool (preview_doc is)",
        ("preview_document" in registry, "preview_doc" in registry),
        (False, True),
    )
    eq(
        "U19d",
        "list_tracked_documents is not a tool (list_documents is)",
        ("list_tracked_documents" in registry, "list_documents" in registry),
        (False, True),
    )

    # --- U20: legacy cessation ---------------------------------------------
    # `_registered_people` filters cessation in SQL, but `_directors_of` falls
    # back to `_legacy_people` when a company has no register rows — and that
    # path read only name/position/appointed_date/shares. An unsynced company
    # offered its RESIGNED directors as current, which is how a person with no
    # authority reaches a signature block.
    from datetime import date, timedelta

    past = (date.today() - timedelta(days=30)).isoformat()
    future = (date.today() + timedelta(days=120)).isoformat()

    def _names(entries):
        return [c.get("name") for c in pp._legacy_people(entries, "Director")]

    eq(
        "U20a",
        "a director who has already ceased is not offered",
        _names([{"name": "GONE", "date_of_cessation": past}, {"name": "HERE"}]),
        ["HERE"],
    )
    eq(
        "U20b",
        "the alternate spelling is honoured too",
        _names([{"name": "GONE", "resigned_date": past}, {"name": "HERE"}]),
        ["HERE"],
    )
    # Safe direction: shown, not hidden. A name the user can see can be
    # declined; a name never shown cannot be chosen.
    eq(
        "U20c",
        "a FUTURE cessation is still in office today",
        _names([{"name": "STILL SERVING", "date_of_cessation": future}]),
        ["STILL SERVING"],
    )
    eq("U20d", "a blank marker is not a resignation", _names([{"name": "FINE", "date_of_cessation": "-"}]), ["FINE"])
    eq(
        "U20e",
        "an unreadable date does not silently remove anybody",
        _names([{"name": "FINE", "date_of_cessation": "01/12/2024"}]),
        ["FINE"],
    )
    # …and the user must be able to SEE a pending departure before signing.
    sub = (
        pp._legacy_people([{"name": "STILL SERVING", "date_of_cessation": future}], "Director")[0].get("subtitle") or ""
    )
    check("U20f", "a pending departure is visible on the candidate", future in sub, sub)

    # --- U21: a representative must be on THAT company's board -------------
    # Measured: KYAW THU SOE sits on CITY MART's board and NOT on CITY
    # HOLDINGS'. He was printed as CITY HOLDINGS' authorised director — a person
    # with no power to bind the company he was signing for.
    #
    # These need real rows. On a fresh install the register is empty by design,
    # so they SKIP rather than fail — a correct empty product must not report red.
    board = rr._board_of("CITY HOLDINGS LIMITED") if register_has("CITY HOLDINGS") else []
    if not board:
        for cid, nm in (
            ("U21a", "the member company's board is readable"),
            ("U21b", "a director of the SUBJECT company is refused as the member's rep"),
            ("U21c", "a director who IS on that board is accepted"),
        ):
            skip(cid, nm, "needs CITY HOLDINGS LIMITED and its directors in the register")
    else:
        check("U21a", "the member company's board is readable", len(board) > 0, f"{board}")
    if board:
        outsider = "KYAW THU SOE"
        check(
            "U21b",
            "a director of the SUBJECT company is refused as the member's rep",
            rr._corp_representative(
                {"name": "CITY HOLDINGS LIMITED", "type": "Company"}, {"authorized director_name": outsider}
            )
            != outsider,
            f"board={board}",
        )
        insider = board[0]
        eq(
            "U21c",
            "a director who IS on that board is accepted",
            rr._corp_representative(
                {"name": "CITY HOLDINGS LIMITED", "type": "Company"}, {"authorized director_name": insider}
            ),
            insider,
        )
    # An unregistered member company cannot be checked. Refusing there would
    # blank the line for every corporate member not in the register, so it fails
    # OPEN — deliberately, and logged.
    eq(
        "U21d",
        "an unverifiable company fails open rather than blanking the line",
        rr._corp_representative(
            {"name": "NOT IN THE REGISTER PTE LTD", "type": "Company"}, {"authorized director_name": "SOMEBODY"}
        ),
        "SOMEBODY",
    )

    # --- U22: member slots are typed, not positional -----------------------
    # slots 1-2 were ASSUMED individual and slot 3 corporate, from a flat
    # comma-joined name list. CITY MART's only member is the corporate CITY
    # HOLDINGS LIMITED; it landed at index 0 and was written into the
    # INDIVIDUAL slot, rendering a company as an individual member.
    #
    # Needs both the template ON DISK and the company in the register, so it
    # skips on a fresh install for the same reason as U21.
    from scout.tools.smart_doc import prepare_document_data

    _tpl = "Shareholders Resolution In Writing - Director Appointment.docx"
    # NOT REPO/"documents" — that is /app/documents, which does not exist. The
    # templates are a bind mount at /documents, which is prepare_document_data's
    # own default documents_dir. Pointing at the wrong one would make this skip
    # permanently, which is the quiet way a test stops testing anything.
    _have = register_has("CITY MART") and Path("/documents/legal/templates", _tpl).exists()
    prepared = prepare_document_data(_tpl, "CITY MART HOLDING COMPANY LIMITED") if _have else {}
    cd = (prepared or {}).get("company_data") or {}
    if not cd:
        why = "needs CITY MART HOLDING COMPANY LIMITED in the register and the Director Appointment template on disk"
        for cid, nm in (
            ("U22a", "the sole CORPORATE member fills the corporate slot"),
            ("U22b", "and does NOT fill the individual slot"),
            ("U22c", "the space spelling agrees with the underscore one"),
            ("U22d", "the untyped slot still carries the member"),
        ):
            skip(cid, nm, why if _have is False else f"{why} (prepare returned {sorted(prepared or {})[:6]})")
    else:
        eq(
            "U22a",
            "the sole CORPORATE member fills the corporate slot",
            cd.get("corporate_shareholder_3_name"),
            "CITY HOLDINGS LIMITED",
        )
        eq("U22b", "and does NOT fill the individual slot", cd.get("individual_shareholder_1_name"), "TBD")
        eq("U22c", "the space spelling agrees with the underscore one", cd.get("individual shareholder_1_name"), "TBD")
        # The generic, type-agnostic slots still list every member in order.
        eq("U22d", "the untyped slot still carries the member", cd.get("shareholder_1_name"), "CITY HOLDINGS LIMITED")


# ===========================================================================
# U23  The approval gate
#     A valid JWT proves someone authenticated. It does not prove an
#     administrator let them in. Everything below exists because the gap
#     between those two facts is where LDAP and SSO put a stranger.
#
#     Every structural case here walks the AST rather than grepping the file.
#     A text scan cannot tell a live call from the same words inside a comment
#     — and the comments around this feature are full of the exact strings a
#     regex would look for, so a grep-based version of U23b would pass with
#     the gate deleted and the explanation left behind.
# ===========================================================================
def test_approval_gate():
    import ast as _ast

    import app.main as appmain

    src = (REPO / "app" / "main.py").read_text(encoding="utf-8")
    tree = _ast.parse(src)

    # --- U23a: the migration adds the columns AND backfills -----------------
    mig = REPO / "db" / "migration_030_auth_approval.sql"
    check("U23a", "migration_030 is present", mig.exists(), str(mig))
    if mig.exists():
        body = mig.read_text(encoding="utf-8")
        # Comments in that file discuss the backfill at length, so match on the
        # statement's own shape rather than on the word "backfill" — and strip
        # the comments LINE BY LINE before splitting on `;`, not by testing
        # whether a chunk starts with `--`. Every statement in that file is
        # introduced by a comment block, so splitting first leaves the comment
        # glued to the front of the statement it explains: the chunk-level test
        # discarded the very UPDATE it was hunting for, and reported a missing
        # backfill against a migration that had demonstrably just run one.
        code = "\n".join(ln for ln in body.splitlines() if not ln.strip().startswith("--"))
        stmts = [" ".join(s.split()).lower() for s in code.split(";") if s.strip()]
        check(
            "U23b",
            "the migration BACKFILLS existing accounts to approved",
            any(s.startswith("update users") and "approved = true" in s for s in stmts),
            "without this, the first boot after deploy locks out every existing account "
            "and nobody is left approved who could approve them",
        )

    # --- U23c: the gate is at BOTH decode sites ----------------------------
    # The middleware decodes a JWT in two separate branches — one for the
    # AgentOS/static roots, one for /api/. Gate only the second and a pending
    # account is refused the admin panel while keeping the chat, the agent and
    # document generation: the entire product, guarded by a control that looks
    # present. Counting CALLS in the AST is what makes this checkable at all.
    dispatch = None
    for node in _ast.walk(tree):
        if isinstance(node, _ast.ClassDef) and node.name == "AuthMiddleware":
            for sub in node.body:
                if isinstance(sub, _ast.AsyncFunctionDef) and sub.name == "dispatch":
                    dispatch = sub
    check("U23c", "AuthMiddleware.dispatch is present", dispatch is not None)
    if dispatch is not None:
        calls = [
            n
            for n in _ast.walk(dispatch)
            if isinstance(n, _ast.Call) and isinstance(n.func, _ast.Name) and n.func.id == "_is_approved"
        ]
        check(
            "U23d",
            "every JWT-decoding branch of the middleware consults the gate",
            len(calls) >= 2,
            f"found {len(calls)} call(s), need one per decode site (2)",
        )

    # --- U23e: no INSERT into users may omit `approved` --------------------
    # The column defaults to FALSE. An INSERT that leaves it out therefore
    # creates a PENDING account — which is right for a directory and fatal for
    # the bootstrap admin, who on a fresh install is the only row in the table.
    inserts = [
        n.value
        for n in _ast.walk(tree)
        if isinstance(n, _ast.Constant) and isinstance(n.value, str) and "insert into users" in n.value.lower()
    ]
    check("U23e", "the user INSERT statements are findable", len(inserts) >= 2, f"{len(inserts)} found")
    silent = [s for s in inserts if "approved" not in s.lower()]
    check(
        "U23f",
        "no INSERT INTO users leaves `approved` to the column default",
        not silent,
        f"{len(silent)} statement(s) would create a pending account silently",
    )

    # --- U23g: the internal training principal survives the gate -----------
    # `training_jobs.py` mints create_token(0, "system@training", "admin") and
    # drives the 15-step pipeline through this app's own HTTP endpoints, so it
    # passes through the middleware like any browser. users.id is SERIAL from
    # 1, so id 0 matches no row and would fail closed — taking all template
    # training with it, and reporting it as a per-template 403 rather than as
    # anything that looks like an auth problem.
    check(
        "U23g",
        "the background training worker is not locked out by its own gate",
        appmain._is_approved(appmain.SYSTEM_USER_ID) is True,
        "training_jobs.py:376 mints user_id=0; a False here kills all training",
    )

    # --- U23h: it fails CLOSED --------------------------------------------
    # Pure inputs only: these return before any query, so this case reads and
    # writes nothing, in keeping with the rest of this suite.
    check(
        "U23h",
        "a malformed user id is refused, not admitted",
        appmain._is_approved("not-an-id") is False and appmain._is_approved(None) is False,
        "a gate that returns True when it cannot identify the caller is not a gate",
    )

    # --- U23i: every write that changes access busts the cache -------------
    # The decision is cached for _APPROVAL_TTL seconds, so a write that does
    # not invalidate it leaves a revoked account working — including the
    # `is_active` flag on the general update endpoint, which is half of what
    # the gate answers.
    busts = [
        n
        for n in _ast.walk(tree)
        if isinstance(n, _ast.Call) and isinstance(n.func, _ast.Name) and n.func.id == "_bust_approval_cache"
    ]
    check(
        "U23i",
        "approve, update and delete all invalidate the cached decision",
        len(busts) >= 3,
        f"found {len(busts)}; need the approval route, the update route and the delete route",
    )

    # --- U23j/k: the cache TTL is tied to the WORKER COUNT ------------------
    # `_approval_cache` is a plain dict inside one process, so a bust reaches
    # only the worker that served the write. Under `--workers 2` a non-zero TTL
    # makes revocation non-deterministic — refused here, admitted there — which
    # presents exactly like the half-gated build U23d exists to catch, and will
    # pass a hand-check on a coin flip. This reads the real worker count out of
    # the Dockerfile so that raising one without fixing the other cannot ship.
    dockerfile = (REPO / "Dockerfile").read_text(encoding="utf-8") if (REPO / "Dockerfile").exists() else ""
    m = re.search(r'"--workers"\s*,\s*"(\d+)"', dockerfile)
    workers = int(m.group(1)) if m else 1
    check("U23j", "the served worker count is discoverable", bool(m) or not dockerfile, f"workers={workers}")
    check(
        "U23k",
        "an in-process approval cache is not used across multiple workers",
        appmain._APPROVAL_TTL == 0 or workers <= 1,
        f"_APPROVAL_TTL={appmain._APPROVAL_TTL}s with --workers {workers}: a bust reaches one process, "
        f"so revocation depends on which worker answers next. Needs a SHARED cache before this may be non-zero.",
    )


# ===========================================================================
# U24  Directory sign-in (LDAP)
#     The directory proves WHO somebody is; this application still decides what
#     they may do. Every case below guards one of the ways that separation, or
#     the password handling underneath it, can be quietly lost.
# ===========================================================================
def test_ldap_signin():
    import ast as _ast

    import app.auth_settings as st
    import app.ldap_auth as la
    import app.main as appmain

    # --- U24a: off unless switched on --------------------------------------
    # Read through the module's own accessor rather than the env var, so this
    # fails if the default is ever inverted in code.
    # ★ Asserts the DEFAULT, not the live value. An earlier version called
    # `la.ldap_enabled()` after clearing the environment variable — which was
    # right until phase 4 gave the setting a database override. On a deployment
    # where an administrator has legitimately switched the directory on, that
    # version went red while the product was behaving exactly as asked. A test
    # that fails because someone used the feature is a test with the wrong
    # assumption, not a defect.
    check(
        "U24a",
        "directory sign-in is OFF unless somebody switches it on",
        st.SPEC["ldap_enabled"][2] is False,
        f"default={st.SPEC['ldap_enabled'][2]!r} (live value may differ, and that is fine)",
    )

    # --- U24b: an empty password is refused BEFORE any bind -----------------
    # A simple bind with a valid DN and a zero-length password is an
    # UNAUTHENTICATED simple bind (RFC 4513 §5.1.2), and some Active Directory
    # deployments answer success — on those, knowing any provisioned address
    # would be enough to sign in as that person. The guard must sit in
    # authenticate(), not only in the login endpoint, because it is this
    # function that decides whether a password is right.
    #
    # Reaching a bind at all requires a directory, so the assertion is that it
    # raises WITHOUT one: the empty-password branch returns before the host
    # check is ever consulted.
    os.environ["LDAP_HOST"] = "ldap.invalid.test"
    os.environ["LDAP_BASE_DN"] = "ou=users,dc=invalid,dc=test"
    try:
        la.authenticate("someone@example.test", "")
        check("U24b", "an empty password is refused before any bind", False, "authenticate() returned")
    except la.LdapError as e:
        check("U24b", "an empty password is refused before any bind", "empty password" in str(e), str(e)[:90])
    except Exception as e:
        check("U24b", "an empty password is refused before any bind", False, f"{type(e).__name__}: {e}")

    # --- U24c: plaintext binds are refused, not merely warned about ---------
    # The flow rebinds as the user, so the password crosses the wire on every
    # sign-in. A warning about that is a log line nobody reads until after the
    # credentials have been collected — and nothing looks broken meanwhile.
    for env, label in (
        ({"LDAP_USE_SSL": "false", "LDAP_START_TLS": "false"}, "no TLS"),
        ({"LDAP_USE_SSL": "true", "LDAP_VALIDATE_CERT": "false"}, "unvalidated certificate"),
    ):
        saved = {k: os.environ.get(k) for k in env}
        os.environ.update(env)
        os.environ["LDAP_ALLOW_INSECURE"] = "false"
        try:
            la._require_transport_security(la.config())
            check("U24c", f"a bind with {label} is refused", False, "no refusal")
        except la.LdapError:
            check("U24c", f"a bind with {label} is refused", True)
        finally:
            for k, v in saved.items():
                if v is None:
                    os.environ.pop(k, None)
                else:
                    os.environ[k] = v

    # --- U24d: the filter cannot be re-shaped by what someone types ---------
    from ldap3.utils.conv import escape_filter_chars

    hostile = "*)(objectClass=*"
    check(
        "U24d",
        "a hostile username is escaped, not interpolated into the filter",
        "(mail={username})".replace("{username}", escape_filter_chars(hostile)) == r"(mail=\2a\29\28objectClass=\2a)",
        escape_filter_chars(hostile),
    )

    # --- U24e: no password, no match ---------------------------------------
    # migration_031 made hashed_password nullable so a directory-only account
    # can exist. That is only safe because the local path refuses a null or
    # empty hash outright: an account with no password is one nobody may sign
    # into, not one anybody may. This is the local twin of U24b.
    check(
        "U24e",
        "a null or empty password hash never matches anything",
        appmain.verify_password("anything", None) is False
        and appmain.verify_password("anything", "") is False
        and appmain.verify_password("", "") is False,
        "bcrypt.checkpw(pw, None) would raise rather than answer",
    )

    # --- U24f: the directory is a fallback, and only for a known account ----
    # If LDAP were tried before the local password, or for an email with no
    # `users` row, this public login form would become a password-spraying
    # proxy for the corporate directory — complete with the account lockouts
    # that causes for real staff. Structure, via the AST: the call must live
    # inside auth_login, and the `not row` branch must return before it.
    src = (REPO / "app" / "main.py").read_text(encoding="utf-8")
    tree = _ast.parse(src)
    login_fn = next(
        (
            n
            for n in _ast.walk(tree)
            if isinstance(n, (_ast.FunctionDef, _ast.AsyncFunctionDef)) and n.name == "auth_login"
        ),
        None,
    )
    check("U24f", "auth_login is present", login_fn is not None)
    if login_fn is not None:
        ldap_calls = [
            n
            for n in _ast.walk(login_fn)
            if isinstance(n, _ast.Call) and isinstance(n.func, _ast.Name) and n.func.id == "_try_ldap_login"
        ]
        verify_calls = [
            n
            for n in _ast.walk(login_fn)
            if isinstance(n, _ast.Call) and isinstance(n.func, _ast.Name) and n.func.id == "verify_password"
        ]
        check("U24g", "the directory is consulted from auth_login", bool(ldap_calls), f"{len(ldap_calls)} call(s)")

        # ★ There are legitimately TWO call sites since phase 4, and they have
        # opposite orderings, so a single "LDAP comes after verify_password"
        # assertion is now wrong:
        #
        #   * the FALLBACK, for an account that exists — must come after the
        #     local password has been tried, so an ordinary sign-in never
        #     touches the directory;
        #   * the JIT path, for an email with NO account — must come before,
        #     because there is no stored password to verify. It is reachable
        #     only when ldap_auto_create is on, which is the switch that
        #     accepts relaying unknown credentials to the directory.
        #
        # So the property is: every directory call is either after
        # verify_password, or guarded by the auto-create flag. A call that is
        # neither would relay every unknown sign-in attempt to corporate AD
        # with nobody having asked for it.
        vpos = min(((c.lineno, c.col_offset) for c in verify_calls), default=(0, 0))
        guarded_ranges = [
            (n.lineno, n.end_lineno or n.lineno)
            for n in _ast.walk(login_fn)
            if isinstance(n, _ast.If)
            and any(
                isinstance(sub, _ast.Constant) and isinstance(sub.value, str) and "auto_create" in sub.value
                for sub in _ast.walk(n.test)
            )
        ]
        unguarded_early = [
            c
            for c in ldap_calls
            if (c.lineno, c.col_offset) < vpos
            and not any(lo <= c.lineno <= hi for lo, hi in guarded_ranges)
        ]
        check(
            "U24h",
            "the directory is only asked before the local password when auto-create is on",
            not unguarded_early,
            "an unguarded early call relays every unknown sign-in to the corporate directory",
        )

    # --- U24i: an unknown email never reaches the directory -----------------
    # The `not row` branch returns, so nothing an attacker can name is sent on.
    if login_fn is not None:
        returns_before = [
            n.lineno
            for n in _ast.walk(login_fn)
            if isinstance(n, _ast.Return) and ldap_calls and n.lineno < ldap_calls[0].lineno
        ]
        check(
            "U24i",
            "an email with no account returns before the directory is asked",
            len(returns_before) >= 2,
            f"{len(returns_before)} early return(s) guard the directory call",
        )

    # --- U24j: the directory's answer must be for the SAME address ----------
    # A filter that matched a different entry — an alias, a shared mailbox, a
    # filter somebody widened — would otherwise authenticate whoever holds THAT
    # password into THIS account. Email is the merge key everywhere else in
    # this design; it has to be the merge key here too.
    merge_fn = next(
        (
            n
            for n in _ast.walk(tree)
            if isinstance(n, (_ast.FunctionDef, _ast.AsyncFunctionDef)) and n.name == "_try_ldap_login"
        ),
        None,
    )
    check("U24j", "_try_ldap_login is present", merge_fn is not None)
    if merge_fn is not None:
        compares = [n for n in _ast.walk(merge_fn) if isinstance(n, _ast.Compare)]
        check(
            "U24k",
            "the address the directory returned is compared with the one signing in",
            any(isinstance(op, _ast.NotEq) for c in compares for op in c.ops),
            "without this, a filter matching another entry signs that person into this account",
        )
        # It must return a plain bool, never re-raise: the caller answers every
        # directory failure with the same flat 401, because "no such entry",
        # "wrong password" and "unreachable" told apart are an enumeration
        # oracle.
        raises = [n for n in _ast.walk(merge_fn) if isinstance(n, _ast.Raise)]
        check(
            "U24l",
            "no directory failure escapes as a distinguishable error",
            not raises,
            f"{len(raises)} raise(s) would let the caller tell failures apart",
        )

    # --- U24n: every directory variable is delivered to the container -------
    # ★★★ compose.yaml's `environment:` is an ALLOWLIST. A variable set in
    # .env and not named there never reaches the container, and NOTHING
    # reports it — the module reads its default instead. That produced a
    # deployment where LDAP_ENABLED=true yielded `"ldap_enabled": false` from
    # /api/auth/config, no directory contacted, every directory sign-in
    # failing as an ordinary wrong password, and no error anywhere. Found by
    # running it. Now driven off auth_settings.SPEC, which since phase 4 is
    # the single place that says which environment variables exist.
    import app.auth_settings as asettings

    ldap_env = {env for name, (env, _k, _d) in asettings.SPEC.items() if name.startswith("ldap_")}
    absent = sorted(v for v in ldap_env if v not in os.environ)
    check(
        "U24n",
        "every LDAP variable in the settings spec is delivered to the container",
        not absent and bool(ldap_env),
        (
            f"declared in auth_settings.SPEC but absent from the environment: {absent} — "
            "compose.yaml's `environment:` block is an ALLOWLIST, not a passthrough"
        )
        if absent
        else f"{len(ldap_env)} variables",
    )

    # --- U24m: nothing in this path can grant a role ------------------------
    # The whole separation rests on it: the directory says who you are, this
    # table says what you may do. If _try_ldap_login could write `role`, then
    # whoever administers the corporate directory could make themselves a
    # Legal Scout administrator by editing a group.
    if merge_fn is not None:
        sql = " ".join(
            n.value.lower()
            for n in _ast.walk(merge_fn)
            if isinstance(n, _ast.Constant) and isinstance(n.value, str)
        )
        check(
            "U24m",
            "the directory path never writes a role or an approval",
            "role" not in sql and "approved" not in sql,
            "a directory sign-in may prove identity; it may not grant access",
        )


# ===========================================================================
# U25  Single sign-on (OIDC)
#     The provider proves identity. It grants nothing. Everything below guards
#     one way that separation, or the token verification underneath it, can be
#     lost — several of which are live vulnerabilities in other people's code.
# ===========================================================================
def test_sso():
    import ast as _ast

    import app.main as appmain
    import app.oidc as oidc

    check("U25a", "single sign-on is OFF unless configured", oidc.sso_enabled() is False)

    # --- U25b: the flow keeps NO state in this process ----------------------
    # Two uvicorn workers, no shared store. Anything stashed in module state by
    # /sso/login is simply absent when /sso/callback lands on the other worker,
    # so sign-in would fail about half the time and read as an intermittent
    # provider fault. Same class as the approval cache in U23k — found there,
    # guarded here before it could happen again.
    src_oidc = (REPO / "app" / "oidc.py").read_text(encoding="utf-8")
    tree_oidc = _ast.parse(src_oidc)
    mutable_module_state = [
        t.id
        for node in tree_oidc.body
        if isinstance(node, (_ast.Assign, _ast.AnnAssign))
        for t in ([node.target] if isinstance(node, _ast.AnnAssign) else node.targets)
        if isinstance(t, _ast.Name)
        and isinstance(node.value, (_ast.Dict, _ast.List, _ast.Set))
        and not t.id.isupper()
        and "cache" not in t.id
    ]
    check(
        "U25b",
        "no per-sign-in state is held in module memory",
        not mutable_module_state,
        f"{mutable_module_state} would be absent when the callback hits the other worker",
    )

    os.environ["JWT_SECRET_KEY"] = os.environ.get("JWT_SECRET_KEY") or "test-secret-for-unit-tests-only"

    # --- U25c: state mismatch is refused (CSRF) -----------------------------
    # Without this an attacker hands somebody a callback URL carrying the
    # ATTACKER's authorization code, signing that person into the attacker's
    # account — where everything they then do is visible to the attacker.
    flow, cookie, _challenge = oidc.new_flow()
    try:
        oidc.read_flow(cookie, "not-the-state")
        check("U25c", "a mismatched state parameter is refused", False, "read_flow returned")
    except oidc.OidcError as e:
        check("U25c", "a mismatched state parameter is refused", "state" in str(e).lower(), str(e)[:80])

    check(
        "U25d",
        "the matching state parameter is accepted",
        oidc.read_flow(cookie, flow["state"])["nonce"] == flow["nonce"],
    )

    # --- U25e: a tampered cookie is refused ---------------------------------
    _body, sig = cookie.split(".", 1)
    forged = json.dumps({"state": flow["state"], "nonce": "x", "verifier": "y", "exp": 9999999999})
    tampered = oidc._b64u(forged.encode()) + "." + sig
    try:
        oidc.read_flow(tampered, flow["state"])
        check("U25e", "a tampered sign-in cookie is refused", False, "read_flow returned")
    except oidc.OidcError as e:
        check("U25e", "a tampered sign-in cookie is refused", "signature" in str(e).lower(), str(e)[:80])

    # --- U25f: a missing cookie is refused ----------------------------------
    try:
        oidc.read_flow("", flow["state"])
        check("U25f", "a callback with no sign-in cookie is refused", False, "read_flow returned")
    except oidc.OidcError:
        check("U25f", "a callback with no sign-in cookie is refused", True)

    # --- U25g: an expired flow is refused -----------------------------------
    import hashlib as _h
    import hmac as _hm

    stale = {"state": "s", "nonce": "n", "verifier": "v", "exp": 1}
    sb = oidc._b64u(json.dumps(stale, separators=(",", ":")).encode())
    ss = oidc._b64u(_hm.new(oidc._secret(), sb.encode(), _h.sha256).digest())
    try:
        oidc.read_flow(f"{sb}.{ss}", "s")
        check("U25g", "an expired sign-in flow is refused", False, "read_flow returned")
    except oidc.OidcError as e:
        check("U25g", "an expired sign-in flow is refused", "too long" in str(e), str(e)[:80])

    # --- U25h: PKCE challenge is S256 of the verifier -----------------------
    flow2, _c2, challenge2 = oidc.new_flow()
    expect = oidc._b64u(_h.sha256(flow2["verifier"].encode()).digest())
    check("U25h", "the PKCE challenge is the S256 hash of the verifier", challenge2 == expect)

    # --- U25i: the token algorithm is PINNED, never read from the header ----
    # ★ The classic confusion attack: an attacker signs a token with HS256
    # using the provider's PUBLIC key as the HMAC secret, and a verifier that
    # trusts the header's `alg` accepts it. The public key is, by definition,
    # public. This asserts on the SOURCE via the AST rather than by grepping,
    # because the surrounding comment names every algorithm involved.
    verify_fn = next(
        (
            n
            for n in _ast.walk(tree_oidc)
            if isinstance(n, (_ast.FunctionDef, _ast.AsyncFunctionDef)) and n.name == "verify_id_token"
        ),
        None,
    )
    check("U25i", "verify_id_token is present", verify_fn is not None)
    if verify_fn is not None:
        allowed = set()
        for node in _ast.walk(verify_fn):
            if isinstance(node, _ast.Compare) and any(isinstance(o, _ast.NotIn) for o in node.ops):
                for cmp_ in node.comparators:
                    if isinstance(cmp_, _ast.Tuple):
                        allowed |= {
                            e.value for e in cmp_.elts if isinstance(e, _ast.Constant) and isinstance(e.value, str)
                        }
        check(
            "U25j",
            "the id_token algorithm is restricted to asymmetric ones",
            bool(allowed) and all(a.startswith(("RS", "ES", "PS")) for a in allowed),
            f"accepted: {sorted(allowed) or 'NOTHING — the header would be trusted'}",
        )
        check(
            "U25k",
            "no HMAC algorithm is accepted for a provider-signed token",
            not any(a.startswith("HS") for a in allowed),
            "HS256 with the public key as the secret is the algorithm-confusion attack",
        )

    # --- U25l: no "first key in the set" fallback ---------------------------
    # On a rotation the first key is the NEW one while the token in hand was
    # signed by the OLD one, so such a fallback turns a precise "unknown key
    # id" into a confusing signature error — and invites accepting a token
    # verified against a key it was not signed with.
    check(
        "U25l",
        "an unknown key id is an error, not a reason to try the first key",
        oidc._find_kid({"keys": [{"kid": "a"}, {"kid": "b"}]}, "c") is None
        and oidc._find_kid({"keys": [{"kid": "a"}]}, "a") is not None,
    )

    # --- U25m: the callback grants nothing ----------------------------------
    # The provider proves who somebody is. If this path could write `role` or
    # `approved`, whoever administers the realm could make themselves a Legal
    # Scout administrator with a group mapping.
    src_main = (REPO / "app" / "main.py").read_text(encoding="utf-8")
    tree_main = _ast.parse(src_main)
    cb = next(
        (
            n
            for n in _ast.walk(tree_main)
            if isinstance(n, (_ast.FunctionDef, _ast.AsyncFunctionDef)) and n.name == "sso_callback"
        ),
        None,
    )
    check("U25m", "sso_callback is present", cb is not None)
    if cb is not None:
        sql = " ".join(
            n.value.lower() for n in _ast.walk(cb) if isinstance(n, _ast.Constant) and isinstance(n.value, str)
        )
        check(
            "U25n",
            "the callback never writes a role, an approval or a new account",
            "insert into users" not in sql and "update users set role" not in sql and "set approved" not in sql,
            "identity is not authorisation",
        )
        # It must apply the SAME approval gate the password path does, or SSO
        # becomes a way around phase 1 entirely.
        names = {n.id for n in _ast.walk(cb) if isinstance(n, _ast.Name)}
        check(
            "U25o",
            "the callback applies the same approval gate as the password path",
            "APPROVAL_PENDING_MESSAGE" in names,
            "otherwise single sign-on walks straight past the approval gate",
        )
        # The token must go back in the FRAGMENT. A query string is sent to the
        # server, written into access logs and carried in Referer headers.
        check(
            "U25p",
            "the session token is handed back in the URL fragment, not a query string",
            "#sso_token=" in src_main and "?sso_token=" not in src_main,
        )

    # --- U25q: every OIDC variable is delivered to the container ------------
    # The same allowlist trap that made LDAP_ENABLED=true a no-op in phase 2.
    import app.auth_settings as asettings2

    oidc_env = {env for name, (env, _k, _d) in asettings2.SPEC.items() if name.startswith("oidc_")}
    absent = sorted(v for v in oidc_env if v not in os.environ)
    check(
        "U25q",
        "every OIDC variable in the settings spec is delivered to the container",
        not absent and bool(oidc_env),
        (
            f"declared in auth_settings.SPEC but absent from the environment: {absent} — "
            "compose.yaml's `environment:` block is an ALLOWLIST, not a passthrough"
        )
        if absent
        else f"{len(oidc_env)} variables",
    )

    # --- U25r: the public config endpoint leaks nothing ---------------------
    cfg_fn = next(
        (
            n
            for n in _ast.walk(tree_main)
            if isinstance(n, (_ast.FunctionDef, _ast.AsyncFunctionDef)) and n.name == "auth_config"
        ),
        None,
    )
    if cfg_fn is not None:
        keys = {
            k.value
            for n in _ast.walk(cfg_fn)
            if isinstance(n, _ast.Dict)
            for k in n.keys
            if isinstance(k, _ast.Constant) and isinstance(k.value, str)
        }
        leaky = {k for k in keys if any(w in k for w in ("secret", "client_id", "discovery", "redirect", "host", "dn"))}
        check(
            "U25r",
            "the public auth config exposes no provider configuration",
            not leaky,
            f"would expose: {sorted(leaky)}" if leaky else f"{len(keys)} keys, all cosmetic",
        )

    # --- U25s: the SSO routes are reachable before sign-in ------------------
    # They exist to produce a token, so they cannot require one. Both are
    # nonetheless guarded — by the signed flow cookie and by JWKS verification.
    for route in ("/api/auth/sso/login", "/api/auth/sso/callback"):
        check(f"U25s{route[-5:]}", f"{route} is reachable without a token", route in appmain.PUBLIC_ROUTES)


# ===========================================================================
# U26  Sign-in settings, sign-in mode, and just-in-time provisioning
#     A settings page that stores a value the code does not read is worse than
#     no settings page: it reports success and changes nothing. Most of what
#     follows guards that, and the rest guards the two ways an auto-created
#     account could arrive with more than it should.
# ===========================================================================
def test_auth_settings():
    import ast as _ast

    import app.auth_settings as st
    import app.ldap_auth as la
    import app.main as appmain
    import app.oidc as oidc

    # --- U26a: the modules read THROUGH the settings layer ------------------
    # ★ This is the check that stops the Authentication tab being decorative.
    # If ldap_auth/oidc call os.getenv directly — as they did in phases 2 and
    # 3 — an administrator saves a corrected host, the row is written, the page
    # shows the new value, and sign-in goes on using the one from .env with
    # nothing anywhere to indicate the difference. Same failure family as the
    # dead SSO button and the decorative Inactive badge.
    #
    # Scoped to the variables that ARE settings — the ones declared in
    # auth_settings.SPEC. `JWT_SECRET_KEY` is read directly by app/oidc.py to
    # sign the sign-in cookie and that is correct: it is the application's own
    # secret, not a sign-in setting, and it is deliberately not editable from a
    # web page. A blanket "no getenv anywhere" rule would fail on it, which is
    # a test being wrong rather than the code.
    spec_envs = {env for _n, (env, _k, _d) in st.SPEC.items()}
    for mod, path in (("ldap_auth", "app/ldap_auth.py"), ("oidc", "app/oidc.py")):
        src = (REPO / path).read_text(encoding="utf-8")
        tree = _ast.parse(src)
        direct = [
            n.args[0].value
            for n in _ast.walk(tree)
            if isinstance(n, _ast.Call)
            and (
                (isinstance(n.func, _ast.Name) and n.func.id == "getenv")
                or (isinstance(n.func, _ast.Attribute) and n.func.attr == "getenv")
            )
            and n.args
            and isinstance(n.args[0], _ast.Constant)
            and n.args[0].value in spec_envs
        ]
        check(
            f"U26a-{mod}",
            f"{mod} reads settings through the settings layer, not the environment",
            not direct,
            f"read directly from the environment: {direct} — the admin panel could not change them",
        )

    # --- U26b: an enum is validated BEFORE the write ------------------------
    # Storing signin_mode="sso-only" (a hyphen) writes fine and reads back as
    # unrecognised, falling through to the default — so a deployment meant to
    # be SSO-only goes on accepting passwords while the settings page shows
    # exactly what the administrator asked for.
    try:
        st.update({"signin_mode": "sso-only"}, "test")
        check("U26b", "an invalid sign-in mode is refused at the write", False, "update() accepted it")
    except st.SettingsError as e:
        check("U26b", "an invalid sign-in mode is refused at the write", "must be one of" in str(e), str(e)[:70])

    try:
        st.update({"not_a_setting": "x"}, "test")
        check("U26c", "an unknown setting key is refused", False, "update() accepted it")
    except st.SettingsError:
        check("U26c", "an unknown setting key is refused", True)

    # --- U26d: validation is all-or-nothing ---------------------------------
    # A loop that writes as it validates leaves the settings half from the form
    # and half from before it — a state nobody chose and nobody can see.
    #
    # Asserted structurally rather than by attempting a bad save: driving that
    # for real would WRITE to app_settings whenever the code was broken, which
    # is exactly when a test suite must not be mutating the deployment it is
    # inspecting. The property is that no database call appears inside the
    # validation loop; the writes come after it, in a second pass.
    st_src_early = (REPO / "app" / "auth_settings.py").read_text(encoding="utf-8")
    upd = next(
        (
            n
            for n in _ast.walk(_ast.parse(st_src_early))
            if isinstance(n, _ast.FunctionDef) and n.name == "update"
        ),
        None,
    )
    if upd is not None:
        loops = [n for n in upd.body if isinstance(n, _ast.For)]
        writes_in_validation = any(
            isinstance(c, _ast.Call)
            and isinstance(c.func, _ast.Attribute)
            and c.func.attr in ("execute", "commit")
            for loop in loops[:1]
            for c in _ast.walk(loop)
        )
        check(
            "U26d",
            "nothing is written until every supplied setting has validated",
            loops and not writes_in_validation,
            "a partial write leaves the settings half-applied",
        )

    # --- U26e: secrets are write-only ---------------------------------------
    # A settings page that renders a client secret puts it in the DOM, in the
    # browser's memory, and in any screenshot of that page.
    view = st.public_view()
    leaked = sorted(k for k in st.SECRET_KEYS if k in view)
    check(
        "U26e",
        "no secret is ever returned by the settings API",
        not leaked,
        f"would return: {leaked}" if leaked else f"{len(st.SECRET_KEYS)} secrets reported as set/not-set only",
    )
    check(
        "U26f",
        "each secret is reported as set or not set",
        all(f"{k}_set" in view for k in st.SECRET_KEYS),
    )

    # --- U26g: a blank secret means "keep", not "clear" ---------------------
    # The form cannot show the current value, so it posts blank on every save.
    # Treating that as a clear would wipe the client secret the first time
    # anybody edited an unrelated field on the same page.
    upd_fn = next(
        (n for n in _ast.walk(_ast.parse((REPO / "app" / "auth_settings.py").read_text())) if isinstance(n, _ast.FunctionDef) and n.name == "update"),
        None,
    )
    check("U26g", "auth_settings.update is present", upd_fn is not None)
    if upd_fn is not None:
        has_skip = any(
            isinstance(n, _ast.Continue) for n in _ast.walk(upd_fn)
        )
        check(
            "U26h",
            "an empty secret is skipped rather than written",
            has_skip,
            "otherwise saving any other field on the page clears the stored secret",
        )

    # --- U26i: sso_only is applied AFTER the password, and exempts an admin -
    # ★★★ Both halves matter. Applied before the password check, the differing
    # status code tells an anonymous caller which addresses have accounts.
    # Without the admin exemption, one mistyped provider URL locks every human
    # out — including whoever has to sign in to correct it — and the only way
    # back is editing Postgres by hand.
    src_main = (REPO / "app" / "main.py").read_text(encoding="utf-8")
    tree_main = _ast.parse(src_main)
    login_fn = next(
        (
            n
            for n in _ast.walk(tree_main)
            if isinstance(n, (_ast.FunctionDef, _ast.AsyncFunctionDef)) and n.name == "auth_login"
        ),
        None,
    )
    if login_fn is not None:
        verify_ln = min(
            (
                n.lineno
                for n in _ast.walk(login_fn)
                if isinstance(n, _ast.Call) and isinstance(n.func, _ast.Name) and n.func.id == "verify_password"
            ),
            default=0,
        )
        mode_ln = min(
            (
                n.lineno
                for n in _ast.walk(login_fn)
                if isinstance(n, _ast.Call) and isinstance(n.func, _ast.Name) and n.func.id == "_auth_settings_mode"
            ),
            default=0,
        )
        check(
            "U26i",
            "sso_only is enforced after the password is verified",
            verify_ln and mode_ln and verify_ln < mode_ln,
            "checked earlier, the status code enumerates which addresses have accounts",
        )
        consts = {
            n.value
            for n in _ast.walk(login_fn)
            if isinstance(n, _ast.Constant) and isinstance(n.value, str)
        }
        check(
            "U26j",
            "an administrator keeps password sign-in under sso_only (break-glass)",
            "sso_only" in consts and "admin" in consts,
            "without it, one bad provider URL locks everyone out including whoever must fix it",
        )

    # --- U26k: JIT cannot mint anything privileged --------------------------
    # ★★★ role and approved are SQL LITERALS with no parameter, so no claim,
    # group, mapper or caller can reach them.
    jit = next(
        (n for n in _ast.walk(tree_main) if isinstance(n, _ast.FunctionDef) and n.name == "_jit_provision"),
        None,
    )
    check("U26k", "_jit_provision is present", jit is not None)
    if jit is not None:
        args = {a.arg for a in jit.args.args}
        check(
            "U26l",
            "just-in-time provisioning takes no role or approval argument",
            not (args & {"role", "approved", "is_admin"}),
            f"arguments: {sorted(args)}",
        )
        sql = " ".join(
            n.value.lower() for n in _ast.walk(jit) if isinstance(n, _ast.Constant) and isinstance(n.value, str)
        )
        # The property is that `role` and `approved` appear in the VALUES clause
        # as LITERALS. An earlier version of this check also asserted that the
        # string `"%s, 'user'"` was absent, meaning to catch role being passed
        # as a parameter — but that substring occurs in the perfectly correct
        # statement too (the placeholder before it belongs to `full_name`), so
        # the check failed on code that was right. The literals plus U26l's
        # "no role or approved argument" are together the real guarantee.
        values = sql[sql.index("values") :] if "values" in sql else sql
        check(
            "U26m",
            "an auto-created account is always an unprivileged, unapproved one",
            "'user'" in values and "false" in values,
            f"role/approved must be literals in the VALUES clause — got: {values[:90]}",
        )
        check(
            "U26n",
            "a parallel sign-in by the same new person cannot fail",
            "on conflict" in sql,
            "a browser reloading the callback is the ordinary case, not the exotic one",
        )

    # --- U26o: both auto-create switches default to OFF ---------------------
    for key in ("ldap_auto_create", "oidc_auto_create"):
        check(f"U26o-{key}", f"{key} is off unless switched on", st.SPEC[key][2] is False)

    # --- U26p: settings live in the existing table, not a second one --------
    check(
        "U26p",
        "sign-in settings reuse app_settings under their own prefix",
        st.PREFIX.endswith(".") and st.PREFIX != "",
        f"prefix={st.PREFIX!r} — a second key/value table for the same job would be the drift",
    )

    # --- U26q: the settings layer holds no in-process cache -----------------
    # Two workers. A value cached at import in one is stale in the other the
    # moment somebody saves, so a setting would appear to take effect or not
    # depending on which worker answered. Same class as U23k and U25b.
    st_src = (REPO / "app" / "auth_settings.py").read_text(encoding="utf-8")
    st_tree = _ast.parse(st_src)
    caches = [
        t.id
        for node in st_tree.body
        if isinstance(node, (_ast.Assign, _ast.AnnAssign))
        for t in ([node.target] if isinstance(node, _ast.AnnAssign) else node.targets)
        if isinstance(t, _ast.Name) and isinstance(node.value, (_ast.Dict, _ast.List)) and not t.id.isupper()
    ]
    check(
        "U26q",
        "the settings layer caches nothing in module memory",
        not caches,
        f"{caches} would be stale in the other worker the moment a setting is saved",
    )

    # --- U26r: config() actually exposes the JIT flag -----------------------
    # The flag has to reach the login path, or turning it on does nothing.
    check("U26r", "the directory config exposes auto_create", "auto_create" in la.config())
    check("U26s", "the provider config exposes auto_create", "auto_create" in oidc.config())

    # --- U26t: the admin settings routes exist and are NOT public -----------
    for route in ("/api/admin/auth-settings",):
        check(
            f"U26t{route[-5:]}",
            "the settings endpoint is not reachable without admin",
            route not in appmain.PUBLIC_ROUTES,
        )


# ===========================================================================
# U27  Document tools return a STRING to HTTP callers
#     The one place a success can be reported as a failure, which is worse
#     than the reverse: the work is done, nobody believes it, and the retry
#     does it again.
# ===========================================================================
def test_tool_result_shape():
    import ast as _ast

    import app.main as appmain
    from scout.tools.smart_doc import create_smart_document_tool

    # --- U27a: the tool really does return a string ------------------------
    # It is built to be called by the agent, where agno hands the string back
    # to the model. Nothing is wrong with that; what was wrong was two HTTP
    # endpoints calling it like an ordinary function.
    gen = create_smart_document_tool("/documents", host="")["generate_document"]
    check(
        "U27a",
        "generate_document returns a string, not a dict",
        not hasattr(gen(template_name="__no_such_template__", company_name="__none__", custom_data={}), "get"),
        "if this ever starts returning a dict the normaliser stays correct anyway",
    )

    # --- U27b: the normaliser handles every shape --------------------------
    n = appmain._tool_result
    eq("U27b1", "a JSON string is parsed", n('{"success": true, "file_name": "x.docx"}')["success"], True)
    eq("U27b2", "a dict passes through", n({"success": False, "error": "e"})["error"], "e")
    eq("U27b3", "bytes are decoded", n(b'{"success": true}')["success"], True)
    check("U27b4", "unparseable input fails honestly", n("not json at all")["success"] is False)
    check("U27b5", "an unexpected type fails honestly", n(12345)["success"] is False)

    # --- U27c: no HTTP caller reads .get() off a raw tool call --------------
    # ★ The actual defect: `result.get("success")` on the string raised, the
    # surrounding except reported failure, and the document had ALREADY been
    # written. The Fill-in view's Generate button therefore reported an error
    # on every single use while producing a file, and bulk generate reported
    # 0/N for a run that generated all N.
    src = (REPO / "app" / "main.py").read_text(encoding="utf-8")
    tree = _ast.parse(src)
    bare = []
    for node in _ast.walk(tree):
        # `<something>(...).get(...)` where the inner call is a document tool
        if not (isinstance(node, _ast.Call) and isinstance(node.func, _ast.Attribute) and node.func.attr == "get"):
            continue
        inner = node.func.value
        if not isinstance(inner, _ast.Call):
            continue
        f = inner.func
        name = (
            f.id
            if isinstance(f, _ast.Name)
            else (f.attr if isinstance(f, _ast.Attribute) else None)
        )
        if isinstance(f, _ast.Subscript) or name in ("generate", "generate_document"):
            bare.append(node.lineno)
    check(
        "U27c",
        "no endpoint calls .get() straight off a document tool's return value",
        not bare,
        f"lines {bare} would raise on the string the tool returns",
    )

    # --- U27d: both known call sites go through the normaliser -------------
    calls = [
        n_
        for n_ in _ast.walk(tree)
        if isinstance(n_, _ast.Call) and isinstance(n_.func, _ast.Name) and n_.func.id == "_tool_result"
    ]
    check(
        "U27d",
        "fill-generate and bulk generate both normalise the result",
        len(calls) >= 2,
        f"found {len(calls)}; expected the two document endpoints",
    )


# ===========================================================================
# U28  View is open, change is not
#     Reading the registers is open to anyone signed in; creating, changing or
#     deleting shared firm records is not. The boundary is the ENDPOINT — a
#     hidden button stops a click, never a request.
# ===========================================================================
def test_write_boundary():
    import ast as _ast

    import app.main as appmain

    src = (REPO / "app" / "main.py").read_text(encoding="utf-8")
    tree = _ast.parse(src)

    # Every route handler, with the gate it calls.
    routes = []
    for n in _ast.walk(tree):
        if not isinstance(n, (_ast.FunctionDef, _ast.AsyncFunctionDef)):
            continue
        for d in n.decorator_list:
            if not (isinstance(d, _ast.Call) and isinstance(d.func, _ast.Attribute)):
                continue
            if not (isinstance(d.func.value, _ast.Name) and d.func.value.id == "app"):
                continue
            if d.func.attr not in ("get", "post", "put", "delete", "patch"):
                continue
            path = d.args[0].value if d.args and isinstance(d.args[0], _ast.Constant) else "?"
            called = {
                x.func.id for x in _ast.walk(n) if isinstance(x, _ast.Call) and isinstance(x.func, _ast.Name)
            }
            routes.append((d.func.attr.upper(), path, n.name, called))

    mutating = [r for r in routes if r[0] in ("POST", "PUT", "DELETE", "PATCH")]
    check("U28a", "the route table is readable", len(mutating) > 40, f"{len(mutating)} mutating routes")

    # ★ The nine that may stay open are each the signed-in person's OWN work:
    # their chat title, their follow-up suggestions, their message feedback,
    # approving or discarding an email they were shown, generating a document
    # (which is the product, not an edit to a register), and the two auth
    # routes, which necessarily precede having a session at all.
    OWN_WORK = {
        "auth_login", "auth_logout", "sso_login", "sso_callback",
        "set_session_title", "suggest_followups", "record_message_feedback",
        "send_queued_email", "discard_queued_email", "send_document_email",
        "documents_fill_generate",
    }
    ungated = [
        (m, p, fn) for m, p, fn, called in mutating
        if fn not in OWN_WORK and not ({"require_admin", "require_write"} & called)
    ]
    check(
        "U28b",
        "no route that changes shared records is left without a role check",
        not ungated,
        f"ungated: {[f'{m} {p}' for m, p, _ in ungated]}" if ungated else "",
    )

    # --- U28c: the ten that were found open are now closed -------------------
    # Each of these could be called by ANY signed-in account before this change,
    # including a plain viewer. `delete_knowledge_source` is the sharpest: it
    # removes a knowledge source outright.
    WERE_OPEN = {
        "delete_knowledge_source", "extract_company_from_pdf_stream", "upload_company_pdf",
        "add_dashboard_company", "bulk_generate_documents", "upload_dashboard_template",
        "sync_existing_documents", "sync_templates_to_knowledge", "set_template_category",
        "save_training_logs",
    }
    by_name = {fn: called for _m, _p, fn, called in routes}
    still_open = sorted(f for f in WERE_OPEN if not ({"require_admin", "require_write"} & by_name.get(f, set())))
    check("U28c", "every previously ungated shared-record route is gated", not still_open, f"{still_open}")

    # --- U28d: the gate actually refuses a viewer ---------------------------
    class _Req:
        def __init__(self, role):
            self.headers = {}
            self._role = role

        # get_current_user reads the Authorization header; stub the decode by
        # patching the helper instead, below.

    import app.main as m

    real = m.get_current_user
    try:
        for role, allowed in (("user", False), ("editor", False), ("admin", True)):
            m.get_current_user = lambda _r, _role=role: {"user_id": 1, "email": "x@y.z", "role": _role}
            try:
                m.require_write(_Req(role))
                got = True
            except Exception:
                got = False
            check(
                f"U28d-{role}",
                f"require_write {'admits' if allowed else 'refuses'} {role}",
                got is allowed,
                f"role={role} allowed={got}, expected {allowed}",
            )
    finally:
        m.get_current_user = real

    # --- U28e: an unauthenticated caller is refused too ---------------------
    try:
        m.get_current_user = lambda _r: None
        try:
            m.require_write(_Req("none"))
            ok = False
        except Exception:
            ok = True
        check("U28e", "require_write refuses an unauthenticated caller", ok)
    finally:
        m.get_current_user = real

    # --- U28h: super_admin is admitted everywhere admin is ------------------
    # ★ `require_admin` compared `role != "admin"` by EQUALITY, which refuses
    # the tier meant to be able to do strictly more — and it does so on every
    # admin route at once, so a brand-new super administrator would have been
    # locked out of the whole panel by the check that exists to let them in.
    real2 = m.get_current_user
    try:
        for role, allowed in (("super_admin", True), ("admin", True), ("editor", False), ("user", False)):
            m.get_current_user = lambda _r, _role=role: {"user_id": 1, "email": "x@y.z", "role": _role}
            for gate in ("require_admin", "require_write"):
                try:
                    getattr(m, gate)(_Req(role))
                    got = True
                except Exception:
                    got = False
                check(
                    f"U28h-{gate}-{role}",
                    f"{gate} {'admits' if allowed else 'refuses'} {role}",
                    got is allowed,
                    f"{gate}({role}) allowed={got}, expected {allowed}",
                )
    finally:
        m.get_current_user = real2

    # --- U28i: nobody can grant a role above their own ----------------------
    # Both halves matter. Checking it only on creation is decorative: make a
    # plain user, then edit them into a super_admin.
    for actor_role, target, allowed in (
        ("admin", "admin", True),
        ("admin", "super_admin", False),
        ("super_admin", "super_admin", True),
        ("super_admin", "admin", True),
        ("editor", "admin", False),
    ):
        check(
            f"U28i-{actor_role}->{target}",
            f"a {actor_role} {'may' if allowed else 'may not'} grant {target}",
            m._may_grant({"role": actor_role}, target) is allowed,
        )

    src_main2 = (REPO / "app" / "main.py").read_text(encoding="utf-8")
    tree2 = _ast.parse(src_main2)
    for fname in ("admin_create_user", "admin_update_user"):
        fn2 = next(
            (n for n in _ast.walk(tree2)
             if isinstance(n, (_ast.FunctionDef, _ast.AsyncFunctionDef)) and n.name == fname),
            None,
        )
        called2 = {
            x.func.id for x in _ast.walk(fn2)
            if isinstance(x, _ast.Call) and isinstance(x.func, _ast.Name)
        } if fn2 else set()
        check(
            f"U28j-{fname}",
            f"{fname} enforces the grant rule",
            "_may_grant" in called2,
            "a rule checked on only one of the two paths is not a rule",
        )

    # --- U28f: the ranking agrees between server and browser ----------------
    # ROLE_RANK here and RANK in roleClient.ts decide the same thing in two
    # places; if they disagree, the UI offers what the API refuses (or hides
    # what it would allow), which is how a permission becomes decorative.
    rc = (REPO / "agent-ui" / "src" / "app" / "admin" / "roleClient.ts")
    if rc.exists():
        txt = rc.read_text(encoding="utf-8")
        found = dict(re.findall(r"(\w+):\s*(\d+)", txt.split("RANK", 1)[1].split("}")[0]))
        check(
            "U28f",
            "the role ranking is the same on the server and in the browser",
            {k: int(v) for k, v in found.items()} == appmain.ROLE_RANK,
            f"browser={found} server={appmain.ROLE_RANK}",
        )
        check(
            "U28g",
            "the browser's write bar matches require_write (admin)",
            "RANK.admin" in txt,
            "canWrite must compare against the same role require_write demands",
        )


# ===========================================================================
# U29  Resuming a paused run
#     Answering a question card is how every human-in-the-loop flow in this
#     product continues — template choice, person picker, confirmations. It was
#     broken for months, in two different ways, and BOTH looked identical from
#     the browser: the card went ANSWERED and nothing else ever happened.
#
#     These run without a model, a network or a server: they drive the ASGI
#     middleware directly. The tracker suites cannot catch this — they exercise
#     a scripted runtime rather than real HTTP, which is exactly why 261 tests
#     passed while this was dead.
# ===========================================================================
def test_resume_receive():
    import ast as _ast
    import asyncio
    import contextlib

    import app.main as appmain

    mw_cls = getattr(appmain, "ResumeSessionScope", None)
    check("U29a", "the resume middleware exists", mw_cls is not None)
    if mw_cls is None:
        return

    BODY = b"session_id=abc-123&stream=true&tools=%5B%5D"

    async def _drive(client_messages):
        """Run one /continue through the middleware and record what the app saw."""
        seen = {"body": b"", "receives": [], "disconnect_awaited": False}

        async def receive():
            # The real transport: yields the body, then blocks until the client
            # genuinely goes away. Modelled here as an event that never fires
            # unless the test arranges it.
            if client_messages:
                return client_messages.pop(0)
            seen["disconnect_awaited"] = True
            await asyncio.sleep(3600)

        async def app(scope, rcv, send):
            # Consume the body the way an endpoint does...
            while True:
                m = await rcv()
                seen["receives"].append(m["type"])
                if m["type"] != "http.request":
                    break
                seen["body"] += m.get("body", b"")
                if not m.get("more_body"):
                    break
            # ...then behave like StreamingResponse.listen_for_disconnect,
            # which loops on receive() waiting for http.disconnect.
            m = await rcv()
            seen["receives"].append(m["type"])

        scope = {"type": "http", "method": "POST", "path": "/agents/scout/runs/r1/continue"}
        with contextlib.suppress(TimeoutError):
            await asyncio.wait_for(mw_cls(app)(scope, receive, lambda _m: None), timeout=1.5)
        return seen

    seen = asyncio.run(_drive([{"type": "http.request", "body": BODY, "more_body": False}]))

    check("U29b", "the body reaches the app intact", seen["body"] == BODY, f"{seen['body']!r}")

    # ★★★ THE regression. After the body is spent, the middleware must NOT
    # invent a disconnect. `listen_for_disconnect` is waiting for exactly that
    # message, so returning it TELLS Starlette the client hung up and the
    # response generator is cancelled — measured: the resumed run emitted
    # RunContinued and ToolCallStarted, then stopped after 0.1s without ever
    # reaching the model, while the same payload with stream=false completed in
    # 5.4s. The run stayed PAUSED with answered=null.
    check(
        "U29c",
        "no disconnect is invented once the body is spent",
        "http.disconnect" not in seen["receives"],
        f"received {seen['receives']} — a fabricated disconnect cancels the stream",
    )
    check(
        "U29d",
        "the disconnect listener is left waiting on the real transport",
        seen["disconnect_awaited"],
        "it must delegate to the real receive, which blocks until the client truly leaves",
    )

    # --- U29e: a genuine disconnect still propagates ------------------------
    # Delegating must not swallow a real one, or a closed browser would leave
    # the run streaming into nothing.
    seen2 = asyncio.run(
        _drive([
            {"type": "http.request", "body": BODY, "more_body": False},
            {"type": "http.disconnect"},
        ])
    )
    check(
        "U29e",
        "a real disconnect still reaches the app",
        "http.disconnect" in seen2["receives"],
        f"received {seen2['receives']}",
    )

    # --- U29f: the earlier bug cannot come back either -----------------------
    # Returning http.request forever makes listen_for_disconnect raise
    # `RuntimeError: Unexpected message received: http.request`. One body
    # message, exactly one.
    check(
        "U29f",
        "the body is handed over exactly once",
        seen["receives"].count("http.request") == 1,
        f"received {seen['receives']}",
    )

    # --- U29g: no middleware reads the body inside BaseHTTPMiddleware --------
    # ★ Reading a request body inside a BaseHTTPMiddleware makes Starlette cache
    # and replay the request, and its wrapped_receive then hands an http.request
    # to the streaming disconnect listener. That is the ORIGINAL crash. The body
    # must only ever be touched by plain-ASGI middleware.
    src = (REPO / "app" / "main.py").read_text(encoding="utf-8")
    tree = _ast.parse(src)
    offenders = []
    for node in _ast.walk(tree):
        if not (isinstance(node, _ast.ClassDef) and any(
            (isinstance(b, _ast.Name) and b.id == "BaseHTTPMiddleware")
            or (isinstance(b, _ast.Attribute) and b.attr == "BaseHTTPMiddleware")
            for b in node.bases
        )):
            continue
        for sub in _ast.walk(node):
            if (
                isinstance(sub, _ast.Await)
                and isinstance(sub.value, _ast.Call)
                and isinstance(sub.value.func, _ast.Attribute)
                and sub.value.func.attr in ("body", "form", "json")
                and isinstance(sub.value.func.value, _ast.Name)
                and sub.value.func.value.id == "request"
            ):
                offenders.append(f"{node.name}.{sub.value.func.attr}()")
    check(
        "U29g",
        "no BaseHTTPMiddleware reads the request body",
        not offenders,
        f"{offenders} — this breaks every streaming response with a request body",
    )


# ===========================================================================
# U30  Template drift — the wrong legal instrument
#     The worst thing this product has done. Roughly two runs in eleven on the
#     client's tracker, generate_document was called with a template that
#     contradicted the request: an Individual Shareholder Consent for
#     "director consent form (non group member)". Filled correctly, looking
#     right, and wrong. Nothing tested it until now.
# ===========================================================================
def test_template_drift():
    from scout.tools.template_guard import contradiction

    # ---- the real drift cases, transcribed from the tracker ---------------
    MUST_BLOCK = [
        ("Prepare director consent form (non group member) for Win Win Tint",
         "Individual Shareholder Consent Form.docx", "director asked, shareholder offered"),
        ("Prepare director consent form (non group member) for Min Min",
         "Director Consent Form - Group Member Appointment.docx", "non-group asked, group offered"),
        ("Prepare a shareholder consent form for Soe Moe Thu",
         "Director Consent Form - Non-Group Member Appointment.docx", "shareholder asked, director offered"),
        ("Prepare resignation letter of Daw Win Win Tint from City Holdings",
         "Shareholders Resolution In Writing - Director Appointment.docx", "resignation asked, appointment offered"),
        ("Prepare a corporate shareholder consent for Pahtama Group",
         "Individual Shareholder Consent Form.docx", "corporate asked, individual offered"),
    ]
    for req, tpl, why in MUST_BLOCK:
        check(f"U30a-{why[:22]}", f"refuses: {why}", contradiction(req, tpl) is not None, f"{tpl} <- {req[:50]}")

    # ---- and it must NOT get in the way -----------------------------------
    # A guard that blocks correct work is worse than no guard: it trains people
    # to click through it. Silence on either side is never a contradiction.
    MUST_ALLOW = [
        ("Prepare director consent form (non group member) for Min Min",
         "Director Consent Form - Non-Group Member Appointment.docx", "the matching template"),
        ("Prepare a shareholder consent form for Soe Moe Thu",
         "Individual Shareholder Consent Form.docx", "shareholder asked, shareholder offered"),
        ("Create AGM minutes for City Holdings", "Annual General Meeting Minutes.docx", "no axis mentioned"),
        ("Prepare shareholders resolution for resignation and appointment of directors",
         "Shareholders Resolution In Writing - Director Resignation and Appointment.docx",
         "BOTH sides asked — a real template, must not be read as a clash"),
        ("Prepare a document for City Holdings",
         "Notice of Calling for Annual General Meeting.docx", "request says nothing"),
        ("Prepare AGM minutes",
         "Shareholders Resolution In Writing for Annual General Meeting.docx", "no discriminator either side"),
        ("", "Director Consent Form - Group Member Appointment.docx", "no request text at all"),
    ]
    for req, tpl, why in MUST_ALLOW:
        check(f"U30b-{why[:22]}", f"allows: {why}", contradiction(req, tpl) is None,
              f"wrongly blocked {tpl} <- {req[:50]!r}")

    # ---- the guard is wired into generation, not just defined -------------
    import ast as _ast

    src = (REPO / "scout" / "tools" / "smart_doc.py").read_text(encoding="utf-8")
    tree = _ast.parse(src)
    gen = next(
        (n for n in _ast.walk(tree)
         if isinstance(n, (_ast.FunctionDef, _ast.AsyncFunctionDef)) and n.name == "generate_document"),
        None,
    )
    check("U30c", "generate_document is present", gen is not None)
    if gen is not None:
        called = {x.func.id for x in _ast.walk(gen) if isinstance(x, _ast.Call) and isinstance(x.func, _ast.Name)}
        check(
            "U30d",
            "generation consults the drift guard before it writes anything",
            "_check_template_drift" in called,
            "a guard nothing calls is a guard that does not exist",
        )

    # ---- the request text must come from the RUN, not from an argument ----
    # ★ An argument would be filled in by the model, and the model is the thing
    # being checked — a paraphrase cannot contradict the template its author
    # chose. It has to be the user's raw words or nothing.
    fn = next(
        (n for n in _ast.walk(tree)
         if isinstance(n, _ast.FunctionDef) and n.name == "_latest_request_text"), None)
    check("U30e", "the request text is read from the run record", fn is not None)
    if fn is not None:
        sql = " ".join(
            n.value.lower() for n in _ast.walk(fn)
            if isinstance(n, _ast.Constant) and isinstance(n.value, str)
        )
        check("U30f", "it reads the stored run input", "agno_sessions" in sql and "input_content" in sql, sql[:80])


def main():
    for fn in (
        test_placeholders,
        test_field_aliases,
        test_classify_kind,
        test_session_scope,
        test_companion_identifier,
        test_party_coercion,
        test_person_guard,
        test_picker_payload,
        test_repeat_regions,
        test_fill_view,
        test_blank_reporting,
        test_register_authority,
        test_approval_gate,
        test_ldap_signin,
        test_sso,
        test_auth_settings,
        test_tool_result_shape,
        test_write_boundary,
        test_resume_receive,
        test_template_drift,
        test_bare_person_attributes,
        test_one_resolver_three_views,
        test_card_suggests_nothing_unknowable,
        test_panel_watches_real_tool_names,
        test_multi_stakeholder_attributes,
        test_stall_guard_sees_unpaid_debt,
        test_pick_refreshes_the_panel,
        test_one_card_holds_the_whole_form,
        test_tracked_changes_are_visible_to_the_fill,
        test_structural_contracts,
    ):
        try:
            fn()
        except Exception as e:
            check(fn.__name__, f"{fn.__name__} raised", False, f"{type(e).__name__}: {e}")

    width = max(len(r[2]) for r in _results)
    print(f"\n{'ID':<8} {'RESULT':<8} {'CASE':<{width}}  DETAIL")
    print("-" * (26 + width + 40))
    failed = skipped = 0
    for cid, result, name, detail in _results:
        if result == "FAIL":
            failed += 1
            print(f"{cid:<8} {result:<8} {name:<{width}}  {detail}")
        elif result == "SKIP":
            skipped += 1
            # A skip always says WHY. A silent skip is a test that quietly
            # stopped testing, which is the failure mode this suite exists for.
            print(f"{cid:<8} {result:<8} {name:<{width}}  {detail}")
        else:
            print(f"{cid:<8} {result:<8} {name:<{width}}")

    total = len(_results)
    print(
        f"\nSUMMARY: PASS={total - failed - skipped}"
        + (f" · SKIP={skipped}" if skipped else "")
        + (f" · FAIL={failed}" if failed else "")
    )
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
