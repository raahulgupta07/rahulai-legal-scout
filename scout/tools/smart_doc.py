"""
Smart Document Workflow Tool
===========================

Complete workflow for legal document generation:
1. Analyze template - extract required fields
2. Check data - validate company data vs required fields
3. Ask for missing info - if needed
4. Validate filled document
5. Generate final document
"""

# ★ Required on python 3.9: `_lookup_recent_generation` is annotated
# `dict[str, Any] | None`, and `GenericAlias | None` is evaluated at def time
# without it — `TypeError: unsupported operand type(s) for |`. The container is
# 3.12 so this never bit there; it made the whole module unimportable on the
# system python, which is what `tracker_fill.py` runs on.
from __future__ import annotations

import hashlib
import json as _json
import logging
import re
import threading
import time
from contextvars import ContextVar
from datetime import datetime
from functools import wraps
from pathlib import Path
from typing import Any

from agno.run import RunContext
from docx import Document

from scout.tools.document_tracker import record_document
from scout.tools.field_aliases import canonical_field, normalize_field, tokens_match
from scout.tools.placeholders import (
    PLACEHOLDER_PATTERN,
    is_empty_placeholder,
    new_empty_counter,
    placeholder_name,
)
from scout.tools.repeat_regions import _is_corporate
from scout.tools.slot_resolver import (
    _attr_tail,
    collect_slot_requests,
    companion_attribute,
    current_session_scope,
    normalise_mapping,
    resolve_slot,
    session_scope,
    slot_of,
)

LEFT_BLANK = ""

# Set on the fill `data` when the document is for a company being incorporated.
# Read in find_replacement's `source == "db"` branch, which otherwise answers
# company-identity placeholders straight from the register row and never looks
# at `data` at all — so setting data["company"] alone had no effect on the
# finished document.
NEWCO_MARKER = "__new_company_name__"

# Only these are overridden. Everything else a template pulls from the register
# is left as it is: the consent forms render just the name and the number, and
# blanking the rest would trip the undeclared-blank guard on fields the document
# never shows.
_NEWCO_NAME_FIELDS = frozenset({"company", "company_name", "company_name_english"})
_NEWCO_REG_FIELDS = frozenset({"company_registration_number"})

# Placeholders this fill deliberately replaced with nothing.
#
# `validate_filled_document` re-opens the SAVED document and looks for leftover
# `{{...}}` patterns, so it is structurally blind to a placeholder that was
# replaced with "" — there is no pattern left to find. That is how a Corporate
# Shareholder Consent came out reading "hereinafter referred to as  ("NewCo")
# ... shall invest in  in NewCo" while the result said 13 placeholders,
# unfilled_names: [] and validation_status "Complete": a resolution
# incorporating an unnamed company for an unstated sum, reported as finished.
#
# A ContextVar rather than a parameter because `find_replacement` is called deep
# inside the paragraph/table walk and threading a collector through every frame
# would touch far more code than the defect warrants. Set per generation and
# reset in a finally, so one document's blanks can never be attributed to the
# next — the same leak the effects turn scope guards against.
_blanked_placeholders: ContextVar = ContextVar("smart_doc_blanked", default=None)


def _note_blank(placeholder: str) -> None:
    """Record that `placeholder` was emitted as empty, if anyone is collecting."""
    sink = _blanked_placeholders.get()
    if sink is not None and placeholder:
        sink.add(str(placeholder))


_drift_log = logging.getLogger("legalscout.drift")


def _latest_request_text(session_id: str) -> str:
    """The user's own words for this conversation's most recent request.

    Read from the run record rather than passed in as an argument, and that is
    deliberate: an argument would be filled in by the MODEL, and the model is
    the thing whose choice is being checked. A paraphrase of the request cannot
    contradict the template the paraphraser picked — it is the raw text or
    nothing.

    Returns "" for any failure, and every caller treats "" as "no opinion", so
    a database hiccup can never block a document.
    """
    if not session_id:
        return ""
    try:
        from db.connection import get_db_conn

        conn = get_db_conn()
        try:
            cur = conn.cursor()
            # The most recent run's input, newest last in the runs array.
            cur.execute(
                "SELECT r->'input'->>'input_content' "
                "FROM ai.agno_sessions s, jsonb_array_elements(s.runs) WITH ORDINALITY AS t(r, i) "
                "WHERE s.session_id = %s AND r->'input'->>'input_content' IS NOT NULL "
                "ORDER BY i DESC LIMIT 1",
                (session_id,),
            )
            row = cur.fetchone()
            cur.close()
            return (row[0] or "") if row else ""
        finally:
            conn.close()
    except Exception as e:
        _drift_log.warning(f"[DRIFT] could not read the request text for {session_id}: {e}")
        return ""


def _check_template_drift(session_id: str, template_name: str) -> dict | None:
    """Block a template that contradicts the request. None means proceed.

    Shaped like the tool's other refusals — `success: False` with a message and
    a question — so the agent surfaces it as a clarification rather than an
    error, and the user gets a choice instead of a wrong document.
    """
    try:
        from scout.tools.template_guard import contradiction

        request_text = _latest_request_text(session_id)
        clash = contradiction(request_text, template_name)
    except Exception as e:
        # A guard that breaks must not break generation.
        _drift_log.warning(f"[DRIFT] check skipped: {e}")
        return None

    if not clash:
        return None

    _drift_log.warning(
        f"[DRIFT] refused {template_name!r}: request asked for "
        f"{clash['requested']!r} but the template is {clash['template']!r} ({clash['axis']})"
    )
    return {
        "success": False,
        "error": "template_mismatch",
        "clarification_needed": True,
        "message": clash["message"],
        "requested": clash["requested"],
        "template_offered": template_name,
        "agent_instruction": (
            "STOP. Do not generate this document. The template you chose states the opposite of "
            "what the user asked for, on a distinction where the two are different legal "
            "instruments. Ask the user which they want — quote both options — and generate only "
            "after they answer. Do not re-call generate_document with the same template."
        ),
    }


def _session_of(run_context: Any) -> str:
    """The conversation a document tool was called in.

    agno injects `run_context` into any tool declaring it and strips it from the
    schema the model sees, so this identifier cannot be forged by the model. It
    scopes the picker-selection read-back in slot_resolver: without it, a person
    chosen in one chat is reused in another chat for the same company.

    Empty for non-chat callers (the Fill-in view), which is correct — they pass
    their values explicitly and must not inherit any conversation's picks.
    """
    return str(getattr(run_context, "session_id", "") or "").strip()


# --- Duplicate-generation guard -------------------------------------------
# generate_document had no idea it had just run. Three things re-fire it with
# identical inputs: the frontend's auto-"continue" nudge re-running the tool on
# a silently-stopped run, a user double-clicking Generate, and browser retries.
# Production produced the same Shareholders Resolution twice 19 seconds apart —
# two near-identical .docx files, neither marked authoritative, for a legal
# filing. So a successful generation is remembered briefly and an identical
# repeat returns the SAME file instead of writing a new one.
#
# 120s covers the realistic accidental-repeat window (a resumed run, a retry, an
# impatient second click) while staying far below any legitimate "regenerate the
# same document again" request, which a user only makes minutes later and which
# almost always carries changed data anyway.
_DUPLICATE_WINDOW_SECONDS = 120
# Bounded so a long-lived worker cannot accumulate entries forever; the oldest
# entry is evicted once the cap is hit. Well above the number of documents any
# one session generates inside a 2-minute window.
_DUPLICATE_CACHE_MAX = 64
# key -> (monotonic timestamp, previous successful result dict)
_recent_generations: dict[str, tuple[float, dict[str, Any]]] = {}
# Tool calls arrive on Agno's worker threads, so the dict is touched
# concurrently.
_recent_generations_lock = threading.Lock()


def _generation_key(template_name: str, company_name: str, data: dict[str, Any]) -> str:
    """Identity of a generation: template + company + the data actually filled in.

    Hashing the resolved `data` (DB values already merged with custom_data) —
    not the raw custom_data argument — is what makes this safe: any genuinely
    different input, including a single changed signer, produces a different key
    and therefore always writes a new document. sort_keys makes dict ordering
    irrelevant; default=str keeps dates/Decimals from raising here.
    """
    try:
        blob = _json.dumps(data, sort_keys=True, default=str)
    except (TypeError, ValueError):
        # Un-hashable data means we cannot prove two calls are identical, so
        # fall back to a unique key — i.e. never suppress the generation.
        blob = repr(object())
    digest = hashlib.sha256(blob.encode("utf-8", "replace")).hexdigest()
    return f"{template_name}|{company_name}|{digest}"


def _lookup_recent_generation(key: str) -> dict[str, Any] | None:
    """Return the previous result for `key` if it is still inside the window."""
    now = time.monotonic()
    with _recent_generations_lock:
        # Drop anything expired while we hold the lock — keeps the cache from
        # holding stale results and doubles as cheap maintenance.
        for stale in [k for k, (ts, _) in _recent_generations.items() if now - ts > _DUPLICATE_WINDOW_SECONDS]:
            del _recent_generations[stale]
        entry = _recent_generations.get(key)
        return dict(entry[1]) if entry else None


def _remember_generation(key: str, result: dict[str, Any]) -> None:
    """Cache a SUCCESSFUL generation only — a failure must always be retryable."""
    with _recent_generations_lock:
        if len(_recent_generations) >= _DUPLICATE_CACHE_MAX and key not in _recent_generations:
            oldest = min(_recent_generations, key=lambda k: _recent_generations[k][0])
            del _recent_generations[oldest]
        _recent_generations[key] = (time.monotonic(), dict(result))


def extract_placeholders_from_template(template_path: Path) -> dict[str, Any]:
    """
    Extract all placeholders from a Word template.
    Returns required fields and their locations.
    """
    doc = Document(str(template_path))

    placeholders = {}
    empty_counter = new_empty_counter()

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        matches = PLACEHOLDER_PATTERN.findall(text)
        for match in matches:
            placeholder = placeholder_name(match, empty_counter)
            if placeholder:
                placeholders[placeholder.lower()] = {
                    "field": placeholder,
                    "location": f"paragraph {idx + 1}",
                    "sample_text": text[:100] if text else "",
                }

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text
                matches = PLACEHOLDER_PATTERN.findall(text)
                for match in matches:
                    placeholder = placeholder_name(match, empty_counter)
                    if placeholder:
                        placeholders[placeholder.lower()] = {
                            "field": placeholder,
                            "location": f"table {table_idx + 1}, row {row_idx + 1}, cell {cell_idx + 1}",
                            "sample_text": text[:100] if text else "",
                        }

    return {
        "template": template_path.name,
        "total_placeholders": len(placeholders),
        "fields": list(placeholders.keys()),
        "details": placeholders,
    }


# The placeholder values the pipeline uses to mean "nothing here yet". Compared
# case-insensitively: `prepare_document_data` writes "TBD" for a party position
# the register does not fill, and a TBD is precisely a field that still needs an
# answer, not one that is satisfied.
_EMPTY_VALUES = {"", "tbd", "[tbd]", "[tbd - needs input]", "none", "n/a"}


def _has_value(company_data: dict, *keys: str) -> bool:
    """Whether any spelling of a field carries real data in `company_data`."""
    for key in keys:
        if not key:
            continue
        value = company_data.get(key)
        if value is None:
            continue
        if isinstance(value, (list, dict)):
            if value:
                return True
            continue
        if str(value).strip().lower() not in _EMPTY_VALUES:
            return True
    return False


def validate_data_vs_template(required_fields: list, company_data: dict) -> dict[str, Any]:
    """
    Validate that company data has all required fields for template.
    Returns what's available vs what's missing.
    """
    available_columns = [
        col.lower().replace(" ", "_").replace("-", "_")
        for col in company_data.get("columns", list(company_data.keys()))
    ]

    available_fields = []
    missing_fields = []

    # Fields that genuinely have a default — `find_replacement` returns a real
    # value for these without anyone being asked, so they are legitimately
    # "available" while absent from the company record.
    #
    # This list used to hold twelve more names on the claim that they "always
    # have defaults". They do not. Measured through `find_replacement`:
    # `auditor_name`, `auditor_fee`, `financial_year_end_date`,
    # `next_financial_year_end_date` and `shareholder_name` return the literal
    # string "TBD" (from a set the resolver itself calls KNOWN_USER_INPUT), and
    # `pronoun`, `date`, `nric`, `identification_type`,
    # `authorized_director_name`, `resigning_director_name` and
    # `resigning_director_identification_number` return None, which leaves the
    # RAW PLACEHOLDER in the document with no highlight at all.
    #
    # Either way the field was exempted from the value check, counted as
    # satisfied, and therefore never put to the user — the same defect as the
    # `new_company_name` case, hidden behind a list that asserted the opposite.
    # `tracker_fill` now pins the claim: every member must resolve to a real
    # value, so the two cannot drift apart again.
    # `address` left this set on 2026-08-27. It was here because a bare
    # `address` placeholder always resolved — from the company record's own
    # `address` alias. That alias is now refused for a person's field, which is
    # the point of the fix, so `address` no longer has a default and claiming it
    # does is what kept it out of the missing list and therefore never asked.
    DEFAULT_FIELDS = {"meeting_location"}

    # Alias mapping for matching
    FIELD_ALIASES = {
        "company": "company_name",
        "date": "meeting_date",
        "meeting_location": "registered_office",
        "director_name": "directors",
        "authorized_director_name": "directors",
        "resigning_director_name": "directors",
        "address": "registered_office",
        "nric": "nrc_passport",
        "identification_type": "nrc_passport",
        "shareholder_name": "individual_shareholder_1_name",
    }

    for field in required_fields:
        field_normalized = field.lower().replace(" ", "_").replace("-", "_")

        found = False
        # Direct match
        for col in available_columns:
            if field_normalized in col or col in field_normalized:
                found = True
                break

        # Alias match
        if not found and field_normalized in FIELD_ALIASES:
            alias = FIELD_ALIASES[field_normalized]
            for col in available_columns:
                if alias in col or col in alias:
                    found = True
                    break

        # Has default
        if not found and field_normalized in DEFAULT_FIELDS:
            found = True

        # A NAME match is not data. `found` above is decided purely by spelling,
        # and the match is a SUBSTRING test, so `new_company_name` matched the
        # `company_name` column and was reported as "available / DB" while its
        # value was None — inflating coverage AND, because a resolved field is
        # never asked about, shipping the document with a blank new-company
        # name. A field is only available when it actually carries a value.
        #
        # DEFAULT_FIELDS are exempt: those are filled from defaults further down
        # the pipeline, so they are legitimately available while still empty here.
        if found and field_normalized not in DEFAULT_FIELDS and not _has_value(company_data, field, field_normalized):
            found = False

        if found:
            available_fields.append(field)
        else:
            missing_fields.append(field)

    return {
        "required_fields": required_fields,
        "available_fields": available_fields,
        "missing_fields": missing_fields,
        "available_columns": available_columns,
        "is_complete": len(missing_fields) == 0,
    }


# --- Declared document state ----------------------------------------------
# The document panel used to RECONSTRUCT its state by parsing six different
# result keys and inferring status from `success`/`ready_to_generate`/
# `file_name`. Every tool that shaped its result differently became a panel bug
# (three were hand-fixed downstream before this). The tools know the answer, so
# they now say it: one `document_state` key, identical in shape everywhere,
# built by the single helper below so the shape cannot drift between tools.
#
# Nothing here queries or recomputes — every argument is something the caller
# already had in hand.

_TBD_VALUE_RE = re.compile(r"^\s*(tbd|\[tbd[^\]]*\]|n/?a|-{1,2})\s*$", re.IGNORECASE)


def _norm_key(key: Any) -> str:
    return str(key).strip().lower().replace(" ", "_").replace("-", "_")


def _resolved_values(
    required_fields: list | None,
    data: dict[str, Any],
    template_name: str | None,
    company_name: str | None,
    document_id: Any = None,
) -> tuple[dict[str, str], list[str]]:
    """Resolve every field the template renders, ONCE, the way generation does.

    Three pieces of code used to decide independently whether a field was
    filled: generation (via `find_replacement`), the Fill-in view (also
    `find_replacement`), and the Fields panel — which asked `_resolve_from_data`
    alone and so could not see the People register at all.

    The result on 2026-08-27, on a build where the register fix was already
    live: the question card correctly stopped asking for nationality and date of
    birth, while the panel beside it still showed both as PENDING and displayed
    the COMPANY's registered office under the director's `address`. Nothing was
    stale — the panel was computing a different, wrong answer.

    One resolver now answers for all three. Returns the resolved values and the
    fields that genuinely came back blank.
    """
    resolved: dict[str, str] = {}
    blank: list[str] = []
    for raw in required_fields or []:
        field = str(raw).strip()
        if not field:
            continue
        try:
            value = find_replacement(
                field,
                data,
                template_name=template_name,
                company_name=company_name,
                document_id=document_id,
            )
        except Exception as e:
            logging.getLogger("legalscout").warning(f"_resolved_values: {field!r} failed: {e}")
            value = None
        text = "" if value is None else str(value).strip()
        if not text or text == LEFT_BLANK or _TBD_VALUE_RE.match(text):
            blank.append(field)
        else:
            resolved[field] = text
    return resolved, blank


def _document_state(
    template: str | None,
    company: str | None,
    required_fields: list | None,
    values: dict[str, Any] | None,
    status: str,
    outstanding: list | None = None,
    file_name: str | None = None,
    download_url: str | None = None,
) -> dict[str, Any]:
    """Build the panel's whole state from data the caller already computed.

    `outstanding` names fields a tool has DECLARED unresolved (missing fields,
    unanswered slots, unfilled placeholders after a fill). That verdict beats a
    value we happen to hold — but a held value is still worth showing as an
    unconfirmed candidate, so it degrades to "tbd" rather than hiding behind
    "pending".
    """
    normalised = {_norm_key(k): v for k, v in (values or {}).items()}
    unresolved = {_norm_key(f) for f in (outstanding or [])}

    fields: list[dict[str, Any]] = []
    seen: set[str] = set()
    for raw in required_fields or []:
        key = str(raw).strip()
        norm = _norm_key(key)
        if not norm or norm in seen:
            continue
        seen.add(norm)

        value = normalised.get(norm)
        text = str(value).strip() if value not in (None, "") else ""

        if norm in unresolved:
            state = "tbd" if text else "pending"
        elif not text:
            state = "pending"
        elif _TBD_VALUE_RE.match(text):
            state = "tbd"
        else:
            state = "filled"

        fields.append({"key": key, "value": text or None, "state": state})

    filled = sum(1 for f in fields if f["state"] == "filled")
    return {
        "template": template,
        "company": company,
        "fields": fields,
        "filled": filled,
        "total": len(fields),
        "status": status,
        "file_name": file_name,
        "download_url": download_url,
    }


def find_company_data(company_name: str, companies_data: dict) -> dict[str, Any]:
    """Find specific company data from companies dict."""
    company_name_lower = company_name.lower().strip()

    for company in companies_data.get("companies", []):
        comp_name = str(company.get("company_name", "")).lower().strip()
        if comp_name == company_name_lower or company_name_lower in comp_name:
            return {"found": True, "data": company}

    return {"found": False, "available_companies": [c.get("company_name") for c in companies_data.get("companies", [])]}


def analyze_template(template_name: str, documents_dir: str = "/documents") -> dict[str, Any]:
    """
    Step 1: Analyze template to extract required fields.
    """
    templates_dir = Path(documents_dir) / "legal" / "templates"

    resolved_name, close = _resolve_template_name(template_name, templates_dir)
    if resolved_name is None:
        return {
            "success": False,
            "error": f"Template not found: {template_name}",
            "close_matches": close,
            "available_templates": [f.name for f in templates_dir.glob("*.docx")] if templates_dir.exists() else [],
            "agent_instruction": (
                "Do NOT guess a different template. "
                + (
                    "Ask the user which of these they mean with ask_questions: " + "; ".join(close)
                    if close
                    else "Pick an exact name from available_templates."
                )
            ),
        }
    template_name = resolved_name
    template_path = templates_dir / template_name

    result = extract_placeholders_from_template(template_path)
    result["success"] = True
    result["template_path"] = str(template_path)

    return result


_NO_COMPANY = {
    "success": False,
    "step": "company_lookup",
    "error": "No company specified and none agreed in this conversation.",
    "agent_instruction": (
        "ASK NOW, IN THIS SAME TURN — ending your turn with no text is forbidden. "
        "Ask which company this document is for via ask_questions, offering the "
        "companies as clickable options. Never guess."
    ),
}


def _resolve_company_arg(company_name, session_id):
    """The company a document tool should act on, or "" when it is unknown.

    NEVER falls through to a lookup with an empty name: `prepare_document_data`
    treats "" as a fuzzy query and returns whichever company happens to match
    first — measured, an empty name resolved to an unrelated client and reported
    success. Making `company_name` optional (so the session binding could supply
    it) put that path within easy reach, so the empty case is refused here.
    """
    named = (company_name or "").strip()
    if named:
        return named
    return _company_from_binding(session_id)


def _company_from_binding(session_id: str) -> str:
    """The company this conversation was already agreed to be about.

    Consulted ONLY when the caller supplied no company name. A conversation
    settles its subject once — at the picker — and every later turn used to
    re-derive it from whatever the model happened to type, which is why the
    agent re-asked which company mid-flow after it had already been answered.

    Returns "" whenever the memory layer is off, nothing is bound, or the bound
    company cannot be named — every caller then behaves exactly as before.
    """
    key = str(session_id or "").strip()
    if not key:
        return ""
    try:
        from scout.memory import RESOLVED, resolve_for_session

        bound = resolve_for_session(key)
        if bound.status != RESOLVED:
            return ""
        return str(getattr(bound, "company_name", "") or "").strip()
    except Exception as e:
        _log_binding_warning(key, e)
        return ""


def _log_binding_warning(key, exc):
    import logging

    logging.getLogger("legalscout").warning(f"session binding lookup failed for {key!r}: {exc}")


# ---------------------------------------------------------------------------
# Numeric presentation
# ---------------------------------------------------------------------------
# The model routinely hands back a bare integer for a figure the user typed with
# separators: asked for "50,000,000 MMK, 10,000 shares, 100% ownership", it
# passed 50000000, 10000 and 100. The document then reads
#
#     10000 Ordinary Shares      50000000 MMK      100
#
# which is not how a Myanmar corporate instrument states a sum. Nothing strips
# the commas — they are simply never sent — so re-applying presentation at fill
# time is the only place this can be corrected once, for every template.
#
# Deliberately conservative. It fires ONLY on a bare integer (or a plain decimal
# for a percentage), and only for a field whose name says what the number means.
# Anything already carrying a separator, a currency word or any other character
# is left exactly as supplied.

# Checked FIRST and unconditionally. These are identifiers that happen to be
# digits: a registration number with thousands separators is a different number,
# and an NRC or a year with a comma in it is simply wrong.
_NOT_A_QUANTITY_RE = re.compile(
    r"registration|regno|reg_no|nrc|passport|phone|mobile|fax|postal|zip"
    r"|year|clause|section|article|serial|invoice|account|licen[cs]e|tin"
)
_MONEY_OR_COUNT_RE = re.compile(
    r"amount|subscription|paid|capital|consideration|value|price|fee|sum"
    r"|shares?|share_count|no_of_share|number_of_share"
)
_PERCENT_RE = re.compile(r"percent|percentage|shareholding|ownership|holding_pct")

_BARE_INT_RE = re.compile(r"^\d{4,}$")

# A bare `date` placeholder, answered under a QUALIFIED key.
#
# Measured on the client's tracker (AGM Minutes): the agent asked
# "meeting_date", the user answered "07 July 2027", and the document came out
# with `Date:` EMPTY — the template's placeholder is `date`, and no tier of
# _resolve_from_data matches "meeting_date" to "date". The same template on the
# next run asked under `date` and filled correctly, so the value was lost purely
# because the model chose an id more specific than the placeholder.
#
# Deliberately a CLOSED list rather than "any key ending in _date". The obvious
# general rule — any key whose tokens include "date" — would let
# `date_of_birth` fill the date of a legal instrument, which is a worse
# document than a blank one.
_QUALIFIED_DATE_RE = re.compile(
    r"^(?:meeting|signing|sign|signature|notice|effective|resolution|consent"
    r"|document|letter|minutes|agm|appointment|resignation)_dates?$"
)
_BARE_PERCENT_RE = re.compile(r"^\d{1,3}(?:\.\d+)?$")


def _present_number(field: str, value: Any) -> Any:
    """Restore the separators / % a figure is written with, when it is safe to.

    Returns `value` untouched unless the field NAMES a quantity and the value is
    a bare number. Identifier-shaped fields are excluded first, so a company
    registration number can never acquire a comma.
    """
    if value is None or isinstance(value, bool):
        return value
    text = str(value).strip()
    if not text:
        return value
    key = canonical_field(field) or normalize_field(field) or str(field).lower()
    if _NOT_A_QUANTITY_RE.search(key):
        return value
    if _PERCENT_RE.search(key) and _BARE_PERCENT_RE.match(text):
        # "100" → "100%". A percentage sign is not decoration in a shareholding
        # table; without it the cell states a share COUNT.
        return f"{text.rstrip('.')}%"
    if _MONEY_OR_COUNT_RE.search(key) and _BARE_INT_RE.match(text):
        return f"{int(text):,}"
    return value


def _present_numbers(data: dict[str, Any]) -> dict[str, Any]:
    """_present_number across a whole data dict."""
    return {k: _present_number(k, v) for k, v in data.items()}


def _normalise_template_key(name: str) -> str:
    """Lowercase, drop the extension, and collapse punctuation to single spaces.

    "Corporate_Shareholder_Consent-Directors Resolution.docx" and
    "corporate shareholder consent - directors resolution" become the same key.
    """
    stem = re.sub(r"\.docx$", "", str(name or "").strip(), flags=re.IGNORECASE)
    return re.sub(r"[^a-z0-9]+", " ", stem.lower()).strip()


def _resolve_template_name(template_name: str, templates_dir: Path) -> tuple[str | None, list[str]]:
    """Map a requested template name onto a real file.

    Returns (filename, []) when it resolves, or (None, close_matches) when it
    does not. A match is returned ONLY when it is unique: two candidates mean
    the caller has not said which document it wants, and picking one is how a
    Director Resignation gets written on an AGM Minutes form.

    Tiers, most literal first: exact filename, the historical
    underscore/space variants, case- and punctuation-insensitive equality,
    then unique prefix, then unique substring.
    """
    if not templates_dir.is_dir():
        return None, []

    candidates = [p.name for p in templates_dir.glob("*.docx") if not p.name.startswith("~$")]

    # Membership in the real listing, NOT Path.is_file(). macOS is
    # case-insensitive, so `(dir / "director resignation letter.docx").is_file()`
    # is True there and this would return the caller's spelling — a name that
    # does not exist inside the Linux container the product actually runs in.
    # The listing gives back the file's true name in every case.
    existing = set(candidates)
    for literal in (
        template_name,
        template_name.replace("_", " "),
        template_name.replace(" ", "_"),
        f"{template_name}.docx",
        f"{template_name.replace('_', ' ')}.docx",
    ):
        if literal and literal in existing:
            return literal, []

    key = _normalise_template_key(template_name)
    if not key:
        return None, []

    by_key = [(c, _normalise_template_key(c)) for c in candidates]

    exact = [c for c, k in by_key if k == key]
    if len(exact) == 1:
        return exact[0], []
    if len(exact) > 1:
        return None, sorted(exact)

    prefix = [c for c, k in by_key if k.startswith(key)]
    if len(prefix) == 1:
        return prefix[0], []
    if len(prefix) > 1:
        return None, sorted(prefix)

    substring = [c for c, k in by_key if key in k]
    if len(substring) == 1:
        return substring[0], []
    return None, sorted(substring)


def prepare_document_data(template_name: str, company_name: str, documents_dir: str = "/documents") -> dict[str, Any]:
    """
    Complete workflow: Analyze → Validate → Prepare data for document.

    Returns:
        - Template analysis (required fields)
        - Company data from database
        - Validation (what's missing)
        - Ready status
    """
    templates_dir = Path(documents_dir) / "legal" / "templates"

    resolved_name, ambiguous = _resolve_template_name(template_name, templates_dir)
    if resolved_name is None:
        # Naming the candidates is the whole point. The old error said only
        # "Template not found: <name>" — and a model given no near match does
        # the worst possible thing, which is guess. Measured: asking for
        # "Corporate Shareholder Consent - Directors Resolution" (a true prefix
        # of the real filename) produced a fully-formed set of ANNUAL GENERAL
        # MEETING MINUTES instead, with none of the requested terms in it.
        return {
            "success": False,
            "error": f"Template not found: {template_name}",
            "step": "template_analysis",
            "close_matches": ambiguous,
            "available_templates": sorted(p.name for p in templates_dir.glob("*.docx") if not p.name.startswith("~$")),
            "agent_instruction": (
                "Do NOT guess a different template and do NOT generate anything. "
                + (
                    "Several templates match that name — ask the user which one with "
                    "ask_questions, using these as the options: " + "; ".join(ambiguous)
                    if ambiguous
                    else "Choose the correct name from available_templates, or ask the "
                    "user with ask_questions if none clearly matches."
                )
            ),
        }
    template_name = resolved_name
    template_path = templates_dir / template_name

    template_analysis = extract_placeholders_from_template(template_path)
    required_fields = template_analysis.get("fields", [])

    # Load company data from database
    company_result = None
    companies_data = {"companies": []}
    try:
        from scout.tools.companies_db import get_all_companies

        db_companies = get_all_companies(limit=200)
        if db_companies:
            companies_list = []
            for c in db_companies:
                director_names = c.get("directors", "")
                shareholder_names = c.get("shareholders", "")
                shareholders_list = c.get("shareholders_list", [])
                dirs_split = [d.strip() for d in director_names.split(",") if d.strip()] if director_names else []
                shs_split = [s.strip() for s in shareholder_names.split(",") if s.strip()] if shareholder_names else []

                sh_shares = {}
                for m in shareholders_list:
                    if isinstance(m, dict):
                        sh_shares[m.get("name", "")] = m.get("shares", "TBD")

                # Member slots used to be assigned by POSITION: slots 1-2 of the
                # flat, comma-joined name list were assumed individual and slot 3
                # assumed corporate. Nothing consulted member type, though the
                # structured `shareholders_list` sitting two lines above carries
                # it. On CITY MART HOLDING COMPANY LIMITED — whose single member
                # is the corporate CITY HOLDINGS LIMITED — that company landed at
                # index 0 and was written into `individual_shareholder_1_name`,
                # so a company was rendered as an individual member. That is
                # legally wrong: a corporate member signs through a named
                # representative, an individual signs personally.
                # Type comes from repeat_regions._is_corporate so this file and
                # the list expander cannot disagree about the same member (real
                # DICA data spells the corporate type "Company", not "corporate").
                # Falls back to the old positional split when a row has no
                # structured list, which older companies do not.
                ind_split = shs_split
                corp_split = shs_split[2:]
                typed_members = [
                    m
                    for m in shareholders_list
                    if isinstance(m, dict) and str(m.get("name") or m.get("full_name") or "").strip()
                ]
                if typed_members:
                    ind_split, corp_split = [], []
                    for m in typed_members:
                        nm = str(m.get("name") or m.get("full_name") or "").strip()
                        (corp_split if _is_corporate(m) else ind_split).append(nm)

                companies_list.append(
                    {
                        "company_name": c.get("name", ""),
                        "company": c.get("name", ""),
                        "company_name_english": c.get("name", ""),
                        "company_registration_number": c.get("registration_number", ""),
                        "registered_office": c.get("address", ""),
                        "registered_office_address": c.get("address", ""),
                        "company_address": c.get("address", ""),
                        "address": c.get("address", ""),
                        "meeting_location": c.get("address", ""),
                        "directors": director_names,
                        "director_name": dirs_split[0] if dirs_split else "TBD",
                        # `authorized_director_name` / `authorized director_name` are
                        # NOT seeded here any more. They are the representative slot
                        # — the person signing on behalf of a CORPORATE MEMBER — and
                        # seeding them from this company's directors[0] was wrong
                        # twice over.
                        #
                        # It named the wrong board. A corporate member is signed for
                        # by ITS OWN directors, so for CM FOODS (sole member CITY
                        # MART HOLDING) the value came from CM FOODS' register, not
                        # CITY MART's.
                        #
                        # And it suppressed the question. resolve_slot treats a value
                        # present in `data` as an answer, so collect_slot_requests
                        # stopped asking: measured on CM FOODS, the requests went
                        # from ['new_director', 'representative'] with empty data to
                        # ['new_director'] once the company defaults were merged in.
                        # The user was never offered the choice, and directors[0] of
                        # the wrong company signed. That is the same shape as the
                        # slot_resolver fallback that was removed for guessing.
                        #
                        # Unseeded, the slot resolves to nothing, the picker is
                        # raised against the MEMBER company's board, and
                        # repeat_regions._corp_representative vets whatever comes
                        # back against that board before printing it.
                        "individual_shareholder_1_name": ind_split[0] if len(ind_split) > 0 else "TBD",
                        "individual_shareholder_2_name": ind_split[1] if len(ind_split) > 1 else "TBD",
                        "individual shareholder_1_name": ind_split[0] if len(ind_split) > 0 else "TBD",
                        "individual shareholder_2_name": ind_split[1] if len(ind_split) > 1 else "TBD",
                        "corporate_shareholder_3_name": corp_split[0] if len(corp_split) > 0 else "TBD",
                        "corporate shareholder_3_name": corp_split[0] if len(corp_split) > 0 else "TBD",
                        "shareholder_1_name": shs_split[0] if len(shs_split) > 0 else "TBD",
                        "shareholder_2_name": shs_split[1] if len(shs_split) > 1 else "TBD",
                        "shareholder_1_shares": sh_shares.get(shs_split[0], "TBD") if shs_split else "TBD",
                        "shareholder_2_shares": sh_shares.get(shs_split[1], "TBD") if len(shs_split) > 1 else "TBD",
                        "status": c.get("status", ""),
                        "company_type": c.get("company_type", ""),
                        "total_shares": c.get("total_shares", ""),
                        "currency": c.get("currency", ""),
                        "financial_year_end_date": c.get("financial_year_end_date", ""),
                        "next_financial_year_end_date": c.get("next_financial_year_end_date", ""),
                        "auditor_name": c.get("auditor_name", ""),
                        "auditor_fee": c.get("auditor_fee", ""),
                        **(c.get("custom_fields") or {}),
                    }
                )
            companies_data = {"success": True, "companies": companies_list}
            company_result = find_company_data(company_name, companies_data)
    except Exception as e:
        logging.getLogger("legalscout").warning(f"Failed to load company data from DB: {e}")

    if not company_result:
        return {
            "success": False,
            "error": "No company data in database. Add companies from the admin panel.",
            "step": "data_loading",
            "required_fields": required_fields,
        }

    if not company_result.get("found"):
        return {
            "success": False,
            "error": f"Company '{company_name}' not found in database",
            "step": "company_lookup",
            "available_companies": company_result.get("available_companies", []),
            "required_fields": required_fields,
        }

    company_data = company_result.get("data", {})
    validation = validate_data_vs_template(required_fields, company_data)

    # Get template info from database (trained data)
    template_info = {}
    try:
        from scout.tools.template_analyzer import get_template_info

        template_data = get_template_info(template_name)
        if template_data:
            template_info = {
                "purpose": template_data.get("purpose", ""),
                "when_to_use": template_data.get("when_to_use", ""),
                "how_to_use": template_data.get("how_to_use", []),
            }
    except Exception as e:
        logging.getLogger("legalscout").warning(f"Failed to load template info: {e}")

    normalized_company_data = {}
    for key, value in company_data.items():
        key_normalized = str(key).lower().replace(" ", "_").replace("-", "_")
        normalized_company_data[key_normalized] = value

    slot_requests = collect_slot_requests(
        _get_field_mapping(template_name) or {},
        required_fields,
        normalized_company_data,
        company_name=company_name,
    )

    # What the panel calls outstanding is decided by the SAME resolver that
    # generation uses, not by `validate_template_data` — which compares field
    # names against company columns and so cannot see the People register at
    # all. Left as it was, the panel reported a director's nationality, NRC and
    # date of birth as pending while generation filled all three from the
    # register, and the question card had already stopped asking for them.
    _resolved_map, _resolved_blank = _resolved_values(
        required_fields, normalized_company_data, template_name, company_name
    )
    _slot_placeholders = [r["placeholder"] for r in slot_requests]
    _slot_canon = {canonical_field(p) for p in _slot_placeholders}
    outstanding = [f for f in _resolved_blank if canonical_field(f) not in _slot_canon] + _slot_placeholders
    ready = not outstanding

    return {
        "success": True,
        "step": "complete",
        "slot_requests": slot_requests,
        "unresolved_slots": [r["placeholder"] for r in slot_requests],
        "document_state": _document_state(
            template=template_name,
            company=company_name,
            required_fields=required_fields,
            # The RESOLVED values, not the raw company record. That record
            # publishes a bare `address` alias holding the company's registered
            # office, which the panel then showed against the DIRECTOR's home
            # address — the one reading the resolver had just refused.
            values=_resolved_map,
            # A slot the user has not picked, or a field the template needs and
            # the record does not carry, is input we are waiting on — not an
            # error and not "ready".
            status="ready" if ready else "awaiting-input",
            outstanding=outstanding,
        ),
        "template_analysis": {
            "template": template_name,
            "required_fields": required_fields,
            "total_fields": len(required_fields),
            **template_info,
        },
        "company_data": company_data,
        "normalized_data": normalized_company_data,
        "validation": validation,
        "ready_to_generate": ready,
        "message": _prepare_message(validation, slot_requests),
    }


def _prepare_message(validation: dict, slot_requests: list) -> str:
    """Separate 'ask the picker for a party' from 'ask the user for a value'."""
    parts = []
    if slot_requests:
        pickers = ", ".join(sorted({f"{r['picker']}({r['kind']})" for r in slot_requests}))
        parts.append(f"Choose the parties for these roles with the person pickers: {pickers}")
    if validation.get("missing_fields"):
        parts.append(f"Missing fields: {', '.join(validation['missing_fields'])}")
    return " | ".join(parts) if parts else "Ready to generate document"


def validate_filled_document(document_path: Path, all_placeholders: list | None = None) -> dict[str, Any]:
    """
    Step 4: Validate that a filled document has no unfilled placeholders.

    Args:
        document_path: Path to the filled document
        all_placeholders: List of all placeholders from the original template
    """
    if not document_path.exists():
        return {"success": False, "error": "Document not found"}

    doc = Document(str(document_path))

    # Find unfilled placeholders in the filled document
    unfilled_placeholders = []
    unfilled_locations = []

    for paragraph in doc.paragraphs:
        text = paragraph.text
        matches = PLACEHOLDER_PATTERN.findall(text)
        for match in matches:
            placeholder = placeholder_name(match)
            if placeholder and placeholder.lower() not in ["", " ", "null", "none"]:
                unfilled_placeholders.append(placeholder)
                unfilled_locations.append(f"Paragraph: {text[:50]}...")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                matches = PLACEHOLDER_PATTERN.findall(text)
                for match in matches:
                    placeholder = placeholder_name(match)
                    if placeholder and placeholder.lower() not in ["", " ", "null", "none"]:
                        unfilled_placeholders.append(placeholder)
                        unfilled_locations.append(f"Table: {text[:50]}...")

    # Calculate filled placeholders
    if all_placeholders:
        unfilled_set = {p.lower() for p in unfilled_placeholders}
        filled_placeholders = [p for p in all_placeholders if p.lower() not in unfilled_set]
        total_placeholders = len(all_placeholders)
        total_filled = len(filled_placeholders)
    else:
        # Fallback if all_placeholders not provided
        filled_placeholders = []
        total_placeholders = len(unfilled_placeholders)
        total_filled = 0

    unique_unfilled = list(set(unfilled_placeholders))
    return {
        "success": True,
        "is_valid": len(unique_unfilled) == 0,
        "all_placeholders": all_placeholders or [],
        "filled_placeholders": filled_placeholders,
        "unfilled_placeholders": unfilled_locations,
        "unfilled_names": unique_unfilled,
        "total_placeholders": total_placeholders,
        "total_filled": total_filled,
        "total_unfilled": len(unique_unfilled),
        "message": f"Filled {total_filled}/{total_placeholders} fields"
        if all_placeholders
        else (
            "Document is valid - all fields filled"
            if len(unfilled_placeholders) == 0
            else f"Warning: {len(unfilled_placeholders)} unfilled placeholders found"
        ),
    }


def create_smart_document_tool(documents_dir: str = "/documents", host: str = ""):
    """Create the smart document workflow tool."""

    def analyze_template_tool(template_name: str) -> dict[str, Any]:
        """Analyze a template to see what fields are required."""
        return analyze_template(template_name, documents_dir)

    def prepare_document(template_name: str, company_name: str = "", run_context: RunContext = None) -> dict[str, Any]:
        """Complete workflow: analyze, validate, prepare data."""
        session_id = _session_of(run_context)
        with session_scope(session_id):
            resolved = _resolve_company_arg(company_name, session_id)
            if not resolved:
                return dict(_NO_COMPANY)
            return prepare_document_data(template_name, resolved, documents_dir)

    def generate_document(
        template_name: str,
        company_name: str = "",
        custom_data: dict | None = None,
        run_context: RunContext = None,
    ) -> dict[str, Any]:
        """Generate document with validation workflow."""
        session_id = _session_of(run_context)
        with session_scope(session_id):
            resolved = _resolve_company_arg(company_name, session_id)
            if not resolved:
                return dict(_NO_COMPANY)

            # ★★★ Refuse a template that CONTRADICTS what was asked for.
            #
            # The worst thing this tool has done is generate the wrong legal
            # instrument — an Individual Shareholder Consent for a request that
            # said "director consent form (non group member)". Roughly two runs
            # in eleven on the client's tracker. It is filled correctly and
            # looks right, which is exactly why nobody catches it.
            #
            # The check is one-sided on purpose: it does not choose a template,
            # it only blocks one that states the opposite of the request on an
            # axis where the two are different instruments. Silence on either
            # side is never a contradiction, so an ambiguous request behaves
            # exactly as before.
            drift = _check_template_drift(session_id, template_name)
            if drift:
                return drift

            return _generate_document(template_name, resolved, custom_data)

    def _generate_document(template_name: str, company_name: str, custom_data: dict | None = None) -> dict[str, Any]:
        # Collect the placeholders this fill empties, so validation can report
        # them. Reset in the finally at the end of the generation so a blank
        # from one document is never attributed to the next.
        _blank_token = _blanked_placeholders.set(set())
        try:
            return _generate_document_inner(template_name, company_name, custom_data)
        finally:
            _blanked_placeholders.reset(_blank_token)

    def _generate_document_inner(
        template_name: str, company_name: str, custom_data: dict | None = None
    ) -> dict[str, Any]:
        result = prepare_document_data(template_name, company_name, documents_dir)

        if not result.get("success"):
            return result

        # prepare_document_data resolves the requested name onto a real file
        # (a prefix, a case variant, different punctuation). Everything below
        # keys off the template NAME — the field mapping, the DB field
        # classification, and the path the .docx is opened from — so adopt the
        # resolved name here. Without this the request succeeds analysis and
        # then dies on "Template not found" one screen later.
        template_name = (result.get("template_analysis") or {}).get("template") or template_name

        # The goal goes on the server the moment it is known. Measured on
        # session 3be25e3c: between a pause and its resume the agent lost track
        # of what it was making and started over. Nothing remembered.
        _session = current_session_scope()
        try:
            from scout.tools.task_memory import remember_task

            remember_task(_session, template_name, company_name, custom_data)
        except Exception as _e:
            logging.getLogger("legalscout").warning(f"[TASK] record skipped: {_e}")

        mapping = _get_field_mapping(template_name) or {}

        # Values collected EARLIER in this conversation are merged back in.
        #
        # `custom_data` does not accumulate between generate_document calls, so
        # anything the model omits from this call would be blank in the
        # document — the defect behind 1.2.14 and 1.2.18. Restating that rule in
        # prose did not hold; carrying the values server-side removes the
        # requirement instead. The CALLER still wins on every key it sends, so a
        # correction is never overwritten by a stale remembered value.
        try:
            from scout.tools.task_memory import recall_task

            _remembered = (recall_task(_session) or {}).get("collected") or {}
        except Exception:
            _remembered = {}
        if _remembered:
            # An EMPTY caller value does not beat a remembered one.
            #
            # The caller winning on every key it sends was right for
            # corrections and wrong for omissions, and the model omits by
            # sending "" rather than by leaving the key out. Measured on the
            # flow a real user takes — answer the cards, then send "now
            # generate" as a separate message — `generate_document` was called
            # with:
            #
            #     {"date": "", "date_of_birth": "", "phone": "", ...}
            #
            # while active_task.collected held {"date": "07 July 2027"} from the
            # card the user had just filled in. The empty string won and the
            # consent form came out with no date. 2 runs in 5.
            #
            # `_explicitly_supplied` still treats PRESENCE as the signal, so a
            # deliberate "leave the effective date blank" is still honoured —
            # that path never had a remembered value to lose, because
            # merge_collected only ever stores what the user actually typed
            # (task_memory.py: `str(v or "").strip()`). The one case this
            # changes is a user who supplies a value and later blanks the same
            # field in the same conversation; they get their earlier value back
            # and must say so again. That is a far smaller loss than silently
            # dropping what they typed, which is the client's first logged
            # defect.
            _caller = custom_data or {}
            _kept = {
                k: v for k, v in _caller.items() if str(v or "").strip() or not str(_remembered.get(k) or "").strip()
            }
            custom_data = {**_remembered, **_kept}

        # A slot names a PERSON the user must choose. prepare_document_data
        # pre-fills those placeholders from the company record (directors[0] and
        # friends), which would make an unanswered slot look already resolved —
        # the picker would then never be asked, and a selection the user did make
        # would be silently overwritten by the first director on file. So only an
        # explicit selection counts here: drop the auto-filled slot placeholders
        # and let custom_data (which carries the picker's answer) speak alone.
        slot_data = {
            key: value for key, value in (result.get("normalized_data") or {}).items() if not slot_of(mapping.get(key))
        }
        if custom_data:
            slot_data.update(custom_data)
        slot_requests = collect_slot_requests(
            mapping,
            result.get("template_analysis", {}).get("required_fields", []),
            slot_data,
            company_name=company_name,
        )
        if slot_requests:
            # `message` is rendered verbatim to the user in the document panel
            # (agent-ui useArtifact.ts), so it has to read as plain English.
            # The plumbing the model needs — which picker, which key to pass
            # back — goes in agent_instruction, which the panel does not show.
            def _role(request: dict) -> str:
                role = (request["kind"] or "person").replace("_", " ")
                scope = request["candidates_from"]
                return f"{role} for {scope}" if scope else role

            return {
                "success": False,
                "error": "Need party selection for role slots",
                "slot_requests": slot_requests,
                "unresolved_slots": [r["placeholder"] for r in slot_requests],
                # success=False here means "waiting on the user", not "failed".
                # Declaring the state is what stops the panel inferring an error
                # pane from the flag alone.
                "document_state": _document_state(
                    template=template_name,
                    company=company_name,
                    required_fields=result.get("template_analysis", {}).get("required_fields", []),
                    values=_resolved_values(
                        result.get("template_analysis", {}).get("required_fields", []),
                        slot_data,
                        template_name,
                        company_name,
                    )[0],
                    status="awaiting-input",
                    outstanding=[r["placeholder"] for r in slot_requests],
                ),
                # Deliberately does not name a place ("the card in the chat"):
                # the same result is shown in the chat picker AND in the Fill-in
                # panel, and naming the wrong one reads as a broken instruction.
                "message": "Choose the "
                + ", ".join(_role(r) for r in slot_requests)
                + " before generating this document.",
                "agent_instruction": "ACT NOW, IN THIS SAME TURN — do not end your turn empty. Call the first lookup tool immediately, then its picker: "
                + ", ".join(
                    f"{r['lookup_tool']} → {r['picker']} for {r['kind']}"
                    + (f" of {r['candidates_from']}" if r["candidates_from"] else "")
                    for r in slot_requests
                )
                + ". Once the user has picked, call generate_document again with "
                + "custom_data={"
                + ", ".join(f'"{r["placeholder"]}": "<chosen name>"' for r in slot_requests)
                + "} PLUS every other value the user has already given you in this "
                + "conversation (company names, sectors, amounts, share counts, "
                + "percentages, dates). custom_data does NOT accumulate between calls — "
                + "whatever you omit is blank in the document.",
            }

        # Smart default: If 90%+ fields available, proceed automatically
        validation = result.get("validation", {})
        # Use available_fields (correct key) instead of matched_fields
        total_fields = len(validation.get("available_fields", [])) + len(validation.get("missing_fields", []))
        matched_fields = len(validation.get("available_fields", []))

        # Calculate percentage
        matched_fields / total_fields if total_fields > 0 else 0

        # Always auto-proceed — use defaults/TBD for missing fields

        # --- Load field classification from DB ---
        field_classification = None
        try:
            from scout.tools.template_analyzer import get_template_from_db

            tpl_data = get_template_from_db(template_name)
            if tpl_data and isinstance(tpl_data.get("fields"), dict):
                field_classification = tpl_data["fields"].get("field_classification")
        except Exception as e:
            logging.getLogger("legalscout").warning(f"Failed to load field classification for '{template_name}': {e}")

        # If classification exists, separate user_input vs db fields
        #
        # This gate used to read `if field_classification and not custom_data`,
        # which skipped it ENTIRELY the moment any custom_data was present. That
        # is exactly the second call the party-slot guard above asks for: it
        # tells the model to "call generate_document again with
        # custom_data={<slot>: <chosen name>}", the model complies with the
        # party names ALONE, custom_data becomes truthy — and every remaining
        # user-input field slips past unchecked and resolves to LEFT_BLANK.
        #
        # The document that comes out is not visibly broken. No {{placeholder}},
        # no TBD, just a resolution incorporating a company that is never named,
        # for a sum never stated: "hereinafter referred to as  ("NewCo")".
        # Measured at roughly 1 run in 4 on the Corporate Shareholder Consent,
        # on 1.2.12 as well as 1.2.13 — it presents as model flakiness, but the
        # hole is here, and whether it opens depends only on whether the model
        # happened to re-send the other values it had already been given.
        #
        # So the check now runs ALWAYS, and asks the honest question: is this
        # field resolvable from the DB record MERGED WITH custom_data, or did
        # the caller deliberately set it? Emptiness of custom_data was never
        # what mattered.
        if field_classification:
            user_input_fields = field_classification.get("user_input_fields", [])
            db_fields_list = field_classification.get("db_fields", [])
            field_descriptions = field_classification.get("field_descriptions", {})

            # Only ask for fields the template actually renders
            template_fields = {
                canonical_field(f) for f in result.get("template_analysis", {}).get("required_fields", [])
            }

            # Check which user_input_fields are actually missing from data
            normalized_data = result.get("normalized_data", {})
            # The DB record is the floor; custom_data is what the caller has
            # collected on top of it. A field is missing only if NEITHER holds
            # it and the caller did not deliberately blank it.
            merged_data = dict(normalized_data)
            if custom_data:
                merged_data.update(custom_data)
            # Resolve the whole template ONCE, through the same resolver
            # generation uses. `_resolve_from_data` alone used to answer this
            # question, and it cannot see the People register — so a director's
            # nationality, NRC and date of birth were reported missing while
            # generation filled all three, and the panel showed the COMPANY's
            # office under the director's `address` because the raw company
            # record publishes a bare `address` alias.
            required_all = list(result.get("template_analysis", {}).get("required_fields", []))
            # No document row exists yet at this point in the flow — the
            # selections are read back by company and session scope instead.
            resolved_map, resolved_blank = _resolved_values(
                required_all,
                merged_data,
                template_name,
                company_name,
            )
            blank_canon = {canonical_field(f) for f in resolved_blank}

            missing_user_fields = []
            for uf in user_input_fields:
                if template_fields and canonical_field(uf) not in template_fields:
                    continue
                if canonical_field(uf) not in blank_canon:
                    continue
                if _explicitly_supplied(uf, custom_data):
                    continue
                missing_user_fields.append(uf)

            # Auto-filled DB fields
            db_fields_filled = []
            for df in db_fields_list:
                df_norm = df.lower().replace(" ", "_").replace("-", "_")
                if normalized_data.get(df_norm):
                    db_fields_filled.append(df)

            # Generate smart defaults for missing fields
            field_defaults = {}
            for mf in missing_user_fields:
                mf_lower = mf.lower()
                if "location" in mf_lower or "venue" in mf_lower:
                    field_defaults[mf] = normalized_data.get(
                        "registered_office_address", normalized_data.get("registered_office", "TBD")
                    )
                else:
                    field_defaults[mf] = "TBD"

            # If there are missing user input fields, return them for the chat to show a form
            if missing_user_fields:
                return {
                    "success": False,
                    "error": "Need user input for some fields",
                    "user_input_fields": missing_user_fields,
                    "field_descriptions": {f: field_descriptions.get(f, "") for f in missing_user_fields},
                    "field_defaults": field_defaults,
                    "db_fields_filled": db_fields_filled,
                    "static_text_warnings": field_classification.get("static_text_warnings", []),
                    "message": "These fields need your input (company data is already filled)",
                    # Without this the model reads success=False as a failure and
                    # either stops or retries the identical call. Naming the next
                    # action — and that the NEXT call must carry every value it
                    # already holds — is what closes the loop.
                    "agent_instruction": (
                        "ACT NOW, IN THIS SAME TURN — do not end your turn empty and do "
                        "not repeat this call unchanged. If the user already gave any of "
                        "these values earlier in the conversation, use them. For the rest, "
                        "call ask_questions with one question per field (use "
                        '"input_type": "date" for any date). Then call generate_document '
                        "again with custom_data carrying EVERY value you hold — the party "
                        "selections you passed before AND these fields. custom_data does "
                        "not accumulate between calls: anything you leave out of the next "
                        "call is BLANK in the document, and a blank legal term is not "
                        "visible in the finished file. Missing now: " + ", ".join(missing_user_fields)
                    ),
                    "document_state": _document_state(
                        template=template_name,
                        company=company_name,
                        required_fields=result.get("template_analysis", {}).get("required_fields", []),
                        # The smart defaults are candidates, not answers: they
                        # show as "tbd" because every one of these keys is also
                        # listed as outstanding.
                        #
                        # `resolved_map`, NOT `normalized_data`. The raw company
                        # record publishes a bare `address` alias holding the
                        # COMPANY's registered office, so passing it here put
                        # that office against the director's home address in the
                        # panel — the one field the resolver had just refused to
                        # answer that way. What the panel shows is now exactly
                        # what generation will write.
                        values={**resolved_map, **field_defaults},
                        status="awaiting-input",
                        outstanding=missing_user_fields,
                    ),
                }

        # Auto-generate defaults for missing fields
        if not result.get("ready_to_generate") and not custom_data:
            custom_data = {}
            for field in validation.get("missing_fields", []):
                field_lower = field.lower()
                if "location" in field_lower or "venue" in field_lower:
                    custom_data[field] = result.get("normalized_data", {}).get("registered_office_address", "TBD")
                else:
                    custom_data[field] = "TBD"
            result["message"] = f"Auto-filling {len(validation.get('missing_fields', []))} fields with defaults"

        template_path = Path(documents_dir) / "legal" / "templates" / template_name
        if not template_path.exists():
            return {"success": False, "error": "Template not found"}

        data = result.get("normalized_data", {})
        # Company-identity scalars the model must never overwrite via custom_data.
        PROTECTED_FIELDS = {
            "company_registration_number",
            "company_name",
            "company_name_english",
            "status",
            "company_type",
            "registered_office",
            "registered_office_address",
        }
        # Party lists (directors/members/shareholders) ARE legitimate custom_data —
        # repeat_regions expands the "Present:" list and signing blocks from them
        # (repeat_regions.py:_parties_for_family). Blanket-stripping them was
        # dropping the caller's attendee/member selection and forcing a silent DB
        # fallback (finding F3). Only block a SCALAR overwrite of these company
        # lists; a real list passes through untouched.
        PROTECTED_UNLESS_LIST = {"directors", "members", "shareholders"}
        if custom_data:
            safe_custom = {}
            for k, v in custom_data.items():
                kl = k.lower()
                if kl in PROTECTED_FIELDS:
                    continue
                if kl in PROTECTED_UNLESS_LIST and not isinstance(v, list):
                    continue
                safe_custom[k] = v
            data.update(safe_custom)

        # Restore separators BEFORE the dedupe key is computed, so the hash sees
        # the same values that reach the document and a re-send of "50000000"
        # is recognised as the same generation as "50,000,000".
        data = _present_numbers(data)

        # ---- A company that does not exist yet -----------------------------
        #
        # Measured against the client's testing tracker (case "Setup 1"): the
        # agent correctly ASKED "What is the name of the new company?", the user
        # answered "ZENITH ORCHID VENTURES LIMITED" — and the consent form came
        # out reading
        #
        #     To: CITY HOLDINGS LIMITED
        #     Company Registration Number: 119510619
        #
        # The answer was collected and then silently discarded: `company_name`
        # is in PROTECTED_FIELDS, so custom_data carrying it is stripped above.
        # The form then kept the register company used for the lookup, which is
        # worse than a blank — a signed consent would attach the director to a
        # DIFFERENT client entity, complete with that entity's real
        # registration number.
        #
        # PROTECTED_FIELDS stays as it is: it stops the model overwriting
        # company identity on an ordinary document, which is a real protection.
        # The escape hatch is explicit instead — a dedicated `new_company_name`
        # key, which the model must pass deliberately.
        #
        # ★ The override applies ONLY when the template does not itself render a
        #   separate `new_company_name`. Measured placeholders:
        #     Director Consent (Non-Group) → company, company_registration_number
        #     Individual Shareholder Consent → company_name
        #     Corporate Shareholder Consent → company_name, company_registration_number,
        #                                     new_company_name
        #   The last one names TWO companies — the parent passing the resolution
        #   and the NewCo being formed. Overwriting `company_name` there would
        #   replace the parent with the NewCo and invert the whole instrument.
        #   `canonical_field` keeps the two keys distinct (verified), so this
        #   test is sound.
        _new_company = ""
        for _k in ("new_company_name", "incorporating_company_name"):
            _v = (custom_data or {}).get(_k)
            if isinstance(_v, str) and _v.strip():
                _new_company = _v.strip()
                break

        if _new_company:
            _tpl_fields = {
                canonical_field(f) for f in (result.get("template_analysis", {}).get("required_fields") or [])
            }
            if canonical_field("new_company_name") not in _tpl_fields:
                for _key in ("company", "company_name", "company_name_english"):
                    if canonical_field(_key) in _tpl_fields or _key in data:
                        data[_key] = _new_company
                # A company being incorporated has no number yet. The client's
                # tracker asks for this exact wording; a real number belonging to
                # some other company is the failure being fixed.
                #
                # The supplied value is read out of custom_data DIRECTLY rather
                # than trusted to be in `data`: company_registration_number is a
                # PROTECTED_FIELD, so a caller-supplied number was stripped
                # upstream and `data` still holds the REGISTER's number. Testing
                # "was it supplied" and then leaving `data` alone would keep the
                # wrong company's number — the exact defect being fixed here.
                _supplied_reg = None
                _reg_key = canonical_field("company_registration_number")
                for _k, _v in (custom_data or {}).items():
                    if canonical_field(_k) == _reg_key:
                        _supplied_reg = _v
                        break
                data["company_registration_number"] = (
                    str(_supplied_reg) if _supplied_reg is not None else "TBC upon incorporation"
                )
                # find_replacement resolves company identity from the DB row,
                # not from `data` — see its `source == "db"` branch. The marker
                # is how the decision reaches it.
                data[NEWCO_MARKER] = _new_company
                logging.getLogger("legalscout").info(
                    f"[NEWCO] '{template_name}' bound to unincorporated company "
                    f"'{_new_company}' (registration number TBC)"
                )

        # Idempotency check sits here, after `data` is fully resolved and before
        # anything is written: this is the last point where the inputs are known
        # and the first side effect (saving the .docx, S3 upload, record_document)
        # has not happened yet.
        #
        # LIMITATION: the cache is per-process, and the API runs under
        # `uvicorn --workers 2`. A duplicate that lands on the other worker is
        # NOT caught. This is a mitigation for the common case (same session,
        # same worker, retry/double-click), not a distributed lock — a real
        # guarantee needs a DB uniqueness constraint or advisory lock.
        dedupe_key = _generation_key(template_name, company_name, data)
        previous = _lookup_recent_generation(dedupe_key)
        if previous:
            previous["reused"] = True
            previous["message"] = (
                f"A matching {template_name} for {company_name} with identical data was just "
                f"generated (within the last {_DUPLICATE_WINDOW_SECONDS} seconds), so the existing "
                f"file {previous.get('file_name', '')} is returned instead of creating a duplicate. "
                "Tell the user this is the same document, not a new one. Change the data if a "
                "different version is needed."
            )
            return previous

        filled_doc = fill_template_with_validation(
            template_path, data, template_name=template_name, company_name=company_name
        )

        # A blank nobody asked for does not get written to disk.
        #
        # `fill_template_with_validation` records every placeholder it replaced
        # with LEFT_BLANK. Downstream, those are folded into `unfilled_names`,
        # `is_valid` goes False and the summary says "Partial" — but the file was
        # still saved and the download link still handed over. For a legal
        # instrument that is the wrong default: the finished .docx carries no
        # {{placeholder}} and no TBD, so a blank term is invisible on the page.
        # The measured example is a resolution reading "hereinafter referred to
        # as  ("NewCo")" — an incorporation with no company named.
        #
        # A blank is only acceptable when it was CHOSEN. `_explicitly_supplied`
        # keys on presence, so a caller that passes {"effective_date": ""} has
        # said "leave it blank" and that still generates. A field simply never
        # supplied has not been decided, and this refuses before anything is
        # written — no orphan file, no link to an incomplete document.
        _blank_now = sorted(_blanked_placeholders.get() or set())
        _undeclared = [b for b in _blank_now if not _explicitly_supplied(b, custom_data)]
        if _undeclared:
            logging.getLogger("legalscout").warning(
                f"refused to generate '{template_name}' for '{company_name}': "
                + f"{len(_undeclared)} undeclared blank field(s): "
                + ", ".join(_undeclared)
            )
            return {
                "success": False,
                "error": "Undeclared blank fields",
                "blank_fields": _undeclared,
                "document_state": _document_state(
                    template=template_name,
                    company=company_name,
                    required_fields=result.get("template_analysis", {}).get("required_fields", []),
                    values=_resolved_values(
                        result.get("template_analysis", {}).get("required_fields", []),
                        data,
                        template_name,
                        company_name,
                    )[0],
                    status="awaiting-input",
                    outstanding=_undeclared,
                ),
                "message": "These fields would be blank in the document: "
                + ", ".join(_undeclared)
                + ". They need a value before it can be generated.",
                "agent_instruction": (
                    "ACT NOW, IN THIS SAME TURN — nothing was written and there is no "
                    "file. These fields would print as EMPTY SPACE in the finished "
                    "document, where nobody would notice them: "
                    + ", ".join(_undeclared)
                    + ". If the user already gave these values earlier in the "
                    "conversation, call generate_document again with custom_data "
                    "carrying EVERY value you hold — custom_data does not accumulate "
                    "between calls. Otherwise ask for them with ask_questions (use "
                    '"input_type": "date" for any date). If the user says a field '
                    "should genuinely stay blank, pass it explicitly as an empty "
                    "string in custom_data and generation will proceed."
                ),
            }

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        # Clean company name - remove registration number, parentheses, etc.
        import re

        company_clean = re.sub(r"\([^)]*\)", "", company_name)  # Remove (Registration No: ...)
        company_clean = re.sub(r"\s+", " ", company_clean).strip()  # Clean extra spaces
        company_safe = (
            company_clean.replace(" ", "_").replace("/", "_").replace("'", "").replace("(", "").replace(")", "")
        )

        # Sanitize template name (remove spaces and special chars for URL-safe filename)
        template_safe = (
            template_name.replace(".docx", "")
            .replace(" ", "_")  # Replace spaces with underscores
            .replace("/", "_")
            .replace("\\", "_")
            .replace("(", "")
            .replace(")", "")
            .replace("[", "")
            .replace("]", "")
            .replace("'", "")
            .replace('"', "")
            .replace(",", "")
            .replace(":", "")
            .replace(";", "")
        )

        output_filename = f"{template_safe}_{company_safe}_{timestamp}.docx"
        output_dir = Path(documents_dir) / "legal" / "output"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / output_filename

        filled_doc.save(str(output_path))

        # A generated document is THE side effect of this product — a file that
        # did not exist, produced for a named client. Recorded after the save
        # succeeds so the ledger cannot claim a document that was never written.
        # `record` returns None and never raises when the flag is off or no turn
        # is open, so this is inert until the ledger is switched on.
        try:
            from scout.effects import record

            record(
                kind="document",
                # "external", not "insert": the effect is a FILE written to disk
                # (and pushed to S3), not a row. `op` is a closed set — an
                # invalid value raises inside build_effect and `record` swallows
                # it, so a wrong op here is a ledger that silently stays empty.
                op="external",
                target_table=None,
                target_label=output_filename,
                after={"company_name": company_name, "template_name": template_name, "path": str(output_path)},
                tool_name="generate_document",
                diff=False,
            )
        except Exception as e:
            logging.getLogger("legalscout").warning(f"effects record failed for '{output_path}': {e}")

        # Upload to S3 (background)
        try:
            from app.s3_storage import s3_upload_async

            s3_upload_async(str(output_path))
        except Exception as e:
            logging.getLogger("legalscout").warning(f"S3 upload failed for '{output_path}': {e}")

        # Extract placeholders from template BEFORE validation
        template_analysis = extract_placeholders_from_template(template_path)
        all_template_placeholders = template_analysis.get("fields", [])

        validation = validate_filled_document(output_path, all_template_placeholders)

        # Fold in the fields the fill emptied. Without this they are invisible:
        # `validate_filled_document` can only see placeholders that were LEFT
        # RAW, so a field replaced with "" reads as filled and the document is
        # reported Complete with its terms blank.
        _blanked = sorted(_blanked_placeholders.get() or set())
        if _blanked:
            _names = list(validation.get("unfilled_names") or [])
            _ph = list(validation.get("unfilled_placeholders") or [])
            for _b in _blanked:
                if _b not in _names:
                    _names.append(_b)
                if _b not in _ph:
                    _ph.append(_b)
            _filled = [
                f
                for f in (validation.get("filled_placeholders") or [])
                if f.lower() not in {b.lower() for b in _blanked}
            ]
            validation["unfilled_names"] = _names
            validation["unfilled_placeholders"] = _ph
            validation["filled_placeholders"] = _filled
            validation["total_unfilled"] = len(_names)
            validation["total_filled"] = len(_filled)
            validation["blanked_placeholders"] = _blanked
            validation["is_valid"] = False
            logging.getLogger("legalscout").warning(
                f"'{output_filename}' generated with {len(_blanked)} EMPTY field(s): " + ", ".join(_blanked)
            )

        preview = generate_preview(filled_doc)

        doc_url = f"{host}/documents/legal/output/{output_filename}"
        google_viewer_url = f"https://docs.google.com/gview?embedded=1&url={doc_url}"

        record_document(
            template_name=template_name,
            company_name=company_name,
            file_name=output_filename,
            file_path=str(output_path),
            download_url=doc_url,
            preview_url=google_viewer_url,
            validation_result=validation,
            custom_data=data,
        )

        payload = {
            "success": True,
            "step": "generated",
            "reused": False,
            "file_name": output_filename,
            "file_path": str(output_path),
            "download_url": doc_url,
            "validation": validation,
            "preview": preview,
            "message": f"Document created successfully for {company_name}",
            # Same numbers the validation_summary reports, expressed per field:
            # every template placeholder, the value it received, and whether the
            # filled .docx still shows it unresolved.
            "document_state": _document_state(
                template=template_name,
                company=company_name,
                required_fields=all_template_placeholders,
                values=_resolved_values(all_template_placeholders, data, template_name, company_name)[0],
                status="generated",
                outstanding=validation.get("unfilled_names", []),
                file_name=output_filename,
                download_url=doc_url,
            ),
            "warnings": validation.get("unfilled_placeholders", []) if not validation.get("is_valid") else [],
            # Validation summary for user
            "validation_summary": {
                "total_placeholders": validation.get("total_placeholders", 0),
                "filled_from_data": validation.get("total_filled", 0),
                "unfilled": validation.get("total_unfilled", 0),
                "unfilled_fields": validation.get("unfilled_names", []),
                "validation_status": "Complete" if validation.get("is_valid") else "Partial - some fields may be empty",
            },
            # Template info from training
            "template_info": {
                "purpose": result.get("template_analysis", {}).get("purpose", ""),
                "when_to_use": result.get("template_analysis", {}).get("when_to_use", ""),
                "how_to_use": result.get("template_analysis", {}).get("how_to_use", []),
            },
            # Told to the model, never shown to the user.
            #
            # This is the single most common place the agent goes quiet. Measured
            # on the live app: the document is written, the tool returns, and the
            # turn ends with 0 characters of visible text and ~200 characters of
            # invisible reasoning ("I've received confirmation to proceed and am
            # now preparing to call the document generator…"). The user sees an
            # empty bubble; the frontend covers it by sending "continue", which
            # costs a whole extra model round-trip.
            #
            # A rule in the system prompt was not enough — this is the text the
            # model reads immediately before deciding whether to stop, so the
            # demand belongs here.
            "agent_instruction": (
                "The document EXISTS now. Do not call this tool again. "
                "Reply to the user IN THIS SAME TURN — ending the turn without text "
                "is a bug, not brevity. One or two sentences: what was generated, for "
                f"which company, and the download link `{doc_url}` on a single line. "
                + (
                    "Then name the fields still unfilled and ask whether to fill them: "
                    + ", ".join(validation.get("unfilled_names", [])[:6])
                    if validation.get("unfilled_names")
                    else "Then offer the obvious next step — emailing it, or the companion document."
                )
            ),
        }

        # The task is done: forget it, so a later unrelated turn in this same
        # conversation is not told a finished document is still pending.
        try:
            from scout.tools.task_memory import clear_task

            clear_task(_session)
        except Exception as _e:
            logging.getLogger("legalscout").warning(f"[TASK] clear skipped: {_e}")

        # Only successes are remembered — a failed run must stay retryable.
        _remember_generation(dedupe_key, payload)
        return payload

    # Combined tool: Analyze + Prepare + Generate in one call
    def create_document(
        template_name: str,
        company_name: str,
        custom_data: dict | None = None,
        run_context: RunContext = None,
    ) -> dict:
        """
        Complete document creation in one call.
        Combines analyze, prepare, and generate.

        Args:
            template_name: Name of the template file
            company_name: Name of the company
            custom_data: Optional custom data to override database data

        Returns:
            Full result with download URL and validation summary
        """
        with session_scope(_session_of(run_context)):
            return _generate_document(template_name, company_name, custom_data)

    def preview_doc(
        template_name: str,
        company_name: str,
        custom_data: dict | None = None,
        run_context: RunContext = None,
    ) -> dict[str, Any]:
        """Preview document - shows what will be filled without saving."""
        with session_scope(_session_of(run_context)):
            return _preview_doc(template_name, company_name, custom_data)

    def _preview_doc(template_name: str, company_name: str, custom_data: dict | None = None) -> dict[str, Any]:
        result = prepare_document_data(template_name, company_name, documents_dir)

        if not result.get("success"):
            return result

        # Same reason as in _generate_document_inner: the name the caller used
        # may be a prefix or a case variant, and everything downstream keys off
        # the name. A preview built from an unresolved name would disagree with
        # the document that generation then produces.
        template_name = (result.get("template_analysis") or {}).get("template") or template_name

        # Record the task HERE too, not only at generation.
        #
        # `merge_collected` (what ask_questions calls to persist a card answer)
        # is UPDATE-only — it needs a row to attach to. In the ordinary flow the
        # preview comes first and the questions are asked against it, so if only
        # generate_document recorded the task, answers given before the first
        # generate call had nowhere to go and were silently dropped. Measured:
        # the AGM Minutes meeting date survived only 2 runs in 3 without this.
        try:
            from scout.tools.task_memory import remember_task

            remember_task(current_session_scope(), template_name, company_name, custom_data)
        except Exception as _e:
            logging.getLogger("legalscout").warning(f"[TASK] preview record skipped: {_e}")

        validation = result.get("validation", {})
        # Use available_fields (correct key) instead of matched_fields
        total_fields = len(validation.get("available_fields", [])) + len(validation.get("missing_fields", []))
        matched_fields = len(validation.get("available_fields", []))

        field_coverage = matched_fields / total_fields if total_fields > 0 else 0

        # Build preview data
        preview_data = result.get("normalized_data", {}).copy()

        # Add custom data
        if custom_data:
            preview_data.update(custom_data)
        # Same presentation the generated document gets — a preview that shows
        # "50,000,000" for a file that will say "50000000" is worse than no
        # preview, because the user approves the version they were shown.
        preview_data = _present_numbers(preview_data)

        # Add TBD for missing fields
        for field in validation.get("missing_fields", []):
            if field not in preview_data:
                preview_data[field] = "[TBD - needs input]"

        # Get template info from database
        from scout.tools.template_analyzer import get_all_templates_from_db

        templates_db = get_all_templates_from_db()
        template_info = None
        for t in templates_db:
            if t.get("name") == template_name:
                template_info = t
                break

        # Generate table format preview
        preview_lines = []

        # Header
        preview_lines.append("📄 **DOCUMENT PREVIEW**")
        preview_lines.append("")

        # Template & Company info
        preview_lines.append(f"**Template:** {template_name}")
        preview_lines.append(f"**Company:** {company_name}")
        preview_lines.append(f"**Data Coverage:** {int(field_coverage * 100)}%")
        preview_lines.append("")

        # Template purpose (if available)
        if template_info and template_info.get("purpose"):
            preview_lines.append(f"📋 **Purpose:** {template_info['purpose']}")
            preview_lines.append("")

        # Data table
        preview_lines.append("### 📊 Data to be filled")
        preview_lines.append("")
        preview_lines.append("| Field | Value | Source |")
        preview_lines.append("|-------|-------|--------|")

        # Show matched fields (available_fields is the correct key)
        for field in validation.get("available_fields", []):
            value = preview_data.get(field, "[empty]")
            # Truncate long values
            if len(str(value)) > 40:
                value = str(value)[:37] + "..."
            preview_lines.append(f"| {field} | {value} | ✅ DB |")

        # Show missing fields
        for field in validation.get("missing_fields", []):
            value = preview_data.get(field, "[TBD]")
            preview_lines.append(f"| {field} | {value} | ⚠️ Needs |")

        preview_lines.append("")
        preview_lines.append("---")
        # The lettered a)/b) menu that used to live here was the very grammar the
        # system prompt bans — and it came from tool output, so no amount of
        # prompt discipline could suppress it. The choice is now handed to
        # `ask_questions`, which renders clickable chips.
        short_name = template_name.replace(".docx", "").replace("_", " ")
        if len(short_name) > 40:
            short_name = short_name[:37] + "..."

        return {
            "success": True,
            "preview": "\n".join(preview_lines),
            "template_name": template_name,
            "company_name": company_name,
            "field_coverage": f"{int(field_coverage * 100)}%",
            "matched_fields": validation.get("available_fields", []),  # Use correct key
            "missing_fields": validation.get("missing_fields", []),
            "data": preview_data,
            "document_state": _document_state(
                template=template_name,
                company=company_name,
                required_fields=result.get("template_analysis", {}).get("required_fields", []),
                values=_resolved_values(
                    result.get("template_analysis", {}).get("required_fields", []),
                    preview_data,
                    template_name,
                    company_name,
                )[0],
                # A preview that still has missing fields or unpicked parties is
                # waiting on the user; otherwise it is ready to generate.
                status="ready" if result.get("ready_to_generate") else "awaiting-input",
                # Same resolver as everywhere else. `validation.missing_fields`
                # is name-matching against company columns and reports a
                # register-held attribute as missing.
                outstanding=_resolved_values(
                    result.get("template_analysis", {}).get("required_fields", []),
                    preview_data,
                    template_name,
                    company_name,
                )[1]
                + list(result.get("unresolved_slots", []) or []),
            ),
            "needs_approval": True,
            # Told to the model, not shown to the user: how to ask for approval.
            #
            # "Show the preview" was being read as optional. Measured on the live
            # app, this tool returned and the turn ended with 0 characters of
            # visible text and 268 characters of invisible reasoning — the user
            # got an empty bubble where the preview should have been, and the
            # frontend had to nudge the run forward with a synthetic "continue".
            # So the demand for text is now explicit and first.
            "agent_instruction": (
                "WRITE THE PREVIEW OUT TO THE USER NOW, IN THIS SAME TURN. A turn that "
                "ends here with no text is a bug — the user is left looking at an empty "
                "reply while waiting to approve. Summarise what will be filled and what "
                "is still missing, THEN call ask_questions with ONE question — "
                f'"Generate {short_name} for {company_name} now?" — and the options '
                '["Yes, generate it", "No, change the data first"]. '
                "Never write a lettered a)/b) menu in prose."
            ),
        }

    def _as_json(fn):
        """Return the tool's result as a JSON string rather than a dict.

        Agno stringifies a dict return with `str()`, which produces Python repr
        — single quotes, True/False/None. The browser then has to parse Python,
        and a value containing an apostrophe or a curly quote breaks the parse
        and blanks the document panel. The generate_document preview text
        contains all three.

        JSON is also what the model reads back, so this makes the tool result
        unambiguous on both sides. `default=str` keeps dates and Decimals from
        raising mid-serialisation.
        """

        @wraps(fn)
        def wrapper(*args, **kwargs):
            result = fn(*args, **kwargs)
            if isinstance(result, str):
                return result  # already serialised by the callee
            try:
                return _json.dumps(result, default=str)
            except (TypeError, ValueError) as e:
                logging.getLogger("legalscout").warning(f"[SMART_DOC] {fn.__name__} result not JSON-serialisable: {e}")
                return _json.dumps({"success": False, "error": f"Result could not be serialised: {e}"})

        return wrapper

    # Only the tools whose results the document panel reads are wrapped. The
    # rest are unchanged so this cannot ripple through the agent's tool surface.
    return {
        "analyze_template": analyze_template_tool,
        "prepare_document": _as_json(prepare_document),
        "generate_document": _as_json(generate_document),
        "create_document": _as_json(create_document),
        "preview_document": _as_json(preview_doc),
    }


def _fill_paragraph_highlighted(
    paragraph,
    data: dict[str, Any],
    placeholder_pattern=None,
    template_name: str | None = None,
    company_name: str | None = None,
    empty_counter=None,
):
    """Fill placeholders in a paragraph, highlighting replaced text in yellow."""
    from copy import deepcopy

    from docx.oxml.ns import qn
    from lxml import etree

    text = paragraph.text
    matches = list((placeholder_pattern or PLACEHOLDER_PATTERN).finditer(text))
    if not matches:
        return

    # Get the formatting from the first run (to preserve font, size, etc.)
    base_run = paragraph.runs[0] if paragraph.runs else None

    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""

    # Build segments: [("text", highlight_color), ...]
    # highlight_color: None=no highlight, "yellow"=filled, "red"=TBD/unfilled
    segments = []
    last_end = 0
    for match in matches:
        placeholder = placeholder_name((match.group(1), match.group(2), match.group(3)), empty_counter).lower()
        replacement = (
            find_replacement(placeholder, data, template_name=template_name, company_name=company_name)
            if placeholder
            else None
        )

        if match.start() > last_end:
            segments.append((text[last_end : match.start()], None))

        if replacement == LEFT_BLANK and replacement is not None:
            segments.append((LEFT_BLANK, None))
        elif replacement:
            is_tbd = str(replacement).strip().upper() == "TBD"
            segments.append((str(replacement), "red" if is_tbd else "yellow"))
        else:
            segments.append((text[match.start() : match.end()], None))

        last_end = match.end()

    if last_end < len(text):
        segments.append((text[last_end:], None))

    # Add runs for each segment
    p_element = paragraph._element
    # Remove all existing run elements
    for r in list(p_element.findall(qn("w:r"))):
        p_element.remove(r)

    for seg_text, highlight_color in segments:
        new_run = paragraph.add_run(seg_text)
        # Copy formatting from base run if available
        if base_run and base_run._element.find(qn("w:rPr")) is not None:
            new_rpr = deepcopy(base_run._element.find(qn("w:rPr")))
            # Remove any existing highlight from copied format
            for existing_hl in new_rpr.findall(qn("w:highlight")):
                new_rpr.remove(existing_hl)
            new_run._element.insert(0, new_rpr)

        # Add highlight: yellow=filled from data, red=TBD/unfilled
        if highlight_color:
            rpr = new_run._element.get_or_add_rPr()
            highlight = etree.SubElement(rpr, qn("w:highlight"))
            highlight.set(qn("w:val"), highlight_color)  # "yellow" for filled, "red" for TBD


def fill_template_with_validation(
    template_path: Path, data: dict[str, Any], template_name: str | None = None, company_name: str | None = None
) -> Document:
    """Fill template with data, highlighting filled values in yellow."""
    doc = Document(str(template_path))
    empty_counter = new_empty_counter()

    # Dynamic list expansion: grow/shrink Present lists, appointed-director
    # lists and signing-table row groups to the real party count BEFORE fill.
    # Returns synthetic {token: value} pairs the highlighter fills like any slot.
    try:
        from scout.tools.repeat_regions import expand_repeat_regions

        synth = expand_repeat_regions(doc, data, template_name=template_name, company_name=company_name)
        if synth:
            data = {**data, **synth}
    except Exception as e:
        logging.getLogger("legalscout").warning(f"repeat-region expansion skipped: {e}")

    for paragraph in doc.paragraphs:
        _fill_paragraph_highlighted(
            paragraph,
            data,
            PLACEHOLDER_PATTERN,
            template_name=template_name,
            company_name=company_name,
            empty_counter=empty_counter,
        )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _fill_paragraph_highlighted(
                        paragraph,
                        data,
                        PLACEHOLDER_PATTERN,
                        template_name=template_name,
                        company_name=company_name,
                        empty_counter=empty_counter,
                    )

    return doc


def _company_own_addresses(company_name: str | None) -> set[str]:
    """The company's OWN addresses, normalised for comparison.

    `prepare_document_data` publishes the company record under several aliases,
    and one of them is a bare `address`. A template that writes `[address]`
    meaning the DIRECTOR's home therefore resolved it on the exact-match tier —
    the first line of `_resolve_from_data`, before any companion lookup — and a
    consent form went out stating that the director resides at the company's
    registered office, phone number and legal@ mailbox included.

    Matching on the VALUE is what makes this safe: it removes the company's own
    address from consideration for a person's field without guessing which key
    it arrived under, and leaves an address the user actually typed alone.
    """
    if not company_name:
        return set()
    try:
        row = _get_company_from_db(company_name)
    except Exception:
        return set()
    if not row or row.get("ambiguous"):
        return set()
    out = set()
    for column in ("registered_office_address", "principal_place_of_business", "address", "company_address"):
        value = row.get(column)
        if value and str(value).strip():
            out.add(re.sub(r"\s+", " ", str(value)).strip().lower())
    return out


def _get_company_field(company_row: dict, column_path: str) -> str | None:
    """Resolve a DB column path from company data.
    Handles: 'company_name_english', 'members[0].name', 'directors[1].position'"""
    if not company_row or not column_path:
        return None

    # Handle array access: members[0].name
    import re

    array_match = re.match(r"(\w+)\[(\d+)\]\.(\w+)", column_path)
    if array_match:
        field, idx, subfield = array_match.groups()
        arr = company_row.get(field, [])
        if isinstance(arr, list) and int(idx) < len(arr):
            item = arr[int(idx)]
            if isinstance(item, dict):
                return str(item.get(subfield, ""))
        return None

    # Direct field access
    return str(company_row.get(column_path, "")) if company_row.get(column_path) else None


def _get_field_mapping(template_name: str) -> dict | None:
    """Load learned field_mapping from DB for a template."""
    conn = None
    try:
        from db.connection import get_db_conn

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT field_mapping FROM templates WHERE name = %s", (template_name,))
        row = cur.fetchone()
        cur.close()
        if row and row[0]:
            raw = row[0] if isinstance(row[0], dict) else __import__("json").loads(row[0])
            return normalise_mapping(raw)
    except Exception as e:
        logging.getLogger("legalscout").warning(f"Failed to load field mapping for '{template_name}': {e}")
    finally:
        if conn:
            conn.close()
    return None


def _company_row_to_dict(row) -> dict:
    """Map a companies table row to the resolver's field dict."""
    return {
        "company_name_english": row[0],
        "company_registration_number": row[1],
        "registered_office_address": row[2],
        "principal_place_of_business": row[3],
        "status": row[4],
        "company_type": row[5],
        "directors": row[6] if isinstance(row[6], list) else [],
        "members": row[7] if isinstance(row[7], list) else [],
        "total_shares_issued": row[8],
        "currency_of_share_capital": row[9],
        "date_of_last_annual_return": str(row[10]) if row[10] else None,
        "financial_year_end_date": str(row[11]) if row[11] else None,
        "ultimate_holding_company_name": row[12],
        "next_financial_year_end_date": str(row[13]) if row[13] else None,
        "auditor_name": row[14],
        "auditor_fee": row[15],
    }


def _get_company_from_db(company_name: str) -> dict | None:
    """Get full company record directly from companies table.

    Matches exact name first, then prefix, then substring. When more than one
    company matches at the chosen tier the result is flagged ambiguous with the
    candidate names so the caller can ask the user which one is meant.
    """
    conn = None
    try:
        from db.connection import get_db_conn

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(
            """
            SELECT company_name_english, company_registration_number, registered_office_address,
                   principal_place_of_business, status, company_type, directors, members,
                   total_shares_issued, currency_of_share_capital, date_of_last_annual_return,
                   financial_year_end_date, ultimate_holding_company_name,
                   next_financial_year_end_date, auditor_name, auditor_fee
            FROM companies WHERE company_name_english ILIKE %s
            ORDER BY company_name_english
        """,
            (f"%{company_name}%",),
        )
        rows = cur.fetchall()
        cur.close()

        needle = str(company_name or "").strip().lower()
        exact = [r for r in rows if str(r[0] or "").strip().lower() == needle]
        prefix = [r for r in rows if str(r[0] or "").strip().lower().startswith(needle)]
        tier = exact or prefix or list(rows)

        if not tier:
            return None
        if len(tier) > 1:
            return {
                "ambiguous": True,
                "matches": [r[0] for r in tier],
                "message": f"'{company_name}' matches {len(tier)} companies",
            }
        return _company_row_to_dict(tier[0])
    except Exception as e:
        logging.getLogger("legalscout").warning(f"Failed to load company '{company_name}' from DB: {e}")
    finally:
        if conn:
            conn.close()
    return None


def _resolve_from_data(placeholder: str, data: dict[str, Any], strict: bool = False) -> str | None:
    """Resolve a placeholder against supplied data: exact, alias, then token set.

    `strict` drops the last tier — the token-SUBSET match. That tier answers a
    placeholder from any key whose tokens contain it, which is right for
    abbreviations and wrong across entities: `{address}` is a subset of
    `{registered, office, address}`, so a bare `address` meaning the DIRECTOR's
    home was silently filled with the COMPANY's registered office, on a form
    whose own trained description reads "Residential address of the appointed
    director". A blank is visible to whoever proof-reads; a confident wrong
    address is not.
    """
    placeholder_norm = normalize_field(placeholder)
    placeholder_canonical = canonical_field(placeholder)

    direct = data.get(placeholder) or data.get(placeholder_norm)
    if direct and str(direct) != "TBD":
        return str(direct)

    normalized_items = []
    for key, value in data.items():
        if not value or str(value) == "TBD":
            continue
        normalized_items.append((normalize_field(key), canonical_field(key), str(value)))

    for key_norm, _key_canonical, value in normalized_items:
        if key_norm == placeholder_norm:
            return value

    for _key_norm, key_canonical, value in normalized_items:
        if key_canonical == placeholder_canonical:
            return value

    if not strict:
        for key_norm, _key_canonical, value in normalized_items:
            if tokens_match(placeholder_norm, key_norm):
                return value

    # Last tier: a bare `date` placeholder answered under a qualified key.
    # Only fires when EXACTLY ONE qualified date was supplied — two of them
    # (a meeting date and a signing date) mean the caller distinguished them,
    # and picking either would put one date where the other belongs.
    if placeholder_canonical == "date":
        qualified = [value for key_norm, _kc, value in normalized_items if _QUALIFIED_DATE_RE.match(key_norm)]
        if len(qualified) == 1:
            return qualified[0]

    return None


def _explicitly_supplied(placeholder: str, custom_data: dict[str, Any] | None) -> bool:
    """True when the caller deliberately set this field — even to an empty value.

    PRESENCE is the signal here, not truthiness. "Leave the effective date
    blank if unconfirmed" is a real answer, and a check that only looked at the
    value would treat it as still missing and ask again forever.

    Key matching mirrors `_resolve_from_data` so that `new_company_name`,
    `New Company Name` and `newCompanyName` are recognised as the same field.
    """
    if not custom_data:
        return False
    placeholder_norm = normalize_field(placeholder)
    placeholder_canonical = canonical_field(placeholder)
    for key in custom_data:
        if key == placeholder:
            return True
        key_norm = normalize_field(key)
        if key_norm == placeholder_norm:
            return True
        if canonical_field(key) == placeholder_canonical:
            return True
        if tokens_match(placeholder_norm, key_norm):
            return True
    return False


def find_replacement(
    placeholder: str,
    data: dict[str, Any],
    template_name: str | None = None,
    company_name: str | None = None,
    document_id: Any = None,
) -> str | None:
    """Find replacement value using learned field_mapping (if available) or fallback to data lookup."""
    placeholder_norm = normalize_field(placeholder)

    if is_empty_placeholder(placeholder_norm):
        return _resolve_from_data(placeholder, data)

    # === LEARNED MAPPING (Priority 1) ===
    if template_name:
        mapping = _get_field_mapping(template_name)
        if mapping:
            # Try exact match and normalized match
            config = mapping.get(placeholder) or mapping.get(placeholder_norm)
            if config:
                source = config.get("source", "")
                db_column = config.get("db_column", "")
                default = config.get("default", "TBD")

                if source == "slot" and slot_of(config):
                    resolved = resolve_slot(
                        placeholder,
                        config,
                        data,
                        company_name=company_name,
                        document_id=document_id,
                        blank=LEFT_BLANK,
                    )
                    if isinstance(resolved, list):
                        return ", ".join(str(item) for item in resolved)
                    return resolved

                if source == "db" and db_column:
                    # A company being incorporated has no register row, and the
                    # row that WAS looked up belongs to somebody else. Answering
                    # from it put another client's name and real registration
                    # number on a consent form (the tracker's "Setup 1" case).
                    _newco = data.get(NEWCO_MARKER) if isinstance(data, dict) else None
                    if _newco:
                        _canon = canonical_field(placeholder)
                        if _canon in _NEWCO_NAME_FIELDS:
                            return str(_newco)
                        if _canon in _NEWCO_REG_FIELDS:
                            supplied = _resolve_from_data(placeholder, data)
                            return supplied or "TBC upon incorporation"
                    # Get value from company DB
                    company_row = _get_company_from_db(company_name) if company_name else None
                    if company_row and not company_row.get("ambiguous"):
                        val = _get_company_field(company_row, db_column)
                        if val:
                            return val

                elif source == "user_input":
                    # A person's own attributes — NRC, nationality, date of
                    # birth, home address — belong to the person a sibling slot
                    # already resolved to, not to the document. Training
                    # classifies every one of them `user_input` because none is
                    # a column on `companies`, which decoupled them from the
                    # name slot beside them: the picker chose the person, the
                    # attribute was asked as free text, and the two could
                    # disagree or come out blank while the register held the
                    # answer all along.
                    #
                    # ORDER: a value the caller SUPPLIED still wins — the user's
                    # own answer outranks the register. But it must win by
                    # naming the field, not by a loose token-subset match
                    # against a company column, so a person attribute reads the
                    # supplied data in `strict` mode. Left loose, `address`
                    # matched `registered_office_address` and the company's
                    # office was filled in as the director's home BEFORE this
                    # companion lookup was ever reached.
                    attr_target = _attr_tail(placeholder)
                    person_attr = attr_target is not None

                    resolved = _resolve_from_data(placeholder, data, strict=person_attr)
                    # The company's OWN address is not the director's home. It
                    # arrives in `data` under a bare `address` alias and so wins
                    # the exact-match tier above; discarding it here is what
                    # sends the placeholder on to the register, and then to the
                    # user, instead of quietly filling in the wrong building.
                    if (
                        resolved
                        and attr_target
                        and attr_target[1] == "residential_address"
                        and re.sub(r"\s+", " ", str(resolved)).strip().lower()
                        in _company_own_addresses(company_name)
                    ):
                        resolved = None
                    if resolved:
                        return resolved

                    companion = companion_attribute(
                        placeholder,
                        mapping,
                        data,
                        company_name=company_name,
                        document_id=document_id,
                    )
                    if companion:
                        return companion

                    # Not a person attribute after all, or the register has
                    # nothing on file: fall back to the loose match the strict
                    # pass skipped, so nothing that used to resolve stops.
                    # A person attribute deliberately does NOT come back here —
                    # for those, blank-and-ask is the correct outcome.
                    if not person_attr:
                        resolved = _resolve_from_data(placeholder, data)
                        if resolved:
                            return resolved

                    if default and default != "today" and str(default).strip():
                        return str(default)
                    # Nothing supplied, no companion, no default. Emitting "" is
                    # what the document needs (a raw {{placeholder}} in a legal
                    # instrument is worse), but it must be REPORTED, not hidden.
                    _note_blank(placeholder)
                    return LEFT_BLANK

    # === DATA LOOKUP (Priority 2 — fallback) ===
    resolved = _resolve_from_data(placeholder, data)
    if resolved:
        return resolved

    # === SMART DEFAULTS (Priority 3) ===
    if "location" in placeholder_norm or "venue" in placeholder_norm:
        return data.get("registered_office", data.get("registered_office_address", "TBD"))

    # Final fallback — any unfilled placeholder becomes TBD so it is not left raw in the doc
    KNOWN_USER_INPUT = {
        "financial_year_end_date",
        "next_financial_year_end_date",
        "auditor_name",
        "auditor_fee",
        "auditor_remuneration",
        "meeting_time",
        "shareholder_name",
    }
    if placeholder_norm in KNOWN_USER_INPUT or "financial_year" in placeholder_norm or "auditor" in placeholder_norm:
        return "TBD"

    return None


def generate_preview(doc: Document) -> str:
    """Generate preview of filled document."""
    preview_lines = ["Document Preview:"]

    for paragraph in doc.paragraphs[:10]:
        text = paragraph.text.strip()
        if text and len(text) > 10:
            preview_lines.append(f"  • {text[:100]}")

    return "\n".join(preview_lines)
