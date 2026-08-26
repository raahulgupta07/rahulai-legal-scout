"""
Fill-view — the whole document, with blanks you click to fill
=============================================================

The legal team asked to see the *entire* document with the missing fields shown
in place, and to fill each blank by choosing from the data we already hold —
"instead of writing, show relevant information and we can just choose and click."

This module renders a template into an ordered list of blocks the frontend draws
inline:

    {"type": "text",  "text": "..."}                     literal document text
    {"type": "break"}                                     paragraph break
    {"type": "blank", "key", "label", "value", "filled",
                       "candidates": [{"label","value","source"}], "kind"}

`value` is what the register already resolves for that slot; `filled` says
whether it is settled. `candidates` are the click-to-pick options drawn from the
company register and the people register. Free-text is always allowed on top.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from docx import Document

from scout.tools.placeholders import PLACEHOLDER_PATTERN, is_empty_placeholder, new_empty_counter, placeholder_name

_log = logging.getLogger("legalscout")

_PERSON_RE = re.compile(
    r"director|shareholder|member|signator|signer|appointed|representative|chairperson|chairman|attendee|person|name",
    re.IGNORECASE,
)
_DATE_RE = re.compile(r"date", re.IGNORECASE)
_PRONOUN_RE = re.compile(r"pronoun", re.IGNORECASE)
_LOCATION_RE = re.compile(r"location|venue|place", re.IGNORECASE)
_AUDITOR_RE = re.compile(r"auditor", re.IGNORECASE)


# A value that is only a placeholder MASK (DD/MM/YYYY, __/__/____, XX-XX) is not
# a real value — it must still count as an outstanding blank (finding F6).
_MASK_RE = re.compile(r"^[dmyx_\-/.\s]+$", re.IGNORECASE)
_CORP_WORDS = ("company", "limited", "ltd", "holdings", "group", "co.,", "pcl", "plc", "corporation", "corp")


def _is_mask(value: str) -> bool:
    v = (value or "").strip()
    return bool(v) and bool(_MASK_RE.match(v)) and any(ch.isalpha() or ch == "_" for ch in v)


def _looks_corporate(name: str) -> bool:
    n = (name or "").lower()
    return any(w in n for w in _CORP_WORDS)


def _slot_party_type(placeholder: str) -> str | None:
    """'corporate' / 'individual' when a person slot names its party type, else None."""
    p = placeholder.lower()
    if "corporate" in p or "company" in p:
        return "corporate"
    if "individual" in p:
        return "individual"
    return None


# A company's name is not a person's name. `_PERSON_RE` matches on the bare word
# "name", so "new_company_name" and "company_name" were classified as people and
# offered a picker full of directors — for a company that does not exist yet.
_COMPANY_NAME_RE = re.compile(r"company[_ ]?name|name[_ ]?of[_ ]?(the[_ ]?)?company|business[_ ]?name", re.IGNORECASE)


def _kind(placeholder: str) -> str:
    p = placeholder.lower()
    if _PRONOUN_RE.search(p):
        return "pronoun"
    if _COMPANY_NAME_RE.search(p):
        return "text"
    if _DATE_RE.search(p):
        return "date"
    if _LOCATION_RE.search(p):
        return "location"
    if _AUDITOR_RE.search(p):
        return "auditor"
    if _PERSON_RE.search(p):
        return "person"
    return "text"


def _people_register_names() -> list[dict]:
    """Names from the standalone People register (directors & individual shareholders)."""
    out = []
    conn = None
    try:
        from db.connection import get_db_conn

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute("SELECT full_name, nrc_passport_no FROM people ORDER BY full_name ASC LIMIT 500")
        for name, nrc in cur.fetchall():
            if name:
                out.append({"label": name, "value": name, "source": "People register", "nrc": nrc or ""})
        cur.close()
    except Exception as e:
        _log.warning(f"fill_view: people register read failed: {e}")
    finally:
        if conn is not None:
            conn.close()
    return out


def _company_people(company_row: dict) -> tuple[list[dict], list[dict]]:
    directors, shareholders = [], []
    for d in company_row.get("directors_list") or []:
        if isinstance(d, dict) and d.get("name"):
            directors.append({"label": d["name"], "value": d["name"], "source": "Directors"})
    for m in company_row.get("shareholders_list") or []:
        if isinstance(m, dict) and m.get("name"):
            shareholders.append({"label": m["name"], "value": m["name"], "source": "Shareholders"})
    return directors, shareholders


def _dedup(cands: list[dict]) -> list[dict]:
    seen, out = set(), []
    for c in cands:
        v = (c.get("value") or "").strip()
        if v and v.lower() not in seen:
            seen.add(v.lower())
            out.append(c)
    return out


def _candidates(kind: str, placeholder: str, company_row: dict) -> list[dict]:
    if kind == "pronoun":
        # A corporate subject is "it / the Company", not he/she/they (finding F8).
        return [
            {"label": "He / His", "value": "he", "source": "Pronoun"},
            {"label": "She / Her", "value": "she", "source": "Pronoun"},
            {"label": "They / Their", "value": "they", "source": "Pronoun"},
            {"label": "It / Its (the Company)", "value": "it", "source": "Pronoun"},
        ]
    if kind == "location":
        addr = company_row.get("address") or company_row.get("registered_office") or ""
        return [{"label": addr, "value": addr, "source": "Registered office"}] if addr else []
    if kind == "auditor":
        if "fee" in placeholder.lower():
            return []  # a fee is free-text, not the auditor's name
        a = company_row.get("auditor_name") or ""
        return [{"label": a, "value": a, "source": "Company register"}] if a else []
    if kind == "person":
        directors, shareholders = _company_people(company_row)
        people = _people_register_names()
        p = placeholder.lower()
        # order the most relevant source first
        if "director" in p or "appointed" in p or "representative" in p or "chairperson" in p:
            ordered = directors + shareholders + people
        elif "shareholder" in p or "member" in p:
            ordered = shareholders + directors + people
        else:
            ordered = people + directors + shareholders
        ordered = _dedup(ordered)
        # An `individual ...` slot must not offer corporate entities, and a
        # `corporate ...` slot must not offer individuals (finding F7).
        ptype = _slot_party_type(placeholder)
        if ptype == "individual":
            ordered = [c for c in ordered if not _looks_corporate(c.get("value", ""))]
        elif ptype == "corporate":
            ordered = [c for c in ordered if _looks_corporate(c.get("value", ""))]
        return ordered
    return []  # date / free text


def _resolve_value(placeholder: str, normalized_data: dict, template_name: str, company_name: str) -> str | None:
    try:
        from scout.tools.smart_doc import LEFT_BLANK, find_replacement

        v = find_replacement(placeholder, normalized_data, template_name=template_name, company_name=company_name)
        if v is None or v == LEFT_BLANK:
            return None
        v = str(v).strip()
        if not v or v.upper() == "TBD":
            return None
        return v
    except Exception as e:
        _log.warning(f"fill_view: resolve failed for {placeholder!r}: {e}")
        return None


# Placeholder tails that describe a person rather than name one. Each maps to
# the People-register column that answers it.
_PERSON_ATTRS = (
    ("identification_number", "nrc_passport_no"),
    ("identification_no", "nrc_passport_no"),
    ("nrc_passport", "nrc_passport_no"),
    ("nrc_no", "nrc_passport_no"),
    ("nrc", "nrc_passport_no"),
    ("passport", "nrc_passport_no"),
    ("id_number", "nrc_passport_no"),
    ("father_name", "father_name"),
    ("fathers_name", "father_name"),
    ("nationality", "nationality"),
    ("date_of_birth", "date_of_birth"),
    ("occupation", "business_occupation"),
    ("residential_address", "residential_address"),
    ("address", "residential_address"),
)

_ID_TYPE_RE = re.compile(r"identification[_ ]?type|id[_ ]?type", re.IGNORECASE)


def _person_attribute(placeholder: str) -> tuple[str, str] | None:
    """Split 'resigning_director_identification_number' into role + column.

    A document names a person once and then refers to their NRC, father's name
    or address in the very next clause. Those companion placeholders share the
    role prefix, so once the person is known the rest follow — asking for them
    separately is how a name and an NRC end up belonging to different people.
    """
    p = placeholder.lower().strip()
    for tail, column in _PERSON_ATTRS:
        if p.endswith(("_" + tail, " " + tail)):
            role = p[: -(len(tail) + 1)].strip(" _")
            return (role, column) if role else None
    return None


def _lookup_person(name: str) -> dict:
    """Fetch one People-register row by name, for companion attributes."""
    if not name:
        return {}
    conn = None
    try:
        from db.connection import get_db_conn

        conn = get_db_conn()
        cur = conn.cursor()
        cur.execute(
            "SELECT nrc_passport_no, father_name, nationality, date_of_birth, "
            "business_occupation, residential_address FROM people "
            "WHERE upper(regexp_replace(trim(full_name), '\\s+', ' ', 'g')) = "
            "upper(regexp_replace(trim(%s), '\\s+', ' ', 'g')) LIMIT 1",
            (name,),
        )
        row = cur.fetchone()
        cur.close()
        if not row:
            return {}
        return {
            "nrc_passport_no": row[0],
            "father_name": row[1],
            "nationality": row[2],
            "date_of_birth": row[3].isoformat() if row[3] else None,
            "business_occupation": row[4],
            "residential_address": row[5],
        }
    except Exception as e:
        _log.warning(f"fill_view: person lookup failed for {name!r}: {e}")
        return {}
    finally:
        if conn:
            conn.close()


_RR_TOKEN_RE = re.compile(r"^__rr_(\d+)__$")


def _blank_label(key: str) -> str:
    """Human label for a blank.

    expand_repeat_regions rewrites fixed party slots into the real party count
    and names the new ones `__rr_1__`, `__rr_2__`… Those are plumbing, not
    something to show a lawyer — titlecasing one gives "Rr 1". A filled blank
    shows its value rather than its label, so this only surfaces in the picker
    header and the tooltip, but it should still read as English.
    """
    m = _RR_TOKEN_RE.match(key)
    if m:
        return f"Party {m.group(1)}"
    return key.replace("_", " ").strip().title()


def _emit_paragraph_text(
    text: str,
    blocks: list,
    blanks: dict,
    normalized_data,
    template_name,
    company_name,
    company_row,
    counter,
    people_by_role=None,
):
    last = 0
    for m in PLACEHOLDER_PATTERN.finditer(text):
        if m.start() > last:
            blocks.append({"type": "text", "text": text[last : m.start()]})
        name = placeholder_name((m.group(1), m.group(2), m.group(3)), counter)
        if not name or is_empty_placeholder(name):
            blocks.append({"type": "text", "text": m.group(0)})
            last = m.end()
            continue
        key = name.strip()
        value = _resolve_value(key, normalized_data, template_name, company_name)
        kind = _kind(key)
        # A bare placeholder mask (DD/MM/YYYY) is not a real value (finding F6);
        # a corporate entity auto-filled into an `individual` person slot (or vice
        # versa) is the wrong party type (finding F7) — treat both as unfilled so
        # the block stays outstanding and the user picks a correct value.
        if value is not None and kind == "date" and _is_mask(value):
            value = None
        if value is not None and kind == "person":
            ptype = _slot_party_type(key)
            if (ptype == "individual" and _looks_corporate(value)) or (
                ptype == "corporate" and not _looks_corporate(value)
            ):
                value = None
        # Remember who filled a person slot, so this person's NRC / father's
        # name / address can follow them into the companion placeholders.
        if people_by_role is not None and kind == "person" and value:
            role = key.lower().rsplit("_name", 1)[0].strip(" _")
            people_by_role.setdefault(role, value)

        # Companion attribute: derive it from the person already chosen for the
        # same role rather than leaving a blank the user must match by hand.
        if value is None and people_by_role is not None:
            attr = _person_attribute(key)
            if attr:
                role, column = attr
                who = people_by_role.get(role)
                if who:
                    value = _lookup_person(who).get(column) or None
            elif _ID_TYPE_RE.search(key):
                # "Identification Type" is answered by the identifier itself:
                # a Myanmar NRC is formatted 12/LATHANA(N)001520.
                for who in people_by_role.values():
                    ident = _lookup_person(who).get("nrc_passport_no")
                    if ident:
                        value = "N.R.C" if "/" in ident else "Passport"
                        break

        # One key, one answer. A template names `new_director_name` in the
        # proposal clause and again in the resolution clause; both must read the
        # same person, and the user must only choose once. Re-emitting the block
        # already held for this key keeps every occurrence in place while tying
        # them to a single value — including a value only the later occurrence
        # could resolve, which backfills the earlier one in place.
        prev = blanks.get(key)
        if prev is not None:
            if prev["value"] is None and value is not None:
                prev["value"] = value
                prev["filled"] = True
            blocks.append(prev)
            last = m.end()
            continue

        blk = {
            "type": "blank",
            "key": key,
            "label": _blank_label(key),
            "value": value,
            "filled": value is not None,
            "kind": kind,
            "candidates": _candidates(kind, key, company_row),
        }
        blocks.append(blk)
        blanks[key] = blk
        last = m.end()
    if last < len(text):
        blocks.append({"type": "text", "text": text[last:]})


_ALIGN_NAMES = {0: "left", 1: "center", 2: "right", 3: "justify"}


def _paragraph_style(para) -> dict:
    """Capture the formatting the panel needs to look like the Word document.

    Only what changes layout is carried across: alignment, heading level, bold /
    italic / underline where the WHOLE paragraph shares it, indent depth and
    list marker. Run-level mixed formatting is deliberately not modelled — the
    blanks are interactive controls, and chasing per-run spans would fight them
    for very little gain.
    """
    text = (para.text or "").strip()
    style_name = ""
    try:
        style_name = (para.style.name or "") if para.style is not None else ""
    except Exception:
        style_name = ""

    align = None
    try:
        if para.alignment is not None:
            align = _ALIGN_NAMES.get(int(para.alignment))
    except Exception:
        align = None

    runs = [r for r in para.runs if (r.text or "").strip()]
    all_bold = bool(runs) and all(bool(r.bold) for r in runs)
    all_italic = bool(runs) and all(bool(r.italic) for r in runs)
    all_underline = bool(runs) and all(bool(r.underline) for r in runs)

    lower = style_name.lower()
    heading = 0
    if lower.startswith("heading"):
        tail = lower.replace("heading", "").strip()
        heading = int(tail) if tail.isdigit() else 1
    elif lower in ("title", "subtitle"):
        heading = 1

    indent = 0
    try:
        left = para.paragraph_format.left_indent
        if left is not None:
            # Word stores EMUs; ~457200 EMU is the usual half-inch tab stop.
            indent = max(0, min(4, int(left / 457200)))
    except Exception:
        indent = 0

    return {
        "align": align,
        "bold": all_bold,
        "italic": all_italic,
        "underline": all_underline,
        "heading": heading,
        "list": "bullet" if "list bullet" in lower else ("number" if "list number" in lower else None),
        "indent": indent,
        "empty": not text,
    }


def build_fill_view(template_name: str, company_name: str, documents_dir: str = "/documents") -> dict[str, Any]:
    """Render `template_name` for `company_name` into fill-view blocks."""
    templates_dir = Path(documents_dir) / "legal" / "templates"
    template_path = templates_dir / template_name
    if not template_path.exists():
        alt = templates_dir / template_name.replace("_", " ")
        template_path = alt if alt.exists() else template_path
    if not template_path.exists():
        return {"success": False, "error": f"Template not found: {template_name}"}

    # Resolve company data the same way generation does.
    normalized_data: dict = {}
    company_row: dict = {}
    try:
        from scout.tools.smart_doc import prepare_document_data

        prep = prepare_document_data(template_name, company_name, documents_dir)
        normalized_data = prep.get("normalized_data", {}) or {}
    except Exception as e:
        _log.warning(f"fill_view: prepare failed: {e}")
    try:
        from scout.tools.companies_db import get_all_companies

        tgt = (company_name or "").strip().lower()
        for c in get_all_companies(limit=300):
            nm = (c.get("name") or "").strip().lower()
            if nm == tgt or (tgt and tgt in nm):
                company_row = c
                break
    except Exception as e:
        _log.warning(f"fill_view: company lookup failed: {e}")

    doc = Document(str(template_path))

    # Same dynamic list expansion generation runs, for the same reason: the
    # template hard-codes fixed party slots (individual shareholder_1..2,
    # corporate shareholder_3) and the real company rarely has exactly those.
    # Without this the preview offers a `corporate shareholder_3_name` blank
    # with zero candidates for a company that has no corporate members, and
    # disagrees with the document that is finally generated.
    try:
        from scout.tools.repeat_regions import expand_repeat_regions

        synth = expand_repeat_regions(doc, normalized_data, template_name=template_name, company_name=company_name)
        if synth:
            # The expander leaves synthetic [__rr_N__] tokens behind; merged
            # into the data they resolve through find_replacement like any
            # other slot, so the expanded lines read as real names.
            normalized_data = {**normalized_data, **synth}
    except Exception as e:
        _log.warning(f"fill_view: repeat-region expansion skipped: {e}")

    counter = new_empty_counter()
    blocks: list = []
    blanks: dict = {}
    # role prefix -> the person filling it, e.g. {"resigning_director": "WIN WIN TINT"}.
    # Populated as the document is walked, so a later NRC placeholder for the
    # same role resolves to that same person instead of standing empty.
    people_by_role: dict[str, str] = {}

    # Paragraphs are emitted as structured units rather than a flat text stream,
    # so the panel can lay the document out the way Word does — centred titles,
    # bold headings, indented clauses — instead of one undifferentiated wall of
    # text with chips in it. A lawyer proof-reads shape as much as wording.
    for para in doc.paragraphs:
        style = _paragraph_style(para)
        blocks.append({"type": "para_start", **style})
        _emit_paragraph_text(
            para.text,
            blocks,
            blanks,
            normalized_data,
            template_name,
            company_name,
            company_row,
            counter,
            people_by_role,
        )
        blocks.append({"type": "para_end"})

    for table in doc.tables:
        blocks.append({"type": "table_start", "columns": len(table.columns)})
        for row in table.rows:
            blocks.append({"type": "row_start"})
            for cell in row.cells:
                blocks.append({"type": "cell_start"})
                # A cell can hold several paragraphs; keep them separate so a
                # signature block does not collapse onto one line.
                for cpara in cell.paragraphs:
                    style = _paragraph_style(cpara)
                    blocks.append({"type": "para_start", **style})
                    _emit_paragraph_text(
                        cpara.text,
                        blocks,
                        blanks,
                        normalized_data,
                        template_name,
                        company_name,
                        company_row,
                        counter,
                        people_by_role,
                    )
                    blocks.append({"type": "para_end"})
                blocks.append({"type": "cell_end"})
            blocks.append({"type": "row_end"})
        blocks.append({"type": "table_end"})

    blank_list = list(blanks.values())
    return {
        "success": True,
        "template_name": template_name,
        "company_name": company_name,
        "blocks": blocks,
        "blanks": blank_list,
        "total_blanks": len(blank_list),
        "outstanding": sum(1 for b in blank_list if not b["filled"]),
    }
