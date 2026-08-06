"""
Repeat Regions — dynamic (unbounded) list expansion for legal templates
=======================================================================

The client's templates hard-code a fixed number of list slots:

    Present:
      [individual shareholder_1_name] (Shareholder)
      [individual shareholder_2_name] (Shareholder)
      [corporate shareholder_3_name] (Shareholder)

and fixed signing-table rows (one row-group per signatory). Companies in the
group have a variable number of members, attendees and signatories, so these
blocks must grow or shrink to fit the real party list.

This module detects those repeating regions in a `.docx` and rewrites them to
the actual number of parties BEFORE the normal placeholder fill runs. It never
touches single, scattered references (e.g. the Chairperson line, which names
one shareholder) — only contiguous list blocks and signing-table row groups.

Design
------
* Detection is conservative: a paragraph/row is a "list unit" only when its
  placeholders all belong to one party family (member / appointed-director /
  signing-director) and the remaining text is just a short role tag.
* Expansion clones the last unit of a block (deep XML copy, formatting kept)
  and re-points its placeholder to a fresh synthetic key; contraction deletes
  surplus units. Values are returned as a `{synthetic_key: value}` dict that the
  caller merges into `data`, so the normal highlighter fills and colours them
  exactly like any other slot.
* Any failure is swallowed and leaves the document untouched — a template that
  does not match these patterns generates exactly as before.
"""

from __future__ import annotations

import copy
import logging
import re
from typing import Any, Callable

from docx.oxml.ns import qn

from scout.tools.placeholders import PLACEHOLDER_PATTERN

_log = logging.getLogger("legalscout")

# Families of list placeholders we know how to expand.
_MEMBER_RE = re.compile(r"(individual|corporate)?\s*shareholder|member", re.I)
_APPOINTED_RE = re.compile(r"appointed[_ ]?director|new[_ ]?director", re.I)
_SIGN_DIR_RE = re.compile(r"^director$|^director[_ ]", re.I)
_ROLE_TAG_RE = re.compile(r"^[\s ()/,.\-–—:]*$")  # what may remain around a slot
_TO_SIGN_RE = re.compile(r"to\s*sign|signature|signed", re.I)
_CORP_SIGN_RE = re.compile(
    # "Represented by its authorized director" is the wording the firm's own
    # signature blocks use; without it a corporate signatory group classified as
    # individual and got the plain block with no representative line.
    r"authoris?ed\s+representative|authoriz?ed\s+representative|on\s+behalf"
    r"|its\s+representative|represented\s+by",
    re.I,
)
_REP_SLOT_RE = re.compile(r"authoris?ed[_ ]?director|authoriz?ed[_ ]?director|representative", re.I)


def _clean(text: str) -> str:
    return (text or "").replace(" ", " ").strip()


def _placeholders(text: str) -> list[str]:
    out = []
    for m in PLACEHOLDER_PATTERN.finditer(text or ""):
        name = (m.group(1) or m.group(2) or m.group(3) or "").strip()
        if name:
            out.append(name)
    return out


def _family(placeholder: str) -> str | None:
    """Classify a placeholder into a party family, or None if not a list slot."""
    p = _clean(placeholder).lower()
    if _APPOINTED_RE.search(p):
        return "appointed_director"
    if _MEMBER_RE.search(p):
        return "member"
    if _SIGN_DIR_RE.search(p):
        return "signing_director"
    return None


def _tail_attr(placeholder: str) -> str:
    p = _clean(placeholder).lower()
    if "nrc" in p or "passport" in p:
        return "nrc"
    if "name" in p:  # check before "share" — "shareholder_1_name" ends in name
        return "name"
    if "percent" in p or "%" in p:
        return "percentage"
    # a bare share count, but NOT the word "shareholder"
    if re.search(r"\bshares?\b", p) and "shareholder" not in p and "share holder" not in p:
        return "shares"
    return "name"


def _unit_family(text: str) -> str | None:
    """A paragraph/cell is a list unit only when EVERY placeholder shares one
    family and the surrounding text is just a short role tag."""
    phs = _placeholders(text)
    if not phs:
        return None
    fams = {_family(p) for p in phs}
    if len(fams) != 1 or None in fams:
        return None
    fam = fams.pop()
    # strip out placeholders, whatever remains must be a short role tag
    residue = PLACEHOLDER_PATTERN.sub("", text or "")
    residue = _clean(residue)
    # allow "(Shareholder)", "(Director)", "Signed by its authorized representative"
    if len(residue) > 60:
        return None
    return fam


# --------------------------------------------------------------------------- #
# Rendering                                                                    #
# --------------------------------------------------------------------------- #

def _member_display(member: dict) -> str:
    return _clean(member.get("name") or member.get("full_name") or member.get("value") or "")


def _is_corporate(member: dict) -> bool:
    t = str(member.get("type") or member.get("party_type") or "").lower().strip()
    if t:
        if "corp" in t or "company" in t or "entity" in t or t == "c":
            return True
        if "indiv" in t or "person" in t or t == "i":
            return False
        # unknown label -> fall through to the name heuristic
    name = _member_display(member).lower()
    return any(w in name for w in ("company", "limited", "ltd", "holdings", "group", "co.,", "pcl", "plc"))


def _render_value(item: Any, tail: str) -> str:
    """Render a party item for the given tail attribute."""
    if isinstance(item, dict):
        if tail == "nrc":
            return _clean(item.get("nrc") or item.get("nrc_passport_no") or item.get("identifier") or item.get("name") or "")
        if tail in ("shares", "percentage"):
            return _clean(str(item.get(tail) or item.get("shares") or ""))
        return _member_display(item)
    return _clean(str(item))


# --------------------------------------------------------------------------- #
# Party-list resolution                                                        #
# --------------------------------------------------------------------------- #

def _members_of(company_name: str | None) -> list[dict]:
    if not company_name:
        return []
    try:
        from scout.tools.companies_db import get_all_companies
        target = _clean(company_name).lower()
        for c in get_all_companies(limit=300):
            if _clean(c.get("name")).lower() == target or target in _clean(c.get("name")).lower():
                mems = c.get("shareholders_list") or []
                return [m for m in mems if isinstance(m, dict) and _member_display(m)]
    except Exception as e:  # noqa: BLE001
        _log.warning(f"repeat_regions: members lookup failed for {company_name!r}: {e}")
    return []


def _board_of(company_name: str) -> list[str]:
    """Current director names of `company_name`, cleaned for comparison.

    Empty when the company is not in the register — which is NOT the same as
    "has no directors", so callers must not treat empty as grounds to reject."""
    try:
        from scout.tools.slot_resolver import corporate_shareholder_directors
        out = []
        for c in corporate_shareholder_directors(company_name) or []:
            name = (c.get("name") or c.get("full_name") or "") if isinstance(c, dict) else str(c)
            name = _clean(name)
            if name:
                out.append(name)
        return out
    except Exception as e:  # noqa: BLE001
        _log.warning(f"repeat_regions: board lookup failed for {company_name!r}: {e}")
        return []


def _on_board(name: str, board: list[str]) -> bool:
    """Case- and spacing-insensitive membership test."""
    n = " ".join(_clean(name).casefold().split())
    return any(n == " ".join(_clean(b).casefold().split()) for b in board)


def _corp_representative(party: dict, data: dict) -> str:
    """Resolve the authorised representative for a CORPORATE signatory.

    Priority: the rep carried on the party itself → a representative the user
    picked for the corporate shareholder (carried in data) → the corporate
    member's own first director from the register. The DB-fallback member path
    supplies no representative, which left the "signed by its authorized
    representative" row with a BLANK name (finding F4).

    ★ Every candidate is checked against THAT COMPANY'S OWN CURRENT BOARD before
    it is printed. A corporate member signs through one of its own directors, so
    a name that is not on its board has no authority to bind it. Two ways this
    went wrong in practice, both measured:
      * a director of the SUBJECT company signed for the member company — the
        register lists KYAW THU SOE on CITY MART's board and not on CITY
        HOLDINGS', yet he was printed as CITY HOLDINGS' authorised director;
      * a RESIGNED director stayed on the line, because a stale answer carried
        in `data` outranked the register that had already ended them.
    The board list also excludes resigned people, so this guard closes both."""
    board = _board_of(_member_display(party) if isinstance(party, dict) else str(party))

    def _vet(name: str, source: str) -> str:
        """Return `name` only if the register does not contradict it."""
        if not name:
            return ""
        if not board:
            # The member company is not in the register, so there is nothing to
            # check against. Refusing here would blank the line for every
            # unregistered corporate member; fail open, but say so.
            _log.warning(
                "repeat_regions: cannot verify representative %r (%s) — %r is not in "
                "the register, so its board is unknown",
                name, source, _member_display(party) if isinstance(party, dict) else party)
            return name
        if _on_board(name, board):
            return name
        _log.warning(
            "repeat_regions: DROPPED representative %r (%s) — not a current director of %r "
            "(board: %s). A name off the board cannot bind the company.",
            name, source, _member_display(party) if isinstance(party, dict) else party, board)
        return ""

    if isinstance(party, dict):
        rep = _clean(party.get("representative") or party.get("representative_name") or "")
        if rep:
            vetted = _vet(rep, "carried on the party")
            if vetted:
                return vetted
    data = data or {}
    # `corporate_shareholder_3_name` is deliberately NOT consulted: that is the
    # corporate MEMBER's own name, and a company cannot be its own authorised
    # director. A blank rep line is a visible gap; the member's name sitting on
    # it reads as a real signatory.
    #
    # ★ The SAME slot arrives under two spellings, and either one can be the
    # real answer:
    #   `authorized_director_name`  — smart_doc's per-company default, seeded
    #                                 from directors[0]; ALSO where some picker
    #                                 answers land.
    #   `authorized director_name`  — the raw placeholder key, which is where
    #                                 this signatory picker's answer lands.
    # Observed on one real run: space = 'PHYOE MIN KYAW' (chosen), underscore =
    # 'KYAW THU SOE' (directors[0]). Reading either spelling first is wrong half
    # the time, so pick on CONTENT: a candidate equal to `director_name` is the
    # company default that was never answered, and a candidate that differs from
    # it is somebody's actual choice.
    #
    # This is a heuristic standing in for a real fix — the two spellings should
    # be normalised to one key where custom_data is assembled, at which point
    # this collapses back to a plain ordered lookup.
    default_first_director = _clean(str(data.get("director_name") or ""))
    candidates = []
    for k in ("representative", "representative_name", "corporate_representative",
              "corporate_shareholder_representative", "authorized_director",
              "authorised_director", "authorized director_name",
              "authorised director_name", "authorized_director_name",
              "authorised_director_name"):
        v = data.get(k)
        if v and str(v).strip() and str(v).strip().upper() != "TBD":
            candidates.append(_clean(str(v)))
    for v in candidates:
        if not default_first_director or v != default_first_director:
            vetted = _vet(v, "chosen / carried in data")
            if vetted:
                return vetted
    # Every candidate matched the default. It may genuinely be the right person;
    # what it is NOT is grounds to go guess a different one from register order.
    for v in candidates:
        vetted = _vet(v, "matches the per-company default")
        if vetted:
            return vetted
    # Last resort: the member company's own first director. This is a POSITIONAL
    # GUESS — register order, not anybody's choice — and slot_resolver was fixed
    # to stop counting it as an answered question, so reaching it means nobody
    # was asked. It is at least somebody with authority to bind the company,
    # which a name off the board is not. `board` is already vetted and excludes
    # resigned directors, so no second lookup is needed.
    if board:
        _log.warning(
            "repeat_regions: falling back to register order for %r's representative (%r) — "
            "nobody chose one",
            _member_display(party) if isinstance(party, dict) else party, board[0])
        return board[0]
    return ""


def _list_from_data(data: dict, keys: list[str]) -> list[dict]:
    """Pull an explicit list a picker/agent supplied under any of `keys`."""
    for k in keys:
        v = data.get(k)
        if isinstance(v, list) and v:
            out = []
            for it in v:
                if isinstance(it, dict) and _member_display(it):
                    out.append(it)
                elif isinstance(it, str) and _clean(it):
                    out.append({"name": _clean(it)})
            if out:
                return out
    return []


def _parties_for_family(family: str, tail: str, data: dict, company_name: str | None) -> list[dict]:
    """The ordered party list a region should expand to."""
    if family == "member":
        parties = _list_from_data(data, ["members", "attendees", "shareholders_selected", "present_members"])
        return parties or _members_of(company_name)
    if family == "appointed_director":
        return _list_from_data(
            data, ["appointed_directors", "new_directors", "appointed_director_list", "directors_appointed"]
        )
    if family == "signing_director":
        return _list_from_data(
            data, ["signing_directors", "signatories", "representative_directors", "directors_to_sign"]
        )
    return []


# --------------------------------------------------------------------------- #
# Low-level docx helpers                                                       #
# --------------------------------------------------------------------------- #

def _set_unit_placeholder(element_text_setter: Callable[[str], None], token: str) -> None:
    element_text_setter(token)


def _para_set_text(paragraph, text: str) -> None:
    """Replace a paragraph's text with `text`, keeping the first run's format."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _para_render_template(orig_text: str, token: str) -> str:
    """Keep the unit's static role tag, swap the (first) placeholder for token."""
    # Replace the FIRST placeholder occurrence with token, drop any others.
    replaced = {"done": False}

    def _sub(m):
        if replaced["done"]:
            return ""
        replaced["done"] = True
        return token

    return PLACEHOLDER_PATTERN.sub(_sub, orig_text)


# --------------------------------------------------------------------------- #
# Paragraph-block expansion                                                    #
# --------------------------------------------------------------------------- #

def _expand_paragraph_blocks(doc, data, company_name, synth, counter):
    paras = doc.paragraphs
    i = 0
    n = len(paras)
    while i < n:
        fam = _unit_family(paras[i].text)
        if not fam:
            i += 1
            continue
        # maximal run of consecutive units of the SAME family
        j = i
        while j < n and _unit_family(paras[j].text) == fam:
            j += 1
        block = paras[i:j]
        _rewrite_paragraph_block(block, fam, data, company_name, synth, counter)
        # The rewrite only ever adds paragraphs directly after block[-1] or removes
        # paragraphs from inside the block, so every index change happens at or
        # before j. Re-anchoring by the size delta therefore lands exactly on the
        # first paragraph the block did not cover — correct whether it grew
        # (delta > 0) or shrank (delta < 0, where the old `i = j` skipped as many
        # unscanned paragraphs as were deleted).
        paras = doc.paragraphs
        delta = len(paras) - n
        n = len(paras)
        # A block always keeps at least one unit (an empty party list returns
        # early), so j + delta > i; max() only guards the loop against a future
        # rewrite that could delete a whole block.
        i = max(i + 1, j + delta)
    return synth


def _rewrite_paragraph_block(block, family, data, company_name, synth, counter):
    if not block:
        return
    tail = _tail_attr(_placeholders(block[0].text)[0])
    parties = _parties_for_family(family, tail, data, company_name)
    if not parties:
        return  # nothing to expand to — leave template as-is
    template_text = block[0].text  # e.g. "[slot] (Shareholder)"
    last = block[-1]

    # 1) fill existing units in place
    for idx, unit in enumerate(block):
        if idx < len(parties):
            key = f"__rr_{next(counter)}__"
            synth[key] = _render_value(parties[idx], tail)
            _para_set_text(unit, _para_render_template(template_text, f"[{key}]"))

    # 2) overflow parties -> clone the last unit's XML after it
    if len(parties) > len(block):
        anchor = last._p
        from docx.text.paragraph import Paragraph
        for idx in range(len(block), len(parties)):
            key = f"__rr_{next(counter)}__"
            synth[key] = _render_value(parties[idx], tail)
            new_p = copy.deepcopy(block[-1]._p)
            anchor.addnext(new_p)
            anchor = new_p
            _para_set_text(Paragraph(new_p, block[-1]._parent), _para_render_template(template_text, f"[{key}]"))

    # 3) surplus template slots -> remove the paragraph entirely (no blank lines)
    if len(parties) < len(block):
        for unit in block[len(parties):]:
            unit._p.getparent().remove(unit._p)


# --------------------------------------------------------------------------- #
# Signing-table expansion                                                      #
# --------------------------------------------------------------------------- #

def _row_family(row) -> str | None:
    for cell in row.cells:
        fam = _unit_family(cell.text)
        if fam:
            return fam
    return None


def _expand_signing_tables(doc, data, company_name, synth, counter):
    for table in doc.tables:
        try:
            _expand_one_table(table, data, company_name, synth, counter)
        except Exception as e:  # noqa: BLE001
            _log.warning(f"repeat_regions: table expansion skipped: {e}")


def _row_group_is_corporate(trs, table) -> bool:
    from docx.table import _Row
    for tr in trs:
        row = _Row(tr, table)
        for cell in row.cells:
            if _CORP_SIGN_RE.search(cell.text or ""):
                return True
    return False


def _prev_paragraph_text(tbl, limit: int = 3) -> str:
    """Text of the paragraphs immediately preceding a table.

    Stops at the previous table so a cue never leaks in from further up."""
    from docx.text.paragraph import Paragraph

    out, el, seen = [], tbl.getprevious(), 0
    while el is not None and seen < limit:
        if el.tag == qn("w:tbl"):
            break
        if el.tag == qn("w:p"):
            t = _clean(Paragraph(el, None).text)
            if t:
                out.append(t)
                seen += 1
        el = el.getprevious()
    return " ".join(out)


def _signing_rows_to_scan(table) -> range | None:
    """Which rows may hold signatory units, or None if this is not a signature
    table.

    Two shapes exist in the firm's templates:
      * a header row that says "…to sign…", units below it;
      * NO header row — the "Members to sign if they agree…" line is an ordinary
        PARAGRAPH above the table and row 0 is already a signatory unit.

    Only the first was recognised, so the second was skipped whole: both slots
    fell through to the flat per-company fill and a company with ONE corporate
    member was rendered signing twice — once via its representative, once again
    on the individual line."""
    rows = table.rows
    if not rows:
        return None
    header = " ".join(c.text for c in rows[0].cells)
    if _TO_SIGN_RE.search(header or ""):
        return range(1, len(rows))
    # Headerless shape. Require BOTH that row 0 is itself a unit and that a
    # nearby preceding paragraph carries the signing cue, so this cannot claim
    # an ordinary data table that happens to mention a shareholder.
    if _row_family(rows[0]) and _TO_SIGN_RE.search(_prev_paragraph_text(table._tbl)):
        return range(0, len(rows))
    return None


def _expand_one_table(table, data, company_name, synth, counter):
    rows = table.rows
    scan = _signing_rows_to_scan(table)
    if scan is None:
        return  # only touch signature tables

    # anchor rows: rows whose first data cell is a list unit
    anchors = []  # (row_index, family)
    for ri in scan:
        fam = _row_family(rows[ri])
        if fam:
            anchors.append((ri, fam))
    if not anchors:
        return

    family = anchors[0][1]
    parties = _parties_for_family(family, "name", data, company_name)
    if not parties:
        return

    # A signatory "unit" = anchor row through the row before the next anchor
    # (captures the trailing "Date:" / spacer rows). Capture each unit's rows as
    # deep-copied XML templates, classified individual vs corporate.
    units = []  # {"trs": [tr,...], "corporate": bool}
    for k, (ri, _fam) in enumerate(anchors):
        end = anchors[k + 1][0] if k + 1 < len(anchors) else len(rows)
        trs = [rows[r]._tr for r in range(ri, end)]
        units.append({"trs": trs, "corporate": _row_group_is_corporate(trs, table)})

    indiv_tpl = next((u["trs"] for u in units if not u["corporate"]), units[0]["trs"])
    corp_tpl = next((u["trs"] for u in units if u["corporate"]), None)

    # insert cloned units right after the row preceding the first unit
    # (normally the header row); fall back to the first unit row itself.
    first_tr = units[0]["trs"][0]
    insert_after = first_tr.getprevious()
    if insert_after is None:
        insert_after = first_tr

    # remove all original unit rows
    for u in units:
        for tr in u["trs"]:
            parent = tr.getparent()
            if parent is not None:
                parent.remove(tr)

    from docx.table import _Row
    for party in parties:
        corporate = _is_corporate(party)
        tpl = corp_tpl if (corporate and corp_tpl is not None) else indiv_tpl
        cloned = [copy.deepcopy(tr) for tr in tpl]
        for tr in cloned:
            insert_after.addnext(tr)
            insert_after = tr
        # fill the signatory name into the anchor cell
        key = f"__rr_{next(counter)}__"
        synth[key] = _render_value(party, "name")
        new_row = _Row(cloned[0], table)
        _rewrite_cell_placeholder(_first_unit_cell(new_row), f"[{key}]")
        # corporate: also fill the representative-director placeholder if present
        if corporate:
            rep = _corp_representative(party if isinstance(party, dict) else {"name": str(party)}, data)
            if rep:
                rkey = f"__rr_{next(counter)}__"
                synth[rkey] = rep
                for tr in cloned:
                    _fill_rep_placeholder(_Row(tr, table), f"[{rkey}]")


def _first_unit_cell(row):
    for cell in row.cells:
        if _unit_family(cell.text):
            return cell
    return row.cells[0]


def _fill_rep_placeholder(row, token: str):
    """In a corporate signatory row, point the representative-director slot
    (e.g. [authorized director_name]) at the chosen representative."""
    for cell in row.cells:
        for para in cell.paragraphs:
            phs = _placeholders(para.text)
            if phs and any(_REP_SLOT_RE.search(_clean(p)) for p in phs):
                _para_set_text(para, _para_render_template(para.text, token))
                return True
    return False


def _rewrite_cell_placeholder(cell, token: str):
    """Swap the first placeholder in a cell for token, keeping surrounding text
    (leading blank lines etc.). Operates paragraph by paragraph."""
    done = False
    for para in cell.paragraphs:
        if done:
            # clear any further placeholder paragraphs in this cell
            if _placeholders(para.text):
                _para_set_text(para, PLACEHOLDER_PATTERN.sub("", para.text))
            continue
        if _placeholders(para.text):
            _para_set_text(para, _para_render_template(para.text, token))
            done = True


# --------------------------------------------------------------------------- #
# Entry point                                                                  #
# --------------------------------------------------------------------------- #

def expand_repeat_regions(doc, data: dict, template_name: str = None, company_name: str = None) -> dict:
    """Expand list regions in `doc` to match the real party count.

    Returns a dict of synthetic {placeholder_key: value} to merge into `data`
    before the normal fill. Mutates `doc` in place. Never raises.
    """
    synth: dict[str, str] = {}
    try:
        counter = iter(range(1, 100000))
        _expand_paragraph_blocks(doc, data, company_name, synth, counter)
        _expand_signing_tables(doc, data, company_name, synth, counter)
    except Exception as e:  # noqa: BLE001
        _log.warning(f"repeat_regions: expansion aborted for {template_name!r}: {e}")
    return synth
